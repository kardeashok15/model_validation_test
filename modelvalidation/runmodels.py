from multiprocessing import Process, Pipe
import sys
from inspect import Traceback
import re
from django.http import HttpResponse, Http404
import smtplib
from sklearn.neural_network import MLPClassifier
from sklearn.ensemble import VotingClassifier, AdaBoostClassifier, RandomForestClassifier, GradientBoostingClassifier, BaggingClassifier
from sklearn.neighbors import KNeighborsClassifier, KNeighborsRegressor
from sklearn import svm
from sklearn.linear_model import LogisticRegression
from statsmodels.stats.outliers_influence import variance_inflation_factor
from typing import Reversible
from sklearn import preprocessing
from sklearn.feature_selection import SelectPercentile, chi2, RFE
from sklearn.preprocessing import StandardScaler, label_binarize
from sklearn.model_selection import train_test_split, RandomizedSearchCV, GridSearchCV, StratifiedKFold, cross_val_score
from sklearn.metrics import classification_report, roc_curve, auc, roc_auc_score, confusion_matrix, recall_score, precision_score, accuracy_score
from bubbly.bubbly import bubbleplot
import markupsafe
import plotly_express as px
import joypy
from django.core.files.storage import FileSystemStorage
from django.conf import settings
import math
from io import StringIO
from statsmodels import robust
import plotly.graph_objs as go
from plotly.offline import plot
import plotly.offline as py
from pandas.plotting import parallel_coordinates
from pandas import plotting
import matplotlib.pyplot as plt
from django.shortcuts import redirect, render
from django.http import JsonResponse
from flask import Markup
import pandas as pd
import numpy as np
import os
from pathlib import Path
import json
# for visualizations
import seaborn as sns
import matplotlib
import xgboost as xgb
from outliers import smirnov_grubbs as grubbs
from scipy.stats import ks_2samp
from scipy import stats
from scipy.stats import randint as sp_randint

# generate random floating point values
from numpy.random import seed
from numpy.random import rand
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
import xlsxwriter
from datetime import date
import traceback
matplotlib.use('Agg')

# for modeling


BASE_DIR = Path(__file__).resolve().parent.parent
# Create your views here.
user_name = "user1"
param_file_path = os.path.join(BASE_DIR, 'static\param_files\\')
file_path = os.path.join(BASE_DIR, 'static\csv_files\\')
param_file_name = "paramfile_"+user_name
file_name = "user1"
app_url = "http://3.131.88.246:8000/modelval/"
font_files = os.path.join(BASE_DIR, 'static\\fonts\\')

csv_file_name = "csvfile_"+user_name


class lstTestModelPerf:
    testName: str
    testResult: str
    testResult_dict: dict


def GBC():
    try:
        csv_file_name = "csvfile_"+user_name
        savefile_x_final = file_path + csv_file_name + "_x_model.csv"
        df = pd.read_csv(savefile_x_final)
        targetVarFile = file_path + csv_file_name + "_targetVar.txt"
        file1 = open(targetVarFile, "r")  # write mode
        targetVar = file1.read()
        file1.close()
        # a variable pdf
        pdf = FPDF()
        # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
        y_model = df[targetVar]
        x_model = df.drop(targetVar, axis=1)
        X_train, X_test, y_train, y_test = train_test_split(
            x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
        X_train, X_val, y_train, y_val = train_test_split(
            X_train, y_train, test_size=0.22222222222222224, random_state=321)

        gbc = GradientBoostingClassifier()
        gbc.fit(X_train, y_train)

        # Test the model
        pred_gbc_val = gbc.predict(X_val)
        pred_gbc_prob_val = gbc.predict_proba(X_val)[:, 1]

        pred_gbc_train = gbc.predict(X_train)
        pred_gbc_prob_train = gbc.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_gbc_train))
        print(classification_report(y_val, pred_gbc_val))

        evaluate_model(y_val, pred_gbc_prob_val, y_train,
                       pred_gbc_prob_train, file_name+"GBC_NT_roc1.png")

        evaluate_model(y_val, pred_gbc_val, y_train, pred_gbc_train,
                       file_name+"GBC_NT_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data
        drawConfMatrix(y_train, pred_gbc_train, file_name +
                       "_GBC_NT_train_data.png")

        drawConfMatrix(y_val, pred_gbc_val, file_name +
                       "_GBC_NT_val_data.png")

        NT_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_gbc_prob_val, pred_gbc_val)
        NT_roc_file1 = "\\static\\media\\" + file_name+"GBC_NT_roc1.png"
        NT_roc_file2 = "\\static\\media\\" + file_name+"GBC_NT_roc2.png"

        # Add a page
        pdf.add_page()
        document = Document()
        pdf = exportPdf(10, 10, pdf, document, os.path.join(
            BASE_DIR, 'static\media\\', file_name+"GBC_NT_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, 'static\media\\', file_name+"GBC_NT_roc2.png"),
            "ROC Curve Test data", "Gradient Boosting - No Parameters Tuning")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_GBC_NT_train_data.png"),
            "Confusion Matrix Validation data", os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_GBC_NT_val_data.png"),
            "Confusion Matrix Test data", "Gradient Boosting - No Parameters Tuning")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, NT_auc_ginin_output, "Gradient Boosting - No Parameters Tuning")

        paramFiles = param_file_path + param_file_name + "_GB_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        seed(1)
     # Estimator for use in random search
        estimator = GradientBoostingClassifier()

        # Create the random search model
        gbc_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                        scoring='roc_auc', cv=5,
                                        n_iter=100, verbose=1, random_state=50)

        # Fit
        gbc_search.fit(X_train, y_train)

        gbc_random = gbc_search.best_estimator_
        # Test the model
        pred_gbc_val0 = gbc_random.predict(X_val)
        pred_gbc_prob_val0 = gbc_random.predict_proba(X_val)[:, 1]

        pred_gbc_train0 = gbc_random.predict(X_train)
        pred_gbc_prob_train0 = gbc_random.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_gbc_train0))
        print(classification_report(y_val, pred_gbc_val0))

        evaluate_model(y_val, pred_gbc_prob_val0, y_train,
                       pred_gbc_prob_train0, file_name+"_GBC_RS_roc1.png")

        evaluate_model(y_val, pred_gbc_val0, y_train,
                       pred_gbc_train0, file_name+"_GBC_RS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_gbc_train0,
                       file_name + "_GBC_RS_train_data.png")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_gbc_val0,
                       file_name + "_GBC_RS_val_data.png")

        RS_roc_file1 = "\\static\\media\\" + file_name + "_GBC_RS_roc1.png"
        RS_roc_file2 = "\\static\\media\\" + file_name + "_GBC_RS_roc2.png"
        RS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_gbc_prob_val0, pred_gbc_val0)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_GBC_RS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_GBC_RS_roc2.png"),
            "ROC Curve Test data", "Gradient Boosting - Random Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_GBC_RS_train_data.png"),
            "Confusion Matrix Validation data", os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_GBC_RS_val_data.png"),
            "Confusion Matrix Test data", "Gradient Boosting - Random Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, RS_auc_ginin_output, "Gradient Boosting - Random Search")

        paramFiles = param_file_path + param_file_name + "_GB_GS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = GradientBoostingClassifier()

        # Create the random search model
        gbc_grid = GridSearchCV(estimator, param_grid, n_jobs=3,
                                scoring='roc_auc', cv=5, verbose=1)

        # Fit
        gbc_grid.fit(X_train, y_train)

        gbc_final = gbc_grid.best_estimator_

        # Test the model
        pred_gbc_val1 = gbc_final.predict(X_val)
        pred_gbc_prob_val1 = gbc_final.predict_proba(X_val)[:, 1]

        pred_gbc_train1 = gbc_final.predict(X_train)
        pred_gbc_prob_train1 = gbc_final.predict_proba(X_train)[:, 1]
        # Get the model performance
        print(classification_report(y_train, pred_gbc_train1))
        print(classification_report(y_val, pred_gbc_val1))

        evaluate_model(y_val, pred_gbc_prob_val1, y_train,
                       pred_gbc_prob_train1, file_name+"_GBC_GS_roc1.png")

        evaluate_model(y_val, pred_gbc_val1, y_train,
                       pred_gbc_train1, file_name+"_GBC_GS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_gbc_train1,
                       file_name + "_GBC_GS_train_data.png")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_gbc_val1,
                       file_name + "_GBC_GS_val_data.png")

        GS_roc_file1 = "\\static\\media\\" + file_name + "_GBC_GS_roc1.png"
        GS_roc_file2 = "\\static\\media\\" + file_name + "_GBC_GS_roc2.png"
        GS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_gbc_prob_val1, pred_gbc_val1)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_GBC_GS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_GBC_GS_roc2.png"),
            "ROC Curve Test data", "Gradient Boosting - Grid Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_GBC_GS_train_data.png"),
            "Confusion Matrix Validation data", os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_GBC_GS_val_data.png"),
            "Confusion Matrix Test data", "Gradient Boosting - Grid Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, GS_auc_ginin_output, "Gradient Boosting - Grid Search")

        pdf.output(os.path.join(
            BASE_DIR, "static\\media\\" + file_name + "_Gradient_Boosting.pdf"))

        context = {'pdfFile': "\\static\\media\\" + file_name + "_Gradient_Boosting.pdf", 'model': 'GBC',  'tableHead': 'Gradient Boosting', 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                   'NT_graphConfMat1': "\\static\\media\\" + file_name + "_GBC_NT_train_data.png", 'NT_graphConfMat2': "\\static\\media\\" + file_name + "_GBC_NT_val_data.png",
                   'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                   'RS_graphConfMat1': "\\static\\media\\" + file_name + "_GBC_RS_train_data.png", 'RS_graphConfMat2': "\\static\\media\\" + file_name + "_GBC_RS_val_data.png",
                   'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                   'GS_graphConfMat1': "\\static\\media\\" + file_name + "_GBC_GS_train_data.png", 'GS_graphConfMat2': "\\static\\media\\" + file_name + "_GBC_GS_val_data.png"}
        print('Bagging Cls done')
    except Exception as e:
        print(e)
        context = {'tableHead': 'Error while processing reuest.', }


def BC(conn):
    try:
        csv_file_name = "csvfile_"+user_name
        savefile_x_final = file_path + csv_file_name + "_x_model.csv"
        df = pd.read_csv(savefile_x_final)
        targetVarFile = file_path + csv_file_name + "_targetVar.txt"
        file1 = open(targetVarFile, "r")  # write mode
        targetVar = file1.read()
        file1.close()
        # a variable pdf
        pdf = FPDF()
        # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
        y_model = df[targetVar]
        x_model = df.drop(targetVar, axis=1)
        X_train, X_test, y_train, y_test = train_test_split(
            x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
        X_train, X_val, y_train, y_val = train_test_split(
            X_train, y_train, test_size=0.22222222222222224, random_state=321)

        bc_model = BaggingClassifier()
        bc_model.fit(X_train, y_train)

        # Test the model
        pred_bc_val = bc_model.predict(X_val)
        pred_bc_prob_val = bc_model.predict_proba(X_val)[:, 1]

        pred_bc_train = bc_model.predict(X_train)
        pred_bc_prob_train = bc_model.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_bc_train))
        print(classification_report(y_val, pred_bc_val))

        evaluate_model(y_val, pred_bc_prob_val, y_train,
                       pred_bc_prob_train, file_name+"BC_NT_roc1.png")

        evaluate_model(y_val, pred_bc_val, y_train,
                       pred_bc_train, file_name+"BC_NT_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data
        drawConfMatrix(y_train, pred_bc_train, file_name +
                       "_BC_NT_train_data.png")

        drawConfMatrix(y_val, pred_bc_val, file_name +
                       "_BC_NT_val_data.png")

        NT_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_bc_prob_val, pred_bc_val)

        # Add a page
        pdf.add_page()
        document = Document()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, 'static\media\\', file_name+"BC_NT_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, 'static\media\\', file_name+"BC_NT_roc2.png"),
            "ROC Curve Test data", "Bagging Classifier - No Parameters Tuning")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_BC_NT_train_data.png"),
            "Confusion Matrix Validation data", os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_BC_NT_val_data.png"),
            "Confusion Matrix Test data", "Bagging Classifier - No Parameters Tuning")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, NT_auc_ginin_output, "Bagging Classifier - No Parameters Tuning")

        NT_roc_file1 = "\\static\\media\\" + file_name+"BC_NT_roc1.png"
        NT_roc_file2 = "\\static\\media\\" + file_name+"BC_NT_roc2.png"

        paramFiles = param_file_path + param_file_name + "_BC_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = BaggingClassifier(random_state=50)

        # Create the random search model
        bc_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                       scoring='roc_auc', cv=5,
                                       n_iter=100, verbose=1, random_state=50)

        # Fit
        bc_search.fit(X_train, y_train)

        bc_random = bc_search.best_estimator_

        # Test the model
        pred_bc_val0 = bc_random.predict(X_val)
        pred_bc_prob_val0 = bc_random.predict_proba(X_val)[:, 1]

        pred_bc_train0 = bc_random.predict(X_train)
        pred_bc_prob_train0 = bc_random.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_bc_train0))
        print(classification_report(y_val, pred_bc_val0))

        evaluate_model(y_val, pred_bc_prob_val0, y_train,
                       pred_bc_prob_train0, file_name+"_BC_RS_roc1.png")

        evaluate_model(y_val, pred_bc_val0, y_train,
                       pred_bc_train0, file_name+"_BC_RS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_bc_train0,
                       file_name + "_BC_RS_train_data.png")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_bc_val0,
                       file_name + "_BC_RS_val_data.png")

        RS_roc_file1 = "\\static\\media\\" + file_name + "_BC_RS_roc1.png"
        RS_roc_file2 = "\\static\\media\\" + file_name + "_BC_RS_roc2.png"
        RS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_bc_prob_val0, pred_bc_val0)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_BC_RS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_BC_RS_roc2.png"),
            "ROC Curve Test data", "Bagging Classifier - Random Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_BC_RS_train_data.png"),
            "Confusion Matrix Validation data", os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_BC_RS_val_data.png"),
            "Confusion Matrix Test data", "Bagging Classifier - Random Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, RS_auc_ginin_output, "Bagging Classifier - Random Search")

        paramFiles = param_file_path + param_file_name + "_BC_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = BaggingClassifier(random_state=50)

        # Create the random search model
        bc_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                       scoring='roc_auc', cv=5,
                                       n_iter=100, verbose=1, random_state=50)

        # Fit
        bc_search.fit(X_train, y_train)

        paramFiles = param_file_path + param_file_name + "_BC_GS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = BaggingClassifier(random_state=50)

        # Create the random search model
        bc_grid = GridSearchCV(estimator, param_grid, n_jobs=3,
                               scoring='roc_auc', cv=5, verbose=1)

        # Fit
        bc_grid.fit(X_train, y_train)

        bc_final = bc_grid.best_estimator_

        pred_bc_val1 = bc_final.predict(X_val)
        pred_bc_prob_val1 = bc_final.predict_proba(X_val)[:, 1]

        pred_bc_train1 = bc_final.predict(X_train)
        pred_bc_prob_train1 = bc_final.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_bc_train1))
        print(classification_report(y_val, pred_bc_val1))

        evaluate_model(y_val, pred_bc_prob_val1, y_train,
                       pred_bc_prob_train1, file_name+"_BC_GS_roc1.png")

        evaluate_model(y_val, pred_bc_val1, y_train,
                       pred_bc_train1, file_name+"_BC_GS_roc2.png")
        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_bc_train1,
                       file_name + "_BC_GS_train_data.png")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_bc_val1,
                       file_name + "_BC_GS_val_data.png")

        GS_roc_file1 = "\\static\\media\\" + file_name + "_BC_GS_roc1.png"
        GS_roc_file2 = "\\static\\media\\" + file_name + "_BC_GS_roc2.png"
        GS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_bc_prob_val1, pred_bc_val1)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_BC_GS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_BC_GS_roc2.png"),
            "ROC Curve Test data", "Bagging Classifier - Grid Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_BC_GS_train_data.png"),
            "Confusion Matrix Validation data", os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_BC_GS_val_data.png"),
            "Confusion Matrix Test data", "Bagging Classifier - Grid Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, GS_auc_ginin_output, "Bagging Classifier - Grid Search")

        pdf.output(os.path.join(
            BASE_DIR, "static\\media\\" + file_name + "_BagCls.pdf"))

        context = {'pdfFile': "\\static\\media\\" + file_name + "_BagCls.pdf", 'model': 'BC',  'tableHead': 'Bagging Classifier', 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                   'NT_graphConfMat1': "\\static\\media\\" + file_name + "_BC_NT_train_data.png", 'NT_graphConfMat2': "\\static\\media\\" + file_name + "_BC_NT_val_data.png",
                   'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                   'RS_graphConfMat1': "\\static\\media\\" + file_name + "_BC_RS_train_data.png", 'RS_graphConfMat2': "\\static\\media\\" + file_name + "_BC_RS_val_data.png",
                   'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                   'GS_graphConfMat1': "\\static\\media\\" + file_name + "_BC_GS_train_data.png", 'GS_graphConfMat2': "\\static\\media\\" + file_name + "_BC_GS_val_data.png"}
        conn.send(context)
        conn.close()
    except Exception as e:
        print(e)


def MLP():
    try:
        savefile_x_final = file_path + csv_file_name + "_x_model.csv"
        if(not os.path.exists(savefile_x_final)):
            return render(request, 'processNotdone.html')
        df = pd.read_csv(savefile_x_final)
        targetVarFile = file_path + csv_file_name + "_targetVar.txt"
        file1 = open(targetVarFile, "r")  # write mode
        targetVar = file1.read()
        file1.close()
        # a variable pdf
        pdf = FPDF()
        # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
        y_model = df[targetVar]
        x_model = df.drop(targetVar, axis=1)
        X_train, X_test, y_train, y_test = train_test_split(
            x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
        X_train, X_val, y_train, y_val = train_test_split(
            X_train, y_train, test_size=0.22222222222222224, random_state=321)

        mlp = MLPClassifier(hidden_layer_sizes=100, alpha=0.0001, random_state=1,
                            learning_rate_init=0.001, max_iter=200)

        mlp.fit(X_train, y_train)
        # Test the model
        pred_mlp_val = mlp.predict(X_val)
        pred_mlp_prob_val = mlp.predict_proba(X_val)[:, 1]

        pred_mlp_train = mlp.predict(X_train)
        pred_mlp_prob_train = mlp.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_mlp_train))
        print(classification_report(y_val, pred_mlp_val))

        evaluate_model(y_val, pred_mlp_prob_val, y_train,
                       pred_mlp_prob_train, file_name+"MLP_NT_roc1.png")

        evaluate_model(y_val, pred_mlp_val, y_train, pred_mlp_train,
                       file_name+"MLP_NT_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data
        drawConfMatrix(y_train, pred_mlp_train, file_name +
                       "_MLP_NT_train_data.png")

        drawConfMatrix(y_val, pred_mlp_val, file_name +
                       "_MLP_NT_val_data.png")

        NT_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_mlp_prob_val, pred_mlp_val)

        NT_roc_file1 = "\\static\\media\\" + file_name+"MLP_NT_roc1.png"
        NT_roc_file2 = "\\static\\media\\" + file_name+"MLP_NT_roc2.png"

        # Add a page
        pdf.add_page()
        document = Document()
        pdf = exportPdf(10, 10, pdf, document, os.path.join(
            BASE_DIR, 'static\media\\', file_name+"MLP_NT_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, 'static\media\\', file_name+"MLP_NT_roc2.png"),
            "ROC Curve Test data", "Multi-Layer Perceptron - No Parameters Tuning")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document, os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_MLP_NT_train_data.png"),
            "Confusion Matrix Validation data", os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_MLP_NT_val_data.png"),
            "Confusion Matrix Test data", "Multi-Layer Perceptron - No Parameters Tuning")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, NT_auc_ginin_output, "Multi-Layer Perceptron - No Parameters Tuning")

        paramFiles = param_file_path + param_file_name + "_MLP_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = MLPClassifier(max_iter=100)
        # Create the random search model
        mlp_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                        scoring='roc_auc', cv=5,
                                        n_iter=100, verbose=1, random_state=50)

        # Fit
        mlp_search.fit(X_train, y_train)

        mlp_search.best_params_

        mlp_random = mlp_search.best_estimator_
        # Test the model
        pred_mlp_val0 = mlp_random.predict(X_val)
        pred_mlp_prob_val0 = mlp_random.predict_proba(X_val)[:, 1]

        pred_mlp_train0 = mlp_random.predict(X_train)
        pred_mlp_prob_train0 = mlp_random.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_mlp_train0))
        print(classification_report(y_val, pred_mlp_val0))

        evaluate_model(y_val, pred_mlp_prob_val0, y_train,
                       pred_mlp_prob_train0, file_name+"_MLP_RS_roc1.png")

        evaluate_model(y_val, pred_mlp_val0, y_train,
                       pred_mlp_train0, file_name+"_MLP_RS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_mlp_train0,
                       file_name + "_MLP_RS_train_data.png")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_mlp_val0,
                       file_name + "_MLP_RS_val_data.png")

        RS_roc_file1 = "\\static\\media\\" + file_name + "_MLP_RS_roc1.png"
        RS_roc_file2 = "\\static\\media\\" + file_name + "_MLP_RS_roc2.png"
        RS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_mlp_prob_val0, pred_mlp_val0)

        # Add a page
        pdf.add_page()
        document = Document()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_MLP_RS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_MLP_RS_roc2.png"),
            "ROC Curve Test data", "Multi-Layer Perceptron - Random Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_MLP_RS_train_data.png"),
            "Confusion Matrix Validation data", os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_MLP_RS_val_data.png"),
            "Confusion Matrix Test data", "Multi-Layer Perceptron - Random Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, RS_auc_ginin_output, "Multi-Layer Perceptron - Random Search")

        paramFiles = param_file_path + param_file_name + "_MLP_GS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = MLPClassifier(max_iter=500)

        # Create the random search model
        mlp_grid = GridSearchCV(estimator, param_grid, n_jobs=3,
                                scoring='roc_auc', cv=5, verbose=1)

        # Fit
        mlp_grid.fit(X_train, y_train)

        mlp_grid.best_params_

        mlp_final = mlp_grid.best_estimator_

        # Test the model
        pred_mlp_val1 = mlp_final.predict(X_val)
        pred_mlp_prob_val1 = mlp_final.predict_proba(X_val)[:, 1]

        pred_mlp_train1 = mlp_final.predict(X_train)
        pred_mlp_prob_train1 = mlp_final.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_mlp_train1))
        print(classification_report(y_val, pred_mlp_val1))

        evaluate_model(y_val, pred_mlp_prob_val1, y_train,
                       pred_mlp_prob_train1, file_name+"_MLP_GS_roc1.png")

        evaluate_model(y_val, pred_mlp_val1, y_train,
                       pred_mlp_train1, file_name+"_MLP_GS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_mlp_train1,
                       file_name + "_MLP_GS_train_data.png")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_mlp_val1,
                       file_name + "_MLP_GS_val_data.png")

        GS_roc_file1 = "\\static\\media\\" + file_name + "_MLP_GS_roc1.png"
        GS_roc_file2 = "\\static\\media\\" + file_name + "_MLP_GS_roc2.png"
        GS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_mlp_prob_val1, pred_mlp_val1)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_MLP_GS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_MLP_GS_roc2.png"),
            "ROC Curve Test data", "Multi-Layer Perceptron - Grid Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_MLP_GS_train_data.png"),
            "Confusion Matrix Validation data", os.path.join(
            BASE_DIR, 'static\media\\', file_name+"_MLP_GS_val_data.png"),
            "Confusion Matrix Test data", "Multi-Layer Perceptron - Grid Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, RS_auc_ginin_output, "Multi-Layer Perceptron - Grid Search")

        pdf.output(os.path.join(
            BASE_DIR, "static\\media\\" + file_name + "_MLP.pdf"))

        print('MLP Completed')
    except Exception as e:
        print('Error is', e)
        print(traceback.print_exc())


def evaluate_model(val_pred, val_probs, train_pred, train_probs, fileName):
    # """Compare machine learning model to baseline performance.
    # Computes statistics and shows ROC curve."""
    csv_file_name = "csvfile_"+user_name
    savefile_x_final = file_path + csv_file_name + "_x_model.csv"
    df = pd.read_csv(savefile_x_final)
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()

    # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
    y_model = df[targetVar]
    x_model = df.drop(targetVar, axis=1)
    X_train, X_test, y_train, y_test = train_test_split(
        x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
    X_train, X_val, y_train, y_val = train_test_split(
        X_train, y_train, test_size=0.22222222222222224, random_state=321)

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    baseline = {}
    baseline['recall'] = recall_score(y_val, [1 for _ in range(len(y_val))])
    baseline['precision'] = precision_score(
        y_val, [1 for _ in range(len(y_val))])
    baseline['roc'] = 0.5

    # Calculate ROC on validation dataset
    val_results = {}
    val_results['recall'] = recall_score(y_val, val_pred)
    val_results['precision'] = precision_score(y_val, val_pred)
    val_results['roc'] = roc_auc_score(y_val, val_probs)

    # Calculate ROC on training dataset
    train_results = {}
    train_results['recall'] = recall_score(y_train, train_pred)
    train_results['precision'] = precision_score(y_train, train_pred)
    train_results['roc'] = roc_auc_score(y_train, train_probs)

    for metric in ['recall', 'precision', 'roc']:
        print(
            f'{metric.capitalize()} Baseline: {round(baseline[metric], 2)} Validation: {round(val_results[metric], 2)} Training: {round(train_results[metric], 2)}')

    # Calculate false positive rates and true positive rates
    base_fpr, base_tpr, _ = roc_curve(y_val, [1 for _ in range(len(y_val))])
    model_fpr, model_tpr, _ = roc_curve(y_val, val_probs)

    # plt.figure(figsize=(8, 6))
    plt.figure(figsize=(10, 5))
    plt.rcParams['font.size'] = 12

    # Plot both curves
    plt.plot(base_fpr, base_tpr, 'b', label='baseline')
    plt.plot(model_fpr, model_tpr, 'r', label='model')
    plt.legend()
    plt.xlabel('False Positive Rate')
    plt.ylabel('True Positive Rate')
    plt.title('ROC Curves')
    plt.tight_layout()
    plt.savefig(os.path.join(
        BASE_DIR, 'static\media\\', fileName))
    plt.close()


def drawConfMatrix(y_val, pred_rf_val, fileName):
    cnf_matrix = confusion_matrix(y_val, pred_rf_val, labels=[0, 1])
    plt.figure(figsize=(10, 5))
    sns.heatmap(pd.DataFrame(cnf_matrix), annot=True,
                cmap="YlGnBu", fmt='g')
    plt.title('Confusion matrix: Validation data')
    plt.ylabel('Actual label')
    plt.xlabel('Predicted label')
    plt.tight_layout()
    plt.savefig(os.path.join(
        BASE_DIR, "static\\media\\" + fileName))
    plt.close()


def test_modelPerfomance(y_val, prob, val):
    # perform the test on model performance

    # perform ks test
    arrlstModelPerf = []
    csv_file_name = "csvfile_"+user_name
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()
    y_val1 = pd.DataFrame(y_val).reset_index()
    pred_prob_val1 = pd.DataFrame(prob).reset_index()
    new_val = pd.concat([y_val1, pred_prob_val1], axis=1).reset_index()
    new_val = new_val.drop(['level_0', 'index'], axis=1)
    new_val.columns = [targetVar, 'Probability']
    prob_to_1 = new_val.loc[new_val.eval(targetVar) ==
                            1, ['Probability']].sort_index()
    prob_to_0 = new_val.loc[new_val.eval(targetVar) ==
                            0, ['Probability']].sort_index()
    prob_to_1 = np.array(prob_to_1).reshape(len(prob_to_1))
    prob_to_0 = np.array(prob_to_0).reshape(len(prob_to_0))

    ks = ks_2samp(prob_to_0, prob_to_1)

    # perform auc and ginin test
    fpr1, tpr1, thresholds = roc_curve(y_val,  prob)
    auc_prob = round(auc(fpr1, tpr1), 3)
    gini_prob = 2*auc_prob-1

    fpr2, tpr2, thresholds = roc_curve(y_val,  val)
    auc_class = round(auc(fpr2, tpr2), 3)
    gini_class = 2*auc_class-1

    # perform MAE test
    abs_error_prob = abs(y_val - prob)
    abs_error_class = abs(y_val - val)

    mae_prob = round(np.mean(abs_error_prob), 4)
    mae_class = round(np.mean(abs_error_class), 4)

    # perform Accuracy test
    accuracy_class = accuracy_score(y_val, val)

    # print('\n')
    # print('Confusion Matrix - Validation:')
    # print(confusion_matrix(y_val, pred_rf_val))
    objtestModelPerf = lstTestModelPerf()
    objtestModelPerf.testName = "Confusion Matrix - Validation:"
    objtestModelPerf.testResult = confusion_matrix(y_val, val)
    arrlstModelPerf.append(objtestModelPerf)

    # print('\n')
    # print('Classification Report - Validation:')
    # print(classification_report(y_val, pred_rf_val))
    objtestModelPerf = lstTestModelPerf()
    objtestModelPerf.testName = "Classification Report - Validation:"
    report = classification_report(
        y_val, val, output_dict=True)
    dfclassification_report = pd.DataFrame(report).transpose()
    dfclassification_json = json.loads(
        dfclassification_report.to_json(orient="index"))
    objtestModelPerf.testResult_dict = dfclassification_json
    objtestModelPerf.testResult = ""
    arrlstModelPerf.append(objtestModelPerf)

    # print('\n')
    # print('Accuracy on classes:', accuracy_class)
    # print('\n')
    objtestModelPerf = lstTestModelPerf()
    objtestModelPerf.testName = "Accuracy on classes:"
    objtestModelPerf.testResult = accuracy_class
    arrlstModelPerf.append(objtestModelPerf)

    # print('KS test:', ks)
    objtestModelPerf = lstTestModelPerf()
    objtestModelPerf.testName = "KS test:"
    objtestModelPerf.testResult = ks
    arrlstModelPerf.append(objtestModelPerf)

    # print('AUC Score on probability:', auc_prob)
    # print('AUC Score on classes:', auc_class)
    objtestModelPerf = lstTestModelPerf()
    objtestModelPerf.testName = "AUC Score on probability:"
    objtestModelPerf.testResult = auc_prob
    arrlstModelPerf.append(objtestModelPerf)
    objtestModelPerf = lstTestModelPerf()
    objtestModelPerf.testName = "AUC Score on classes:"
    objtestModelPerf.testResult = auc_class
    arrlstModelPerf.append(objtestModelPerf)

    # print('GINI Score on probability:', gini_prob)
    # print('GINI Score on classes:', gini_class)
    objtestModelPerf = lstTestModelPerf()
    objtestModelPerf.testName = "GINI Score on probability:"
    objtestModelPerf.testResult = gini_prob
    arrlstModelPerf.append(objtestModelPerf)
    objtestModelPerf = lstTestModelPerf()
    objtestModelPerf.testName = "GINI Score on classes:"
    objtestModelPerf.testResult = gini_class
    arrlstModelPerf.append(objtestModelPerf)

    # print('MAE on probability:', mae_prob)
    # print('MAE on classes:', mae_class)
    objtestModelPerf = lstTestModelPerf()
    objtestModelPerf.testName = "MAE Score on probability:"
    objtestModelPerf.testResult = mae_prob
    arrlstModelPerf.append(objtestModelPerf)
    objtestModelPerf = lstTestModelPerf()
    objtestModelPerf.testName = "MAE Score on classes:"
    objtestModelPerf.testResult = mae_class
    arrlstModelPerf.append(objtestModelPerf)
    # print('arrlstModelPerf')
    # print(arrlstModelPerf)
    return arrlstModelPerf


def exportPdf(x, y, pdf, document, graph, text, graph2, text2, header):

    # set style and size of font
    # that you want in the pdf
    pdf.set_font("Arial", size=15)
    pdf.set_xy(x, y)
    pdf.set_text_color(0.0, 0.0, 0.0)
    pdf.multi_cell(0, 10, header, align='C')

    y += 20.0
    pdf.set_xy(40, y)
    if (os.path.exists(graph)):
        pdf.image(graph,  link='', type='', w=700/5, h=450/5)

    y += 450/5
    pdf.set_xy(x, y)
    pdf.set_font("Arial", size=12)
    pdf.set_text_color(0.0, 0.0, 0.0)
    pdf.multi_cell(0, 10, text, align='C')

    y += 20
    pdf.set_xy(40, y)
    if (os.path.exists(graph2)):
        pdf.image(graph2,  link='', type='', w=700/5, h=450/5)

    y += 450/5
    pdf.set_xy(x, y)
    pdf.set_font("Arial", size=12)
    pdf.set_text_color(0.0, 0.0, 0.0)
    pdf.multi_cell(0, 10, text2, align='C')

    para = document.add_paragraph()
    paragraph_format = para.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(header)
    font = run.font
    font.name = 'Arial'
    font.size = Pt(15)

    if (os.path.exists(graph)):
        document.add_picture(graph, width=Inches(6.0), height=Inches(3.25))

    para = document.add_paragraph()
    paragraph_format = para.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    font = run.font
    font.name = 'Arial'
    font.size = Pt(12)

    if (os.path.exists(graph2)):
        document.add_picture(graph2, width=Inches(6.0), height=Inches(3.25))

    para = document.add_paragraph()
    paragraph_format = para.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text2)
    font = run.font
    font.name = 'Arial'
    font.size = Pt(12)

    return pdf


def exportTestResultPdf(x, y, pdf, result, header):

    # set style and size of font
    # that you want in the pdf
    pdf.set_font("Arial", size=15)
    pdf.set_xy(x, y)
    pdf.set_text_color(0.0, 0.0, 0.0)
    pdf.multi_cell(0, 10, header, align='C')

    y += 20.0

    pdf.set_xy(x, y)
    pdf.set_font("Arial", size=12)
    pdf.set_text_color(0.0, 0.0, 0.0)
    pdf.multi_cell(0, 10, "Perform The Test On Model Performance", align='C')
    pdf.set_font("Arial", size=10)
    pdf.set_text_color(0.0, 0.0, 0.0)
    y += 10
    workbook = xlsxwriter.Workbook(os.path.join(
        BASE_DIR, "static\\media\\" + header+".xlsx"))
    worksheet = workbook.add_worksheet('Data')
    # Start from the first cell. Rows and columns are zero indexed.
    row = 0
    col = 0
    worksheet.write(row, col,     "testName")
    worksheet.write(row, col + 1, "testResult")
    row = 1
    for i in range(len(result)):
        # print(' result[i]', result[i])
        y += 5
        pdf.set_xy(20, y)
        pdf.cell(0, 10, result[i].testName)

        # Iterate over the data and write it out row by row.
        worksheet.write(row, col,     result[i].testName)
        worksheet.write(row, col + 1, str(result[i].testResult))
        row += 1

        if result[i].testName == "Classification Report - Validation:" or result[i].testName == "Confusion Matrix - Validation:":
            if result[i].testName == "Confusion Matrix - Validation:":
                y += 5
                pdf.set_xy(20, y)
                pdf.cell(0, 10, str(result[i].testResult))
            else:

                dictRes = result[i].testResult_dict
                worksheet2 = workbook.add_worksheet('Data2')
                y += 5
                row2 = 0
                col2 = 0
                worksheet2.write(row2, col2, "")
                pdf.set_xy(20, y)
                pdf.cell(0, 10, "")
                pdf.set_xy(50, y)
                pdf.cell(0, 10, "precision")
                worksheet2.write(row2, col2+1, "precision")
                pdf.set_xy(80, y)
                pdf.cell(0, 10, "recall")
                worksheet2.write(row2, col2+2, "recall")
                pdf.set_xy(110, y)
                pdf.cell(0, 10, "f1-score")
                worksheet2.write(row2, col2+3, "f1-score")
                pdf.set_xy(140, y)
                pdf.cell(0, 10, "support")
                worksheet2.write(row2, col2+4, "support")

                for key in dictRes:
                    y += 5
                    row2 += 1
                    col2 = 0
                    pdf.set_xy(20, y)
                    pdf.cell(0, 10, str(key))
                    worksheet2.write(row2, col2, str(key))
                    dictVal = dictRes[key]
                    x2 = 50
                    for keyval in dictVal:
                        pdf.set_xy(x2, y)
                        pdf.cell(0, 10, str(dictVal[keyval]))
                        x2 += 30
                        col2 += 1
                        worksheet2.write(row2, col2, str(dictVal[keyval]))
        else:
            pdf.set_xy(70, y)
            pdf.cell(0, 10, str(result[i].testResult))

        y += 5
        pdf.set_xy(20, y)
        pdf.cell(0, 10, "")
    workbook.close()
    return pdf


def exportTestResultPdfFromExcel(pdf, document,  header):

    # set style and size of font
    # that you want in the pdf
    x, y = 10, 10
    pdf.set_font("Arial", size=15)
    pdf.set_xy(x, y)
    pdf.set_text_color(0.0, 0.0, 0.0)
    pdf.multi_cell(0, 10, header, align='C')

    y += 20.0

    pdf.set_xy(x, y)
    pdf.set_font("Arial", size=12)
    pdf.set_text_color(0.0, 0.0, 0.0)
    pdf.multi_cell(0, 10, "Perform The Test On Model Performance", align='C')
    pdf.set_font("Arial", size=10)
    pdf.set_text_color(0.0, 0.0, 0.0)
    y += 10
    xlsfile_name = os.path.join(
        BASE_DIR, "static\\media\\" + header+".xlsx")
    df = pd.read_excel(xlsfile_name, sheet_name="Data")
    # print('df ', df)

    para = document.add_paragraph()
    paragraph_format = para.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(header)
    font = run.font
    font.name = 'Arial'
    font.size = Pt(15)

    para = document.add_paragraph()
    paragraph_format = para.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Perform The Test On Model Performance")
    font = run.font
    font.name = 'Arial'
    font.size = Pt(12)
    table = document.add_table(rows=1, cols=2)
    table2 = document.add_table(rows=1, cols=5)
    table1 = document.add_table(rows=1, cols=2)
    for index, row in df.iterrows():
        # print(' result[i]', result[i])
        y += 5
        pdf.set_xy(20, y)
        pdf.cell(0, 10, row["testName"])

        if row["testName"] == "Classification Report - Validation:" or row["testName"] == "Confusion Matrix - Validation:":

            row_cells = table.add_row().cells
            row_cells[0].text = row["testName"]
            if row["testName"] == "Confusion Matrix - Validation:":
                y += 5
                pdf.set_xy(20, y)
                pdf.cell(0, 10, row["testResult"])
                row_cells = table.add_row().cells
                row_cells[0].text = row["testResult"]
            else:

                dictRes = pd.read_excel(xlsfile_name, sheet_name="Data2")
                # print('dictRes', dictRes)
                y += 5
                pdf.set_xy(20, y)
                pdf.cell(0, 10, "")
                pdf.set_xy(50, y)
                pdf.cell(0, 10, "precision")
                pdf.set_xy(80, y)
                pdf.cell(0, 10, "recall")
                pdf.set_xy(110, y)
                pdf.cell(0, 10, "f1-score")
                pdf.set_xy(140, y)
                pdf.cell(0, 10, "support")

                hdr_cells = table2.rows[0].cells
                hdr_cells[0].text = ''
                hdr_cells[1].text = 'precision'
                hdr_cells[2].text = 'recall'
                hdr_cells[3].text = 'f1-score'
                hdr_cells[4].text = 'support'

                for index2, row2 in dictRes.iterrows():
                    y += 5
                    pdf.set_xy(20, y)
                    pdf.cell(0, 10, row2[0])
                    x2 = 50
                    pdf.set_xy(x2, y)
                    pdf.cell(0, 10, str(row2[1]))
                    x2 += 30
                    pdf.set_xy(x2, y)
                    pdf.cell(0, 10, str(row2[2]))
                    x2 += 30
                    pdf.set_xy(x2, y)
                    pdf.cell(0, 10, str(row2[3]))
                    x2 += 30
                    pdf.set_xy(x2, y)
                    pdf.cell(0, 10, str(row2[4]))

                    row_cells2 = table2.add_row().cells
                    row_cells2[0].text = str(row2[0])
                    row_cells2[1].text = str(row2[1])
                    row_cells2[2].text = str(row2[2])
                    row_cells2[3].text = str(row2[3])
                    row_cells2[4].text = str(row2[4])
        else:
            pdf.set_xy(70, y)
            pdf.cell(0, 10, row["testResult"])

            row_cells1 = table1.add_row().cells
            row_cells1[0].text = row["testName"]
            row_cells1[1].text = row["testResult"]

        y += 5
        pdf.set_xy(20, y)
        pdf.cell(0, 10, "")

    return pdf


def addCommenttoPdf(x, y, pdf, document, comment, header):

    # set style and size of font
    # that you want in the pdf
    pdf.set_font("Arial", size=15)
    pdf.set_xy(x, y)
    pdf.set_text_color(0.0, 0.0, 0.0)
    pdf.multi_cell(0, 10, header, align='C')

    y += 20.0

    pdf.set_xy(x, y)
    pdf.set_font("Arial", size=12)
    pdf.set_text_color(0.0, 0.0, 0.0)
    # pdf.multi_cell(0, 10, "User Comments", align='C')
    # pdf.set_font("Arial", size=10)
    # pdf.set_text_color(0.0, 0.0, 0.0)
    # y += 10
    pdf.multi_cell(0, 10, comment)

    para = document.add_paragraph()
    paragraph_format = para.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(header)
    font = run.font
    font.name = 'Arial'
    font.size = Pt(15)

    paracomment = document.add_paragraph()
    paragraph_format = paracomment.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paracomment.add_run(comment)
    font = run.font
    font.name = 'Arial'
    font.size = Pt(12)

    return pdf


if __name__ == '__main__':
    # p1 = Process(target=func1)
    # p1.start()
    # p2 = Process(target=func2)
    # p2.start()
    # # This is where I had to add the join() function.
    # p1.join()
    # p2.join()

    # MLP()
    parent_conn, child_conn = Pipe()
    p = Process(target=BC, args=(child_conn,))
    p.start()
    print(parent_conn.recv())   # prints "[42, None, 'hello']"
    p.join()
