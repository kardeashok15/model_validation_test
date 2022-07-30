from inspect import Traceback
import logging
from django.http import HttpResponse, Http404
import smtplib
from sklearn.neural_network import MLPClassifier
from sklearn.ensemble import VotingClassifier, AdaBoostClassifier, RandomForestClassifier, GradientBoostingClassifier, BaggingClassifier
from sklearn.neighbors import KNeighborsClassifier, KNeighborsRegressor
from sklearn import svm
from sklearn.linear_model import LogisticRegression
from statsmodels.stats.outliers_influence import variance_inflation_factor
from typing import Reversible
from sklearn import preprocessing
from sklearn.feature_selection import SelectPercentile, chi2, RFE
from sklearn.preprocessing import StandardScaler, label_binarize
from sklearn.model_selection import train_test_split, RandomizedSearchCV, GridSearchCV, StratifiedKFold, cross_val_score
from sklearn.metrics import classification_report, roc_curve, auc, roc_auc_score, confusion_matrix, recall_score, precision_score, accuracy_score
from bubbly.bubbly import bubbleplot
import plotly_express as px
import joypy
from django.core.files.storage import FileSystemStorage
from django.conf import settings
import math
from io import StringIO
from statsmodels import robust
import plotly.graph_objs as go
from plotly.offline import plot
import plotly.offline as py
from pandas.plotting import parallel_coordinates
from pandas import plotting
import matplotlib.pyplot as plt
from django.shortcuts import redirect, render
from django.http import JsonResponse
import pandas as pd
#import terality as pd
import numpy as np
from .models import lstTestModelPerf
import os
from pathlib import Path
import json
# for visualizations
import seaborn as sns
import matplotlib
import xgboost as xgb
from scipy.stats import ks_2samp
from scipy import stats
from scipy.stats import randint as sp_randint

# generate random floating point values
from numpy.random import seed
from numpy.random import rand
from fpdf import FPDF, HTMLMixin
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
import xlsxwriter
from datetime import date
import time
import traceback
import shutil
import vaex as vx
matplotlib.use('Agg')
logging.basicConfig(filename='example.log',
                    level=logging.DEBUG)

# for modeling


BASE_DIR = Path(__file__).resolve().parent.parent
# Create your views here.
user_name = "user1"
param_file_path = os.path.join(BASE_DIR, 'static/param_files/')
file_path = os.path.join(BASE_DIR, 'static/csv_files/')
param_file_name = "paramfile_"+user_name
file_name = "user1"
app_url = "http://3.131.88.246:8000/modelval/"
font_files = os.path.join(BASE_DIR, 'static/fonts/')

processingFile_path='static/reportTemplates/processing.csv' 

plot_dir='/static/media/'
plot_dir_view='static/media/'

 
src_files='static/cnfrmsrc_files/' 


def tuneParams(request):
    return render(request, 'tuneParameters.html')


def getParamName(request):
    RF = ['criterion', 'max_features', 'max_depth', 'min_samples_leaf',
          'min_samples_split', 'n_estimators', 'max_leaf_nodes', 'bootstrap']
    RF_GS = ['criterion', 'max_features', 'bootstrap', 'max_depth',
             'min_samples_leaf', 'min_samples_split', 'n_estimators', 'max_leaf_nodes']
    XGB = ['objective', 'colsample_bytree', 'learning_rate',
           'max_depth', 'lambda', 'n_estimators', 'missing', 'seed']
    MLP = ['hidden_layer_sizes', 'activation',
           'solver', 'alpha', 'momentum', 'learning_rate']
    GB = ['loss', 'learning_rate', 'n_estimators', 'criterion', 'min_samples_split',
          'min_samples_leaf', 'max_depth', 'max_features', 'max_leaf_nodes', 'init', 'validation_fraction']
    KNN = ['n_neighbors', 'weights', 'algorithm', 'p']
    SVM = ['C', 'kernel', 'gamma', 'probability', 'class_weight']
    BC = ['base_estimator', 'n_estimators', 'max_samples',
          'max_features', 'bootstrap', 'bootstrap_features', 'oob_score']
    modelName = request.GET['modelName']
    tuneMethod = request.GET['tuneMethod']
    if (modelName == "Random_Forest"):
        paramFiles = param_file_path + param_file_name + "_RF_" + tuneMethod + ".csv"
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            result = df.to_json(orient="records")
            result = json.loads(result)
            data = {
                'params': RF,
                'paramVals': result
            }
        else:
            data = {
                'params': RF,
                'paramVals': []
            }
    elif (modelName == "XGBoost"):
        paramFiles = param_file_path + param_file_name + "_XGB_" + tuneMethod + ".csv"
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            result = df.to_json(orient="records")
            result = json.loads(result)
            data = {
                'params': XGB,
                'paramVals': result
            }
        else:
            data = {
                'params': XGB,
                'paramVals': []
            }
    elif (modelName == "MLP"):
        paramFiles = param_file_path + param_file_name + "_MLP_" + tuneMethod + ".csv"
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            result = df.to_json(orient="records")
            result = json.loads(result)
            data = {
                'params': MLP,
                'paramVals': result
            }
        else:
            data = {
                'params': MLP,
                'paramVals': []
            }
    elif (modelName == "Gradient_Boosting"):
        paramFiles = param_file_path + param_file_name + "_GB_" + tuneMethod + ".csv"
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            result = df.to_json(orient="records")
            result = json.loads(result)
            data = {
                'params': GB,
                'paramVals': result
            }
        else:
            data = {
                'params': GB,
                'paramVals': []
            }
    elif (modelName == "KNN"):
        paramFiles = param_file_path + param_file_name + "_KNN_" + tuneMethod + ".csv"
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            result = df.to_json(orient="records")
            result = json.loads(result)
            data = {
                'params': KNN,
                'paramVals': result
            }
        else:
            data = {
                'params': KNN,
                'paramVals': []
            }
    elif (modelName == "SVM"):
        paramFiles = param_file_path + param_file_name + "_SVM_" + tuneMethod + ".csv"
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            result = df.to_json(orient="records")
            result = json.loads(result)
            data = {
                'params': SVM,
                'paramVals': result
            }
        else:
            data = {
                'params': SVM,
                'paramVals': []
            }
    elif (modelName == "Bagging_Classifier"):
        paramFiles = param_file_path + param_file_name + "_BC_" + tuneMethod + ".csv"
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            result = df.to_json(orient="records")
            result = json.loads(result)
            data = {
                'params': BC,
                'paramVals': result
            }
        else:
            data = {
                'params': BC,
                'paramVals': []
            }
    return JsonResponse(data)


def setParamName(request):
    modelName = request.GET['modelName']
    tuneMethod = request.GET['tuneMethod']
    paramName = request.GET['paramName']
    paramValue = request.GET['paramValue']
    if (modelName == "Random_Forest"):
        paramFiles = param_file_path + param_file_name + "_RF_" + tuneMethod + ".csv"
    elif (modelName == "XGBoost"):
        paramFiles = param_file_path + param_file_name + "_XGB_" + tuneMethod + ".csv"
    elif (modelName == "MLP"):
        paramFiles = param_file_path + param_file_name + "_MLP_" + tuneMethod + ".csv"
    elif (modelName == "Gradient_Boosting"):
        paramFiles = param_file_path + param_file_name + "_GB_" + tuneMethod + ".csv"
    elif (modelName == "KNN"):
        paramFiles = param_file_path + param_file_name + "_KNN_" + tuneMethod + ".csv"
    elif (modelName == "SVM"):
        paramFiles = param_file_path + param_file_name + "_SVM_" + tuneMethod + ".csv"
    elif (modelName == "Bagging_Classifier"):
        paramFiles = param_file_path + param_file_name + "_BC_" + tuneMethod + ".csv"
    if os.path.exists(paramFiles):
        df_old = pd.read_csv(paramFiles)
        if (df_old["paramName"] == paramName).any():
            df_old.loc[df_old.paramName ==
                       paramName, "paramValue"] = paramValue
            df_old.to_csv(paramFiles, index=False)
        else:
            data = [[paramName, paramValue]]
            df_new = pd.DataFrame(
                data, columns=['paramName', 'paramValue'])
            df = pd.concat([df_old, df_new], axis=0)
            df.to_csv(paramFiles, index=False)
    else:
        data = [[paramName, paramValue]]
        df = pd.DataFrame(data, columns=['paramName', 'paramValue'])
        df.to_csv(paramFiles, index=False)

    if os.path.exists(paramFiles):
        df = pd.read_csv(paramFiles)
        result = df.to_json(orient="records")
        result = json.loads(result)
    data = {
        'paramVals': result
    }
    return JsonResponse(data)


def evaluate_model(val_pred, val_probs, train_pred, train_probs, fileName):
    # """Compare machine learning model to baseline performance.
    # Computes statistics and shows ROC curve."""
    csv_file_name = "csvfile_"+user_name
    savefile_x_final = file_path + csv_file_name + "_x_model.csv"
    df = pd.read_csv(savefile_x_final)
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()

    # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
    y_model = df[targetVar]
    x_model = df.drop(targetVar, axis=1)
    X_train, X_test, y_train, y_test = train_test_split(
        x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
    X_train, X_val, y_train, y_val = train_test_split(
        X_train, y_train, test_size=0.22222222222222224, random_state=321)

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    baseline = {}
    baseline['recall'] = recall_score(y_val, [1 for _ in range(len(y_val))])
    baseline['precision'] = precision_score(
        y_val, [1 for _ in range(len(y_val))])
    baseline['roc'] = 0.5

    # Calculate ROC on validation dataset
    val_results = {}
    val_results['recall'] = recall_score(y_val, val_pred)
    val_results['precision'] = precision_score(y_val, val_pred)
    val_results['roc'] = roc_auc_score(y_val, val_probs)

    # Calculate ROC on training dataset
    train_results = {}
    train_results['recall'] = recall_score(y_train, train_pred)
    train_results['precision'] = precision_score(y_train, train_pred)
    train_results['roc'] = roc_auc_score(y_train, train_probs)

    for metric in ['recall', 'precision', 'roc']:
        print(
            f'{metric.capitalize()} Baseline: {round(baseline[metric], 2)} Validation: {round(val_results[metric], 2)} Training: {round(train_results[metric], 2)}')

    # Calculate false positive rates and true positive rates
    base_fpr, base_tpr, _ = roc_curve(y_val, [1 for _ in range(len(y_val))])
    model_fpr, model_tpr, _ = roc_curve(y_val, val_probs)

    # plt.figure(figsize=(8, 6))
    plt.figure(figsize=(10, 5))
    plt.rcParams['font.size'] = 12

    # Plot both curves
    plt.plot(base_fpr, base_tpr, 'b', label='baseline')
    plt.plot(model_fpr, model_tpr, 'r', label='model')
    plt.legend()
    plt.xlabel('False Positive Rate')
    plt.ylabel('True Positive Rate')
    plt.title('ROC Curves')
    plt.tight_layout()
    plt.savefig(os.path.join(
        BASE_DIR, plot_dir_view, fileName))
    plt.close()


def randomForest_NoTunning(request):
    try:
        csv_file_name = "csvfile_"+user_name
        savefile_x_final = file_path + csv_file_name + "_x_model.csv"
        df = pd.read_csv(savefile_x_final)
        targetVarFile = file_path + csv_file_name + "_targetVar.txt"
        file1 = open(targetVarFile, "r")  # write mode
        targetVar = file1.read()
        file1.close()

        # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
        y_model = df[targetVar]
        x_model = df.drop(targetVar, axis=1)
        X_train, X_test, y_train, y_test = train_test_split(
            x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
        X_train, X_val, y_train, y_val = train_test_split(
            X_train, y_train, test_size=0.22222222222222224, random_state=321)

        rf_model = RandomForestClassifier(n_estimators=100, criterion="gini",
                                          random_state=50,
                                          max_features='sqrt',
                                          n_jobs=-1, verbose=1)
        rf_model.fit(X_train, y_train)

        # check the average of the nodes and depth.
        n_nodes = []
        max_depths = []
        for ind_tree in rf_model.estimators_:
            n_nodes.append(ind_tree.tree_.node_count)
            max_depths.append(ind_tree.tree_.max_depth)

        print(f'Average number of nodes {int(np.mean(n_nodes))}')
        print(f'Average maximum depth {int(np.mean(max_depths))}')

        # Test the model
        pred_rf_val = rf_model.predict(X_val)
        pred_rf_prob_val = rf_model.predict_proba(X_val)[:, 1]

        pred_rf_train = rf_model.predict(X_train)
        pred_rf_prob_train = rf_model.predict_proba(X_train)[:, 1]

        # Get the model performance
        print('classification report on training data')
        print(classification_report(y_train, pred_rf_train))
        print('\n')
        print('classification report on validation data')
        print(classification_report(y_val, pred_rf_val))

        evaluate_model(y_val, pred_rf_prob_val, y_train,
                       pred_rf_prob_train, file_name+"_roc1.png")

        evaluate_model(y_val, pred_rf_val, y_train,
                       pred_rf_train, file_name+"_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data
        print('y_train')
        y_train.to_csv('y_trainData.csv', index=False)
        cnf_matrix = confusion_matrix(y_train, pred_rf_train, labels=[0, 1])
        sns.heatmap(pd.DataFrame(cnf_matrix), annot=True,
                    cmap="YlGnBu", fmt='g')
        plt.tight_layout()
        plt.title('Confusion matrix: Training data')
        plt.ylabel('Actual label')
        plt.xlabel('Predicted label')
        plt.savefig(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_RF_train_data.png"))
        plt.close()
        # show the confusion matrix for validation data

        cnf_matrix2 = confusion_matrix(y_val, pred_rf_val, labels=[0, 1])
        sns.heatmap(pd.DataFrame(cnf_matrix2), annot=True,
                    cmap="YlGnBu", fmt='g')
        plt.tight_layout()
        plt.title('Confusion matrix: Validation data')
        plt.ylabel('Actual label')
        plt.xlabel('Predicted label')
        plt.savefig(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_RF_val_data.png"))
        plt.close()
        roc_file1 = plot_dir + file_name + "_roc1.png"
        roc_file2 = plot_dir + file_name + "_roc2.png"
        auc_ginin_output = test_modelPerfomance(
            y_val, pred_rf_prob_val, pred_rf_val)
        context = {'rocgraphpath1': roc_file1, 'rocgraphpath2': roc_file2, 'auc_ginin_output': auc_ginin_output,
                   'graphConfMat1': plot_dir + file_name + "_RF_train_data.png", 'graphConfMat2': plot_dir + file_name + "_RF_val_data.png"}
        return render(request, 'showModelOutput.html', context)
    except Exception as e:
        print(e)
        return render(request, 'error.html')


def randomForest_RS(request):
    csv_file_name = "csvfile_"+user_name
    savefile_x_final = file_path + csv_file_name + "_x_model.csv"
    df = pd.read_csv(savefile_x_final)
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()

    # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
    y_model = df[targetVar]
    x_model = df.drop(targetVar, axis=1)
    X_train, X_test, y_train, y_test = train_test_split(
        x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
    X_train, X_val, y_train, y_val = train_test_split(
        X_train, y_train, test_size=0.22222222222222224, random_state=321)

    paramFiles = param_file_path + param_file_name + "_RF_RS.csv"
    param_grid = {}
    if os.path.exists(paramFiles):
        df = pd.read_csv(paramFiles)
        for index, row in df.iterrows():
            param_grid[row['paramName']] = eval(row['paramValue'])
    # Hyperparameter grid
    # param_grid = {'criterion': ['entropy', 'gini'],
    #               'max_features': ['auto', 'sqrt', 'log2', None],
    #               'max_depth': list(np.linspace(10, 200, 50).astype(int)),
    #               'min_samples_leaf': list(np.linspace(2, 20, 10).astype(int)),
    #               'min_samples_split': [2, 5, 7, 10, 12, 15],
    #               'n_estimators': np.linspace(10, 200, 50).astype(int),
    #               'max_leaf_nodes': list(np.linspace(5, 100, 20).astype(int)),
    #               'bootstrap': [True, False]
    #               }
    # print(param_grid)
    # Estimator for use in random search
    model = RandomForestClassifier()

    # Create the random search model
    rf_search = RandomizedSearchCV(model, param_grid, n_jobs=3,
                                   scoring='roc_auc', cv=5,
                                   n_iter=100, verbose=1, random_state=50)

    # Fit
    rf_search.fit(X_train, y_train)

    rf_search.best_params_

    rf_random = rf_search.best_estimator_

    # Test the model
    pred_rf_val0 = rf_random.predict(X_val)
    pred_rf_prob_val0 = rf_random.predict_proba(X_val)[:, 1]

    pred_rf_train0 = rf_random.predict(X_train)
    pred_rf_prob_train0 = rf_random.predict_proba(X_train)[:, 1]

    # Get the model performance
    print(classification_report(y_train, pred_rf_train0))
    print(classification_report(y_val, pred_rf_val0))

    evaluate_model(y_val, pred_rf_prob_val0, y_train,
                   pred_rf_prob_train0, file_name+"_RF_RS_roc1.png")

    evaluate_model(y_val, pred_rf_val0, y_train,
                   pred_rf_train0, file_name+"_RF_RS_roc2.png")

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    # show the confusion matrix for training data

    # cnf_matrix = confusion_matrix(y_train, pred_rf_train0, labels=[0, 1])
    # plt.figure(figsize=(10, 6))
    # sns.heatmap(pd.DataFrame(cnf_matrix), annot=True,
    #             cmap="YlGnBu", fmt='g')

    # plt.title('Confusion matrix: Training data')
    # plt.ylabel('Actual label')
    # plt.xlabel('Predicted label')
    # plt.tight_layout()
    # plt.savefig(os.path.join(
    #     BASE_DIR, plot_dir_view + file_name + "_RF_RS_train_data.png"))
    # plt.close()
    drawConfMatrix(y_train, pred_rf_train0,
                   file_name + "_RF_RS_train_data.png", "Training")

    # show the confusion matrix for validation data

    # cnf_matrix2 = confusion_matrix(y_val, pred_rf_val0, labels=[0, 1])
    # plt.figure(figsize=(10, 6))
    # sns.heatmap(pd.DataFrame(cnf_matrix2), annot=True,
    #             cmap="YlGnBu", fmt='g')
    # plt.title('Confusion matrix: Validation data')
    # plt.ylabel('Actual label')
    # plt.xlabel('Predicted label')
    # plt.tight_layout()
    # plt.savefig(os.path.join(
    #     BASE_DIR, plot_dir_view + file_name + "_RF_RS_val_data.png"))
    # plt.close()

    drawConfMatrix(y_val, pred_rf_val0, file_name +
                   "_RF_RS_val_data.png", "Validation")

    roc_file1 = plot_dir + file_name + "_roc1.png"
    roc_file2 = plot_dir + file_name + "_roc2.png"
    auc_ginin_output = test_modelPerfomance(
        y_val, pred_rf_prob_val0, pred_rf_val0)

    context = {'rocgraphpath1': roc_file1, 'rocgraphpath2': roc_file2, 'auc_ginin_output': auc_ginin_output,
               'graphConfMat1': plot_dir + file_name + "_RF_RS_train_data.png", 'graphConfMat2': plot_dir + file_name + "_RF_RS_val_data.png"}
    return render(request, 'showModelOutput.html', context)


def randomForest_GS(request):
    csv_file_name = "csvfile_"+user_name
    savefile_x_final = file_path + csv_file_name + "_x_model.csv"
    df = pd.read_csv(savefile_x_final)
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()

    # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
    y_model = df[targetVar]
    x_model = df.drop(targetVar, axis=1)
    X_train, X_test, y_train, y_test = train_test_split(
        x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
    X_train, X_val, y_train, y_val = train_test_split(
        X_train, y_train, test_size=0.22222222222222224, random_state=321)

    paramFiles = param_file_path + param_file_name + "_RF_RS.csv"
    param_grid = {}
    if os.path.exists(paramFiles):
        df = pd.read_csv(paramFiles)
        for index, row in df.iterrows():
            param_grid[row['paramName']] = eval(row['paramValue'])

    model = RandomForestClassifier()

    # Create the random search model
    rf_search = RandomizedSearchCV(model, param_grid, n_jobs=3,
                                   scoring='roc_auc', cv=5,
                                   n_iter=100, verbose=1, random_state=50)

    # Fit
    rf_search.fit(X_train, y_train)

    paramFiles = param_file_path + param_file_name + "_RF_GS.csv"
    param_grid = {}
    if os.path.exists(paramFiles):
        df = pd.read_csv(paramFiles)
        for index, row in df.iterrows():
            param_grid[row['paramName']] = eval(row['paramValue'])
    # Hyperparameter grid
    # param_grid = {'criterion': [rf_search.best_params_['criterion']],
    #               'max_features': [rf_search.best_params_['max_features']],
    #               'bootstrap': [rf_search.best_params_['bootstrap']],

    #               'max_depth': [rf_search.best_params_['max_depth']-5,
    #                             rf_search.best_params_['max_depth'],
    #                             rf_search.best_params_['max_depth']+5, ],

    #               'min_samples_leaf': [rf_search.best_params_['min_samples_leaf'] - 2,
    #                                    rf_search.best_params_[
    #                                        'min_samples_leaf'],
    #                                    rf_search.best_params_['min_samples_leaf'] + 2],

    #               'min_samples_split': [rf_search.best_params_['min_samples_split'] - 2,
    #                                     rf_search.best_params_[
    #                                     'min_samples_split'],
    #                                     rf_search.best_params_['min_samples_split'] + 2],

    #               'n_estimators': [rf_search.best_params_['n_estimators'] - 5,
    #                                rf_search.best_params_['n_estimators'],
    #                                rf_search.best_params_['n_estimators'] + 5],

    #               'max_leaf_nodes': [rf_search.best_params_['max_leaf_nodes'] - 3,
    #                                  rf_search.best_params_['max_leaf_nodes'],
    #                                  rf_search.best_params_['max_leaf_nodes'] + 3]
    #               }
    # print('param_grid')
    # print(param_grid)

    # Estimator for use in random search
    model = RandomForestClassifier()

    # Create the random search model
    rf_grid = GridSearchCV(model, param_grid, n_jobs=3,
                           scoring='roc_auc', cv=5, verbose=1)

    # Fit
    rf_grid.fit(X_train, y_train)

    rf_grid.best_params_

    rf_grid = rf_grid.best_estimator_

    # Test the model
    pred_rf_val1 = rf_grid.predict(X_val)
    pred_rf_prob_val1 = rf_grid.predict_proba(X_val)[:, 1]

    pred_rf_train1 = rf_grid.predict(X_train)
    pred_rf_prob_train1 = rf_grid.predict_proba(X_train)[:, 1]

    # Get the model performance
    print(classification_report(y_train, pred_rf_train1))
    print(classification_report(y_val, pred_rf_val1))

    evaluate_model(y_val, pred_rf_prob_val1, y_train,
                   pred_rf_prob_train1, file_name+"_RF_GS_roc1.png")

    evaluate_model(y_val, pred_rf_val1, y_train,
                   pred_rf_train1, file_name+"_RF_GS_roc2.png")

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    # show the confusion matrix for training data

    cnf_matrix = confusion_matrix(y_train, pred_rf_train1, labels=[0, 1])
    plt.figure(figsize=(10, 6))
    sns.heatmap(pd.DataFrame(cnf_matrix), annot=True,
                cmap="YlGnBu", fmt='g')

    plt.title('Confusion matrix: Training data')
    plt.ylabel('Actual label')
    plt.xlabel('Predicted label')
    plt.tight_layout()
    plt.savefig(os.path.join(
        BASE_DIR, plot_dir_view + file_name + "_RF_GS_train_data.png"))
    plt.close()
    # show the confusion matrix for validation data

    cnf_matrix2 = confusion_matrix(y_val, pred_rf_val1, labels=[0, 1])
    plt.figure(figsize=(10, 6))
    sns.heatmap(pd.DataFrame(cnf_matrix2), annot=True,
                cmap="YlGnBu", fmt='g')
    plt.title('Confusion matrix: Validation data')
    plt.ylabel('Actual label')
    plt.xlabel('Predicted label')
    plt.tight_layout()
    plt.savefig(os.path.join(
        BASE_DIR, plot_dir_view + file_name + "_RF_GS_val_data.png"))
    plt.close()
    roc_file1 = plot_dir + file_name + "_RF_GS_roc1.png"
    roc_file2 = plot_dir + file_name + "_RF_GS_roc2.png"
    auc_ginin_output = test_modelPerfomance(
        y_val, pred_rf_prob_val1, pred_rf_val1)

    context = {'rocgraphpath1': roc_file1, 'rocgraphpath2': roc_file2, 'auc_ginin_output': auc_ginin_output,
               'graphConfMat1': plot_dir + file_name + "_RF_GS_train_data.png", 'graphConfMat2': plot_dir + file_name + "_RF_GS_val_data.png"}
    return render(request, 'showModelOutput.html', context)


def xgBoost_NoTunning():
    csv_file_name = "csvfile_"+user_name
    savefile_x_final = file_path + csv_file_name + "_x_model.csv"
    df = pd.read_csv(savefile_x_final)
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()
    # a variable pdf
    pdf = FPDF()
    # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
    y_model = df[targetVar]
    x_model = df.drop(targetVar, axis=1)
    X_train, X_test, y_train, y_test = train_test_split(
        x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
    X_train, X_val, y_train, y_val = train_test_split(
        X_train, y_train, test_size=0.22222222222222224, random_state=321)

    # Fit XGRegressor to the Training set
    xg_clf = xgb.XGBClassifier()
    xg_clf.fit(X_train, y_train)

    # Test the model
    pred_xgb_val = xg_clf.predict(X_val)
    pred_xgb_prob_val = xg_clf.predict_proba(X_val)[:, 1]

    pred_xgb_train = xg_clf.predict(X_train)
    pred_xgb_prob_train = xg_clf.predict_proba(X_train)[:, 1]

    # Get the model performance
    print(classification_report(y_train, pred_xgb_train))
    print(classification_report(y_val, pred_xgb_val))

    evaluate_model(pred_xgb_val, pred_xgb_prob_val, pred_xgb_train,
                   pred_xgb_prob_train, file_name+"xgboost_NT_roc1.png")

    evaluate_model(pred_xgb_val, pred_xgb_val, pred_xgb_train,
                   pred_xgb_train, file_name+"xgboost_NT_roc2.png")

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    # show the confusion matrix for training data
    drawConfMatrix(y_train, pred_xgb_train, file_name +
                   "_xgboost_NT_train_data.png", "Training")

    drawConfMatrix(y_val, pred_xgb_val, file_name +
                   "_xgboost_NT_val_data.png", "Validation")

    NT_auc_ginin_output = test_modelPerfomance(
        y_val,  pred_xgb_prob_val, pred_xgb_val)

    NT_roc_file1 = plot_dir + file_name+"xgboost_NT_roc1.png"
    NT_roc_file2 = plot_dir + file_name+"xgboost_NT_roc2.png"

    # Add a page
    pdf.add_page()
    document = Document()
    pdf = exportPdf(10, 10, pdf, document,  os.path.join(
        BASE_DIR, plot_dir_view, file_name+"xgboost_NT_roc1.png"),
        "ROC Curve Validation data", os.path.join(
        BASE_DIR, plot_dir_view, file_name+"xgboost_NT_roc2.png"),
        "ROC Curve Test data", "XGBoost - No Parameters Tuning")
    pdf.add_page()
    pdf = exportPdf(10, 10, pdf, document,  os.path.join(
        BASE_DIR, plot_dir_view, file_name+"_xgboost_NT_train_data.png"),
        "Confusion Matrix Training data", os.path.join(
        BASE_DIR, plot_dir_view, file_name+"_xgboost_NT_val_data.png"),
        "Confusion Matrix Validation data", "XGBoost - No Parameters Tuning")
    pdf.add_page()
    pdf = exportTestResultPdf(
        10, 10, pdf, NT_auc_ginin_output, "XGBoost - No Parameters Tuning")


def xgBoost_GS():
    try:
        csv_file_name = "csvfile_"+user_name
        savefile_x_final = file_path + csv_file_name + "_x_model.csv"
        df = pd.read_csv(savefile_x_final)
        targetVarFile = file_path + csv_file_name + "_targetVar.txt"
        file1 = open(targetVarFile, "r")  # write mode
        targetVar = file1.read()
        file1.close()
        # a variable pdf
        pdf = FPDF()
        document = Document()
        # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
        y_model = df[targetVar]
        x_model = df.drop(targetVar, axis=1)
        X_train, X_test, y_train, y_test = train_test_split(
            x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
        X_train, X_val, y_train, y_val = train_test_split(
            X_train, y_train, test_size=0.22222222222222224, random_state=321)
        pdf.add_page()

        paramFiles = param_file_path + param_file_name + "_XGB_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = xgb.XGBClassifier(random_state=50)

        # Create the random search model
        xgb_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                        scoring='roc_auc', cv=5,
                                        n_iter=100, verbose=1, random_state=50)

        # Fit
        logging.debug('inside xgboost GS')
        t1 = time.time()
        xgb_search.fit(X_train, y_train)
        logging.debug('inside xgboost GS took %s', str(time.time()-t1))
        xgb_search.best_params_

        xgb_random = xgb_search.best_estimator_

        paramFiles = param_file_path + param_file_name + "_XGB_GS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        t2 = time.time()
        estimator = xgb.XGBClassifier(random_state=50)

        # Create the random search model
        xgb_grid = GridSearchCV(estimator, param_grid, n_jobs=3,
                                scoring='roc_auc', cv=5, verbose=1)

        # Fit
        xgb_grid.fit(X_train, y_train)
        logging.debug('inside xgboost GS2 took %s', str(time.time()-t2))
        xgb_grid.best_params_

        xgb_final = xgb_grid.best_estimator_

        # Test the model
        pred_xgb_val1 = xgb_final.predict(X_val)
        pred_xgb_prob_val1 = xgb_final.predict_proba(X_val)[:, 1]

        pred_xgb_train1 = xgb_final.predict(X_train)
        pred_xgb_prob_train1 = xgb_final.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_xgb_train1))
        print(classification_report(y_val, pred_xgb_val1))

        evaluate_model(y_val, pred_xgb_prob_val1, y_train,
                       pred_xgb_prob_train1, file_name+"_xgboost_GS_roc1.png")

        evaluate_model(y_val, pred_xgb_val1, y_train,
                       pred_xgb_train1, file_name+"_xgboost_GS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_xgb_train1,
                       file_name + "_xgboost_GS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_xgb_val1,  file_name +
                       "_xgboost_GS_val_data.png", "Validation")

        GS_roc_file1 = plot_dir + file_name + "_xgboost_GS_roc1.png"
        GS_roc_file2 = plot_dir + file_name + "_xgboost_GS_roc2.png"
        GS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_xgb_prob_val1, pred_xgb_val1)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_GS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_GS_roc2.png"),
            "ROC Curve Test data", "XGBoost - Grid Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document, os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_GS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_GS_val_data.png"),
            "Confusion Matrix Validation data", "XGBoost - Grid Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, GS_auc_ginin_output, "XGBoost - Grid Search")

        pdf.output(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_XGBoost.pdf"))

        print('Process XGboost ended')
    except Exception as e:
        print(e)


def xgBoost_RS():
    try:
        csv_file_name = "csvfile_"+user_name
        savefile_x_final = file_path + csv_file_name + "_x_model.csv"
        df = pd.read_csv(savefile_x_final)
        targetVarFile = file_path + csv_file_name + "_targetVar.txt"
        file1 = open(targetVarFile, "r")  # write mode
        targetVar = file1.read()
        file1.close()
        # a variable pdf
        pdf = FPDF()
        document = Document()
        # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
        y_model = df[targetVar]
        x_model = df.drop(targetVar, axis=1)
        X_train, X_test, y_train, y_test = train_test_split(
            x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
        X_train, X_val, y_train, y_val = train_test_split(
            X_train, y_train, test_size=0.22222222222222224, random_state=321)
        pdf.add_page()

        paramFiles = param_file_path + param_file_name + "_XGB_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        logging.debug('inside xgboost RS')
        t1 = time.time()
        estimator = xgb.XGBClassifier(random_state=50)

        # Create the random search model

        xgb_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                        scoring='roc_auc', cv=5,
                                        n_iter=100, verbose=1, random_state=50)

        # Fit
        xgb_search.fit(X_train, y_train)

        logging.debug('XGBoost RS fit took %s', str(time.time()-t1))

        xgb_random = xgb_search.best_estimator_

        # Test the model
        pred_xgb_val0 = xgb_random.predict(X_val)
        pred_xgb_prob_val0 = xgb_random.predict_proba(X_val)[:, 1]

        pred_xgb_train0 = xgb_random.predict(X_train)
        pred_xgb_prob_train0 = xgb_random.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_xgb_train0))
        print(classification_report(y_val, pred_xgb_val0))

        evaluate_model(y_val, pred_xgb_prob_val0, y_train,
                       pred_xgb_prob_train0, file_name+"_xgboost_RS_roc1.png")

        evaluate_model(y_val, pred_xgb_val0, y_train,
                       pred_xgb_train0, file_name+"_xgboost_RS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_xgb_train0,
                       file_name + "_xgboost_RS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_xgb_val0,  file_name +
                       "_xgboost_RS_val_data.png", "Validation")

        RS_roc_file1 = plot_dir + file_name + "_xgboost_RS_roc1.png"
        RS_roc_file2 = plot_dir + file_name + "_xgboost_RS_roc2.png"
        RS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_xgb_prob_val0, pred_xgb_val0)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_RS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_RS_roc2.png"),
            "ROC Curve Test data", "XGBoost - Random Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_RS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_RS_val_data.png"),
            "Confusion Matrix Validation data", "XGBoost - Random Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, RS_auc_ginin_output, "XGBoost - Random Search")

        paramFiles = param_file_path + param_file_name + "_XGB_GS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        t2 = time.time()
        estimator = xgb.XGBClassifier(random_state=50)

        # Create the random search model
        xgb_grid = GridSearchCV(estimator, param_grid, n_jobs=3,
                                scoring='roc_auc', cv=5, verbose=1)

        # Fit
        xgb_grid.fit(X_train, y_train)
        logging.debug('inside xgboost GS2 took %s', str(time.time()-t2))
        xgb_grid.best_params_

        xgb_final = xgb_grid.best_estimator_

        # Test the model
        pred_xgb_val1 = xgb_final.predict(X_val)
        pred_xgb_prob_val1 = xgb_final.predict_proba(X_val)[:, 1]

        pred_xgb_train1 = xgb_final.predict(X_train)
        pred_xgb_prob_train1 = xgb_final.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_xgb_train1))
        print(classification_report(y_val, pred_xgb_val1))

        evaluate_model(y_val, pred_xgb_prob_val1, y_train,
                       pred_xgb_prob_train1, file_name+"_xgboost_GS_roc1.png")

        evaluate_model(y_val, pred_xgb_val1, y_train,
                       pred_xgb_train1, file_name+"_xgboost_GS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_xgb_train1,
                       file_name + "_xgboost_GS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_xgb_val1,  file_name +
                       "_xgboost_GS_val_data.png", "Validation")

        GS_roc_file1 = plot_dir + file_name + "_xgboost_GS_roc1.png"
        GS_roc_file2 = plot_dir + file_name + "_xgboost_GS_roc2.png"
        GS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_xgb_prob_val1, pred_xgb_val1)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_GS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_GS_roc2.png"),
            "ROC Curve Test data", "XGBoost - Grid Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document, os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_GS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_GS_val_data.png"),
            "Confusion Matrix Validation data", "XGBoost - Grid Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, GS_auc_ginin_output, "XGBoost - Grid Search")

        pdf.output(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_XGBoost.pdf"))

        print('Process XGboost ended')
    except Exception as e:
        print(e)


def MLP_NoTunning(request):
    csv_file_name = "csvfile_"+user_name
    savefile_x_final = file_path + csv_file_name + "_x_model.csv"
    df = pd.read_csv(savefile_x_final)
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()

    # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
    y_model = df[targetVar]
    x_model = df.drop(targetVar, axis=1)
    X_train, X_test, y_train, y_test = train_test_split(
        x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
    X_train, X_val, y_train, y_val = train_test_split(
        X_train, y_train, test_size=0.22222222222222224, random_state=321)

    mlp = MLPClassifier(hidden_layer_sizes=100, alpha=0.0001, random_state=1,
                        learning_rate_init=0.001, max_iter=200)

    mlp.fit(X_train, y_train)
    # Test the model
    pred_mlp_val = mlp.predict(X_val)
    pred_mlp_prob_val = mlp.predict_proba(X_val)[:, 1]

    pred_mlp_train = mlp.predict(X_train)
    pred_mlp_prob_train = mlp.predict_proba(X_train)[:, 1]

    # Get the model performance
    print(classification_report(y_train, pred_mlp_train))
    print(classification_report(y_val, pred_mlp_val))

    evaluate_model(y_val, pred_mlp_prob_val, y_train,
                   pred_mlp_prob_train, file_name+"MLPNoTune_roc1.png")

    evaluate_model(y_val, pred_mlp_val, y_train, pred_mlp_train,
                   file_name+"MLPNoTune_roc2.png")

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    # show the confusion matrix for training data
    drawConfMatrix(y_train, pred_mlp_train, file_name +
                   "_MLPNoTune_train_data.png", "Training")

    drawConfMatrix(y_val, pred_mlp_val, file_name +
                   "_MLPNoTune_val_data.png", "Validation")

    auc_ginin_output = test_modelPerfomance(
        y_val,  pred_mlp_prob_val, pred_mlp_val)

    roc_file1 = plot_dir + file_name+"MLPNoTune_roc1.png"
    roc_file2 = plot_dir + file_name+"MLPNoTune_roc2.png"

    context = {'rocgraphpath1': roc_file1, 'rocgraphpath2': roc_file2, 'auc_ginin_output': auc_ginin_output,
               'graphConfMat1': plot_dir + file_name + "_MLPNoTune_train_data.png", 'graphConfMat2': plot_dir + file_name + "_MLPNoTune_val_data.png"}
    return render(request, 'showModelOutput.html', context)


def MLP_RS(request):
    csv_file_name = "csvfile_"+user_name
    savefile_x_final = file_path + csv_file_name + "_x_model.csv"
    df = pd.read_csv(savefile_x_final)
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()

    # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
    y_model = df[targetVar]
    x_model = df.drop(targetVar, axis=1)
    X_train, X_test, y_train, y_test = train_test_split(
        x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
    X_train, X_val, y_train, y_val = train_test_split(
        X_train, y_train, test_size=0.22222222222222224, random_state=321)

    paramFiles = param_file_path + param_file_name + "_MLP_RS.csv"
    param_grid = {}
    if os.path.exists(paramFiles):
        df = pd.read_csv(paramFiles)
        for index, row in df.iterrows():
            param_grid[row['paramName']] = eval(row['paramValue'])

   # Estimator for use in random search
    estimator = MLPClassifier(max_iter=100)
    # Create the random search model
    mlp_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                    scoring='roc_auc', cv=5,
                                    n_iter=100, verbose=1, random_state=50)

    # Fit
    mlp_search.fit(X_train, y_train)

    mlp_search.best_params_

    mlp_random = mlp_search.best_estimator_
    # Test the model
    pred_mlp_val0 = mlp_random.predict(X_val)
    pred_mlp_prob_val0 = mlp_random.predict_proba(X_val)[:, 1]

    pred_mlp_train0 = mlp_random.predict(X_train)
    pred_mlp_prob_train0 = mlp_random.predict_proba(X_train)[:, 1]

    # Get the model performance
    print(classification_report(y_train, pred_mlp_train0))
    print(classification_report(y_val, pred_mlp_val0))

    evaluate_model(y_val, pred_mlp_prob_val0, y_train,
                   pred_mlp_prob_train0, file_name+"_MLP_RS_roc1.png")

    evaluate_model(y_val, pred_mlp_val0, y_train,
                   pred_mlp_train0, file_name+"_MLP_RS_roc2.png")

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    # show the confusion matrix for training data

    drawConfMatrix(y_train, pred_mlp_train0,
                   file_name + "_MLP_RS_train_data.png", "Training")

    # show the confusion matrix for validation data

    drawConfMatrix(y_val, pred_mlp_val0,   file_name +
                   "_MLP_RS_val_data.png", "Validation")

    roc_file1 = plot_dir + file_name + "_MLP_RS_roc1.png"
    roc_file2 = plot_dir + file_name + "_MLP_RS_roc2.png"
    auc_ginin_output = test_modelPerfomance(
        y_val,  pred_mlp_prob_val0, pred_mlp_val0)

    context = {'rocgraphpath1': roc_file1, 'rocgraphpath2': roc_file2, 'auc_ginin_output': auc_ginin_output,
               'graphConfMat1': plot_dir + file_name + "_MLP_RS_train_data.png", 'graphConfMat2': plot_dir + file_name + "_MLP_RS_val_data.png"}
    return render(request, 'showModelOutput.html', context)


def GBC_NoTunning(request):
    csv_file_name = "csvfile_"+user_name
    savefile_x_final = file_path + csv_file_name + "_x_model.csv"
    df = pd.read_csv(savefile_x_final)
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()

    # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
    y_model = df[targetVar]
    x_model = df.drop(targetVar, axis=1)
    X_train, X_test, y_train, y_test = train_test_split(
        x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
    X_train, X_val, y_train, y_val = train_test_split(
        X_train, y_train, test_size=0.22222222222222224, random_state=321)

    gbc = GradientBoostingClassifier()
    gbc.fit(X_train, y_train)

    # Test the model
    pred_gbc_val = gbc.predict(X_val)
    pred_gbc_prob_val = gbc.predict_proba(X_val)[:, 1]

    pred_gbc_train = gbc.predict(X_train)
    pred_gbc_prob_train = gbc.predict_proba(X_train)[:, 1]

    # Get the model performance
    print(classification_report(y_train, pred_gbc_train))
    print(classification_report(y_val, pred_gbc_val))

    evaluate_model(y_val, pred_gbc_prob_val, y_train,
                   pred_gbc_prob_train, file_name+"GBC_NT_roc1.png")

    evaluate_model(y_val, pred_gbc_val, y_train, pred_gbc_train,
                   file_name+"GBC_NT_roc2.png")

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    # show the confusion matrix for training data
    drawConfMatrix(y_train, pred_gbc_train, file_name +
                   "_GBC_NT_train_data.png", "Training")

    drawConfMatrix(y_val, pred_gbc_val, file_name +
                   "_GBC_NT_val_data.png", "Validation")

    auc_ginin_output = test_modelPerfomance(
        y_val,  pred_gbc_prob_val, pred_gbc_val)

    roc_file1 = plot_dir + file_name+"GBC_NT_roc1.png"
    roc_file2 = plot_dir + file_name+"GBC_NT_roc2.png"

    context = {'rocgraphpath1': roc_file1, 'rocgraphpath2': roc_file2, 'auc_ginin_output': auc_ginin_output,
               'graphConfMat1': plot_dir + file_name + "_GBC_NT_train_data.png", 'graphConfMat2': plot_dir + file_name + "_GBC_NT_val_data.png"}
    return render(request, 'showModelOutput.html', context)


def MLP_GS(request):
    csv_file_name = "csvfile_"+user_name
    savefile_x_final = file_path + csv_file_name + "_x_model.csv"
    df = pd.read_csv(savefile_x_final)
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()

    # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
    y_model = df[targetVar]
    x_model = df.drop(targetVar, axis=1)
    X_train, X_test, y_train, y_test = train_test_split(
        x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
    X_train, X_val, y_train, y_val = train_test_split(
        X_train, y_train, test_size=0.22222222222222224, random_state=321)

    paramFiles = param_file_path + param_file_name + "_MLP_RS.csv"
    param_grid = {}
    if os.path.exists(paramFiles):
        df = pd.read_csv(paramFiles)
        for index, row in df.iterrows():
            param_grid[row['paramName']] = eval(row['paramValue'])

   # Estimator for use in random search
    estimator = MLPClassifier(max_iter=100)

    # Create the random search model
    mlp_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                    scoring='roc_auc', cv=5,
                                    n_iter=100, verbose=1, random_state=50)

    # Fit
    mlp_search.fit(X_train, y_train)

    paramFiles = param_file_path + param_file_name + "_MLP_GS.csv"
    param_grid = {}
    if os.path.exists(paramFiles):
        df = pd.read_csv(paramFiles)
        for index, row in df.iterrows():
            param_grid[row['paramName']] = eval(row['paramValue'])

    # Estimator for use in random search
    estimator = MLPClassifier(max_iter=500)

    # Create the random search model
    mlp_grid = GridSearchCV(estimator, param_grid, n_jobs=3,
                            scoring='roc_auc', cv=5, verbose=1)

    # Fit
    mlp_grid.fit(X_train, y_train)

    mlp_grid.best_params_

    mlp_final = mlp_grid.best_estimator_

    # Test the model
    pred_mlp_val1 = mlp_final.predict(X_val)
    pred_mlp_prob_val1 = mlp_final.predict_proba(X_val)[:, 1]

    pred_mlp_train1 = mlp_final.predict(X_train)
    pred_mlp_prob_train1 = mlp_final.predict_proba(X_train)[:, 1]

    # Get the model performance
    print(classification_report(y_train, pred_mlp_train1))
    print(classification_report(y_val, pred_mlp_val1))

    evaluate_model(y_val, pred_mlp_prob_val1, y_train,
                   pred_mlp_prob_train1, file_name+"_MLP_GS_roc1.png")

    evaluate_model(y_val, pred_mlp_val1, y_train,
                   pred_mlp_train1, file_name+"_MLP_GS_roc2.png")

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    # show the confusion matrix for training data

    drawConfMatrix(y_train, pred_mlp_train1,
                   file_name + "_MLP_GS_train_data.png", "Training")

    # show the confusion matrix for validation data

    drawConfMatrix(y_val, pred_mlp_val1,   file_name +
                   "_MLP_GS_val_data.png", "Validation")

    roc_file1 = plot_dir + file_name + "_MLP_GS_roc1.png"
    roc_file2 = plot_dir + file_name + "_MLP_GS_roc2.png"
    auc_ginin_output = test_modelPerfomance(
        y_val,  pred_mlp_prob_val1, pred_mlp_val1)

    context = {'rocgraphpath1': roc_file1, 'rocgraphpath2': roc_file2, 'auc_ginin_output': auc_ginin_output,
               'graphConfMat1': plot_dir + file_name + "_MLP_GS_train_data.png", 'graphConfMat2': plot_dir + file_name + "_MLP_GS_val_data.png"}
    return render(request, 'showModelOutput.html', context)


def GBC_RS(request):
    csv_file_name = "csvfile_"+user_name
    savefile_x_final = file_path + csv_file_name + "_x_model.csv"
    df = pd.read_csv(savefile_x_final)
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()

    # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
    y_model = df[targetVar]
    x_model = df.drop(targetVar, axis=1)
    X_train, X_test, y_train, y_test = train_test_split(
        x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
    X_train, X_val, y_train, y_val = train_test_split(
        X_train, y_train, test_size=0.22222222222222224, random_state=321)

    paramFiles = param_file_path + param_file_name + "_GB_RS.csv"
    param_grid = {}
    if os.path.exists(paramFiles):
        df = pd.read_csv(paramFiles)
        for index, row in df.iterrows():
            param_grid[row['paramName']] = eval(row['paramValue'])

    seed(1)
   # Estimator for use in random search
    estimator = GradientBoostingClassifier()

    # Create the random search model
    gbc_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                    scoring='roc_auc', cv=5,
                                    n_iter=100, verbose=1, random_state=50)

    # Fit
    gbc_search.fit(X_train, y_train)

    gbc_random = gbc_search.best_estimator_
    # Test the model
    pred_gbc_val0 = gbc_random.predict(X_val)
    pred_gbc_prob_val0 = gbc_random.predict_proba(X_val)[:, 1]

    pred_gbc_train0 = gbc_random.predict(X_train)
    pred_gbc_prob_train0 = gbc_random.predict_proba(X_train)[:, 1]

    # Get the model performance
    print(classification_report(y_train, pred_gbc_train0))
    print(classification_report(y_val, pred_gbc_val0))

    evaluate_model(y_val, pred_gbc_prob_val0, y_train,
                   pred_gbc_prob_train0, file_name+"_GBC_RS_roc1.png")

    evaluate_model(y_val, pred_gbc_val0, y_train,
                   pred_gbc_train0, file_name+"_GBC_RS_roc2.png")

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    # show the confusion matrix for training data

    drawConfMatrix(y_train, pred_gbc_train0,
                   file_name + "_GBC_RS_train_data.png", "Training")

    # show the confusion matrix for validation data

    drawConfMatrix(y_val, pred_gbc_val0,   file_name +
                   "_GBC_RS_val_data.png", "Validation")

    roc_file1 = plot_dir + file_name + "_GBC_RS_roc1.png"
    roc_file2 = plot_dir + file_name + "_GBC_RS_roc2.png"
    auc_ginin_output = test_modelPerfomance(
        y_val,  pred_gbc_prob_val0, pred_gbc_val0)

    context = {'rocgraphpath1': roc_file1, 'rocgraphpath2': roc_file2, 'auc_ginin_output': auc_ginin_output,
               'graphConfMat1': plot_dir + file_name + "_GBC_RS_train_data.png", 'graphConfMat2': plot_dir + file_name + "_GBC_RS_val_data.png"}
    return render(request, 'showModelOutput.html', context)


def GBC_GS(request):
    csv_file_name = "csvfile_"+user_name
    savefile_x_final = file_path + csv_file_name + "_x_model.csv"
    df = pd.read_csv(savefile_x_final)
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()

    # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
    y_model = df[targetVar]
    x_model = df.drop(targetVar, axis=1)
    X_train, X_test, y_train, y_test = train_test_split(
        x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
    X_train, X_val, y_train, y_val = train_test_split(
        X_train, y_train, test_size=0.22222222222222224, random_state=321)

    paramFiles = param_file_path + param_file_name + "_GB_RS.csv"
    param_grid = {}
    if os.path.exists(paramFiles):
        df = pd.read_csv(paramFiles)
        for index, row in df.iterrows():
            param_grid[row['paramName']] = eval(row['paramValue'])

    # Estimator for use in random search
    estimator = GradientBoostingClassifier()

    # Create the random search model
    gbc_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                    scoring='roc_auc', cv=5,
                                    n_iter=100, verbose=1, random_state=50)

    # Fit
    gbc_search.fit(X_train, y_train)

    paramFiles = param_file_path + param_file_name + "_GB_GS.csv"
    param_grid = {}
    if os.path.exists(paramFiles):
        df = pd.read_csv(paramFiles)
        for index, row in df.iterrows():
            param_grid[row['paramName']] = eval(row['paramValue'])

    # Estimator for use in random search
    estimator = GradientBoostingClassifier()

    # Create the random search model
    gbc_grid = GridSearchCV(estimator, param_grid, n_jobs=3,
                            scoring='roc_auc', cv=5, verbose=1)

    # Fit
    gbc_grid.fit(X_train, y_train)

    gbc_final = gbc_grid.best_estimator_

    # Test the model
    pred_gbc_val1 = gbc_final.predict(X_val)
    pred_gbc_prob_val1 = gbc_final.predict_proba(X_val)[:, 1]

    pred_gbc_train1 = gbc_final.predict(X_train)
    pred_gbc_prob_train1 = gbc_final.predict_proba(X_train)[:, 1]
    # Get the model performance
    print(classification_report(y_train, pred_gbc_train1))
    print(classification_report(y_val, pred_gbc_val1))

    evaluate_model(y_val, pred_gbc_prob_val1, y_train,
                   pred_gbc_prob_train1, file_name+"_GBC_GS_roc1.png")

    evaluate_model(y_val, pred_gbc_val1, y_train,
                   pred_gbc_train1, file_name+"_GBC_GS_roc2.png")

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    # show the confusion matrix for training data

    drawConfMatrix(y_train, pred_gbc_train1,
                   file_name + "_GBC_GS_train_data.png", "Training")

    # show the confusion matrix for validation data

    drawConfMatrix(y_val, pred_gbc_val1,   file_name +
                   "_GBC_GS_val_data.png", "Validation")

    roc_file1 = plot_dir + file_name + "_GBC_GS_roc1.png"
    roc_file2 = plot_dir + file_name + "_GBC_GS_roc2.png"
    auc_ginin_output = test_modelPerfomance(
        y_val,  pred_gbc_prob_val1, pred_gbc_val1)

    context = {'rocgraphpath1': roc_file1, 'rocgraphpath2': roc_file2, 'auc_ginin_output': auc_ginin_output,
               'graphConfMat1': plot_dir + file_name + "_GBC_GS_train_data.png", 'graphConfMat2': plot_dir + file_name + "_GBC_GS_val_data.png"}
    return render(request, 'showModelOutput.html', context)


def KNN_NT(request):
    csv_file_name = "csvfile_"+user_name
    savefile_x_final = file_path + csv_file_name + "_x_model.csv"
    df = pd.read_csv(savefile_x_final)
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()

    # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
    y_model = df[targetVar]
    x_model = df.drop(targetVar, axis=1)
    X_train, X_test, y_train, y_test = train_test_split(
        x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
    X_train, X_val, y_train, y_val = train_test_split(
        X_train, y_train, test_size=0.22222222222222224, random_state=321)

    knn = KNeighborsClassifier()
    knn.fit(X_train, y_train)

    # Test the model
    pred_knn_val = knn.predict(X_val)
    pred_knn_prob_val = knn.predict_proba(X_val)[:, 1]

    pred_knn_train = knn.predict(X_train)
    pred_knn_prob_train = knn.predict_proba(X_train)[:, 1]

    # Get the model performance
    print(classification_report(y_train, pred_knn_train))
    print(classification_report(y_val, pred_knn_val))

    evaluate_model(y_val, pred_knn_prob_val, y_train,
                   pred_knn_prob_train, file_name+"KNN_NT_roc1.png")

    evaluate_model(y_val, pred_knn_val, y_train,
                   pred_knn_train, file_name+"KNN_NT_roc2.png")

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    # show the confusion matrix for training data
    drawConfMatrix(y_train, pred_knn_train, file_name +
                   "_KNN_NT_train_data.png", "Training")

    drawConfMatrix(y_val, pred_knn_val, file_name +
                   "_KNN_NT_val_data.png", "Validation")

    auc_ginin_output = test_modelPerfomance(
        y_val,  pred_knn_prob_val, pred_knn_val)

    roc_file1 = plot_dir + file_name+"KNN_NT_roc1.png"
    roc_file2 = plot_dir + file_name+"KNN_NT_roc2.png"

    context = {'rocgraphpath1': roc_file1, 'rocgraphpath2': roc_file2, 'auc_ginin_output': auc_ginin_output,
               'graphConfMat1': plot_dir + file_name + "_KNN_NT_train_data.png", 'graphConfMat2': plot_dir + file_name + "_KNN_NT_val_data.png"}
    return render(request, 'showModelOutput.html', context)


def KNN_RS(request):
    csv_file_name = "csvfile_"+user_name
    savefile_x_final = file_path + csv_file_name + "_x_model.csv"
    df = pd.read_csv(savefile_x_final)
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()

    # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
    y_model = df[targetVar]
    x_model = df.drop(targetVar, axis=1)
    X_train, X_test, y_train, y_test = train_test_split(
        x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
    X_train, X_val, y_train, y_val = train_test_split(
        X_train, y_train, test_size=0.22222222222222224, random_state=321)

    paramFiles = param_file_path + param_file_name + "_KNN_RS.csv"
    param_grid = {}
    if os.path.exists(paramFiles):
        df = pd.read_csv(paramFiles)
        for index, row in df.iterrows():
            param_grid[row['paramName']] = eval(row['paramValue'])

    # Estimator for use in random search
    estimator = KNeighborsClassifier()

    # Create the random search model
    knn_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                    scoring='roc_auc', cv=5,
                                    n_iter=100, verbose=1, random_state=50)

    # Fit
    knn_search.fit(X_train, y_train)

    knn_random = knn_search.best_estimator_
    # Test the model
    pred_knn_val0 = knn_random.predict(X_val)
    pred_knn_prob_val0 = knn_random.predict_proba(X_val)[:, 1]

    pred_knn_train0 = knn_random.predict(X_train)
    pred_knn_prob_train0 = knn_random.predict_proba(X_train)[:, 1]

    # Get the model performance
    print(classification_report(y_train, pred_knn_train0))
    print(classification_report(y_val, pred_knn_val0))

    evaluate_model(y_val, pred_knn_prob_val0, y_train,
                   pred_knn_prob_train0, file_name+"_KNN_RS_roc1.png")

    evaluate_model(y_val, pred_knn_val0, y_train,
                   pred_knn_train0, file_name+"_KNN_RS_roc2.png")

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    # show the confusion matrix for training data

    drawConfMatrix(y_train, pred_knn_train0,
                   file_name + "_KNN_RS_train_data.png", "Training")

    # show the confusion matrix for validation data

    drawConfMatrix(y_val, pred_knn_val0,   file_name +
                   "_KNN_RS_val_data.png", "Validation")

    roc_file1 = plot_dir + file_name + "_KNN_RS_roc1.png"
    roc_file2 = plot_dir + file_name + "_KNN_RS_roc2.png"
    auc_ginin_output = test_modelPerfomance(
        y_val,  pred_knn_prob_val0, pred_knn_val0)

    context = {'rocgraphpath1': roc_file1, 'rocgraphpath2': roc_file2, 'auc_ginin_output': auc_ginin_output,
               'graphConfMat1': plot_dir + file_name + "_KNN_RS_train_data.png", 'graphConfMat2': plot_dir + file_name + "_KNN_RS_val_data.png"}
    return render(request, 'showModelOutput.html', context)


def KNN_GS(request):
    csv_file_name = "csvfile_"+user_name
    savefile_x_final = file_path + csv_file_name + "_x_model.csv"
    df = pd.read_csv(savefile_x_final)
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()

    # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
    y_model = df[targetVar]
    x_model = df.drop(targetVar, axis=1)
    X_train, X_test, y_train, y_test = train_test_split(
        x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
    X_train, X_val, y_train, y_val = train_test_split(
        X_train, y_train, test_size=0.22222222222222224, random_state=321)

    paramFiles = param_file_path + param_file_name + "_KNN_RS.csv"
    param_grid = {}
    if os.path.exists(paramFiles):
        df = pd.read_csv(paramFiles)
        for index, row in df.iterrows():
            param_grid[row['paramName']] = eval(row['paramValue'])

    # Estimator for use in random search
    estimator = KNeighborsClassifier()

    # Create the random search model
    knn_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                    scoring='roc_auc', cv=5,
                                    n_iter=100, verbose=1, random_state=50)

    # Fit
    knn_search.fit(X_train, y_train)

    paramFiles = param_file_path + param_file_name + "_KNN_GS.csv"
    param_grid = {}
    if os.path.exists(paramFiles):
        df = pd.read_csv(paramFiles)
        for index, row in df.iterrows():
            param_grid[row['paramName']] = eval(row['paramValue'])

    # Estimator for use in random search
    estimator = KNeighborsClassifier()

    # Create the random search model
    knn_grid = GridSearchCV(estimator, param_grid, n_jobs=3,
                            scoring='roc_auc', cv=5, verbose=1)

    # Fit
    knn_grid.fit(X_train, y_train)

    knn_final = knn_grid.best_estimator_

    # Test the model
    pred_knn_val1 = knn_final.predict(X_val)
    pred_knn_prob_val1 = knn_final.predict_proba(X_val)[:, 1]

    pred_knn_train1 = knn_final.predict(X_train)
    pred_knn_prob_train1 = knn_final.predict_proba(X_train)[:, 1]

    # Get the model performance
    print(classification_report(y_train, pred_knn_train1))
    print(classification_report(y_val, pred_knn_val1))

    evaluate_model(y_val, pred_knn_prob_val1, y_train,
                   pred_knn_prob_train1, file_name+"_KNN_GS_roc1.png")

    evaluate_model(y_val, pred_knn_val1, y_train,
                   pred_knn_train1, file_name+"_KNN_GS_roc2.png")

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    # show the confusion matrix for training data

    drawConfMatrix(y_train, pred_knn_train1,
                   file_name + "_KNN_GS_train_data.png", "Training")

    # show the confusion matrix for validation data

    drawConfMatrix(y_val, pred_knn_val1,   file_name +
                   "_KNN_GS_val_data.png", "Validation")

    roc_file1 = plot_dir + file_name + "_KNN_GS_roc1.png"
    roc_file2 = plot_dir + file_name + "_KNN_GS_roc2.png"
    auc_ginin_output = test_modelPerfomance(
        y_val,  pred_knn_prob_val1, pred_knn_val1)

    context = {'rocgraphpath1': roc_file1, 'rocgraphpath2': roc_file2, 'auc_ginin_output': auc_ginin_output,
               'graphConfMat1': plot_dir + file_name + "_KNN_GS_train_data.png", 'graphConfMat2': plot_dir + file_name + "_KNN_GS_val_data.png"}
    return render(request, 'showModelOutput.html', context)


def SVM_NT(request):
    csv_file_name = "csvfile_"+user_name
    savefile_x_final = file_path + csv_file_name + "_x_model.csv"
    df = pd.read_csv(savefile_x_final)
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()

    # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
    y_model = df[targetVar]
    x_model = df.drop(targetVar, axis=1)
    X_train, X_test, y_train, y_test = train_test_split(
        x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
    X_train, X_val, y_train, y_val = train_test_split(
        X_train, y_train, test_size=0.22222222222222224, random_state=321)

    svm_clf = svm.SVC(probability=True)
    svm_clf.fit(X_train, y_train)

    # Test the model
    pred_svm_val = svm_clf.predict(X_val)
    pred_svm_prob_val = svm_clf.predict_proba(X_val)[:, 1]

    pred_svm_train = svm_clf.predict(X_train)
    pred_svm_prob_train = svm_clf.predict_proba(X_train)[:, 1]

    # Get the model performance
    print(classification_report(y_train, pred_svm_train))
    print(classification_report(y_val, pred_svm_val))

    evaluate_model(y_val, pred_svm_prob_val, y_train,
                   pred_svm_prob_train, file_name+"SVM_NT_roc1.png")

    evaluate_model(y_val, pred_svm_val, y_train,
                   pred_svm_train, file_name+"SVM_NT_roc2.png")

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    # show the confusion matrix for training data
    drawConfMatrix(y_train, pred_svm_train, file_name +
                   "_SVM_NT_train_data.png", "Training")

    drawConfMatrix(y_val, pred_svm_val, file_name +
                   "_SVM_NT_val_data.png", "Validation")

    auc_ginin_output = test_modelPerfomance(
        y_val,  pred_svm_prob_val, pred_svm_val)

    roc_file1 = plot_dir + file_name+"SVM_NT_roc1.png"
    roc_file2 = plot_dir + file_name+"SVM_NT_roc2.png"

    context = {'rocgraphpath1': roc_file1, 'rocgraphpath2': roc_file2, 'auc_ginin_output': auc_ginin_output,
               'graphConfMat1': plot_dir + file_name + "_SVM_NT_train_data.png", 'graphConfMat2': plot_dir + file_name + "_SVM_NT_val_data.png"}
    return render(request, 'showModelOutput.html', context)


def SVM_RS(request):
    csv_file_name = "csvfile_"+user_name
    savefile_x_final = file_path + csv_file_name + "_x_model.csv"
    df = pd.read_csv(savefile_x_final)
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()

    # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
    y_model = df[targetVar]
    x_model = df.drop(targetVar, axis=1)
    X_train, X_test, y_train, y_test = train_test_split(
        x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
    X_train, X_val, y_train, y_val = train_test_split(
        X_train, y_train, test_size=0.22222222222222224, random_state=321)

    paramFiles = param_file_path + param_file_name + "_SVM_RS.csv"
    param_grid = {}
    if os.path.exists(paramFiles):
        df = pd.read_csv(paramFiles)
        for index, row in df.iterrows():
            param_grid[row['paramName']] = eval(row['paramValue'])
    param_grid1 = {
        'C':  list(np.arange(0.01, 1, 0.1)),  # list(np.arange(0.01, 1, 0.1)),
        'kernel': ['linear', 'poly', 'rbf', 'sigmoid'],
        'gamma': ['scale', 'auto'],
        'probability': [True],
        'class_weight': [dict, 'balanced', None]
    }

    print('param_grid')
    print(param_grid)
    print('param_grid1')
    print(param_grid1)
    # Estimator for use in random search
    estimator = svm.SVC()

    # Create the random search model
    svm_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                    scoring='roc_auc', cv=5,
                                    n_iter=100, verbose=1, random_state=50)

    # Fit
    svm_search.fit(X_train, y_train)

    svm_random = svm_search.best_estimator_

    # Test the model
    pred_svm_val0 = svm_random.predict(X_val)
    pred_svm_prob_val0 = svm_random.predict_proba(X_val)[:, 1]

    pred_svm_train0 = svm_random.predict(X_train)
    pred_svm_prob_train0 = svm_random.predict_proba(X_train)[:, 1]

    # Get the model performance
    print(classification_report(y_train, pred_svm_train0))
    print(classification_report(y_val, pred_svm_val0))

    evaluate_model(y_val, pred_svm_prob_val0, y_train,
                   pred_svm_prob_train0, file_name+"_SVM_RS_roc1.png")

    evaluate_model(y_val, pred_svm_val0, y_train,
                   pred_svm_train0, file_name+"_SVM_RS_roc2.png")

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    # show the confusion matrix for training data

    drawConfMatrix(y_train, pred_svm_train0,
                   file_name + "_SVM_RS_train_data.png", "Training")

    # show the confusion matrix for validation data

    drawConfMatrix(y_val, pred_svm_val0,   file_name +
                   "_SVM_RS_val_data.png", "Validation")

    roc_file1 = plot_dir + file_name + "_SVM_RS_roc1.png"
    roc_file2 = plot_dir + file_name + "_SVM_RS_roc2.png"
    auc_ginin_output = test_modelPerfomance(
        y_val,  pred_svm_prob_val0, pred_svm_val0)

    context = {'rocgraphpath1': roc_file1, 'rocgraphpath2': roc_file2, 'auc_ginin_output': auc_ginin_output,
               'graphConfMat1': plot_dir + file_name + "_SVM_RS_train_data.png", 'graphConfMat2': plot_dir + file_name + "_SVM_RS_val_data.png"}
    return render(request, 'showModelOutput.html', context)


def SVM_GS(request):
    csv_file_name = "csvfile_"+user_name
    savefile_x_final = file_path + csv_file_name + "_x_model.csv"
    df = pd.read_csv(savefile_x_final)
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()

    # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
    y_model = df[targetVar]
    x_model = df.drop(targetVar, axis=1)
    X_train, X_test, y_train, y_test = train_test_split(
        x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
    X_train, X_val, y_train, y_val = train_test_split(
        X_train, y_train, test_size=0.22222222222222224, random_state=321)

    paramFiles = param_file_path + param_file_name + "_SVM_RS.csv"
    param_grid = {}
    if os.path.exists(paramFiles):
        df = pd.read_csv(paramFiles)
        for index, row in df.iterrows():
            param_grid[row['paramName']] = eval(row['paramValue'])

    # Estimator for use in random search
    estimator = svm.SVC()

    # Create the random search model
    svm_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                    scoring='roc_auc', cv=5,
                                    n_iter=100, verbose=1, random_state=50)

    # Fit
    svm_search.fit(X_train, y_train)

    paramFiles = param_file_path + param_file_name + "_SVM_GS.csv"
    param_grid = {}
    if os.path.exists(paramFiles):
        df = pd.read_csv(paramFiles)
        for index, row in df.iterrows():
            param_grid[row['paramName']] = eval(row['paramValue'])

    # Estimator for use in random search
    estimator = svm.SVC()

    # Create the random search model
    svm_grid = GridSearchCV(estimator, param_grid, n_jobs=3,
                            scoring='roc_auc', cv=5, verbose=1)

    # Fit
    svm_grid.fit(X_train, y_train)

    svm_final = svm_grid.best_estimator_

    # Test the model
    pred_svm_val1 = svm_final.predict(X_val)
    pred_svm_prob_val1 = svm_final.predict_proba(X_val)[:, 1]

    pred_svm_train1 = svm_final.predict(X_train)
    pred_svm_prob_train1 = svm_final.predict_proba(X_train)[:, 1]

    # Get the model performance
    print(classification_report(y_train, pred_svm_train1))
    print(classification_report(y_val, pred_svm_val1))

    evaluate_model(y_val, pred_svm_prob_val1, y_train,
                   pred_svm_prob_train1, file_name+"_SVM_GS_roc1.png")

    evaluate_model(y_val, pred_svm_val1, y_train,
                   pred_svm_train1, file_name+"_SVM_GS_roc2.png")

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    # show the confusion matrix for training data

    drawConfMatrix(y_train, pred_svm_train1,
                   file_name + "_SVM_GS_train_data.png", "Training")

    # show the confusion matrix for validation data

    drawConfMatrix(y_val, pred_svm_val1,   file_name +
                   "_SVM_GS_val_data.png", "Validation")

    roc_file1 = plot_dir + file_name + "_SVM_GS_roc1.png"
    roc_file2 = plot_dir + file_name + "_SVM_GS_roc2.png"
    auc_ginin_output = test_modelPerfomance(
        y_val,  pred_svm_prob_val1, pred_svm_val1)

    context = {'rocgraphpath1': roc_file1, 'rocgraphpath2': roc_file2, 'auc_ginin_output': auc_ginin_output,
               'graphConfMat1': plot_dir + file_name + "_SVM_GS_train_data.png", 'graphConfMat2': plot_dir + file_name + "_SVM_GS_val_data.png"}
    return render(request, 'showModelOutput.html', context)


def BC_NT(request):
    csv_file_name = "csvfile_"+user_name
    savefile_x_final = file_path + csv_file_name + "_x_model.csv"
    df = pd.read_csv(savefile_x_final)
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()

    # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
    y_model = df[targetVar]
    x_model = df.drop(targetVar, axis=1)
    X_train, X_test, y_train, y_test = train_test_split(
        x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
    X_train, X_val, y_train, y_val = train_test_split(
        X_train, y_train, test_size=0.22222222222222224, random_state=321)

    bc_model = BaggingClassifier()
    bc_model.fit(X_train, y_train)

    # Test the model
    pred_bc_val = bc_model.predict(X_val)
    pred_bc_prob_val = bc_model.predict_proba(X_val)[:, 1]

    pred_bc_train = bc_model.predict(X_train)
    pred_bc_prob_train = bc_model.predict_proba(X_train)[:, 1]

    # Get the model performance
    print(classification_report(y_train, pred_bc_train))
    print(classification_report(y_val, pred_bc_val))

    evaluate_model(y_val, pred_bc_prob_val, y_train,
                   pred_bc_prob_train, file_name+"BC_NT_roc1.png")

    evaluate_model(y_val, pred_bc_val, y_train,
                   pred_bc_train, file_name+"BC_NT_roc2.png")

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    # show the confusion matrix for training data
    drawConfMatrix(y_train, pred_bc_train, file_name +
                   "_BC_NT_train_data.png", "Training")

    drawConfMatrix(y_val, pred_bc_val, file_name +
                   "_BC_NT_val_data.png", "Validation")

    auc_ginin_output = test_modelPerfomance(
        y_val,  pred_bc_prob_val, pred_bc_val)

    roc_file1 = plot_dir + file_name+"BC_NT_roc1.png"
    roc_file2 = plot_dir + file_name+"BC_NT_roc2.png"

    context = {'rocgraphpath1': roc_file1, 'rocgraphpath2': roc_file2, 'auc_ginin_output': auc_ginin_output,
               'graphConfMat1': plot_dir + file_name + "_BC_NT_train_data.png", 'graphConfMat2': plot_dir + file_name + "_BC_NT_val_data.png"}
    return render(request, 'showModelOutput.html', context)


def BC_RS(request):
    csv_file_name = "csvfile_"+user_name
    savefile_x_final = file_path + csv_file_name + "_x_model.csv"
    df = pd.read_csv(savefile_x_final)
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()

    # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
    y_model = df[targetVar]
    x_model = df.drop(targetVar, axis=1)
    X_train, X_test, y_train, y_test = train_test_split(
        x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
    X_train, X_val, y_train, y_val = train_test_split(
        X_train, y_train, test_size=0.22222222222222224, random_state=321)

    paramFiles = param_file_path + param_file_name + "_BC_RS.csv"
    param_grid = {}
    if os.path.exists(paramFiles):
        df = pd.read_csv(paramFiles)
        for index, row in df.iterrows():
            param_grid[row['paramName']] = eval(row['paramValue'])

   # Estimator for use in random search
    estimator = BaggingClassifier(random_state=50)

    # Create the random search model
    bc_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                   scoring='roc_auc', cv=5,
                                   n_iter=100, verbose=1, random_state=50)

    # Fit
    bc_search.fit(X_train, y_train)

    bc_random = bc_search.best_estimator_

    # Test the model
    pred_bc_val0 = bc_random.predict(X_val)
    pred_bc_prob_val0 = bc_random.predict_proba(X_val)[:, 1]

    pred_bc_train0 = bc_random.predict(X_train)
    pred_bc_prob_train0 = bc_random.predict_proba(X_train)[:, 1]

    # Get the model performance
    print(classification_report(y_train, pred_bc_train0))
    print(classification_report(y_val, pred_bc_val0))

    evaluate_model(y_val, pred_bc_prob_val0, y_train,
                   pred_bc_prob_train0, file_name+"_BC_RS_roc1.png")

    evaluate_model(y_val, pred_bc_val0, y_train,
                   pred_bc_train0, file_name+"_BC_RS_roc2.png")

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    # show the confusion matrix for training data

    drawConfMatrix(y_train, pred_bc_train0,
                   file_name + "_BC_RS_train_data.png", "Training")

    # show the confusion matrix for validation data

    drawConfMatrix(y_val, pred_bc_val0,   file_name +
                   "_BC_RS_val_data.png", "Validation")

    roc_file1 = plot_dir + file_name + "_BC_RS_roc1.png"
    roc_file2 = plot_dir + file_name + "_BC_RS_roc2.png"
    auc_ginin_output = test_modelPerfomance(
        y_val,  pred_bc_prob_val0, pred_bc_val0)

    context = {'rocgraphpath1': roc_file1, 'rocgraphpath2': roc_file2, 'auc_ginin_output': auc_ginin_output,
               'graphConfMat1': plot_dir + file_name + "_BC_RS_train_data.png", 'graphConfMat2': plot_dir + file_name + "_BC_RS_val_data.png"}
    return render(request, 'showModelOutput.html', context)


def BC_GS(request):
    csv_file_name = "csvfile_"+user_name
    savefile_x_final = file_path + csv_file_name + "_x_model.csv"
    df = pd.read_csv(savefile_x_final)
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()

    # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
    y_model = df[targetVar]
    x_model = df.drop(targetVar, axis=1)
    X_train, X_test, y_train, y_test = train_test_split(
        x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
    X_train, X_val, y_train, y_val = train_test_split(
        X_train, y_train, test_size=0.22222222222222224, random_state=321)

    paramFiles = param_file_path + param_file_name + "_BC_RS.csv"
    param_grid = {}
    if os.path.exists(paramFiles):
        df = pd.read_csv(paramFiles)
        for index, row in df.iterrows():
            param_grid[row['paramName']] = eval(row['paramValue'])

   # Estimator for use in random search
    estimator = BaggingClassifier(random_state=50)

    # Create the random search model
    bc_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                   scoring='roc_auc', cv=5,
                                   n_iter=100, verbose=1, random_state=50)

    # Fit
    bc_search.fit(X_train, y_train)

    paramFiles = param_file_path + param_file_name + "_BC_GS.csv"
    param_grid = {}
    if os.path.exists(paramFiles):
        df = pd.read_csv(paramFiles)
        for index, row in df.iterrows():
            param_grid[row['paramName']] = eval(row['paramValue'])

    # Estimator for use in random search
    estimator = BaggingClassifier(random_state=50)

    # Create the random search model
    bc_grid = GridSearchCV(estimator, param_grid, n_jobs=3,
                           scoring='roc_auc', cv=5, verbose=1)

    # Fit
    bc_grid.fit(X_train, y_train)

    bc_final = bc_grid.best_estimator_

    pred_bc_val1 = bc_final.predict(X_val)
    pred_bc_prob_val1 = bc_final.predict_proba(X_val)[:, 1]

    pred_bc_train1 = bc_final.predict(X_train)
    pred_bc_prob_train1 = bc_final.predict_proba(X_train)[:, 1]

    # Get the model performance
    print(classification_report(y_train, pred_bc_train1))
    print(classification_report(y_val, pred_bc_val1))

    evaluate_model(y_val, pred_bc_prob_val1, y_train,
                   pred_bc_prob_train1, file_name+"_BC_GS_roc1.png")

    evaluate_model(y_val, pred_bc_val1, y_train,
                   pred_bc_train1, file_name+"_BC_GS_roc2.png")

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    # show the confusion matrix for training data

    drawConfMatrix(y_train, pred_bc_train1,
                   file_name + "_BC_GS_train_data.png", "Training")

    # show the confusion matrix for validation data

    drawConfMatrix(y_val, pred_bc_val1,   file_name +
                   "_BC_GS_val_data.png", "Validation")

    roc_file1 = plot_dir + file_name + "_BC_GS_roc1.png"
    roc_file2 = plot_dir + file_name + "_BC_GS_roc2.png"
    auc_ginin_output = test_modelPerfomance(
        y_val,  pred_bc_prob_val1, pred_bc_val1)

    context = {'rocgraphpath1': roc_file1, 'rocgraphpath2': roc_file2, 'auc_ginin_output': auc_ginin_output,
               'graphConfMat1': plot_dir + file_name + "_BC_GS_train_data.png", 'graphConfMat2': plot_dir + file_name + "_BC_GS_val_data.png"}
    return render(request, 'showModelOutput.html', context)


def runModel(request):
    try:
        content = request.POST.get('optModel', False)
        print('content ', content)
        if(content == "RF"):
            return redirect('randomForest')
        elif(content == "XGB"):
            return redirect('xgBoost')
        elif(content == "MLP"):
            return redirect('MLP')
        elif(content == "GBC"):
            return redirect('GBC')
        elif(content == "KNN"):
            return redirect('KNN')
        elif(content == "SVM"):
            return redirect('SVM')
        elif(content == "BC"):
            return redirect('BC')
        return render(request, 'showModelOutputAll.html',    {'pdfFile': "", 'model': '0'})
    except Exception as e:
        print(e)
        return render(request, 'error.html')


def runModelAjax(request):
    try:
        from multiprocessing import Process, Pipe 
        content = request.GET['model']
        print('content is ', content) 
        if(content == "RF"):
            t1 = time.time()
            print('_RF process started')
            randomForestAjax()
            if os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "_RF_NT_roc1.png")) or os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "_RF_GS_roc1.png")):
                NT_roc_file1 = plot_dir + file_name + "_RF_NT_roc1.png"
                NT_roc_file2 = plot_dir + file_name + "_RF_NT_roc2.png"
                RS_roc_file1 = plot_dir + file_name + "_RF_RS_roc1.png"
                RS_roc_file2 = plot_dir + file_name + "_RF_RS_roc2.png"
                GS_roc_file1 = plot_dir + file_name + "_RF_GS_roc1.png"
                GS_roc_file2 = plot_dir + file_name + "_RF_GS_roc2.png"
                NT_auc_ginin_output = exportTestResultLstFromExcel(
                    "Random Forest - No Parameters Tuning")
                RS_auc_ginin_output = exportTestResultLstFromExcel(
                    "Random Forest - Random Search")
                GS_auc_ginin_output = exportTestResultLstFromExcel(
                    "Random Forest - Grid Search")
                context = {'is_data': True, 'pdfFile': plot_dir + file_name + "_Random_Forest.pdf", 'model': 'RF', 'tableHead': 'Random Forest', 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                           'NT_graphConfMat1': plot_dir + file_name + "_RF_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_RF_NT_val_data.png",
                           'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                           'RS_graphConfMat1': plot_dir + file_name + "_RF_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_RF_RS_val_data.png",
                           'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                           'GS_graphConfMat1': plot_dir + file_name + "_RF_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_RF_GS_val_data.png"}
            else:
                context = {'is_data': False}
            print('RF respoanse sent it took ', time.time()-t1)

            return JsonResponse(context)
        elif(content == "XGB"):
            print('xgBoost process started')
            xgBoostAjax()
            if os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "xgboost_NT_roc1.png")) or os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "_xgboost_GS_roc1.png")):
                NT_roc_file1 = plot_dir + file_name + "xgboost_NT_roc1.png"
                NT_roc_file2 = plot_dir + file_name + "xgboost_NT_roc2.png"
                RS_roc_file1 = plot_dir + file_name + "_xgboost_RS_roc1.png"
                RS_roc_file2 = plot_dir + file_name + "_xgboost_RS_roc2.png"
                GS_roc_file1 = plot_dir + file_name + "_xgboost_GS_roc1.png"
                GS_roc_file2 = plot_dir + file_name + "_xgboost_GS_roc2.png"
                NT_auc_ginin_output = exportTestResultLstFromExcel(
                    "XGBoost - No Parameters Tuning")
                RS_auc_ginin_output = exportTestResultLstFromExcel(
                    "XGBoost - Random Search")
                GS_auc_ginin_output = exportTestResultLstFromExcel(
                    "XGBoost - Grid Search")
                context = {'is_data': True, 'pdfFile': plot_dir + file_name + "_XGBoost.pdf", 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                           'NT_graphConfMat1': plot_dir + file_name + "_xgboost_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_xgboost_NT_val_data.png",
                           'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                           'RS_graphConfMat1': plot_dir + file_name + "_xgboost_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_xgboost_RS_val_data.png",
                           'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                           'GS_graphConfMat1': plot_dir + file_name + "_xgboost_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_xgboost_GS_val_data.png"}
            else:
                context = {'is_data': False}
            print('xgBoost respoanse sent')
            return JsonResponse(context)
        elif(content == "MLP"):
            MLPAjax()
            if os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "MLP_NT_roc1.png")) or os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "_MLP_GS_roc1.png")):
                NT_roc_file1 = plot_dir + file_name+"MLP_NT_roc1.png"
                NT_roc_file2 = plot_dir + file_name+"MLP_NT_roc2.png"
                RS_roc_file1 = plot_dir + file_name + "_MLP_RS_roc1.png"
                RS_roc_file2 = plot_dir + file_name + "_MLP_RS_roc2.png"
                GS_roc_file1 = plot_dir + file_name + "_MLP_GS_roc1.png"
                GS_roc_file2 = plot_dir + file_name + "_MLP_GS_roc2.png"
                NT_auc_ginin_output = exportTestResultLstFromExcel(
                    "Multi-Layer Perceptron - No Parameters Tuning")
                RS_auc_ginin_output = exportTestResultLstFromExcel(
                    "Multi-Layer Perceptron - Random Search")
                GS_auc_ginin_output = exportTestResultLstFromExcel(
                    "Multi-Layer Perceptron - Grid Search")
                context = {'is_data': True, 'pdfFile': plot_dir + file_name + "_MLP.pdf", 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                           'NT_graphConfMat1': plot_dir + file_name + "_MLP_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_MLP_NT_val_data.png",
                           'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                           'RS_graphConfMat1': plot_dir + file_name + "_MLP_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_MLP_RS_val_data.png",
                           'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                           'GS_graphConfMat1': plot_dir + file_name + "_MLP_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_MLP_GS_val_data.png"}
            else:
                context = {'is_data': False}
            print('MLP process ended')
            return JsonResponse(context)
        elif(content == "GBC"):
            # p1 = Process(target=GBCAjax)
            # p1.start()
            # p1.join()
            # print('process started GBC')
            GBCAjax()
            if os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "_GBC_NT_roc1.png")) or os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "_GBC_GS_roc1.png")):
                NT_roc_file1 = plot_dir + file_name+"GBC_NT_roc1.png"
                NT_roc_file2 = plot_dir + file_name+"GBC_NT_roc2.png"
                RS_roc_file1 = plot_dir + file_name + "_GBC_RS_roc1.png"
                RS_roc_file2 = plot_dir + file_name + "_GBC_RS_roc2.png"
                GS_roc_file1 = plot_dir + file_name + "_GBC_GS_roc1.png"
                GS_roc_file2 = plot_dir + file_name + "_GBC_GS_roc2.png"
                NT_auc_ginin_output = exportTestResultLstFromExcel(
                    "Gradient Boosting - No Parameters Tuning")
                RS_auc_ginin_output = exportTestResultLstFromExcel(
                    "Gradient Boosting - Random Search")
                GS_auc_ginin_output = exportTestResultLstFromExcel(
                    "Gradient Boosting - Grid Search")
                context = {'is_data': True, 'pdfFile': plot_dir + file_name + "_Gradient_Boosting.pdf", 'model': 'GBC',  'tableHead': 'Gradient Boosting', 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                           'NT_graphConfMat1': plot_dir + file_name + "_GBC_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_GBC_NT_val_data.png",
                           'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                           'RS_graphConfMat1': plot_dir + file_name + "_GBC_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_GBC_RS_val_data.png",
                           'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                           'GS_graphConfMat1': plot_dir + file_name + "_GBC_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_GBC_GS_val_data.png"}

            else:
                context = {'is_data': False}
            print('process ended GBC')
            return JsonResponse(context)

        elif(content == "KNN"):
            t1 = time.time()
            # p1 = Process(target=KNN)
            # p1.start()
            # print('process started KNN')
            # p1.join()
            KNNAjax()
            print('process after join took ', time.time()-t1)
            if os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "KNN_NT_roc1.png")) or os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "_KNN_GS_roc1.png")):
                NT_roc_file1 = plot_dir + file_name+"KNN_NT_roc1.png"
                NT_roc_file2 = plot_dir + file_name+"KNN_NT_roc2.png"
                RS_roc_file1 = plot_dir + file_name + "_KNN_RS_roc1.png"
                RS_roc_file2 = plot_dir + file_name + "_KNN_RS_roc2.png"
                GS_roc_file1 = plot_dir + file_name + "_KNN_GS_roc1.png"
                GS_roc_file2 = plot_dir + file_name + "_KNN_GS_roc2.png"
                t2 = time.time()
                NT_auc_ginin_output = exportTestResultLstFromExcel(
                    "K Nearest Neighbors - No Parameters Tuning")
                RS_auc_ginin_output = exportTestResultLstFromExcel(
                    "K Nearest Neighbors - Random Search")
                GS_auc_ginin_output = exportTestResultLstFromExcel(
                    "K Nearest Neighbors - Grid Search")
                print('process test result took  ', time.time()-t2)
                context = {'is_data': True, 'pdfFile': plot_dir + file_name + "_KNN.pdf", 'model': 'KNN',  'tableHead': 'K Nearest Neighbors', 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                           'NT_graphConfMat1': plot_dir + file_name + "_KNN_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_KNN_NT_val_data.png",
                           'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                           'RS_graphConfMat1': plot_dir + file_name + "_KNN_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_KNN_RS_val_data.png",
                           'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                           'GS_graphConfMat1': plot_dir + file_name + "_KNN_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_KNN_GS_val_data.png"}

            else:
                context = {'is_data': False}
            print('process ended KNN time taken is %s', str(time.time()-t1))
            return JsonResponse(context)
        elif(content == "SVM"):
            # p1 = Process(target=SVMAjax)
            # print('process started SVM')
            # p1.start()
            # p1.join()
            SVMAjax()
            if os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "SVM_NT_roc1.png")) or os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "_SVM_GS_roc1.png")):
                NT_roc_file1 = plot_dir + file_name+"SVM_NT_roc1.png"
                NT_roc_file2 = plot_dir + file_name+"SVM_NT_roc2.png"
                RS_roc_file1 = plot_dir + file_name + "_SVM_RS_roc1.png"
                RS_roc_file2 = plot_dir + file_name + "_SVM_RS_roc2.png"
                GS_roc_file1 = plot_dir + file_name + "_SVM_GS_roc1.png"
                GS_roc_file2 = plot_dir + file_name + "_SVM_GS_roc2.png"
                NT_auc_ginin_output = exportTestResultLstFromExcel(
                    "Support Machine Vector - No Parameters Tuning")
                RS_auc_ginin_output = exportTestResultLstFromExcel(
                    "Support Machine Vector - Random Search")
                GS_auc_ginin_output = exportTestResultLstFromExcel(
                    "Support Machine Vector - Grid Search")
                context = {'is_data': True, 'pdfFile': plot_dir + file_name + "_SVM.pdf", 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                           'NT_graphConfMat1': plot_dir + file_name + "_SVM_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_SVM_NT_val_data.png",
                           'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                           'RS_graphConfMat1': plot_dir + file_name + "_SVM_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_SVM_RS_val_data.png",
                           'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                           'GS_graphConfMat1': plot_dir + file_name + "_SVM_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_SVM_GS_val_data.png"}
            else:
                context = {'is_data': False}
            print('process ended SVM')
            return JsonResponse(context)
        elif(content == "BC"):
            # p1 = Process(target=BCAjax)
            # print('process started BC')
            # p1.start()
            # p1.join()
            BCAjax()
            if os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "BC_NT_roc1.png")) or os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "_BC_GS_roc1.png")):
                NT_roc_file1 = plot_dir + file_name+"BC_NT_roc1.png"
                NT_roc_file2 = plot_dir + file_name+"BC_NT_roc2.png"
                RS_roc_file1 = plot_dir + file_name + "_BC_RS_roc1.png"
                RS_roc_file2 = plot_dir + file_name + "_BC_RS_roc2.png"
                GS_roc_file1 = plot_dir + file_name + "_BC_GS_roc1.png"
                GS_roc_file2 = plot_dir + file_name + "_BC_GS_roc2.png"
                NT_auc_ginin_output = exportTestResultLstFromExcel(
                    "Bagging Classifier - No Parameters Tuning")
                RS_auc_ginin_output = exportTestResultLstFromExcel(
                    "Bagging Classifier - Random Search")
                GS_auc_ginin_output = exportTestResultLstFromExcel(
                    "Bagging Classifier - Grid Search")
                context = {'pdfFile': plot_dir + file_name + "_BagCls.pdf", 'model': 'BC',  'tableHead': 'Bagging Classifier', 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                           'NT_graphConfMat1': plot_dir + file_name + "_BC_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_BC_NT_val_data.png",
                           'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                           'RS_graphConfMat1': plot_dir + file_name + "_BC_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_BC_RS_val_data.png",
                           'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                           'GS_graphConfMat1': plot_dir + file_name + "_BC_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_BC_GS_val_data.png"}
            else:
                context = {'is_data': False}
            print('process ended BC')
            return JsonResponse(context)
        return render(request, 'showModelOutputAll.html',    {'pdfFile': "", 'model': '0'})
    except Exception as e:
        print(e)
        print('traceback ', traceback.print_exc())
        return render(request, 'error.html')


def getPrvChart(request):
    try:
        content = request.GET['model']
        if(content == "RF"):
            if os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "_RF_NT_roc1.png")) and os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "_RF_GS_roc1.png")):
                NT_roc_file1 = plot_dir + file_name + "_RF_NT_roc1.png"
                NT_roc_file2 = plot_dir + file_name + "_RF_NT_roc2.png"
                RS_roc_file1 = plot_dir + file_name + "_RF_RS_roc1.png"
                RS_roc_file2 = plot_dir + file_name + "_RF_RS_roc2.png"
                GS_roc_file1 = plot_dir + file_name + "_RF_GS_roc1.png"
                GS_roc_file2 = plot_dir + file_name + "_RF_GS_roc2.png"
                NT_auc_ginin_output = exportTestResultLstFromExcel(
                    "Random Forest - No Parameters Tuning")
                RS_auc_ginin_output = exportTestResultLstFromExcel(
                    "Random Forest - Random Search")
                GS_auc_ginin_output = exportTestResultLstFromExcel(
                    "Random Forest - Grid Search")
                context = {'is_data': True, 'pdfFile': plot_dir + file_name + "_Random_Forest.pdf", 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                           'NT_graphConfMat1': plot_dir + file_name + "_RF_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_RF_NT_val_data.png",
                           'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                           'RS_graphConfMat1': plot_dir + file_name + "_RF_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_RF_RS_val_data.png",
                           'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                           'GS_graphConfMat1': plot_dir + file_name + "_RF_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_RF_GS_val_data.png"}
            else:
                context = {'is_data': False}

            return JsonResponse(context)
        elif(content == "XGB"):
            if os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "xgboost_NT_roc1.png")) or os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "_xgboost_GS_roc1.png")):
                NT_roc_file1 = plot_dir + file_name + "xgboost_NT_roc1.png"
                NT_roc_file2 = plot_dir + file_name + "xgboost_NT_roc2.png"
                RS_roc_file1 = plot_dir + file_name + "_xgboost_RS_roc1.png"
                RS_roc_file2 = plot_dir + file_name + "_xgboost_RS_roc2.png"
                GS_roc_file1 = plot_dir + file_name + "_xgboost_GS_roc1.png"
                GS_roc_file2 = plot_dir + file_name + "_xgboost_GS_roc2.png"
                NT_auc_ginin_output = exportTestResultLstFromExcel(
                    "XGBoost - No Parameters Tuning")
                RS_auc_ginin_output = exportTestResultLstFromExcel(
                    "XGBoost - Random Search")
                GS_auc_ginin_output = exportTestResultLstFromExcel(
                    "XGBoost - Grid Search")
                context = {'is_data': True, 'pdfFile': plot_dir + file_name + "_Random_Forest.pdf", 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                           'NT_graphConfMat1': plot_dir + file_name + "_xgboost_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_xgboost_NT_val_data.png",
                           'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                           'RS_graphConfMat1': plot_dir + file_name + "_xgboost_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_xgboost_RS_val_data.png",
                           'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                           'GS_graphConfMat1': plot_dir + file_name + "_xgboost_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_xgboost_GS_val_data.png"}
            else:
                context = {'is_data': False}

            return JsonResponse(context)
        elif(content == "MLP"):
            if os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "MLP_NT_roc1.png")) or os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "_MLP_GS_roc1.png")):
                NT_roc_file1 = plot_dir + file_name+"MLP_NT_roc1.png"
                NT_roc_file2 = plot_dir + file_name+"MLP_NT_roc2.png"
                RS_roc_file1 = plot_dir + file_name + "_MLP_RS_roc1.png"
                RS_roc_file2 = plot_dir + file_name + "_MLP_RS_roc2.png"
                GS_roc_file1 = plot_dir + file_name + "_MLP_GS_roc1.png"
                GS_roc_file2 = plot_dir + file_name + "_MLP_GS_roc2.png"
                NT_auc_ginin_output = exportTestResultLstFromExcel(
                    "Multi-Layer Perceptron - No Parameters Tuning")
                RS_auc_ginin_output = exportTestResultLstFromExcel(
                    "Multi-Layer Perceptron - Random Search")
                GS_auc_ginin_output = exportTestResultLstFromExcel(
                    "Multi-Layer Perceptron - Grid Search")
                context = {'is_data': True, 'pdfFile': plot_dir + file_name + "_Random_Forest.pdf", 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                           'NT_graphConfMat1': plot_dir + file_name + "_MLP_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_MLP_NT_val_data.png",
                           'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                           'RS_graphConfMat1': plot_dir + file_name + "_MLP_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_MLP_RS_val_data.png",
                           'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                           'GS_graphConfMat1': plot_dir + file_name + "_MLP_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_MLP_GS_val_data.png"}
            else:
                context = {'is_data': False}

            return JsonResponse(context)
        elif(content == "GBC"):
            if os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "GBC_NT_roc1.png")) or os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "_GBC_GS_roc1.png")):
                NT_roc_file1 = plot_dir + file_name+"GBC_NT_roc1.png"
                NT_roc_file2 = plot_dir + file_name+"GBC_NT_roc2.png"
                RS_roc_file1 = plot_dir + file_name + "_GBC_RS_roc1.png"
                RS_roc_file2 = plot_dir + file_name + "_GBC_RS_roc2.png"
                GS_roc_file1 = plot_dir + file_name + "_GBC_GS_roc1.png"
                GS_roc_file2 = plot_dir + file_name + "_GBC_GS_roc2.png"
                NT_auc_ginin_output = exportTestResultLstFromExcel(
                    "Gradient Boosting - No Parameters Tuning")
                RS_auc_ginin_output = exportTestResultLstFromExcel(
                    "Gradient Boosting - Random Search")
                GS_auc_ginin_output = exportTestResultLstFromExcel(
                    "Gradient Boosting - Grid Search")
                context = {'is_data': True, 'pdfFile': plot_dir + file_name + "_Random_Forest.pdf", 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                           'NT_graphConfMat1': plot_dir + file_name + "_GBC_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_GBC_NT_val_data.png",
                           'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                           'RS_graphConfMat1': plot_dir + file_name + "_GBC_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_GBC_RS_val_data.png",
                           'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                           'GS_graphConfMat1': plot_dir + file_name + "_GBC_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_GBC_GS_val_data.png"}
            else:
                context = {'is_data': False}

            return JsonResponse(context)
        elif(content == "KNN"):
            if os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "KNN_NT_roc1.png")) or os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "_KNN_GS_roc1.png")):
                print('insie if ', plot_dir_view +
                      file_name + "KNN_NT_roc1.png")
                NT_roc_file1 = plot_dir + file_name+"KNN_NT_roc1.png"
                NT_roc_file2 = plot_dir + file_name+"KNN_NT_roc2.png"
                RS_roc_file1 = plot_dir + file_name + "_KNN_RS_roc1.png"
                RS_roc_file2 = plot_dir + file_name + "_KNN_RS_roc2.png"
                GS_roc_file1 = plot_dir + file_name + "_KNN_GS_roc1.png"
                GS_roc_file2 = plot_dir + file_name + "_KNN_GS_roc2.png"
                NT_auc_ginin_output = exportTestResultLstFromExcel(
                    "K Nearest Neighbors - No Parameters Tuning")
                RS_auc_ginin_output = exportTestResultLstFromExcel(
                    "K Nearest Neighbors - Random Search")
                GS_auc_ginin_output = exportTestResultLstFromExcel(
                    "K Nearest Neighbors - Grid Search")
                context = {'is_data': True, 'pdfFile': plot_dir + file_name + "_Random_Forest.pdf", 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                           'NT_graphConfMat1': plot_dir + file_name + "_KNN_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_KNN_NT_val_data.png",
                           'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                           'RS_graphConfMat1': plot_dir + file_name + "_KNN_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_KNN_RS_val_data.png",
                           'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                           'GS_graphConfMat1': plot_dir + file_name + "_KNN_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_KNN_GS_val_data.png"}
            else:
                context = {'is_data': False}

            return JsonResponse(context)
        elif(content == "SVM"):
            if os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "SVM_NT_roc1.png")) and os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "_SVM_GS_roc1.png")):
                NT_roc_file1 = plot_dir + file_name+"SVM_NT_roc1.png"
                NT_roc_file2 = plot_dir + file_name+"SVM_NT_roc2.png"
                RS_roc_file1 = plot_dir + file_name + "_SVM_RS_roc1.png"
                RS_roc_file2 = plot_dir + file_name + "_SVM_RS_roc2.png"
                GS_roc_file1 = plot_dir + file_name + "_SVM_GS_roc1.png"
                GS_roc_file2 = plot_dir + file_name + "_SVM_GS_roc2.png"
                NT_auc_ginin_output = exportTestResultLstFromExcel(
                    "Support Machine Vector - No Parameters Tuning")
                RS_auc_ginin_output = exportTestResultLstFromExcel(
                    "Support Machine Vector - Random Search")
                GS_auc_ginin_output = exportTestResultLstFromExcel(
                    "Support Machine Vector - Grid Search")
                context = {'is_data': True, 'pdfFile': plot_dir + file_name + "_Random_Forest.pdf", 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                           'NT_graphConfMat1': plot_dir + file_name + "_SVM_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_SVM_NT_val_data.png",
                           'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                           'RS_graphConfMat1': plot_dir + file_name + "_SVM_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_SVM_RS_val_data.png",
                           'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                           'GS_graphConfMat1': plot_dir + file_name + "_SVM_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_SVM_GS_val_data.png"}
            else:
                context = {'is_data': False}

            return JsonResponse(context)
        elif(content == "BC"):
            if os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "BC_NT_roc1.png")) or os.path.exists(os.path.join(BASE_DIR, plot_dir_view + file_name + "_BC_GS_roc1.png")):
                NT_roc_file1 = plot_dir + file_name+"BC_NT_roc1.png"
                NT_roc_file2 = plot_dir + file_name+"BC_NT_roc2.png"
                RS_roc_file1 = plot_dir + file_name + "_BC_RS_roc1.png"
                RS_roc_file2 = plot_dir + file_name + "_BC_RS_roc2.png"
                GS_roc_file1 = plot_dir + file_name + "_BC_GS_roc1.png"
                GS_roc_file2 = plot_dir + file_name + "_BC_GS_roc2.png"
                NT_auc_ginin_output = exportTestResultLstFromExcel(
                    "Bagging Classifier - No Parameters Tuning")
                RS_auc_ginin_output = exportTestResultLstFromExcel(
                    "Bagging Classifier - Random Search")
                GS_auc_ginin_output = exportTestResultLstFromExcel(
                    "Bagging Classifier - Grid Search")
                context = {'is_data': True, 'pdfFile': plot_dir + file_name + "_Random_Forest.pdf", 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                           'NT_graphConfMat1': plot_dir + file_name + "_BC_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_BC_NT_val_data.png",
                           'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                           'RS_graphConfMat1': plot_dir + file_name + "_BC_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_BC_RS_val_data.png",
                           'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                           'GS_graphConfMat1': plot_dir + file_name + "_BC_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_BC_GS_val_data.png"}
            else:
                context = {'is_data': False}

            return JsonResponse(context)

    except Exception as e:
        print(e)
        return render(request, 'error.html')


def randomForest(request):
    try:
        csv_file_name = "csvfile_"+user_name
        savefile_x_final = file_path + csv_file_name + "_x_model.csv"
        if(not os.path.exists(savefile_x_final)):
            return render(request, 'processNotdone.html')
        df = pd.read_csv(savefile_x_final)
        targetVarFile = file_path + csv_file_name + "_targetVar.txt"
        file1 = open(targetVarFile, "r")  # write mode
        targetVar = file1.read()
        file1.close()
        # a variable pdf
        pdf = FPDF()
        # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
        y_model = df[targetVar]
        x_model = df.drop(targetVar, axis=1)
        X_train, X_test, y_train, y_test = train_test_split(
            x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
        X_train, X_val, y_train, y_val = train_test_split(
            X_train, y_train, test_size=0.22222222222222224, random_state=321)

        rf_model = RandomForestClassifier(n_estimators=100, criterion="gini",
                                          random_state=50,
                                          max_features='sqrt',
                                          n_jobs=-1, verbose=1)
        rf_model.fit(X_train, y_train)

        # check the average of the nodes and depth.
        n_nodes = []
        max_depths = []
        for ind_tree in rf_model.estimators_:
            n_nodes.append(ind_tree.tree_.node_count)
            max_depths.append(ind_tree.tree_.max_depth)

        print(f'Average number of nodes {int(np.mean(n_nodes))}')
        print(f'Average maximum depth {int(np.mean(max_depths))}')

        # Test the model
        pred_rf_val = rf_model.predict(X_val)
        pred_rf_prob_val = rf_model.predict_proba(X_val)[:, 1]

        pred_rf_train = rf_model.predict(X_train)
        pred_rf_prob_train = rf_model.predict_proba(X_train)[:, 1]

        # Get the model performance
        print('classification report on training data')
        print(classification_report(y_train, pred_rf_train))
        print('\n')
        print('classification report on validation data')
        print(classification_report(y_val, pred_rf_val))

        evaluate_model(y_val, pred_rf_prob_val, y_train,
                       pred_rf_prob_train, file_name+"_RF_NT_roc1.png")

        evaluate_model(y_val, pred_rf_val, y_train,
                       pred_rf_train, file_name+"_RF_NT_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data
        drawConfMatrix(y_train, pred_rf_train,
                       file_name + "_RF_NT_train_data.png", "Training")

        # show the confusion matrix for validation data
        drawConfMatrix(y_val, pred_rf_val, file_name +
                       "_RF_NT_val_data.png", "Validation")

        NT_roc_file1 = plot_dir + file_name + "_RF_NT_roc1.png"
        NT_roc_file2 = plot_dir + file_name + "_RF_NT_roc2.png"
        NT_auc_ginin_output = test_modelPerfomance(
            y_val, pred_rf_prob_val, pred_rf_val)

        # Add a page
        pdf.add_page()
        document = Document()
        pdf = exportPdf(10, 10, pdf, document, os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_NT_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_NT_roc2.png"),
            "ROC Curve Test data", "Random Forest - No Parameters Tuning")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_NT_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_NT_val_data.png"),
            "Confusion Matrix Validation data", "Random Forest - No Parameters Tuning")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, NT_auc_ginin_output, "Random Forest - No Parameters Tuning")
        # Random Search
        paramFiles = param_file_path + param_file_name + "_RF_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])
        model = RandomForestClassifier()

        # Create the random search model
        rf_search = RandomizedSearchCV(model, param_grid, n_jobs=3,
                                       scoring='roc_auc', cv=5,
                                       n_iter=100, verbose=1, random_state=50)

        # Fit
        rf_search.fit(X_train, y_train)

        rf_search.best_params_

        rf_random = rf_search.best_estimator_

        # Test the model
        pred_rf_val0 = rf_random.predict(X_val)
        pred_rf_prob_val0 = rf_random.predict_proba(X_val)[:, 1]

        pred_rf_train0 = rf_random.predict(X_train)
        pred_rf_prob_train0 = rf_random.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_rf_train0))
        print(classification_report(y_val, pred_rf_val0))

        evaluate_model(y_val, pred_rf_prob_val0, y_train,
                       pred_rf_prob_train0, file_name+"_RF_RS_roc1.png")

        evaluate_model(y_val, pred_rf_val0, y_train,
                       pred_rf_train0, file_name+"_RF_RS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data
        drawConfMatrix(y_train, pred_rf_train0,
                       file_name + "_RF_RS_train_data.png", "Training")
        drawConfMatrix(y_val, pred_rf_val0, file_name +
                       "_RF_RS_val_data.png", "Validation")

        RS_roc_file1 = plot_dir + file_name + "_RF_RS_roc1.png"
        RS_roc_file2 = plot_dir + file_name + "_RF_RS_roc2.png"
        RS_auc_ginin_output = test_modelPerfomance(
            y_val, pred_rf_prob_val0, pred_rf_val0)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_RS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_RS_roc2.png"),
            "ROC Curve Test data", "Random Forest - Random Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_RS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_RS_val_data.png"),
            "Confusion Matrix Validation data", "Random Forest - Random Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, RS_auc_ginin_output, "Random Forest - Random Search")

        # Grid Search
        paramFiles = param_file_path + param_file_name + "_RF_GS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        model = RandomForestClassifier()

        # Create the random search model
        rf_grid = GridSearchCV(model, param_grid, n_jobs=3,
                               scoring='roc_auc', cv=5, verbose=1)

        # Fit
        rf_grid.fit(X_train, y_train)

        rf_grid.best_params_

        rf_grid = rf_grid.best_estimator_

        # Test the model
        pred_rf_val1 = rf_grid.predict(X_val)
        pred_rf_prob_val1 = rf_grid.predict_proba(X_val)[:, 1]

        pred_rf_train1 = rf_grid.predict(X_train)
        pred_rf_prob_train1 = rf_grid.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_rf_train1))
        print(classification_report(y_val, pred_rf_val1))

        evaluate_model(y_val, pred_rf_prob_val1, y_train,
                       pred_rf_prob_train1, file_name+"_RF_GS_roc1.png")

        evaluate_model(y_val, pred_rf_val1, y_train,
                       pred_rf_train1, file_name+"_RF_GS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        cnf_matrix = confusion_matrix(y_train, pred_rf_train1, labels=[0, 1])
        plt.figure(figsize=(10, 6))
        sns.heatmap(pd.DataFrame(cnf_matrix), annot=True,
                    cmap="YlGnBu", fmt='g')

        plt.title('Confusion matrix: Training data')
        plt.ylabel('Actual label')
        plt.xlabel('Predicted label')
        plt.tight_layout()
        plt.savefig(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_RF_GS_train_data.png"))
        plt.close()
        # show the confusion matrix for validation data

        cnf_matrix2 = confusion_matrix(y_val, pred_rf_val1, labels=[0, 1])
        plt.figure(figsize=(10, 6))
        sns.heatmap(pd.DataFrame(cnf_matrix2), annot=True,
                    cmap="YlGnBu", fmt='g')
        plt.title('Confusion matrix: Validation data')
        plt.ylabel('Actual label')
        plt.xlabel('Predicted label')
        plt.tight_layout()
        plt.savefig(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_RF_GS_val_data.png"))
        plt.close()
        GS_roc_file1 = plot_dir + file_name + "_RF_GS_roc1.png"
        GS_roc_file2 = plot_dir + file_name + "_RF_GS_roc2.png"
        GS_auc_ginin_output = test_modelPerfomance(
            y_val, pred_rf_prob_val1, pred_rf_val1)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_GS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_GS_roc2.png"),
            "ROC Curve Test data", "Random Forest - Grid Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_GS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_GS_val_data.png"),
            "Confusion Matrix Validation data", "Random Forest - Grid Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, GS_auc_ginin_output, "Random Forest - Grid Search")

        pdf.output(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_Random_Forest.pdf"))
        context = {'pdfFile': plot_dir + file_name + "_Random_Forest.pdf", 'model': 'RF', 'tableHead': 'Random Forest', 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                   'NT_graphConfMat1': plot_dir + file_name + "_RF_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_RF_NT_val_data.png",
                   'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                   'RS_graphConfMat1': plot_dir + file_name + "_RF_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_RF_RS_val_data.png",
                   'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                   'GS_graphConfMat1': plot_dir + file_name + "_RF_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_RF_GS_val_data.png"}
        return render(request, 'showModelOutputAll.html', context)
        # return render(request, 'runModels.html', context)
    except Exception as e:
        print(e)
        return render(request, 'error.html')


def randomForestAjax():
    try:
        csv_file_name = "csvfile_"+user_name
        savefile_x_final = file_path + csv_file_name + "_x_model.csv"
        df = pd.read_csv(savefile_x_final)
        targetVarFile = file_path + csv_file_name + "_targetVar.txt"
        file1 = open(targetVarFile, "r")  # write mode
        targetVar = file1.read()
        file1.close()
        # a variable pdf
        pdf = FPDF()
        # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
        y_model = df[targetVar]
        x_model = df.drop(targetVar, axis=1)
        X_train, X_test, y_train, y_test = train_test_split(
            x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
        X_train, X_val, y_train, y_val = train_test_split(
            X_train, y_train, test_size=0.22222222222222224, random_state=321)

        rf_model = RandomForestClassifier(n_estimators=100, criterion="gini",
                                          random_state=50,
                                          max_features='sqrt',
                                          n_jobs=-1, verbose=1)
        rf_model.fit(X_train, y_train)

        # check the average of the nodes and depth.
        n_nodes = []
        max_depths = []
        for ind_tree in rf_model.estimators_:
            n_nodes.append(ind_tree.tree_.node_count)
            max_depths.append(ind_tree.tree_.max_depth)

        print(f'Average number of nodes {int(np.mean(n_nodes))}')
        print(f'Average maximum depth {int(np.mean(max_depths))}')

        # Test the model
        pred_rf_val = rf_model.predict(X_val)
        pred_rf_prob_val = rf_model.predict_proba(X_val)[:, 1]

        pred_rf_train = rf_model.predict(X_train)
        pred_rf_prob_train = rf_model.predict_proba(X_train)[:, 1]

        # Get the model performance
        print('classification report on training data')
        print(classification_report(y_train, pred_rf_train))
        print('\n')
        print('classification report on validation data')
        print(classification_report(y_val, pred_rf_val))

        evaluate_model(y_val, pred_rf_prob_val, y_train,
                       pred_rf_prob_train, file_name+"_RF_NT_roc1.png")

        evaluate_model(y_val, pred_rf_val, y_train,
                       pred_rf_train, file_name+"_RF_NT_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data
        drawConfMatrix(y_train, pred_rf_train,
                       file_name + "_RF_NT_train_data.png", "Training")

        # show the confusion matrix for validation data
        drawConfMatrix(y_val, pred_rf_val, file_name +
                       "_RF_NT_val_data.png", "Validation")

        NT_roc_file1 = plot_dir + file_name + "_RF_NT_roc1.png"
        NT_roc_file2 = plot_dir + file_name + "_RF_NT_roc2.png"
        NT_auc_ginin_output = test_modelPerfomance(
            y_val, pred_rf_prob_val, pred_rf_val)

        # Add a page
        pdf.add_page()
        document = Document()
        pdf = exportPdf(10, 10, pdf, document, os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_NT_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_NT_roc2.png"),
            "ROC Curve Test data", "Random Forest - No Parameters Tuning")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_NT_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_NT_val_data.png"),
            "Confusion Matrix Validation data", "Random Forest - No Parameters Tuning")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, NT_auc_ginin_output, "Random Forest - No Parameters Tuning")
        # Random Search
        paramFiles = param_file_path + param_file_name + "_RF_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])
        model = RandomForestClassifier()

        # Create the random search model
        rf_search = RandomizedSearchCV(model, param_grid, n_jobs=-1,
                                       scoring='roc_auc', cv=5,
                                       n_iter=100, verbose=1, random_state=50)

        # Fit
        rf_search.fit(X_train, y_train)

        rf_search.best_params_

        rf_random = rf_search.best_estimator_

        # Test the model
        pred_rf_val0 = rf_random.predict(X_val)
        pred_rf_prob_val0 = rf_random.predict_proba(X_val)[:, 1]

        pred_rf_train0 = rf_random.predict(X_train)
        pred_rf_prob_train0 = rf_random.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_rf_train0))
        print(classification_report(y_val, pred_rf_val0))

        evaluate_model(y_val, pred_rf_prob_val0, y_train,
                       pred_rf_prob_train0, file_name+"_RF_RS_roc1.png")

        evaluate_model(y_val, pred_rf_val0, y_train,
                       pred_rf_train0, file_name+"_RF_RS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data
        drawConfMatrix(y_train, pred_rf_train0,
                       file_name + "_RF_RS_train_data.png", "Training")
        drawConfMatrix(y_val, pred_rf_val0, file_name +
                       "_RF_RS_val_data.png", "Validation")

        RS_roc_file1 = plot_dir + file_name + "_RF_RS_roc1.png"
        RS_roc_file2 = plot_dir + file_name + "_RF_RS_roc2.png"
        RS_auc_ginin_output = test_modelPerfomance(
            y_val, pred_rf_prob_val0, pred_rf_val0)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_RS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_RS_roc2.png"),
            "ROC Curve Test data", "Random Forest - Random Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_RS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_RS_val_data.png"),
            "Confusion Matrix Validation data", "Random Forest - Random Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, RS_auc_ginin_output, "Random Forest - Random Search")

        # Grid Search
        paramFiles = param_file_path + param_file_name + "_RF_GS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        model = RandomForestClassifier()

        # Create the random search model
        rf_grid = GridSearchCV(model, param_grid, n_jobs=-1,
                               scoring='roc_auc', cv=10, verbose=1)

        # Fit
        rf_grid.fit(X_train, y_train)

        rf_grid.best_params_

        rf_grid = rf_grid.best_estimator_

        # Test the model
        pred_rf_val1 = rf_grid.predict(X_val)
        pred_rf_prob_val1 = rf_grid.predict_proba(X_val)[:, 1]

        pred_rf_train1 = rf_grid.predict(X_train)
        pred_rf_prob_train1 = rf_grid.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_rf_train1))
        print(classification_report(y_val, pred_rf_val1))

        evaluate_model(y_val, pred_rf_prob_val1, y_train,
                       pred_rf_prob_train1, file_name+"_RF_GS_roc1.png")

        evaluate_model(y_val, pred_rf_val1, y_train,
                       pred_rf_train1, file_name+"_RF_GS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        cnf_matrix = confusion_matrix(y_train, pred_rf_train1, labels=[0, 1])
        plt.figure(figsize=(10, 6))
        sns.heatmap(pd.DataFrame(cnf_matrix), annot=True,
                    cmap="YlGnBu", fmt='g')

        plt.title('Confusion matrix: Training data')
        plt.ylabel('Actual label')
        plt.xlabel('Predicted label')
        plt.tight_layout()
        plt.savefig(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_RF_GS_train_data.png"))
        plt.close()
        # show the confusion matrix for validation data

        cnf_matrix2 = confusion_matrix(y_val, pred_rf_val1, labels=[0, 1])
        plt.figure(figsize=(10, 6))
        sns.heatmap(pd.DataFrame(cnf_matrix2), annot=True,
                    cmap="YlGnBu", fmt='g')
        plt.title('Confusion matrix: Validation data')
        plt.ylabel('Actual label')
        plt.xlabel('Predicted label')
        plt.tight_layout()
        plt.savefig(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_RF_GS_val_data.png"))
        plt.close()
        GS_roc_file1 = plot_dir + file_name + "_RF_GS_roc1.png"
        GS_roc_file2 = plot_dir + file_name + "_RF_GS_roc2.png"
        GS_auc_ginin_output = test_modelPerfomance(
            y_val, pred_rf_prob_val1, pred_rf_val1)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_GS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_GS_roc2.png"),
            "ROC Curve Test data", "Random Forest - Grid Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_GS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_RF_GS_val_data.png"),
            "Confusion Matrix Validation data", "Random Forest - Grid Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, GS_auc_ginin_output, "Random Forest - Grid Search")

        pdf.output(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_Random_Forest.pdf"))
        context = {'pdfFile': plot_dir + file_name + "_Random_Forest.pdf", 'model': 'RF', 'tableHead': 'Random Forest', 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                   'NT_graphConfMat1': plot_dir + file_name + "_RF_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_RF_NT_val_data.png",
                   'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                   'RS_graphConfMat1': plot_dir + file_name + "_RF_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_RF_RS_val_data.png",
                   'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                   'GS_graphConfMat1': plot_dir + file_name + "_RF_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_RF_GS_val_data.png"}

        # return render(request, 'runModels.html', context)

    except Exception as e:
        print(e)


def xgBoost(request):
    try:
        csv_file_name = "csvfile_"+user_name
        savefile_x_final = file_path + csv_file_name + "_x_model.csv"
        if(not os.path.exists(savefile_x_final)):
            return render(request, 'processNotdone.html')
        df = pd.read_csv(savefile_x_final)
        targetVarFile = file_path + csv_file_name + "_targetVar.txt"
        file1 = open(targetVarFile, "r")  # write mode
        targetVar = file1.read()
        file1.close()
        # a variable pdf
        pdf = FPDF()
        # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
        y_model = df[targetVar]
        x_model = df.drop(targetVar, axis=1)
        X_train, X_test, y_train, y_test = train_test_split(
            x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
        X_train, X_val, y_train, y_val = train_test_split(
            X_train, y_train, test_size=0.22222222222222224, random_state=321)

        # Fit XGRegressor to the Training set
        xg_clf = xgb.XGBClassifier()
        xg_clf.fit(X_train, y_train)

        # Test the model
        pred_xgb_val = xg_clf.predict(X_val)
        pred_xgb_prob_val = xg_clf.predict_proba(X_val)[:, 1]

        pred_xgb_train = xg_clf.predict(X_train)
        pred_xgb_prob_train = xg_clf.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_xgb_train))
        print(classification_report(y_val, pred_xgb_val))

        evaluate_model(pred_xgb_val, pred_xgb_prob_val, pred_xgb_train,
                       pred_xgb_prob_train, file_name+"xgboost_NT_roc1.png")

        evaluate_model(pred_xgb_val, pred_xgb_val, pred_xgb_train,
                       pred_xgb_train, file_name+"xgboost_NT_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data
        drawConfMatrix(y_train, pred_xgb_train, file_name +
                       "_xgboost_NT_train_data.png", "Training")

        drawConfMatrix(y_val, pred_xgb_val, file_name +
                       "_xgboost_NT_val_data.png", "Validation")

        NT_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_xgb_prob_val, pred_xgb_val)

        NT_roc_file1 = plot_dir + file_name+"xgboost_NT_roc1.png"
        NT_roc_file2 = plot_dir + file_name+"xgboost_NT_roc2.png"

        # Add a page
        pdf.add_page()
        document = Document()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"xgboost_NT_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"xgboost_NT_roc2.png"),
            "ROC Curve Test data", "XGBoost - No Parameters Tuning")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_NT_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_NT_val_data.png"),
            "Confusion Matrix Validation data", "XGBoost - No Parameters Tuning")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, NT_auc_ginin_output, "XGBoost - No Parameters Tuning")

        paramFiles = param_file_path + param_file_name + "_XGB_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = xgb.XGBClassifier(random_state=50)

        # Create the random search model
        xgb_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                        scoring='roc_auc', cv=5,
                                        n_iter=100, verbose=1, random_state=50)

        # Fit
        xgb_search.fit(X_train, y_train)

        xgb_search.best_params_

        xgb_random = xgb_search.best_estimator_

        # Test the model
        pred_xgb_val0 = xgb_random.predict(X_val)
        pred_xgb_prob_val0 = xgb_random.predict_proba(X_val)[:, 1]

        pred_xgb_train0 = xgb_random.predict(X_train)
        pred_xgb_prob_train0 = xgb_random.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_xgb_train0))
        print(classification_report(y_val, pred_xgb_val0))

        evaluate_model(y_val, pred_xgb_prob_val0, y_train,
                       pred_xgb_prob_train0, file_name+"_xgboost_RS_roc1.png")

        evaluate_model(y_val, pred_xgb_val0, y_train,
                       pred_xgb_train0, file_name+"_xgboost_RS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_xgb_train0,
                       file_name + "_xgboost_RS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_xgb_val0,  file_name +
                       "_xgboost_RS_val_data.png", "Validation")

        RS_roc_file1 = plot_dir + file_name + "_xgboost_RS_roc1.png"
        RS_roc_file2 = plot_dir + file_name + "_xgboost_RS_roc2.png"
        RS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_xgb_prob_val0, pred_xgb_val0)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_RS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_RS_roc2.png"),
            "ROC Curve Test data", "XGBoost - Random Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_RS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_RS_val_data.png"),
            "Confusion Matrix Validation data", "XGBoost - Random Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, RS_auc_ginin_output, "XGBoost - Random Search")

        paramFiles = param_file_path + param_file_name + "_XGB_GS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = xgb.XGBClassifier(random_state=50)

        # Create the random search model
        xgb_grid = GridSearchCV(estimator, param_grid, n_jobs=3,
                                scoring='roc_auc', cv=5, verbose=1)

        # Fit
        xgb_grid.fit(X_train, y_train)

        xgb_grid.best_params_

        xgb_final = xgb_grid.best_estimator_

        # Test the model
        pred_xgb_val1 = xgb_final.predict(X_val)
        pred_xgb_prob_val1 = xgb_final.predict_proba(X_val)[:, 1]

        pred_xgb_train1 = xgb_final.predict(X_train)
        pred_xgb_prob_train1 = xgb_final.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_xgb_train1))
        print(classification_report(y_val, pred_xgb_val1))

        evaluate_model(y_val, pred_xgb_prob_val1, y_train,
                       pred_xgb_prob_train1, file_name+"_xgboost_GS_roc1.png")

        evaluate_model(y_val, pred_xgb_val1, y_train,
                       pred_xgb_train1, file_name+"_xgboost_GS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_xgb_train1,
                       file_name + "_xgboost_GS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_xgb_val1,  file_name +
                       "_xgboost_GS_val_data.png", "Validation")

        GS_roc_file1 = plot_dir + file_name + "_xgboost_GS_roc1.png"
        GS_roc_file2 = plot_dir + file_name + "_xgboost_GS_roc2.png"
        GS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_xgb_prob_val1, pred_xgb_val1)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_GS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_GS_roc2.png"),
            "ROC Curve Test data", "XGBoost - Grid Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document, os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_GS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_GS_val_data.png"),
            "Confusion Matrix Validation data", "XGBoost - Grid Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, GS_auc_ginin_output, "XGBoost - Grid Search")

        pdf.output(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_XGBoost.pdf"))

        context = {'pdfFile': plot_dir + file_name + "_XGBoost.pdf", 'model': 'XGB', 'tableHead': 'XGBoost', 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                   'NT_graphConfMat1': plot_dir + file_name + "_xgboost_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_xgboost_NT_val_data.png",
                   'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                   'RS_graphConfMat1': plot_dir + file_name + "_xgboost_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_xgboost_RS_val_data.png",
                   'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                   'GS_graphConfMat1': plot_dir + file_name + "_xgboost_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_xgboost_GS_val_data.png"}
        return render(request, 'showModelOutputAll.html', context)
    except Exception as e:
        print(e)
        return render(request, 'error.html')


def xgBoostAjax():
    try:
        csv_file_name = "csvfile_"+user_name
        savefile_x_final = file_path + csv_file_name + "_x_model.csv"
        df = pd.read_csv(savefile_x_final)
        targetVarFile = file_path + csv_file_name + "_targetVar.txt"
        file1 = open(targetVarFile, "r")  # write mode
        targetVar = file1.read()
        file1.close()
        # a variable pdf
        pdf = FPDF()
        # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
        y_model = df[targetVar]
        x_model = df.drop(targetVar, axis=1)
        X_train, X_test, y_train, y_test = train_test_split(
            x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
        X_train, X_val, y_train, y_val = train_test_split(
            X_train, y_train, test_size=0.22222222222222224, random_state=321)

        # Fit XGRegressor to the Training set
        xg_clf = xgb.XGBClassifier()
        xg_clf.fit(X_train, y_train)

        # Test the model
        pred_xgb_val = xg_clf.predict(X_val)
        pred_xgb_prob_val = xg_clf.predict_proba(X_val)[:, 1]

        pred_xgb_train = xg_clf.predict(X_train)
        pred_xgb_prob_train = xg_clf.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_xgb_train))
        print(classification_report(y_val, pred_xgb_val))

        evaluate_model(pred_xgb_val, pred_xgb_prob_val, pred_xgb_train,
                       pred_xgb_prob_train, file_name+"xgboost_NT_roc1.png")

        evaluate_model(pred_xgb_val, pred_xgb_val, pred_xgb_train,
                       pred_xgb_train, file_name+"xgboost_NT_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data
        drawConfMatrix(y_train, pred_xgb_train, file_name +
                       "_xgboost_NT_train_data.png", "Training")

        drawConfMatrix(y_val, pred_xgb_val, file_name +
                       "_xgboost_NT_val_data.png", "Validation")

        NT_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_xgb_prob_val, pred_xgb_val)

        NT_roc_file1 = plot_dir + file_name+"xgboost_NT_roc1.png"
        NT_roc_file2 = plot_dir + file_name+"xgboost_NT_roc2.png"

        # Add a page
        pdf.add_page()
        document = Document()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"xgboost_NT_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"xgboost_NT_roc2.png"),
            "ROC Curve Test data", "XGBoost - No Parameters Tuning")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_NT_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_NT_val_data.png"),
            "Confusion Matrix Validation data", "XGBoost - No Parameters Tuning")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, NT_auc_ginin_output, "XGBoost - No Parameters Tuning")

        paramFiles = param_file_path + param_file_name + "_XGB_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = xgb.XGBClassifier(random_state=50)

        # Create the random search model
        xgb_search = RandomizedSearchCV(estimator, param_grid, n_jobs=-1,
                                        scoring='roc_auc', cv=5,
                                        n_iter=100, verbose=1, random_state=50)

        # Fit
        xgb_search.fit(X_train, y_train)

        xgb_search.best_params_

        xgb_random = xgb_search.best_estimator_

        # Test the model
        pred_xgb_val0 = xgb_random.predict(X_val)
        pred_xgb_prob_val0 = xgb_random.predict_proba(X_val)[:, 1]

        pred_xgb_train0 = xgb_random.predict(X_train)
        pred_xgb_prob_train0 = xgb_random.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_xgb_train0))
        print(classification_report(y_val, pred_xgb_val0))

        evaluate_model(y_val, pred_xgb_prob_val0, y_train,
                       pred_xgb_prob_train0, file_name+"_xgboost_RS_roc1.png")

        evaluate_model(y_val, pred_xgb_val0, y_train,
                       pred_xgb_train0, file_name+"_xgboost_RS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_xgb_train0,
                       file_name + "_xgboost_RS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_xgb_val0,  file_name +
                       "_xgboost_RS_val_data.png", "Validation")

        RS_roc_file1 = plot_dir + file_name + "_xgboost_RS_roc1.png"
        RS_roc_file2 = plot_dir + file_name + "_xgboost_RS_roc2.png"
        RS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_xgb_prob_val0, pred_xgb_val0)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_RS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_RS_roc2.png"),
            "ROC Curve Test data", "XGBoost - Random Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_RS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_RS_val_data.png"),
            "Confusion Matrix Validation data", "XGBoost - Random Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, RS_auc_ginin_output, "XGBoost - Random Search")

        paramFiles = param_file_path + param_file_name + "_XGB_GS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = xgb.XGBClassifier(random_state=50)

        # Create the random search model
        xgb_grid = GridSearchCV(estimator, param_grid, n_jobs=-1,
                                scoring='roc_auc', cv=5, verbose=1)

        # Fit
        xgb_grid.fit(X_train, y_train)

        xgb_grid.best_params_

        xgb_final = xgb_grid.best_estimator_

        # Test the model
        pred_xgb_val1 = xgb_final.predict(X_val)
        pred_xgb_prob_val1 = xgb_final.predict_proba(X_val)[:, 1]

        pred_xgb_train1 = xgb_final.predict(X_train)
        pred_xgb_prob_train1 = xgb_final.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_xgb_train1))
        print(classification_report(y_val, pred_xgb_val1))

        evaluate_model(y_val, pred_xgb_prob_val1, y_train,
                       pred_xgb_prob_train1, file_name+"_xgboost_GS_roc1.png")

        evaluate_model(y_val, pred_xgb_val1, y_train,
                       pred_xgb_train1, file_name+"_xgboost_GS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_xgb_train1,
                       file_name + "_xgboost_GS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_xgb_val1,  file_name +
                       "_xgboost_GS_val_data.png", "Validation")

        GS_roc_file1 = plot_dir + file_name + "_xgboost_GS_roc1.png"
        GS_roc_file2 = plot_dir + file_name + "_xgboost_GS_roc2.png"
        GS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_xgb_prob_val1, pred_xgb_val1)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_GS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_GS_roc2.png"),
            "ROC Curve Test data", "XGBoost - Grid Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document, os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_GS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_xgboost_GS_val_data.png"),
            "Confusion Matrix Validation data", "XGBoost - Grid Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, GS_auc_ginin_output, "XGBoost - Grid Search")

        pdf.output(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_XGBoost.pdf"))

        context = {'pdfFile': plot_dir + file_name + "_XGBoost.pdf", 'model': 'XGB', 'tableHead': 'XGBoost', 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                   'NT_graphConfMat1': plot_dir + file_name + "_xgboost_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_xgboost_NT_val_data.png",
                   'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                   'RS_graphConfMat1': plot_dir + file_name + "_xgboost_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_xgboost_RS_val_data.png",
                   'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                   'GS_graphConfMat1': plot_dir + file_name + "_xgboost_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_xgboost_GS_val_data.png"}
        print('Process XGboost ended')
    except Exception as e:
        print(e)


def MLP(request):
    try:
        csv_file_name = "csvfile_"+user_name
        savefile_x_final = file_path + csv_file_name + "_x_model.csv"
        if(not os.path.exists(savefile_x_final)):
            return render(request, 'processNotdone.html')
        df = pd.read_csv(savefile_x_final)
        targetVarFile = file_path + csv_file_name + "_targetVar.txt"
        file1 = open(targetVarFile, "r")  # write mode
        targetVar = file1.read()
        file1.close()
        # a variable pdf
        pdf = FPDF()
        # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
        y_model = df[targetVar]
        x_model = df.drop(targetVar, axis=1)
        X_train, X_test, y_train, y_test = train_test_split(
            x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
        X_train, X_val, y_train, y_val = train_test_split(
            X_train, y_train, test_size=0.22222222222222224, random_state=321)

        mlp = MLPClassifier(hidden_layer_sizes=100, alpha=0.0001, random_state=1,
                            learning_rate_init=0.001, max_iter=200)

        mlp.fit(X_train, y_train)
        # Test the model
        pred_mlp_val = mlp.predict(X_val)
        pred_mlp_prob_val = mlp.predict_proba(X_val)[:, 1]

        pred_mlp_train = mlp.predict(X_train)
        pred_mlp_prob_train = mlp.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_mlp_train))
        print(classification_report(y_val, pred_mlp_val))

        evaluate_model(y_val, pred_mlp_prob_val, y_train,
                       pred_mlp_prob_train, file_name+"MLP_NT_roc1.png")

        evaluate_model(y_val, pred_mlp_val, y_train, pred_mlp_train,
                       file_name+"MLP_NT_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data
        drawConfMatrix(y_train, pred_mlp_train, file_name +
                       "_MLP_NT_train_data.png", "Training")

        drawConfMatrix(y_val, pred_mlp_val, file_name +
                       "_MLP_NT_val_data.png", "Validation")

        NT_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_mlp_prob_val, pred_mlp_val)

        NT_roc_file1 = plot_dir + file_name+"MLP_NT_roc1.png"
        NT_roc_file2 = plot_dir + file_name+"MLP_NT_roc2.png"

        # Add a page
        pdf.add_page()
        document = Document()
        pdf = exportPdf(10, 10, pdf, document, os.path.join(
            BASE_DIR, plot_dir_view, file_name+"MLP_NT_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"MLP_NT_roc2.png"),
            "ROC Curve Test data", "Multi-Layer Perceptron - No Parameters Tuning")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document, os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_MLP_NT_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_MLP_NT_val_data.png"),
            "Confusion Matrix Validation data", "Multi-Layer Perceptron - No Parameters Tuning")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, NT_auc_ginin_output, "Multi-Layer Perceptron - No Parameters Tuning")

        paramFiles = param_file_path + param_file_name + "_MLP_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = MLPClassifier(max_iter=100)
        # Create the random search model
        mlp_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                        scoring='roc_auc', cv=5,
                                        n_iter=100, verbose=1, random_state=50)

        # Fit
        mlp_search.fit(X_train, y_train)

        mlp_search.best_params_

        mlp_random = mlp_search.best_estimator_
        # Test the model
        pred_mlp_val0 = mlp_random.predict(X_val)
        pred_mlp_prob_val0 = mlp_random.predict_proba(X_val)[:, 1]

        pred_mlp_train0 = mlp_random.predict(X_train)
        pred_mlp_prob_train0 = mlp_random.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_mlp_train0))
        print(classification_report(y_val, pred_mlp_val0))

        evaluate_model(y_val, pred_mlp_prob_val0, y_train,
                       pred_mlp_prob_train0, file_name+"_MLP_RS_roc1.png")

        evaluate_model(y_val, pred_mlp_val0, y_train,
                       pred_mlp_train0, file_name+"_MLP_RS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_mlp_train0,
                       file_name + "_MLP_RS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_mlp_val0,
                       file_name + "_MLP_RS_val_data.png", "Validation")

        RS_roc_file1 = plot_dir + file_name + "_MLP_RS_roc1.png"
        RS_roc_file2 = plot_dir + file_name + "_MLP_RS_roc2.png"
        RS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_mlp_prob_val0, pred_mlp_val0)

        # Add a page
        pdf.add_page()
        document = Document()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_MLP_RS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_MLP_RS_roc2.png"),
            "ROC Curve Test data", "Multi-Layer Perceptron - Random Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_MLP_RS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_MLP_RS_val_data.png"),
            "Confusion Matrix Validation data", "Multi-Layer Perceptron - Random Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, RS_auc_ginin_output, "Multi-Layer Perceptron - Random Search")

        paramFiles = param_file_path + param_file_name + "_MLP_GS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = MLPClassifier(max_iter=500)

        # Create the random search model
        mlp_grid = GridSearchCV(estimator, param_grid, n_jobs=3,
                                scoring='roc_auc', cv=5, verbose=1)

        # Fit
        mlp_grid.fit(X_train, y_train)

        mlp_grid.best_params_

        mlp_final = mlp_grid.best_estimator_

        # Test the model
        pred_mlp_val1 = mlp_final.predict(X_val)
        pred_mlp_prob_val1 = mlp_final.predict_proba(X_val)[:, 1]

        pred_mlp_train1 = mlp_final.predict(X_train)
        pred_mlp_prob_train1 = mlp_final.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_mlp_train1))
        print(classification_report(y_val, pred_mlp_val1))

        evaluate_model(y_val, pred_mlp_prob_val1, y_train,
                       pred_mlp_prob_train1, file_name+"_MLP_GS_roc1.png")

        evaluate_model(y_val, pred_mlp_val1, y_train,
                       pred_mlp_train1, file_name+"_MLP_GS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_mlp_train1,
                       file_name + "_MLP_GS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_mlp_val1,
                       file_name + "_MLP_GS_val_data.png", "Validation")

        GS_roc_file1 = plot_dir + file_name + "_MLP_GS_roc1.png"
        GS_roc_file2 = plot_dir + file_name + "_MLP_GS_roc2.png"
        GS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_mlp_prob_val1, pred_mlp_val1)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_MLP_GS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_MLP_GS_roc2.png"),
            "ROC Curve Test data", "Multi-Layer Perceptron - Grid Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_MLP_GS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_MLP_GS_val_data.png"),
            "Confusion Matrix Validation data", "Multi-Layer Perceptron - Grid Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, GS_auc_ginin_output, "Multi-Layer Perceptron - Grid Search")

        pdf.output(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_MLP.pdf"))

        context = {'pdfFile': plot_dir + file_name + "_MLP.pdf", 'model': 'MLP', 'tableHead': 'Multi-Layer Perceptron (MLP)', 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                   'NT_graphConfMat1': plot_dir + file_name + "_MLP_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_MLP_NT_val_data.png",
                   'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                   'RS_graphConfMat1': plot_dir + file_name + "_MLP_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_MLP_RS_val_data.png",
                   'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                   'GS_graphConfMat1': plot_dir + file_name + "_MLP_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_MLP_GS_val_data.png"}
        return render(request, 'showModelOutputAll.html', context)
    except Exception as e:
        print(e)
        return render(request, 'error.html')


def MLPAjax():
    try:
        csv_file_name = "csvfile_"+user_name
        savefile_x_final = file_path + csv_file_name + "_x_model.csv"
        df = pd.read_csv(savefile_x_final)
        targetVarFile = file_path + csv_file_name + "_targetVar.txt"
        file1 = open(targetVarFile, "r")  # write mode
        targetVar = file1.read()
        file1.close()
        # a variable pdf
        pdf = FPDF()
        # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
        y_model = df[targetVar]
        x_model = df.drop(targetVar, axis=1)
        X_train, X_test, y_train, y_test = train_test_split(
            x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
        X_train, X_val, y_train, y_val = train_test_split(
            X_train, y_train, test_size=0.22222222222222224, random_state=321)

        mlp = MLPClassifier(hidden_layer_sizes=100, alpha=0.0001, random_state=1,
                            learning_rate_init=0.001, max_iter=200)

        mlp.fit(X_train, y_train)
        # Test the model
        pred_mlp_val = mlp.predict(X_val)
        pred_mlp_prob_val = mlp.predict_proba(X_val)[:, 1]

        pred_mlp_train = mlp.predict(X_train)
        pred_mlp_prob_train = mlp.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_mlp_train))
        print(classification_report(y_val, pred_mlp_val))

        evaluate_model(y_val, pred_mlp_prob_val, y_train,
                       pred_mlp_prob_train, file_name+"MLP_NT_roc1.png")

        evaluate_model(y_val, pred_mlp_val, y_train, pred_mlp_train,
                       file_name+"MLP_NT_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data
        drawConfMatrix(y_train, pred_mlp_train, file_name +
                       "_MLP_NT_train_data.png", "Training")

        drawConfMatrix(y_val, pred_mlp_val, file_name +
                       "_MLP_NT_val_data.png", "Validation")

        NT_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_mlp_prob_val, pred_mlp_val)

        NT_roc_file1 = plot_dir + file_name+"MLP_NT_roc1.png"
        NT_roc_file2 = plot_dir + file_name+"MLP_NT_roc2.png"

        # Add a page
        pdf.add_page()
        document = Document()
        pdf = exportPdf(10, 10, pdf, document, os.path.join(
            BASE_DIR, plot_dir_view, file_name+"MLP_NT_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"MLP_NT_roc2.png"),
            "ROC Curve Test data", "Multi-Layer Perceptron - No Parameters Tuning")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document, os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_MLP_NT_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_MLP_NT_val_data.png"),
            "Confusion Matrix Validation data", "Multi-Layer Perceptron - No Parameters Tuning")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, NT_auc_ginin_output, "Multi-Layer Perceptron - No Parameters Tuning")

        paramFiles = param_file_path + param_file_name + "_MLP_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = MLPClassifier(max_iter=100)
        # Create the random search model
        mlp_search = RandomizedSearchCV(estimator, param_grid, n_jobs=-1,
                                        scoring='roc_auc', cv=5,
                                        n_iter=100, verbose=1, random_state=50)

        # Fit
        mlp_search.fit(X_train, y_train)

        mlp_search.best_params_

        mlp_random = mlp_search.best_estimator_
        # Test the model
        pred_mlp_val0 = mlp_random.predict(X_val)
        pred_mlp_prob_val0 = mlp_random.predict_proba(X_val)[:, 1]

        pred_mlp_train0 = mlp_random.predict(X_train)
        pred_mlp_prob_train0 = mlp_random.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_mlp_train0))
        print(classification_report(y_val, pred_mlp_val0))

        evaluate_model(y_val, pred_mlp_prob_val0, y_train,
                       pred_mlp_prob_train0, file_name+"_MLP_RS_roc1.png")

        evaluate_model(y_val, pred_mlp_val0, y_train,
                       pred_mlp_train0, file_name+"_MLP_RS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_mlp_train0,
                       file_name + "_MLP_RS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_mlp_val0,
                       file_name + "_MLP_RS_val_data.png", "Validation")

        RS_roc_file1 = plot_dir + file_name + "_MLP_RS_roc1.png"
        RS_roc_file2 = plot_dir + file_name + "_MLP_RS_roc2.png"
        RS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_mlp_prob_val0, pred_mlp_val0)

        # Add a page
        pdf.add_page()
        document = Document()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_MLP_RS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_MLP_RS_roc2.png"),
            "ROC Curve Test data", "Multi-Layer Perceptron - Random Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_MLP_RS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_MLP_RS_val_data.png"),
            "Confusion Matrix Validation data", "Multi-Layer Perceptron - Random Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, RS_auc_ginin_output, "Multi-Layer Perceptron - Random Search")

        paramFiles = param_file_path + param_file_name + "_MLP_GS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = MLPClassifier(max_iter=500)

        # Create the random search model
        mlp_grid = GridSearchCV(estimator, param_grid, n_jobs=-1,
                                scoring='roc_auc', cv=5, verbose=1)

        # Fit
        mlp_grid.fit(X_train, y_train)

        mlp_grid.best_params_

        mlp_final = mlp_grid.best_estimator_

        # Test the model
        pred_mlp_val1 = mlp_final.predict(X_val)
        pred_mlp_prob_val1 = mlp_final.predict_proba(X_val)[:, 1]

        pred_mlp_train1 = mlp_final.predict(X_train)
        pred_mlp_prob_train1 = mlp_final.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_mlp_train1))
        print(classification_report(y_val, pred_mlp_val1))

        evaluate_model(y_val, pred_mlp_prob_val1, y_train,
                       pred_mlp_prob_train1, file_name+"_MLP_GS_roc1.png")

        evaluate_model(y_val, pred_mlp_val1, y_train,
                       pred_mlp_train1, file_name+"_MLP_GS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_mlp_train1,
                       file_name + "_MLP_GS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_mlp_val1,
                       file_name + "_MLP_GS_val_data.png", "Validation")

        GS_roc_file1 = plot_dir + file_name + "_MLP_GS_roc1.png"
        GS_roc_file2 = plot_dir + file_name + "_MLP_GS_roc2.png"
        GS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_mlp_prob_val1, pred_mlp_val1)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_MLP_GS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_MLP_GS_roc2.png"),
            "ROC Curve Test data", "Multi-Layer Perceptron - Grid Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_MLP_GS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_MLP_GS_val_data.png"),
            "Confusion Matrix Validation data", "Multi-Layer Perceptron - Grid Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, GS_auc_ginin_output, "Multi-Layer Perceptron - Grid Search")

        pdf.output(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_MLP.pdf"))

        context = {'pdfFile': plot_dir + file_name + "_MLP.pdf", 'model': 'MLP', 'tableHead': 'Multi-Layer Perceptron (MLP)', 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                   'NT_graphConfMat1': plot_dir + file_name + "_MLP_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_MLP_NT_val_data.png",
                   'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                   'RS_graphConfMat1': plot_dir + file_name + "_MLP_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_MLP_RS_val_data.png",
                   'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                   'GS_graphConfMat1': plot_dir + file_name + "_MLP_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_MLP_GS_val_data.png"}

    except Exception as e:
        print(e)


def GBC(request):
    try:
        csv_file_name = "csvfile_"+user_name
        savefile_x_final = file_path + csv_file_name + "_x_model.csv"
        if(not os.path.exists(savefile_x_final)):
            return render(request, 'processNotdone.html')
        df = pd.read_csv(savefile_x_final)
        targetVarFile = file_path + csv_file_name + "_targetVar.txt"
        file1 = open(targetVarFile, "r")  # write mode
        targetVar = file1.read()
        file1.close()
        # a variable pdf
        pdf = FPDF()
        # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
        y_model = df[targetVar]
        x_model = df.drop(targetVar, axis=1)
        X_train, X_test, y_train, y_test = train_test_split(
            x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
        X_train, X_val, y_train, y_val = train_test_split(
            X_train, y_train, test_size=0.22222222222222224, random_state=321)

        gbc = GradientBoostingClassifier()
        gbc.fit(X_train, y_train)

        # Test the model
        pred_gbc_val = gbc.predict(X_val)
        pred_gbc_prob_val = gbc.predict_proba(X_val)[:, 1]

        pred_gbc_train = gbc.predict(X_train)
        pred_gbc_prob_train = gbc.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_gbc_train))
        print(classification_report(y_val, pred_gbc_val))

        evaluate_model(y_val, pred_gbc_prob_val, y_train,
                       pred_gbc_prob_train, file_name+"GBC_NT_roc1.png")

        evaluate_model(y_val, pred_gbc_val, y_train, pred_gbc_train,
                       file_name+"GBC_NT_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data
        drawConfMatrix(y_train, pred_gbc_train, file_name +
                       "_GBC_NT_train_data.png", "Training")

        drawConfMatrix(y_val, pred_gbc_val, file_name +
                       "_GBC_NT_val_data.png", "Validation")

        NT_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_gbc_prob_val, pred_gbc_val)
        NT_roc_file1 = plot_dir + file_name+"GBC_NT_roc1.png"
        NT_roc_file2 = plot_dir + file_name+"GBC_NT_roc2.png"

        # Add a page
        pdf.add_page()
        document = Document()
        pdf = exportPdf(10, 10, pdf, document, os.path.join(
            BASE_DIR, plot_dir_view, file_name+"GBC_NT_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"GBC_NT_roc2.png"),
            "ROC Curve Test data", "Gradient Boosting - No Parameters Tuning")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_GBC_NT_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_GBC_NT_val_data.png"),
            "Confusion Matrix Validation data", "Gradient Boosting - No Parameters Tuning")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, NT_auc_ginin_output, "Gradient Boosting - No Parameters Tuning")

        paramFiles = param_file_path + param_file_name + "_GB_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        seed(1)
     # Estimator for use in random search
        estimator = GradientBoostingClassifier()

        # Create the random search model
        gbc_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                        scoring='roc_auc', cv=5,
                                        n_iter=100, verbose=1, random_state=50)

        # Fit
        gbc_search.fit(X_train, y_train)

        gbc_random = gbc_search.best_estimator_
        # Test the model
        pred_gbc_val0 = gbc_random.predict(X_val)
        pred_gbc_prob_val0 = gbc_random.predict_proba(X_val)[:, 1]

        pred_gbc_train0 = gbc_random.predict(X_train)
        pred_gbc_prob_train0 = gbc_random.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_gbc_train0))
        print(classification_report(y_val, pred_gbc_val0))

        evaluate_model(y_val, pred_gbc_prob_val0, y_train,
                       pred_gbc_prob_train0, file_name+"_GBC_RS_roc1.png")

        evaluate_model(y_val, pred_gbc_val0, y_train,
                       pred_gbc_train0, file_name+"_GBC_RS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_gbc_train0,
                       file_name + "_GBC_RS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_gbc_val0,
                       file_name + "_GBC_RS_val_data.png", "Validation")

        RS_roc_file1 = plot_dir + file_name + "_GBC_RS_roc1.png"
        RS_roc_file2 = plot_dir + file_name + "_GBC_RS_roc2.png"
        RS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_gbc_prob_val0, pred_gbc_val0)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_GBC_RS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_GBC_RS_roc2.png"),
            "ROC Curve Test data", "Gradient Boosting - Random Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_GBC_RS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_GBC_RS_val_data.png"),
            "Confusion Matrix Validation data", "Gradient Boosting - Random Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, RS_auc_ginin_output, "Gradient Boosting - Random Search")

        paramFiles = param_file_path + param_file_name + "_GB_GS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = GradientBoostingClassifier()

        # Create the random search model
        gbc_grid = GridSearchCV(estimator, param_grid, n_jobs=3,
                                scoring='roc_auc', cv=5, verbose=1)

        # Fit
        gbc_grid.fit(X_train, y_train)

        gbc_final = gbc_grid.best_estimator_

        # Test the model
        pred_gbc_val1 = gbc_final.predict(X_val)
        pred_gbc_prob_val1 = gbc_final.predict_proba(X_val)[:, 1]

        pred_gbc_train1 = gbc_final.predict(X_train)
        pred_gbc_prob_train1 = gbc_final.predict_proba(X_train)[:, 1]
        # Get the model performance
        print(classification_report(y_train, pred_gbc_train1))
        print(classification_report(y_val, pred_gbc_val1))

        evaluate_model(y_val, pred_gbc_prob_val1, y_train,
                       pred_gbc_prob_train1, file_name+"_GBC_GS_roc1.png")

        evaluate_model(y_val, pred_gbc_val1, y_train,
                       pred_gbc_train1, file_name+"_GBC_GS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_gbc_train1,
                       file_name + "_GBC_GS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_gbc_val1,
                       file_name + "_GBC_GS_val_data.png", "Validation")

        GS_roc_file1 = plot_dir + file_name + "_GBC_GS_roc1.png"
        GS_roc_file2 = plot_dir + file_name + "_GBC_GS_roc2.png"
        GS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_gbc_prob_val1, pred_gbc_val1)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_GBC_GS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_GBC_GS_roc2.png"),
            "ROC Curve Test data", "Gradient Boosting - Grid Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_GBC_GS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_GBC_GS_val_data.png"),
            "Confusion Matrix Validation data", "Gradient Boosting - Grid Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, GS_auc_ginin_output, "Gradient Boosting - Grid Search")

        pdf.output(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_Gradient_Boosting.pdf"))

        context = {'pdfFile': plot_dir + file_name + "_Gradient_Boosting.pdf", 'model': 'GBC',  'tableHead': 'Gradient Boosting', 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                   'NT_graphConfMat1': plot_dir + file_name + "_GBC_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_GBC_NT_val_data.png",
                   'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                   'RS_graphConfMat1': plot_dir + file_name + "_GBC_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_GBC_RS_val_data.png",
                   'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                   'GS_graphConfMat1': plot_dir + file_name + "_GBC_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_GBC_GS_val_data.png"}
    except Exception as e:
        print(e)
        context = {'tableHead': 'Error while processing reuest.', }
    return render(request, 'showModelOutputAll.html', context)


def GBCAjax():
    try:
        csv_file_name = "csvfile_"+user_name
        savefile_x_final = file_path + csv_file_name + "_x_model.csv"
        df = pd.read_csv(savefile_x_final)
        targetVarFile = file_path + csv_file_name + "_targetVar.txt"
        file1 = open(targetVarFile, "r")  # write mode
        targetVar = file1.read()
        file1.close()
        # a variable pdf
        pdf = FPDF()
        # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
        y_model = df[targetVar]
        x_model = df.drop(targetVar, axis=1)
        X_train, X_test, y_train, y_test = train_test_split(
            x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
        X_train, X_val, y_train, y_val = train_test_split(
            X_train, y_train, test_size=0.22222222222222224, random_state=321)

        gbc = GradientBoostingClassifier()
        gbc.fit(X_train, y_train)

        # Test the model
        pred_gbc_val = gbc.predict(X_val)
        pred_gbc_prob_val = gbc.predict_proba(X_val)[:, 1]

        pred_gbc_train = gbc.predict(X_train)
        pred_gbc_prob_train = gbc.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_gbc_train))
        print(classification_report(y_val, pred_gbc_val))

        evaluate_model(y_val, pred_gbc_prob_val, y_train,
                       pred_gbc_prob_train, file_name+"GBC_NT_roc1.png")

        evaluate_model(y_val, pred_gbc_val, y_train, pred_gbc_train,
                       file_name+"GBC_NT_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data
        drawConfMatrix(y_train, pred_gbc_train, file_name +
                       "_GBC_NT_train_data.png", "Training")

        drawConfMatrix(y_val, pred_gbc_val, file_name +
                       "_GBC_NT_val_data.png", "Validaiton")

        NT_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_gbc_prob_val, pred_gbc_val)
        NT_roc_file1 = plot_dir + file_name+"GBC_NT_roc1.png"
        NT_roc_file2 = plot_dir + file_name+"GBC_NT_roc2.png"

        # Add a page
        pdf.add_page()
        document = Document()
        pdf = exportPdf(10, 10, pdf, document, os.path.join(
            BASE_DIR, plot_dir_view, file_name+"GBC_NT_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"GBC_NT_roc2.png"),
            "ROC Curve Test data", "Gradient Boosting - No Parameters Tuning")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_GBC_NT_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_GBC_NT_val_data.png"),
            "Confusion Matrix Validation data", "Gradient Boosting - No Parameters Tuning")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, NT_auc_ginin_output, "Gradient Boosting - No Parameters Tuning")

        paramFiles = param_file_path + param_file_name + "_GB_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        seed(1)
     # Estimator for use in random search
        estimator = GradientBoostingClassifier()

        # Create the random search model
        gbc_search = RandomizedSearchCV(estimator, param_grid, n_jobs=-1,
                                        scoring='roc_auc', cv=5,
                                        n_iter=100, verbose=1, random_state=50)

        # Fit
        gbc_search.fit(X_train, y_train)

        gbc_random = gbc_search.best_estimator_
        # Test the model
        pred_gbc_val0 = gbc_random.predict(X_val)
        pred_gbc_prob_val0 = gbc_random.predict_proba(X_val)[:, 1]

        pred_gbc_train0 = gbc_random.predict(X_train)
        pred_gbc_prob_train0 = gbc_random.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_gbc_train0))
        print(classification_report(y_val, pred_gbc_val0))

        evaluate_model(y_val, pred_gbc_prob_val0, y_train,
                       pred_gbc_prob_train0, file_name+"_GBC_RS_roc1.png")

        evaluate_model(y_val, pred_gbc_val0, y_train,
                       pred_gbc_train0, file_name+"_GBC_RS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_gbc_train0,
                       file_name + "_GBC_RS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_gbc_val0,
                       file_name + "_GBC_RS_val_data.png", "Validaiton")

        RS_roc_file1 = plot_dir + file_name + "_GBC_RS_roc1.png"
        RS_roc_file2 = plot_dir + file_name + "_GBC_RS_roc2.png"
        RS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_gbc_prob_val0, pred_gbc_val0)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_GBC_RS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_GBC_RS_roc2.png"),
            "ROC Curve Test data", "Gradient Boosting - Random Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_GBC_RS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_GBC_RS_val_data.png"),
            "Confusion Matrix Validation data", "Gradient Boosting - Random Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, RS_auc_ginin_output, "Gradient Boosting - Random Search")

        paramFiles = param_file_path + param_file_name + "_GB_GS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = GradientBoostingClassifier()

        # Create the random search model
        gbc_grid = GridSearchCV(estimator, param_grid, n_jobs=-1,
                                scoring='roc_auc', cv=5, verbose=1)

        # Fit
        gbc_grid.fit(X_train, y_train)

        gbc_final = gbc_grid.best_estimator_

        # Test the model
        pred_gbc_val1 = gbc_final.predict(X_val)
        pred_gbc_prob_val1 = gbc_final.predict_proba(X_val)[:, 1]

        pred_gbc_train1 = gbc_final.predict(X_train)
        pred_gbc_prob_train1 = gbc_final.predict_proba(X_train)[:, 1]
        # Get the model performance
        print(classification_report(y_train, pred_gbc_train1))
        print(classification_report(y_val, pred_gbc_val1))

        evaluate_model(y_val, pred_gbc_prob_val1, y_train,
                       pred_gbc_prob_train1, file_name+"_GBC_GS_roc1.png")

        evaluate_model(y_val, pred_gbc_val1, y_train,
                       pred_gbc_train1, file_name+"_GBC_GS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_gbc_train1,
                       file_name + "_GBC_GS_train_data.png", "Training")

        # show the confusion matrix for validaiton data

        drawConfMatrix(y_val, pred_gbc_val1,
                       file_name + "_GBC_GS_val_data.png", "Validaiton")

        GS_roc_file1 = plot_dir + file_name + "_GBC_GS_roc1.png"
        GS_roc_file2 = plot_dir + file_name + "_GBC_GS_roc2.png"
        GS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_gbc_prob_val1, pred_gbc_val1)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_GBC_GS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_GBC_GS_roc2.png"),
            "ROC Curve Test data", "Gradient Boosting - Grid Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_GBC_GS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_GBC_GS_val_data.png"),
            "Confusion Matrix Validation data", "Gradient Boosting - Grid Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, GS_auc_ginin_output, "Gradient Boosting - Grid Search")

        pdf.output(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_Gradient_Boosting.pdf"))

        context = {'pdfFile': plot_dir + file_name + "_Gradient_Boosting.pdf", 'model': 'GBC',  'tableHead': 'Gradient Boosting', 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                   'NT_graphConfMat1': plot_dir + file_name + "_GBC_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_GBC_NT_val_data.png",
                   'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                   'RS_graphConfMat1': plot_dir + file_name + "_GBC_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_GBC_RS_val_data.png",
                   'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                   'GS_graphConfMat1': plot_dir + file_name + "_GBC_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_GBC_GS_val_data.png"}
        print('Bagging Cls done')
    except Exception as e:
        print(e)
        context = {'tableHead': 'Error while processing reuest.', }


def BCAjax():
    try:
        csv_file_name = "csvfile_"+user_name
        savefile_x_final = file_path + csv_file_name + "_x_model.csv"
        df = pd.read_csv(savefile_x_final)
        targetVarFile = file_path + csv_file_name + "_targetVar.txt"
        file1 = open(targetVarFile, "r")  # write mode
        targetVar = file1.read()
        file1.close()
        # a variable pdf
        pdf = FPDF()
        # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
        y_model = df[targetVar]
        x_model = df.drop(targetVar, axis=1)
        X_train, X_test, y_train, y_test = train_test_split(
            x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
        X_train, X_val, y_train, y_val = train_test_split(
            X_train, y_train, test_size=0.22222222222222224, random_state=321)

        bc_model = BaggingClassifier()
        bc_model.fit(X_train, y_train)

        # Test the model
        pred_bc_val = bc_model.predict(X_val)
        pred_bc_prob_val = bc_model.predict_proba(X_val)[:, 1]

        pred_bc_train = bc_model.predict(X_train)
        pred_bc_prob_train = bc_model.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_bc_train))
        print(classification_report(y_val, pred_bc_val))

        evaluate_model(y_val, pred_bc_prob_val, y_train,
                       pred_bc_prob_train, file_name+"BC_NT_roc1.png")

        evaluate_model(y_val, pred_bc_val, y_train,
                       pred_bc_train, file_name+"BC_NT_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data
        drawConfMatrix(y_train, pred_bc_train, file_name +
                       "BC_NT_train_data.png","Training")

        drawConfMatrix(y_val, pred_bc_val, file_name +
                       "BC_NT_val_data.png","Validation")

        NT_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_bc_prob_val, pred_bc_val)

        # Add a page
        pdf.add_page()
        document = Document()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"BC_NT_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"BC_NT_roc2.png"),
            "ROC Curve Test data", "Bagging Classifier - No Parameters Tuning")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_BC_NT_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_BC_NT_val_data.png"),
            "Confusion Matrix Validation data", "Bagging Classifier - No Parameters Tuning")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, NT_auc_ginin_output, "Bagging Classifier - No Parameters Tuning")

        NT_roc_file1 = plot_dir + file_name+"BC_NT_roc1.png"
        NT_roc_file2 = plot_dir + file_name+"BC_NT_roc2.png"

        paramFiles = param_file_path + param_file_name + "_BC_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = BaggingClassifier(random_state=50)

        # Create the random search model
        bc_search = RandomizedSearchCV(estimator, param_grid, n_jobs=-1,
                                       scoring='roc_auc', cv=5,
                                       n_iter=100, verbose=1, random_state=50)

        # Fit
        bc_search.fit(X_train, y_train)

        bc_random = bc_search.best_estimator_

        # Test the model
        pred_bc_val0 = bc_random.predict(X_val)
        pred_bc_prob_val0 = bc_random.predict_proba(X_val)[:, 1]

        pred_bc_train0 = bc_random.predict(X_train)
        pred_bc_prob_train0 = bc_random.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_bc_train0))
        print(classification_report(y_val, pred_bc_val0))

        evaluate_model(y_val, pred_bc_prob_val0, y_train,
                       pred_bc_prob_train0, file_name+"_BC_RS_roc1.png")

        evaluate_model(y_val, pred_bc_val0, y_train,
                       pred_bc_train0, file_name+"_BC_RS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_bc_train0,
                       file_name + "_BC_RS_train_data.png","Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_bc_val0,
                       file_name + "_BC_RS_val_data.png","Validation")

        RS_roc_file1 = plot_dir + file_name + "_BC_RS_roc1.png"
        RS_roc_file2 = plot_dir + file_name + "_BC_RS_roc2.png"
        RS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_bc_prob_val0, pred_bc_val0)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_BC_RS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_BC_RS_roc2.png"),
            "ROC Curve Test data", "Bagging Classifier - Random Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_BC_RS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_BC_RS_val_data.png"),
            "Confusion Matrix Validation data", "Bagging Classifier - Random Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, RS_auc_ginin_output, "Bagging Classifier - Random Search")

        paramFiles = param_file_path + param_file_name + "_BC_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = BaggingClassifier(random_state=50)

        # Create the random search model
        bc_search = RandomizedSearchCV(estimator, param_grid, n_jobs=-1,
                                       scoring='roc_auc', cv=5,
                                       n_iter=100, verbose=1, random_state=50)

        # Fit
        bc_search.fit(X_train, y_train)

        paramFiles = param_file_path + param_file_name + "_BC_GS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = BaggingClassifier(random_state=50)

        # Create the random search model
        bc_grid = GridSearchCV(estimator, param_grid, n_jobs=3,
                               scoring='roc_auc', cv=5, verbose=1)

        # Fit
        bc_grid.fit(X_train, y_train)

        bc_final = bc_grid.best_estimator_

        pred_bc_val1 = bc_final.predict(X_val)
        pred_bc_prob_val1 = bc_final.predict_proba(X_val)[:, 1]

        pred_bc_train1 = bc_final.predict(X_train)
        pred_bc_prob_train1 = bc_final.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_bc_train1))
        print(classification_report(y_val, pred_bc_val1))

        evaluate_model(y_val, pred_bc_prob_val1, y_train,
                       pred_bc_prob_train1, file_name+"_BC_GS_roc1.png")

        evaluate_model(y_val, pred_bc_val1, y_train,
                       pred_bc_train1, file_name+"_BC_GS_roc2.png")
        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_bc_train1,
                       file_name + "_BC_GS_train_data.png","Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_bc_val1,
                       file_name + "_BC_GS_val_data.png","Validation")

        GS_roc_file1 = plot_dir + file_name + "_BC_GS_roc1.png"
        GS_roc_file2 = plot_dir + file_name + "_BC_GS_roc2.png"
        GS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_bc_prob_val1, pred_bc_val1)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_BC_GS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_BC_GS_roc2.png"),
            "ROC Curve Test data", "Bagging Classifier - Grid Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_BC_GS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_BC_GS_val_data.png"),
            "Confusion Matrix Validation data", "Bagging Classifier - Grid Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, GS_auc_ginin_output, "Bagging Classifier - Grid Search")

        pdf.output(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_BagCls.pdf"))

        context = {'pdfFile': plot_dir + file_name + "_BagCls.pdf", 'model': 'BC',  'tableHead': 'Bagging Classifier', 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                   'NT_graphConfMat1': plot_dir + file_name + "_BC_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_BC_NT_val_data.png",
                   'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                   'RS_graphConfMat1': plot_dir + file_name + "_BC_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_BC_RS_val_data.png",
                   'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                   'GS_graphConfMat1': plot_dir + file_name + "_BC_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_BC_GS_val_data.png"}

    except Exception as e:
        print(e)


def KNN(request):
    try:
        csv_file_name = "csvfile_"+user_name
        savefile_x_final = file_path + csv_file_name + "_x_model.csv"
        if(not os.path.exists(savefile_x_final)):
            return render(request, 'processNotdone.html')
        df = pd.read_csv(savefile_x_final)
        targetVarFile = file_path + csv_file_name + "_targetVar.txt"
        file1 = open(targetVarFile, "r")  # write mode
        targetVar = file1.read()
        file1.close()
        # a variable pdf
        pdf = FPDF()
        # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
        y_model = df[targetVar]
        x_model = df.drop(targetVar, axis=1)
        X_train, X_test, y_train, y_test = train_test_split(
            x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
        X_train, X_val, y_train, y_val = train_test_split(
            X_train, y_train, test_size=0.22222222222222224, random_state=321)

        knn = KNeighborsClassifier()
        knn.fit(X_train, y_train)

        # Test the model
        pred_knn_val = knn.predict(X_val)
        pred_knn_prob_val = knn.predict_proba(X_val)[:, 1]

        pred_knn_train = knn.predict(X_train)
        pred_knn_prob_train = knn.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_knn_train))
        print(classification_report(y_val, pred_knn_val))
        x = 10.0
        y = 10.0
        NT_roc_file1 = plot_dir + file_name+"KNN_NT_roc1.png"
        NT_roc_file2 = plot_dir + file_name+"KNN_NT_roc2.png"

        evaluate_model(y_val, pred_knn_prob_val, y_train,
                       pred_knn_prob_train, file_name+"KNN_NT_roc1.png")

        evaluate_model(y_val, pred_knn_val, y_train,
                       pred_knn_train, file_name+"KNN_NT_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data
        drawConfMatrix(y_train, pred_knn_train, file_name +
                       "_KNN_NT_train_data.png", "Training")

        drawConfMatrix(y_val, pred_knn_val, file_name +
                       "_KNN_NT_val_data.png", "Validation")
        NT_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_knn_prob_val, pred_knn_val)

        # Add a page
        pdf.add_page()
        document = Document()
        pdf = exportPdf(x, y, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"KNN_NT_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"KNN_NT_roc2.png"),
            "ROC Curve Test data", "K Nearest Neighbors - No Parameters Tuning")
        pdf.add_page()
        pdf = exportPdf(x, y, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_KNN_NT_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_KNN_NT_val_data.png"),
            "Confusion Matrix Validation data", "K Nearest Neighbors - No Parameters Tuning")
        pdf.add_page()
        pdf = exportTestResultPdf(
            x, y, pdf, NT_auc_ginin_output, "K Nearest Neighbors - No Parameters Tuning")

        paramFiles = param_file_path + param_file_name + "_KNN_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = KNeighborsClassifier()

        # Create the random search model
        knn_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                        scoring='roc_auc', cv=5,
                                        n_iter=100, verbose=1, random_state=50)

        # Fit
        knn_search.fit(X_train, y_train)

        knn_random = knn_search.best_estimator_
        # Test the model
        pred_knn_val0 = knn_random.predict(X_val)
        pred_knn_prob_val0 = knn_random.predict_proba(X_val)[:, 1]

        pred_knn_train0 = knn_random.predict(X_train)
        pred_knn_prob_train0 = knn_random.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_knn_train0))
        print(classification_report(y_val, pred_knn_val0))

        evaluate_model(y_val, pred_knn_prob_val0, y_train,
                       pred_knn_prob_train0, file_name+"_KNN_RS_roc1.png")

        evaluate_model(y_val, pred_knn_val0, y_train,
                       pred_knn_train0, file_name+"_KNN_RS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_knn_train0,
                       file_name + "_KNN_RS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_knn_val0,
                       file_name + "_KNN_RS_val_data.png", "Validation")

        RS_roc_file1 = plot_dir + file_name + "_KNN_RS_roc1.png"
        RS_roc_file2 = plot_dir + file_name + "_KNN_RS_roc2.png"
        RS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_knn_prob_val0, pred_knn_val0)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(x, y, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_KNN_RS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_KNN_RS_roc2.png"),
            "ROC Curve Test data", "K Nearest Neighbors - Random Search")
        pdf.add_page()
        pdf = exportPdf(x, y, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_KNN_RS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_KNN_RS_val_data.png"),
            "Confusion Matrix Validation data", "K Nearest Neighbors - Random Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            x, y, pdf, RS_auc_ginin_output, "K Nearest Neighbors - Random Search")

        paramFiles = param_file_path + param_file_name + "_KNN_GS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = KNeighborsClassifier()

        # Create the random search model
        knn_grid = GridSearchCV(estimator, param_grid, n_jobs=3,
                                scoring='roc_auc', cv=5, verbose=1)

        # Fit
        knn_grid.fit(X_train, y_train)

        knn_final = knn_grid.best_estimator_

        # Test the model
        pred_knn_val1 = knn_final.predict(X_val)
        pred_knn_prob_val1 = knn_final.predict_proba(X_val)[:, 1]

        pred_knn_train1 = knn_final.predict(X_train)
        pred_knn_prob_train1 = knn_final.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_knn_train1))
        print(classification_report(y_val, pred_knn_val1))

        evaluate_model(y_val, pred_knn_prob_val1, y_train,
                       pred_knn_prob_train1, file_name+"_KNN_GS_roc1.png")

        evaluate_model(y_val, pred_knn_val1, y_train,
                       pred_knn_train1, file_name+"_KNN_GS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_knn_train1,
                       file_name + "_KNN_GS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_knn_val1,
                       file_name + "_KNN_GS_val_data.png", "Validation")

        GS_roc_file1 = plot_dir + file_name + "_KNN_GS_roc1.png"
        GS_roc_file2 = plot_dir + file_name + "_KNN_GS_roc2.png"
        GS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_knn_prob_val1, pred_knn_val1)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(x, y, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_KNN_GS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_KNN_GS_roc2.png"),
            "ROC Curve Test data", "K Nearest Neighbors - Grid Search")
        pdf.add_page()
        pdf = exportPdf(x, y, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_KNN_GS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_KNN_GS_val_data.png", "Validation"),
            "Confusion Matrix Validation data", "K Nearest Neighbors - Grid Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            x, y, pdf, GS_auc_ginin_output, "K Nearest Neighbors - Grid Search")

        pdf.output(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_KNN.pdf"))

        context = {'pdfFile': plot_dir + file_name + "_KNN.pdf", 'model': 'KNN',  'tableHead': 'K Nearest Neighbors', 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                   'NT_graphConfMat1': plot_dir + file_name + "_KNN_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_KNN_NT_val_data.png",
                   'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                   'RS_graphConfMat1': plot_dir + file_name + "_KNN_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_KNN_RS_val_data.png",
                   'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                   'GS_graphConfMat1': plot_dir + file_name + "_KNN_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_KNN_GS_val_data.png"}
        return render(request, 'showModelOutputAll.html', context)
    except Exception as e:
        print(e)
        return render(request, 'error.html')


def KNNAjax():
    try:
        csv_file_name = "csvfile_"+user_name
        savefile_x_final = file_path + csv_file_name + "_x_model.csv"
        df = pd.read_csv(savefile_x_final)
        targetVarFile = file_path + csv_file_name + "_targetVar.txt"
        file1 = open(targetVarFile, "r")  # write mode
        targetVar = file1.read()
        file1.close()
        # a variable pdf
        pdf = FPDF()
        # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
        y_model = df[targetVar]
        x_model = df.drop(targetVar, axis=1)
        X_train, X_test, y_train, y_test = train_test_split(
            x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
        X_train, X_val, y_train, y_val = train_test_split(
            X_train, y_train, test_size=0.22222222222222224, random_state=321)

        knn = KNeighborsClassifier()
        knn.fit(X_train, y_train)

        # Test the model
        pred_knn_val = knn.predict(X_val)
        pred_knn_prob_val = knn.predict_proba(X_val)[:, 1]

        pred_knn_train = knn.predict(X_train)
        pred_knn_prob_train = knn.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_knn_train))
        print(classification_report(y_val, pred_knn_val))
        x = 10.0
        y = 10.0
        NT_roc_file1 = plot_dir + file_name+"KNN_NT_roc1.png"
        NT_roc_file2 = plot_dir + file_name+"KNN_NT_roc2.png"

        evaluate_model(y_val, pred_knn_prob_val, y_train,
                       pred_knn_prob_train, file_name+"KNN_NT_roc1.png")

        evaluate_model(y_val, pred_knn_val, y_train,
                       pred_knn_train, file_name+"KNN_NT_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data
        drawConfMatrix(y_train, pred_knn_train, file_name +
                       "_KNN_NT_train_data.png", "Training")

        drawConfMatrix(y_val, pred_knn_val, file_name +
                       "_KNN_NT_val_data.png", "Validation")
        NT_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_knn_prob_val, pred_knn_val)

        # Add a page
        pdf.add_page()
        document = Document()
        pdf = exportPdf(x, y, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"KNN_NT_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"KNN_NT_roc2.png"),
            "ROC Curve Test data", "K Nearest Neighbors - No Parameters Tuning")
        pdf.add_page()
        pdf = exportPdf(x, y, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_KNN_NT_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_KNN_NT_val_data.png"),
            "Confusion Matrix Validation data", "K Nearest Neighbors - No Parameters Tuning")
        pdf.add_page()
        pdf = exportTestResultPdf(
            x, y, pdf, NT_auc_ginin_output, "K Nearest Neighbors - No Parameters Tuning")

        paramFiles = param_file_path + param_file_name + "_KNN_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = KNeighborsClassifier()

        # Create the random search model
        knn_search = RandomizedSearchCV(estimator, param_grid, n_jobs=-1,
                                        scoring='roc_auc', cv=5,
                                        n_iter=100, verbose=1, random_state=50)

        # Fit
        knn_search.fit(X_train, y_train)

        knn_random = knn_search.best_estimator_
        # Test the model
        pred_knn_val0 = knn_random.predict(X_val)
        pred_knn_prob_val0 = knn_random.predict_proba(X_val)[:, 1]

        pred_knn_train0 = knn_random.predict(X_train)
        pred_knn_prob_train0 = knn_random.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_knn_train0))
        print(classification_report(y_val, pred_knn_val0))

        evaluate_model(y_val, pred_knn_prob_val0, y_train,
                       pred_knn_prob_train0, file_name+"_KNN_RS_roc1.png")

        evaluate_model(y_val, pred_knn_val0, y_train,
                       pred_knn_train0, file_name+"_KNN_RS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_knn_train0,
                       file_name + "_KNN_RS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_knn_val0,
                       file_name + "_KNN_RS_val_data.png", "Validation")

        RS_roc_file1 = plot_dir + file_name + "_KNN_RS_roc1.png"
        RS_roc_file2 = plot_dir + file_name + "_KNN_RS_roc2.png"
        RS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_knn_prob_val0, pred_knn_val0)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(x, y, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_KNN_RS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_KNN_RS_roc2.png"),
            "ROC Curve Test data", "K Nearest Neighbors - Random Search")
        pdf.add_page()
        pdf = exportPdf(x, y, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_KNN_RS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_KNN_RS_val_data.png"),
            "Confusion Matrix Validation data", "K Nearest Neighbors - Random Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            x, y, pdf, RS_auc_ginin_output, "K Nearest Neighbors - Random Search")

        paramFiles = param_file_path + param_file_name + "_KNN_GS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = KNeighborsClassifier()

        # Create the random search model
        knn_grid = GridSearchCV(estimator, param_grid, n_jobs=-1,
                                scoring='roc_auc', cv=5, verbose=1)

        # Fit
        knn_grid.fit(X_train, y_train)

        knn_final = knn_grid.best_estimator_

        # Test the model
        pred_knn_val1 = knn_final.predict(X_val)
        pred_knn_prob_val1 = knn_final.predict_proba(X_val)[:, 1]

        pred_knn_train1 = knn_final.predict(X_train)
        pred_knn_prob_train1 = knn_final.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_knn_train1))
        print(classification_report(y_val, pred_knn_val1))

        evaluate_model(y_val, pred_knn_prob_val1, y_train,
                       pred_knn_prob_train1, file_name+"_KNN_GS_roc1.png")

        evaluate_model(y_val, pred_knn_val1, y_train,
                       pred_knn_train1, file_name+"_KNN_GS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_knn_train1,
                       file_name + "_KNN_GS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_knn_val1,
                       file_name + "_KNN_GS_val_data.png", "Validation")

        GS_roc_file1 = plot_dir + file_name + "_KNN_GS_roc1.png"
        GS_roc_file2 = plot_dir + file_name + "_KNN_GS_roc2.png"
        GS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_knn_prob_val1, pred_knn_val1)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(x, y, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_KNN_GS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_KNN_GS_roc2.png"),
            "ROC Curve Test data", "K Nearest Neighbors - Grid Search")
        pdf.add_page()
        pdf = exportPdf(x, y, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_KNN_GS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_KNN_GS_val_data.png"),
            "Confusion Matrix Validation data", "K Nearest Neighbors - Grid Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            x, y, pdf, GS_auc_ginin_output, "K Nearest Neighbors - Grid Search")

        pdf.output(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_KNN.pdf"))

        NT_roc_file1 = plot_dir + file_name+"KNN_NT_roc1.png"
        NT_roc_file2 = plot_dir + file_name+"KNN_NT_roc2.png"
        RS_roc_file1 = plot_dir + file_name + "_KNN_RS_roc1.png"
        RS_roc_file2 = plot_dir + file_name + "_KNN_RS_roc2.png"
        GS_roc_file1 = plot_dir + file_name + "_KNN_GS_roc1.png"
        GS_roc_file2 = plot_dir + file_name + "_KNN_GS_roc2.png"
        NT_auc_ginin_output = exportTestResultLstFromExcel(
            "K Nearest Neighbors - No Parameters Tuning")
        RS_auc_ginin_output = exportTestResultLstFromExcel(
            "K Nearest Neighbors - Random Search")
        GS_auc_ginin_output = exportTestResultLstFromExcel(
            "K Nearest Neighbors - Grid Search")
        context = {'is_data': True, 'pdfFile': plot_dir + file_name + "_KNN.pdf", 'model': 'KNN',  'tableHead': 'K Nearest Neighbors', 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                   'NT_graphConfMat1': plot_dir + file_name + "_KNN_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_KNN_NT_val_data.png",
                   'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                   'RS_graphConfMat1': plot_dir + file_name + "_KNN_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_KNN_RS_val_data.png",
                   'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                   'GS_graphConfMat1': plot_dir + file_name + "_KNN_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_KNN_GS_val_data.png"}

        print('KNN Proces ended')

    except Exception as e:
        print(e)


def SVM(request):
    try:
        csv_file_name = "csvfile_"+user_name
        savefile_x_final = file_path + csv_file_name + "_x_model.csv"
        if(not os.path.exists(savefile_x_final)):
            return render(request, 'processNotdone.html')
        df = pd.read_csv(savefile_x_final)
        targetVarFile = file_path + csv_file_name + "_targetVar.txt"
        file1 = open(targetVarFile, "r")  # write mode
        targetVar = file1.read()
        file1.close()
        pdf = FPDF()
        # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
        y_model = df[targetVar]
        x_model = df.drop(targetVar, axis=1)
        X_train, X_test, y_train, y_test = train_test_split(
            x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
        X_train, X_val, y_train, y_val = train_test_split(
            X_train, y_train, test_size=0.22222222222222224, random_state=321)

        svm_clf = svm.SVC(probability=True)
        svm_clf.fit(X_train, y_train)

        # Test the model
        pred_svm_val = svm_clf.predict(X_val)
        pred_svm_prob_val = svm_clf.predict_proba(X_val)[:, 1]

        pred_svm_train = svm_clf.predict(X_train)
        pred_svm_prob_train = svm_clf.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_svm_train))
        print(classification_report(y_val, pred_svm_val))

        evaluate_model(y_val, pred_svm_prob_val, y_train,
                       pred_svm_prob_train, file_name+"SVM_NT_roc1.png")

        evaluate_model(y_val, pred_svm_val, y_train,
                       pred_svm_train, file_name+"SVM_NT_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data
        drawConfMatrix(y_train, pred_svm_train, file_name +
                       "_SVM_NT_train_data.png", "Training")

        drawConfMatrix(y_val, pred_svm_val, file_name +
                       "_SVM_NT_val_data.png", "Validation")

        NT_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_svm_prob_val, pred_svm_val)

        # Add a page
        pdf.add_page()
        document = Document()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"SVM_NT_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"SVM_NT_roc2.png"),
            "ROC Curve Test data", "Support Machine Vector - No Parameters Tuning")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_SVM_NT_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_SVM_NT_val_data.png"),
            "Confusion Matrix Validation data", "Support Machine Vector - No Parameters Tuning")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, NT_auc_ginin_output, "Support Machine Vector - No Parameters Tuning")

        NT_roc_file1 = plot_dir + file_name+"SVM_NT_roc1.png"
        NT_roc_file2 = plot_dir + file_name+"SVM_NT_roc2.png"

        paramFiles = param_file_path + param_file_name + "_SVM_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = svm.SVC()

        # Create the random search model
        svm_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                        scoring='roc_auc', cv=5,
                                        n_iter=100, verbose=1, random_state=50)

        # Fit
        svm_search.fit(X_train, y_train)

        svm_random = svm_search.best_estimator_

        # Test the model
        pred_svm_val0 = svm_random.predict(X_val)
        pred_svm_prob_val0 = svm_random.predict_proba(X_val)[:, 1]

        pred_svm_train0 = svm_random.predict(X_train)
        pred_svm_prob_train0 = svm_random.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_svm_train0))
        print(classification_report(y_val, pred_svm_val0))

        evaluate_model(y_val, pred_svm_prob_val0, y_train,
                       pred_svm_prob_train0, file_name+"_SVM_RS_roc1.png")

        evaluate_model(y_val, pred_svm_val0, y_train,
                       pred_svm_train0, file_name+"_SVM_RS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_svm_train0,
                       file_name + "_SVM_RS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_svm_val0,
                       file_name + "_SVM_RS_val_data.png", "Validation")

        RS_roc_file1 = plot_dir + file_name + "_SVM_RS_roc1.png"
        RS_roc_file2 = plot_dir + file_name + "_SVM_RS_roc2.png"
        RS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_svm_prob_val0, pred_svm_val0)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_SVM_RS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_SVM_RS_roc2.png"),
            "ROC Curve Test data", "Support Machine Vector - Random Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_SVM_RS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_SVM_RS_val_data.png"),
            "Confusion Matrix Validation data", "Support Machine Vector - Random Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, RS_auc_ginin_output, "Support Machine Vector - Random Search")

        paramFiles = param_file_path + param_file_name + "_SVM_GS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = svm.SVC()

        # Create the random search model
        svm_grid = GridSearchCV(estimator, param_grid, n_jobs=3,
                                scoring='roc_auc', cv=5, verbose=1)

        # Fit
        svm_grid.fit(X_train, y_train)

        svm_final = svm_grid.best_estimator_

        # Test the model
        pred_svm_val1 = svm_final.predict(X_val)
        pred_svm_prob_val1 = svm_final.predict_proba(X_val)[:, 1]

        pred_svm_train1 = svm_final.predict(X_train)
        pred_svm_prob_train1 = svm_final.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_svm_train1))
        print(classification_report(y_val, pred_svm_val1))

        evaluate_model(y_val, pred_svm_prob_val1, y_train,
                       pred_svm_prob_train1, file_name+"_SVM_GS_roc1.png")

        evaluate_model(y_val, pred_svm_val1, y_train,
                       pred_svm_train1, file_name+"_SVM_GS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_svm_train1,
                       file_name + "_SVM_GS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_svm_val1,
                       file_name + "_SVM_GS_val_data.png", "Validation")

        GS_roc_file1 = plot_dir + file_name + "_SVM_GS_roc1.png"
        GS_roc_file2 = plot_dir + file_name + "_SVM_GS_roc2.png"
        GS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_svm_prob_val1, pred_svm_val1)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_SVM_GS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_SVM_GS_roc2.png"),
            "ROC Curve Test data", "Support Machine Vector - Grid Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_SVM_GS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_SVM_GS_val_data.png"),
            "Confusion Matrix Validation data", "Support Machine Vector - Grid Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, GS_auc_ginin_output, "Support Machine Vector - Grid Search")

        pdf.output(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_SVM_Output.pdf"))
        context = {'pdfFile': plot_dir + file_name + "_SVM_Output.pdf", 'model': 'SVM',  'tableHead': 'Support Machine Vector', 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                   'NT_graphConfMat1': plot_dir + file_name + "_SVM_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_SVM_NT_val_data.png",
                   'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                   'RS_graphConfMat1': plot_dir + file_name + "_SVM_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_SVM_RS_val_data.png",
                   'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                   'GS_graphConfMat1': plot_dir + file_name + "_SVM_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_SVM_GS_val_data.png"}
        return render(request, 'showModelOutputAll.html', context)
    except Exception as e:
        print(e)
        return render(request, 'error.html')


def SVMAjax():
    try:
        print('inside SVMAjax')
        csv_file_name = "csvfile_"+user_name
        savefile_x_final = file_path + csv_file_name + "_x_model.csv"
        df = pd.read_csv(savefile_x_final)
        targetVarFile = file_path + csv_file_name + "_targetVar.txt"
        file1 = open(targetVarFile, "r")  # write mode
        targetVar = file1.read()
        file1.close()
        pdf = FPDF()
        # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
        y_model = df[targetVar]
        x_model = df.drop(targetVar, axis=1)
        X_train, X_test, y_train, y_test = train_test_split(
            x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
        X_train, X_val, y_train, y_val = train_test_split(
            X_train, y_train, test_size=0.22222222222222224, random_state=321)

        svm_clf = svm.SVC(probability=True)
        svm_clf.fit(X_train, y_train)

        # Test the model
        pred_svm_val = svm_clf.predict(X_val)
        pred_svm_prob_val = svm_clf.predict_proba(X_val)[:, 1]

        pred_svm_train = svm_clf.predict(X_train)
        pred_svm_prob_train = svm_clf.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_svm_train))
        print(classification_report(y_val, pred_svm_val))

        evaluate_model(y_val, pred_svm_prob_val, y_train,
                       pred_svm_prob_train, file_name+"SVM_NT_roc1.png")

        evaluate_model(y_val, pred_svm_val, y_train,
                       pred_svm_train, file_name+"SVM_NT_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data
        drawConfMatrix(y_train, pred_svm_train, file_name +
                       "_SVM_NT_train_data.png", "Training")

        drawConfMatrix(y_val, pred_svm_val, file_name +
                       "_SVM_NT_val_data.png", "Validation")

        NT_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_svm_prob_val, pred_svm_val)

        # Add a page
        pdf.add_page()
        document = Document()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"SVM_NT_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"SVM_NT_roc2.png"),
            "ROC Curve Test data", "Support Machine Vector - No Parameters Tuning")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_SVM_NT_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_SVM_NT_val_data.png"),
            "Confusion Matrix Validation data", "Support Machine Vector - No Parameters Tuning")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, NT_auc_ginin_output, "Support Machine Vector - No Parameters Tuning")

        NT_roc_file1 = plot_dir + file_name+"SVM_NT_roc1.png"
        NT_roc_file2 = plot_dir + file_name+"SVM_NT_roc2.png"

        print('SVM NT done')
        paramFiles = param_file_path + param_file_name + "_SVM_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        print('SVM RS Started')
        print('param_grid')
        print(param_grid)
        # Estimator for use in random search
        estimator = svm.SVC()

        # Create the random search model
        svm_search = RandomizedSearchCV(estimator, param_grid, n_jobs=-1,
                                        scoring='roc_auc', cv=5,
                                        n_iter=100, verbose=1, random_state=50)

        # Fit
        svm_search.fit(X_train, y_train)

        svm_random = svm_search.best_estimator_

        # Test the model
        pred_svm_val0 = svm_random.predict(X_val)
        pred_svm_prob_val0 = svm_random.predict_proba(X_val)[:, 1]

        pred_svm_train0 = svm_random.predict(X_train)
        pred_svm_prob_train0 = svm_random.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_svm_train0))
        print(classification_report(y_val, pred_svm_val0))

        evaluate_model(y_val, pred_svm_prob_val0, y_train,
                       pred_svm_prob_train0, file_name+"_SVM_RS_roc1.png")

        evaluate_model(y_val, pred_svm_val0, y_train,
                       pred_svm_train0, file_name+"_SVM_RS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_svm_train0,
                       file_name + "_SVM_RS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_svm_val0,
                       file_name + "_SVM_RS_val_data.png", "Validation")

        RS_roc_file1 = plot_dir + file_name + "_SVM_RS_roc1.png"
        RS_roc_file2 = plot_dir + file_name + "_SVM_RS_roc2.png"
        RS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_svm_prob_val0, pred_svm_val0)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_SVM_RS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_SVM_RS_roc2.png"),
            "ROC Curve Test data", "Support Machine Vector - Random Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_SVM_RS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_SVM_RS_val_data.png"),
            "Confusion Matrix Validation data", "Support Machine Vector - Random Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, RS_auc_ginin_output, "Support Machine Vector - Random Search")

        paramFiles = param_file_path + param_file_name + "_SVM_GS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = svm.SVC()

        # Create the random search model
        svm_grid = GridSearchCV(estimator, param_grid, n_jobs=-1,
                                scoring='roc_auc', cv=5, verbose=1)

        # Fit
        svm_grid.fit(X_train, y_train)

        svm_final = svm_grid.best_estimator_

        # Test the model
        pred_svm_val1 = svm_final.predict(X_val)
        pred_svm_prob_val1 = svm_final.predict_proba(X_val)[:, 1]

        pred_svm_train1 = svm_final.predict(X_train)
        pred_svm_prob_train1 = svm_final.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_svm_train1))
        print(classification_report(y_val, pred_svm_val1))

        evaluate_model(y_val, pred_svm_prob_val1, y_train,
                       pred_svm_prob_train1, file_name+"_SVM_GS_roc1.png")

        evaluate_model(y_val, pred_svm_val1, y_train,
                       pred_svm_train1, file_name+"_SVM_GS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_svm_train1,
                       file_name + "_SVM_GS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_svm_val1,
                       file_name + "_SVM_GS_val_data.png", "Validation")

        GS_roc_file1 = plot_dir + file_name + "_SVM_GS_roc1.png"
        GS_roc_file2 = plot_dir + file_name + "_SVM_GS_roc2.png"
        GS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_svm_prob_val1, pred_svm_val1)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_SVM_GS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_SVM_GS_roc2.png"),
            "ROC Curve Test data", "Support Machine Vector - Grid Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_SVM_GS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_SVM_GS_val_data.png"),
            "Confusion Matrix Validation data", "Support Machine Vector - Grid Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, GS_auc_ginin_output, "Support Machine Vector - Grid Search")

        pdf.output(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_SVM_Output.pdf"))
        context = {'pdfFile': plot_dir + file_name + "_SVM_Output.pdf", 'model': 'SVM',  'tableHead': 'Support Machine Vector', 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                   'NT_graphConfMat1': plot_dir + file_name + "_SVM_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_SVM_NT_val_data.png",
                   'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                   'RS_graphConfMat1': plot_dir + file_name + "_SVM_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_SVM_RS_val_data.png",
                   'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                   'GS_graphConfMat1': plot_dir + file_name + "_SVM_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_SVM_GS_val_data.png"}

    except Exception as e:
        print(e)


def BC(request):
    try:
        csv_file_name = "csvfile_"+user_name
        savefile_x_final = file_path + csv_file_name + "_x_model.csv"
        if(not os.path.exists(savefile_x_final)):
            return render(request, 'processNotdone.html')
        df = pd.read_csv(savefile_x_final)
        targetVarFile = file_path + csv_file_name + "_targetVar.txt"
        file1 = open(targetVarFile, "r")  # write mode
        targetVar = file1.read()
        file1.close()
        # a variable pdf
        pdf = FPDF()
        # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
        y_model = df[targetVar]
        x_model = df.drop(targetVar, axis=1)
        X_train, X_test, y_train, y_test = train_test_split(
            x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
        X_train, X_val, y_train, y_val = train_test_split(
            X_train, y_train, test_size=0.22222222222222224, random_state=321)

        bc_model = BaggingClassifier()
        bc_model.fit(X_train, y_train)

        # Test the model
        pred_bc_val = bc_model.predict(X_val)
        pred_bc_prob_val = bc_model.predict_proba(X_val)[:, 1]

        pred_bc_train = bc_model.predict(X_train)
        pred_bc_prob_train = bc_model.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_bc_train))
        print(classification_report(y_val, pred_bc_val))

        evaluate_model(y_val, pred_bc_prob_val, y_train,
                       pred_bc_prob_train, file_name+"BC_NT_roc1.png")

        evaluate_model(y_val, pred_bc_val, y_train,
                       pred_bc_train, file_name+"BC_NT_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data
        drawConfMatrix(y_train, pred_bc_train, file_name +
                       "_BC_NT_train_data.png", "Training")

        drawConfMatrix(y_val, pred_bc_val, file_name +
                       "_BC_NT_val_data.png", "Validation")

        NT_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_bc_prob_val, pred_bc_val)

        # Add a page
        pdf.add_page()
        document = Document()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"BC_NT_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"BC_NT_roc2.png"),
            "ROC Curve Test data", "Bagging Classifier - No Parameters Tuning")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_BC_NT_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_BC_NT_val_data.png"),
            "Confusion Matrix Validation data", "Bagging Classifier - No Parameters Tuning")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, NT_auc_ginin_output, "Bagging Classifier - No Parameters Tuning")

        NT_roc_file1 = plot_dir + file_name+"BC_NT_roc1.png"
        NT_roc_file2 = plot_dir + file_name+"BC_NT_roc2.png"

        paramFiles = param_file_path + param_file_name + "_BC_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = BaggingClassifier(random_state=50)

        # Create the random search model
        bc_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                       scoring='roc_auc', cv=5,
                                       n_iter=100, verbose=1, random_state=50)

        # Fit
        bc_search.fit(X_train, y_train)

        bc_random = bc_search.best_estimator_

        # Test the model
        pred_bc_val0 = bc_random.predict(X_val)
        pred_bc_prob_val0 = bc_random.predict_proba(X_val)[:, 1]

        pred_bc_train0 = bc_random.predict(X_train)
        pred_bc_prob_train0 = bc_random.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_bc_train0))
        print(classification_report(y_val, pred_bc_val0))

        evaluate_model(y_val, pred_bc_prob_val0, y_train,
                       pred_bc_prob_train0, file_name+"_BC_RS_roc1.png")

        evaluate_model(y_val, pred_bc_val0, y_train,
                       pred_bc_train0, file_name+"_BC_RS_roc2.png")

        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_bc_train0,
                       file_name + "_BC_RS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_bc_val0,
                       file_name + "_BC_RS_val_data.png", "Validation")

        RS_roc_file1 = plot_dir + file_name + "_BC_RS_roc1.png"
        RS_roc_file2 = plot_dir + file_name + "_BC_RS_roc2.png"
        RS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_bc_prob_val0, pred_bc_val0)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_BC_RS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_BC_RS_roc2.png"),
            "ROC Curve Test data", "Bagging Classifier - Random Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_BC_RS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_BC_RS_val_data.png"),
            "Confusion Matrix Validation data", "Bagging Classifier - Random Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, RS_auc_ginin_output, "Bagging Classifier - Random Search")

        paramFiles = param_file_path + param_file_name + "_BC_RS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = BaggingClassifier(random_state=50)

        # Create the random search model
        bc_search = RandomizedSearchCV(estimator, param_grid, n_jobs=3,
                                       scoring='roc_auc', cv=5,
                                       n_iter=100, verbose=1, random_state=50)

        # Fit
        bc_search.fit(X_train, y_train)

        paramFiles = param_file_path + param_file_name + "_BC_GS.csv"
        param_grid = {}
        if os.path.exists(paramFiles):
            df = pd.read_csv(paramFiles)
            for index, row in df.iterrows():
                param_grid[row['paramName']] = eval(row['paramValue'])

        # Estimator for use in random search
        estimator = BaggingClassifier(random_state=50)

        # Create the random search model
        bc_grid = GridSearchCV(estimator, param_grid, n_jobs=3,
                               scoring='roc_auc', cv=5, verbose=1)

        # Fit
        bc_grid.fit(X_train, y_train)

        bc_final = bc_grid.best_estimator_

        pred_bc_val1 = bc_final.predict(X_val)
        pred_bc_prob_val1 = bc_final.predict_proba(X_val)[:, 1]

        pred_bc_train1 = bc_final.predict(X_train)
        pred_bc_prob_train1 = bc_final.predict_proba(X_train)[:, 1]

        # Get the model performance
        print(classification_report(y_train, pred_bc_train1))
        print(classification_report(y_val, pred_bc_val1))

        evaluate_model(y_val, pred_bc_prob_val1, y_train,
                       pred_bc_prob_train1, file_name+"_BC_GS_roc1.png")

        evaluate_model(y_val, pred_bc_val1, y_train,
                       pred_bc_train1, file_name+"_BC_GS_roc2.png")
        # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
        # show the confusion matrix for training data

        drawConfMatrix(y_train, pred_bc_train1,
                       file_name + "_BC_GS_train_data.png", "Training")

        # show the confusion matrix for validation data

        drawConfMatrix(y_val, pred_bc_val1,
                       file_name + "_BC_GS_val_data.png", "Validation")

        GS_roc_file1 = plot_dir + file_name + "_BC_GS_roc1.png"
        GS_roc_file2 = plot_dir + file_name + "_BC_GS_roc2.png"
        GS_auc_ginin_output = test_modelPerfomance(
            y_val,  pred_bc_prob_val1, pred_bc_val1)

        # Add a page
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_BC_GS_roc1.png"),
            "ROC Curve Validation data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_BC_GS_roc2.png"),
            "ROC Curve Test data", "Bagging Classifier - Grid Search")
        pdf.add_page()
        pdf = exportPdf(10, 10, pdf, document,  os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_BC_GS_train_data.png"),
            "Confusion Matrix Training data", os.path.join(
            BASE_DIR, plot_dir_view, file_name+"_BC_GS_val_data.png"),
            "Confusion Matrix Validation data", "Bagging Classifier - Grid Search")
        pdf.add_page()
        pdf = exportTestResultPdf(
            10, 10, pdf, GS_auc_ginin_output, "Bagging Classifier - Grid Search")

        pdf.output(os.path.join(
            BASE_DIR, plot_dir_view + file_name + "_BagCls.pdf"))

        context = {'pdfFile': plot_dir + file_name + "_BagCls.pdf", 'model': 'BC',  'tableHead': 'Bagging Classifier', 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
                   'NT_graphConfMat1': plot_dir + file_name + "_BC_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_BC_NT_val_data.png",
                   'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
                   'RS_graphConfMat1': plot_dir + file_name + "_BC_RS_train_data.png", 'RS_graphConfMat2': plot_dir + file_name + "_BC_RS_val_data.png",
                   'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
                   'GS_graphConfMat1': plot_dir + file_name + "_BC_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_BC_GS_val_data.png"}
        return render(request, 'showModelOutputAll.html', context)
    except Exception as e:
        print(e)
        return render(request, 'error.html')


def test_modelPerfomance(y_val, prob, val):
    # perform the test on model performance

    # perform ks test
    arrlstModelPerf = []
    csv_file_name = "csvfile_"+user_name
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()
    y_val1 = pd.DataFrame(y_val).reset_index()
    pred_prob_val1 = pd.DataFrame(prob).reset_index()
    new_val = pd.concat([y_val1, pred_prob_val1], axis=1).reset_index()
    new_val = new_val.drop(['level_0', 'index'], axis=1)
    new_val.columns = [targetVar, 'Probability']
    prob_to_1 = new_val.loc[new_val.eval(targetVar) ==
                            1, ['Probability']].sort_index()
    prob_to_0 = new_val.loc[new_val.eval(targetVar) ==
                            0, ['Probability']].sort_index()
    prob_to_1 = np.array(prob_to_1).reshape(len(prob_to_1))
    prob_to_0 = np.array(prob_to_0).reshape(len(prob_to_0))

    ks = ks_2samp(prob_to_0, prob_to_1)

    # perform auc and ginin test
    fpr1, tpr1, thresholds = roc_curve(y_val,  prob)
    auc_prob = round(auc(fpr1, tpr1), 3)
    gini_prob = 2*auc_prob-1

    fpr2, tpr2, thresholds = roc_curve(y_val,  val)
    auc_class = round(auc(fpr2, tpr2), 3)
    gini_class = 2*auc_class-1

    # perform MAE test
    abs_error_prob = abs(y_val - prob)
    abs_error_class = abs(y_val - val)

    mae_prob = round(np.mean(abs_error_prob), 4)
    mae_class = round(np.mean(abs_error_class), 4)

    # perform Accuracy test
    accuracy_class = accuracy_score(y_val, val)

    # print('\n')
    # print('Confusion Matrix - Validation:')
    # print(confusion_matrix(y_val, pred_rf_val))
    objtestModelPerf = lstTestModelPerf()
    objtestModelPerf.testName = "Confusion Matrix - Validation:"
    objtestModelPerf.testResult = confusion_matrix(y_val, val)
    arrlstModelPerf.append(objtestModelPerf)

    # print('\n')
    # print('Classification Report - Validation:')
    # print(classification_report(y_val, pred_rf_val))
    objtestModelPerf = lstTestModelPerf()
    objtestModelPerf.testName = "Classification Report - Validation:"
    report = classification_report(
        y_val, val, output_dict=True)
    dfclassification_report = pd.DataFrame(report).transpose()
    dfclassification_json = json.loads(
        dfclassification_report.to_json(orient="index"))
    objtestModelPerf.testResult_dict = dfclassification_json
    objtestModelPerf.testResult = ""
    arrlstModelPerf.append(objtestModelPerf)

    # print('\n')
    # print('Accuracy on classes:', accuracy_class)
    # print('\n')
    objtestModelPerf = lstTestModelPerf()
    objtestModelPerf.testName = "Accuracy on classes:"
    objtestModelPerf.testResult = accuracy_class
    arrlstModelPerf.append(objtestModelPerf)

    # print('KS test:', ks)
    objtestModelPerf = lstTestModelPerf()
    objtestModelPerf.testName = "KS test:"
    objtestModelPerf.testResult = ks
    arrlstModelPerf.append(objtestModelPerf)

    # print('AUC Score on probability:', auc_prob)
    # print('AUC Score on classes:', auc_class)
    objtestModelPerf = lstTestModelPerf()
    objtestModelPerf.testName = "AUC Score on probability:"
    objtestModelPerf.testResult = auc_prob
    arrlstModelPerf.append(objtestModelPerf)
    objtestModelPerf = lstTestModelPerf()
    objtestModelPerf.testName = "AUC Score on classes:"
    objtestModelPerf.testResult = auc_class
    arrlstModelPerf.append(objtestModelPerf)

    # print('GINI Score on probability:', gini_prob)
    # print('GINI Score on classes:', gini_class)
    objtestModelPerf = lstTestModelPerf()
    objtestModelPerf.testName = "GINI Score on probability:"
    objtestModelPerf.testResult = gini_prob
    arrlstModelPerf.append(objtestModelPerf)
    objtestModelPerf = lstTestModelPerf()
    objtestModelPerf.testName = "GINI Score on classes:"
    objtestModelPerf.testResult = gini_class
    arrlstModelPerf.append(objtestModelPerf)

    # print('MAE on probability:', mae_prob)
    # print('MAE on classes:', mae_class)
    objtestModelPerf = lstTestModelPerf()
    objtestModelPerf.testName = "MAE Score on probability:"
    objtestModelPerf.testResult = mae_prob
    arrlstModelPerf.append(objtestModelPerf)
    objtestModelPerf = lstTestModelPerf()
    objtestModelPerf.testName = "MAE Score on classes:"
    objtestModelPerf.testResult = mae_class
    arrlstModelPerf.append(objtestModelPerf)
    # print('arrlstModelPerf')
    # print(arrlstModelPerf)
    return arrlstModelPerf


def drawConfMatrix(y_val, pred_rf_val, fileName, strLbl):
    cnf_matrix = confusion_matrix(y_val, pred_rf_val, labels=[0, 1])
    plt.figure(figsize=(12, 9))
    sns.heatmap(pd.DataFrame(cnf_matrix), annot=True,
                cmap="YlGnBu", fmt='g')
    plt.title('Confusion matrix: '+strLbl+' data')
    plt.ylabel('Actual label')
    plt.xlabel('Predicted label')
    plt.tight_layout()
    plt.savefig(os.path.join(
        BASE_DIR, plot_dir_view + fileName))
    plt.close()


def randomForest_ajax(request):
    csv_file_name = "csvfile_"+user_name
    savefile_x_final = file_path + csv_file_name + "_x_model.csv"
    df = pd.read_csv(savefile_x_final)
    targetVarFile = file_path + csv_file_name + "_targetVar.txt"
    file1 = open(targetVarFile, "r")  # write mode
    targetVar = file1.read()
    file1.close()

    # split the dastaset into train, validation and test with the ratio 0.7, 0.2 and 0.1
    y_model = df[targetVar]
    x_model = df.drop(targetVar, axis=1)
    X_train, X_test, y_train, y_test = train_test_split(
        x_model, y_model, test_size=0.1, random_state=321)  # Predictor and target variables
    X_train, X_val, y_train, y_val = train_test_split(
        X_train, y_train, test_size=0.22222222222222224, random_state=321)

    rf_model = RandomForestClassifier(n_estimators=100, criterion="gini",
                                      random_state=50,
                                      max_features='sqrt',
                                      n_jobs=-1, verbose=1)
    rf_model.fit(X_train, y_train)

    # check the average of the nodes and depth.
    n_nodes = []
    max_depths = []
    for ind_tree in rf_model.estimators_:
        n_nodes.append(ind_tree.tree_.node_count)
        max_depths.append(ind_tree.tree_.max_depth)

    print(f'Average number of nodes {int(np.mean(n_nodes))}')
    print(f'Average maximum depth {int(np.mean(max_depths))}')

    # Test the model
    pred_rf_val = rf_model.predict(X_val)
    pred_rf_prob_val = rf_model.predict_proba(X_val)[:, 1]

    pred_rf_train = rf_model.predict(X_train)
    pred_rf_prob_train = rf_model.predict_proba(X_train)[:, 1]

    # Get the model performance
    print('classification report on training data')
    print(classification_report(y_train, pred_rf_train))
    print('\n')
    print('classification report on validation data')
    print(classification_report(y_val, pred_rf_val))

    evaluate_model(y_val, pred_rf_prob_val, y_train,
                   pred_rf_prob_train, file_name+"_RF_NT_roc1.png")

    evaluate_model(y_val, pred_rf_val, y_train,
                   pred_rf_train, file_name+"_RF_NT_roc2.png")

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    # show the confusion matrix for training data
    drawConfMatrix(y_train, pred_rf_train, file_name +
                   "_RF_NT_train_data.png", "Training")

    # show the confusion matrix for validation data
    drawConfMatrix(y_val, pred_rf_val, file_name +
                   "_RF_NT_val_data.png", "Validation")

    NT_roc_file1 = plot_dir + file_name + "_RF_NT_roc1.png"
    NT_roc_file2 = plot_dir + file_name + "_RF_NT_roc2.png"
    NT_auc_ginin_output = test_modelPerfomance(
        y_val, pred_rf_prob_val, pred_rf_val)

    # Random Search
    paramFiles = param_file_path + param_file_name + "_RF_RS.csv"
    param_grid = {}
    if os.path.exists(paramFiles):
        df = pd.read_csv(paramFiles)
        for index, row in df.iterrows():
            param_grid[row['paramName']] = eval(row['paramValue'])
    model = RandomForestClassifier()

    # Create the random search model
    rf_search = RandomizedSearchCV(model, param_grid, n_jobs=3,
                                   scoring='roc_auc', cv=5,
                                   n_iter=100, verbose=1, random_state=50)

    # Fit
    rf_search.fit(X_train, y_train)

    rf_search.best_params_

    rf_random = rf_search.best_estimator_

    # Test the model
    pred_rf_val0 = rf_random.predict(X_val)
    pred_rf_prob_val0 = rf_random.predict_proba(X_val)[:, 1]

    pred_rf_train0 = rf_random.predict(X_train)
    pred_rf_prob_train0 = rf_random.predict_proba(X_train)[:, 1]

    # Get the model performance
    print(classification_report(y_train, pred_rf_train0))
    print(classification_report(y_val, pred_rf_val0))

    evaluate_model(y_val, pred_rf_prob_val0, y_train,
                   pred_rf_prob_train0, file_name+"_RF_RS_roc1.png")

    evaluate_model(y_val, pred_rf_val0, y_train,
                   pred_rf_train0, file_name+"_RF_RS_roc2.png")

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    # show the confusion matrix for training data
    drawConfMatrix(y_train, pred_rf_train0,
                   file_name + "_RF_RS_train_data.png", "Training")
    drawConfMatrix(y_val, pred_rf_val0, file_name +
                   "_RF_RS_val_data.png", "Validation")

    RS_roc_file1 = plot_dir + file_name + "_RF_RS_roc1.png"
    RS_roc_file2 = plot_dir + file_name + "_RF_RS_roc2.png"
    RS_auc_ginin_output = test_modelPerfomance(
        y_val, pred_rf_prob_val0, pred_rf_val0)

    # Grid Search
    paramFiles = param_file_path + param_file_name + "_RF_GS.csv"
    param_grid = {}
    if os.path.exists(paramFiles):
        df = pd.read_csv(paramFiles)
        for index, row in df.iterrows():
            param_grid[row['paramName']] = eval(row['paramValue'])

    # Estimator for use in random search
    model = RandomForestClassifier()

    # Create the random search model
    rf_grid = GridSearchCV(model, param_grid, n_jobs=3,
                           scoring='roc_auc', cv=5, verbose=1)

    # Fit
    rf_grid.fit(X_train, y_train)

    rf_grid.best_params_

    rf_grid = rf_grid.best_estimator_

    # Test the model
    pred_rf_val1 = rf_grid.predict(X_val)
    pred_rf_prob_val1 = rf_grid.predict_proba(X_val)[:, 1]

    pred_rf_train1 = rf_grid.predict(X_train)
    pred_rf_prob_train1 = rf_grid.predict_proba(X_train)[:, 1]

    # Get the model performance
    print(classification_report(y_train, pred_rf_train1))
    print(classification_report(y_val, pred_rf_val1))

    evaluate_model(y_val, pred_rf_prob_val1, y_train,
                   pred_rf_prob_train1, file_name+"_RF_GS_roc1.png")

    evaluate_model(y_val, pred_rf_val1, y_train,
                   pred_rf_train1, file_name+"_RF_GS_roc2.png")

    # baseline, ROC=0.5 with equal chance to classify for any randome selected target account
    # show the confusion matrix for training data

    cnf_matrix = confusion_matrix(y_train, pred_rf_train1, labels=[0, 1])
    plt.figure(figsize=(10, 6))
    sns.heatmap(pd.DataFrame(cnf_matrix), annot=True,
                cmap="YlGnBu", fmt='g')

    plt.title('Confusion matrix: Training data')
    plt.ylabel('Actual label')
    plt.xlabel('Predicted label')
    plt.tight_layout()
    plt.savefig(os.path.join(
        BASE_DIR, plot_dir_view + file_name + "_RF_GS_train_data.png"))
    plt.close()
    # show the confusion matrix for validation data

    cnf_matrix2 = confusion_matrix(y_val, pred_rf_val1, labels=[0, 1])
    plt.figure(figsize=(10, 6))
    sns.heatmap(pd.DataFrame(cnf_matrix2), annot=True,
                cmap="YlGnBu", fmt='g')
    plt.title('Confusion matrix: Validation data')
    plt.ylabel('Actual label')
    plt.xlabel('Predicted label')
    plt.tight_layout()
    plt.savefig(os.path.join(
        BASE_DIR, plot_dir_view + file_name + "_RF_GS_val_data.png"))
    plt.close()
    GS_roc_file1 = plot_dir + file_name + "_RF_GS_roc1.png"
    GS_roc_file2 = plot_dir + file_name + "_RF_GS_roc2.png"
    GS_auc_ginin_output = test_modelPerfomance(
        y_val, pred_rf_prob_val1, pred_rf_val1)

    context = {'tableHead': 'Random Forest', 'NT_rocgraphpath1': NT_roc_file1, 'NT_rocgraphpath2': NT_roc_file2, 'NT_auc_ginin_output': NT_auc_ginin_output,
               'NT_graphConfMat1': plot_dir + file_name + "_RF_NT_train_data.png", 'NT_graphConfMat2': plot_dir + file_name + "_RF_NT_val_data.png",
               'RS_rocgraphpath1': RS_roc_file1, 'RS_rocgraphpath2': RS_roc_file2, 'RS_auc_ginin_output': RS_auc_ginin_output,
               'RS_graphConfMat1': plot_dir + file_name + "_RF_RS_train_data.png", 'RF_graphConfMat2': plot_dir + file_name + "_RF_RS_val_data.png",
               'GS_rocgraphpath1': GS_roc_file1, 'GS_rocgraphpath2': GS_roc_file2, 'GS_auc_ginin_output': GS_auc_ginin_output,
               'GS_graphConfMat1': plot_dir + file_name + "_RF_GS_train_data.png", 'GS_graphConfMat2': plot_dir + file_name + "_RF_GS_val_data.png"}
    return render(request, 'showModelOutputAll.html', context)


def sendMail(request):
    sender = 'n.bawaskar@prescio.com'
    receivers = ['n.bawaskar@prescio.com']

    message = """From: From Person <n.bawaskar@prescio.com>
    To: To Person <n.bawaskar@prescio.com>
    Subject: SMTP e-mail test

    This is a test e-mail message.
    """

    try:
        # smtpObj = smtplib.SMTP('mail.prescio.com')
        # smtpObj.sendmail(sender, receivers, message)
        # print("Successfully sent email")

        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText
        mail_content = """Hello,
        This is a simple mail. There is only text, no attachments are there The mail is sent using Python SMTP library.
        Thank You
        """
        # The mail addresses and password
        sender_address = 'n.bawaskar@prescio.com'
        sender_pass = ''
        receiver_address = 'nilesh.bawaskar86@gmail.com'
        # Setup the MIME
        message = MIMEMultipart()
        message['From'] = sender_address
        message['To'] = receiver_address
        # The subject line
        message['Subject'] = 'A test mail sent by Python. It has an attachment.'
        # The body and the attachments for the mail
        message.attach(MIMEText(mail_content, 'plain'))
        # Create SMTP session for sending the mail
        session = smtplib.SMTP('mail.prescio.com')  # use gmail with port
        session.starttls()  # enable security
        # login with mail_id and password
        session.login(sender_address, sender_pass)
        text = message.as_string()
        session.sendmail(sender_address, receiver_address, text)
        session.quit()
        print('Mail Sent')
    except Exception as e:
        print(e)
        print("Error: unable to send email")
    return render(request, 'showModelOutputAll.html')


def sendGMail(request):
    try:
        # smtpObj = smtplib.SMTP('mail.prescio.com')
        # smtpObj.sendmail(sender, receivers, message)
        # print("Successfully sent email")

        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText
        mail_content = """Hello,
        This is a simple mail. There is only text, no attachments are there The mail is sent using Python SMTP library.
        Thank You
        """
        # The mail addresses and password
        sender_address = 'modvaladm@gmail.com'
        sender_pass = 'modelVal#12'
        receiver_address = 'nilesh.bawaskar86@gmail.com'
        # Setup the MIME
        message = MIMEMultipart()
        message['From'] = sender_address
        message['To'] = receiver_address
        # The subject line
        message['Subject'] = 'A test mail sent by Python.'

        # The body and the attachments for the mail        message.attach(MIMEText(mail_content, 'plain'))

        # Create SMTP session for sending the mail
        session = smtplib.SMTP('smtp.gmail.com', 587)  # use gmail with port
        session.starttls()  # enable security
        # login with mail_id and password
        session.login(sender_address, sender_pass)
        text = message.as_string()
        session.sendmail(sender_address, receiver_address, text)
        session.quit()
        print('Mail Sent')
    except Exception as e:
        print(e)
        print("Error: unable to send email")
    return render(request, 'showModelOutputAll.html')


def savevalFindingsComment(request):
    comment = request.GET.get('comment', 'False')
    validationFindings = file_path + file_name + "_validationFindings.csv"
    if os.path.exists(validationFindings) and comment != "False":
        df_old = pd.read_csv(validationFindings)
        df_old.at[0, "Comment"] = comment
        df_old.to_csv(validationFindings, index=False)
        del df_old
    return JsonResponse({'is_taken': True})


def sendDevloperMail(request):
    try:
        import smtplib
        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText
        emailId = request.GET.get('emailId', 'False')
        validationFindings = file_path + file_name + "_validationFindings.csv"
        if os.path.exists(validationFindings) and emailId != "False":
            df_old = pd.read_csv(validationFindings)
            if ((df_old["EmailId"] == emailId)).any():
                print('emailsent')
                data = {'is_taken': True}
            else:
                df_old.at[0, "EmailId"] = emailId
                df_old.to_csv(validationFindings, index=False)
                mail_content = """Hello,
                Please click link below to responde the model validation findings.
                """+app_url + """valFindingsResp
                Thank You
                """
                # The mail addresses and password
                sender_address = 'modvaladm@gmail.com'
                sender_pass = 'modelVal#12'
                receiver_address = emailId
                # Setup the MIME
                message = MIMEMultipart()
                message['From'] = sender_address
                message['To'] = receiver_address
                # The subject line
                message['Subject'] = 'Model validation findings.'

                # The body and the attachments for the mail
                message.attach(MIMEText(mail_content, 'plain'))

                # Create SMTP session for sending the mail
                # use gmail with port
                session = smtplib.SMTP('smtp.gmail.com', 587)
                session.starttls()  # enable security
                # login with mail_id and password
                session.login(sender_address, sender_pass)
                text = message.as_string()
                session.sendmail(sender_address, receiver_address, text)
                session.quit()
                print('Mail Sent')
                data = {'is_taken': True}
    except Exception as e:
        print(e)
        print("Error: unable to send email")
        data = {'is_taken': False}
    return JsonResponse(data)


def sendCommentsMail(request):
    try:
        import smtplib
        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText
        emailId = request.GET.get('emailId', 'False')
        validationFindings = file_path + file_name + "_DocComments.csv"
        if os.path.exists(validationFindings) and emailId != "False":
            df_old = pd.read_csv(validationFindings)
            if ((df_old["EmailId"] == emailId)).any():
                print('emailsent')
                data = {'is_taken': True}
            else:
                df_old.at[0, "EmailId"] = emailId
                df_old.to_csv(validationFindings, index=False)
                mail_content = """Hello,
                Please click link below to responde the comments added against documet(s).
                """+app_url + """docCommResp
                Thank You
                """
                # The mail addresses and password
                sender_address = 'modvaladm@gmail.com'
                sender_pass = 'modelVal#12'
                receiver_address = emailId
                # Setup the MIME
                message = MIMEMultipart()
                message['From'] = sender_address
                message['To'] = receiver_address
                # The subject line
                message['Subject'] = 'Model validation.'

                # The body and the attachments for the mail
                message.attach(MIMEText(mail_content, 'plain'))

                # Create SMTP session for sending the mail
                # use gmail with port
                session = smtplib.SMTP('smtp.gmail.com', 587)
                session.starttls()  # enable security
                # login with mail_id and password
                session.login(sender_address, sender_pass)
                text = message.as_string()
                session.sendmail(sender_address, receiver_address, text)
                session.quit()
                print('Mail Sent')
                data = {'is_taken': True}
    except Exception as e:
        print(e)
        print("Error: unable to send email")
        data = {'is_taken': False}
    return JsonResponse(data)

def exportPdf(x, y, pdf, document, graph, text, graph2, text2, header):

    # set style and size of font
    # that you want in the pdf
    pdf.set_font("Arial", size=15)
    pdf.set_xy(x, y)
    pdf.set_text_color(0.0, 0.0, 0.0)
    pdf.multi_cell(0, 10, header, align='C')

    y += 20.0
    pdf.set_xy(40, y)
    if (os.path.exists(graph)):
        pdf.image(graph,  link='', type='', w=700/5, h=450/5)

    y += 450/5
    pdf.set_xy(x, y)
    pdf.set_font("Arial", size=12)
    pdf.set_text_color(0.0, 0.0, 0.0)
    pdf.multi_cell(0, 10, text, align='C')

    if(text2 !="ROC Curve Test data"):
        y += 20
        pdf.set_xy(40, y)
        if (os.path.exists(graph2)):
            pdf.image(graph2,  link='', type='', w=700/5, h=450/5)

        y += 450/5
        pdf.set_xy(x, y)
        pdf.set_font("Arial", size=12)
        pdf.set_text_color(0.0, 0.0, 0.0)
        pdf.multi_cell(0, 10, text2, align='C')

    para = document.add_paragraph()
    paragraph_format = para.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(header)
    font = run.font
    font.name = 'Arial'
    font.size = Pt(15)

    if (os.path.exists(graph)):
        document.add_picture(graph, width=Inches(6.0), height=Inches(3.25))

    para = document.add_paragraph()
    paragraph_format = para.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    font = run.font
    font.name = 'Arial'
    font.size = Pt(12)

    if (os.path.exists(graph2)):
        document.add_picture(graph2, width=Inches(6.0), height=Inches(3.25))

    para = document.add_paragraph()
    paragraph_format = para.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text2)
    font = run.font
    font.name = 'Arial'
    font.size = Pt(12)

    return pdf


def exportTestResultPdf(x, y, pdf, result, header):

    # set style and size of font
    # that you want in the pdf
    pdf.set_font("Arial", size=15)
    pdf.set_xy(x, y)
    pdf.set_text_color(0.0, 0.0, 0.0)
    pdf.multi_cell(0, 10, header, align='C')

    y += 20.0

    pdf.set_xy(x, y)
    pdf.set_font("Arial", size=12)
    pdf.set_text_color(0.0, 0.0, 0.0)
    pdf.multi_cell(0, 10, "Perform The Test On Model Performance", align='C')
    pdf.set_font("Arial", size=10)
    pdf.set_text_color(0.0, 0.0, 0.0)
    y += 10
    workbook = xlsxwriter.Workbook(os.path.join(
        BASE_DIR, plot_dir_view + header+".xlsx"))
    worksheet = workbook.add_worksheet('Data')
    # Start from the first cell. Rows and columns are zero indexed.
    row = 0
    col = 0
    worksheet.write(row, col,     "testName")
    worksheet.write(row, col + 1, "testResult")
    row = 1
    for i in range(len(result)):
        # print(' result[i]', result[i])
        y += 5
        pdf.set_xy(20, y)
        pdf.cell(0, 10, result[i].testName)

        # Iterate over the data and write it out row by row.
        worksheet.write(row, col,     result[i].testName)
        worksheet.write(row, col + 1, str(result[i].testResult))
        row += 1

        if result[i].testName == "Classification Report - Validation:" or result[i].testName == "Confusion Matrix - Validation:":
            if result[i].testName == "Confusion Matrix - Validation:":
                y += 5
                pdf.set_xy(20, y)
                pdf.cell(0, 10, str(result[i].testResult))
            else:

                dictRes = result[i].testResult_dict
                worksheet2 = workbook.add_worksheet('Data2')
                y += 5
                row2 = 0
                col2 = 0
                worksheet2.write(row2, col2, "")
                pdf.set_xy(20, y)
                pdf.cell(0, 10, "")
                pdf.set_xy(50, y)
                pdf.cell(0, 10, "precision")
                worksheet2.write(row2, col2+1, "precision")
                pdf.set_xy(80, y)
                pdf.cell(0, 10, "recall")
                worksheet2.write(row2, col2+2, "recall")
                pdf.set_xy(110, y)
                pdf.cell(0, 10, "f1-score")
                worksheet2.write(row2, col2+3, "f1-score")
                pdf.set_xy(140, y)
                pdf.cell(0, 10, "support")
                worksheet2.write(row2, col2+4, "support")

                for key in dictRes:
                    y += 5
                    row2 += 1
                    col2 = 0
                    pdf.set_xy(20, y)
                    pdf.cell(0, 10, str(key))
                    worksheet2.write(row2, col2, str(key))
                    dictVal = dictRes[key]
                    x2 = 50
                    for keyval in dictVal:
                        pdf.set_xy(x2, y)
                        pdf.cell(0, 10, str(dictVal[keyval]))
                        x2 += 30
                        col2 += 1
                        worksheet2.write(row2, col2, str(dictVal[keyval]))
        else:
            pdf.set_xy(70, y)
            pdf.cell(0, 10, str(result[i].testResult))

        y += 5
        pdf.set_xy(20, y)
        pdf.cell(0, 10, "")
    workbook.close()
    return pdf


def exportTestResultPdfFromExcel(pdf, document,  header):

    # set style and size of font
    # that you want in the pdf
    x, y = 10, 10
    pdf.set_font("Arial", size=15)
    pdf.set_xy(x, y)
    pdf.set_text_color(0.0, 0.0, 0.0)
    pdf.multi_cell(0, 10, header, align='C')

    y += 20.0

    pdf.set_xy(x, y)
    pdf.set_font("Arial", size=12)
    pdf.set_text_color(0.0, 0.0, 0.0)
    pdf.multi_cell(0, 10, "Perform The Test On Model Performance", align='C')
    pdf.set_font("Arial", size=10)
    pdf.set_text_color(0.0, 0.0, 0.0)
    y += 10
    xlsfile_name = os.path.join(
        BASE_DIR, plot_dir_view + header+".xlsx")
    df = pd.read_excel(xlsfile_name, sheet_name="Data")
    # print('df ', df)

    para = document.add_paragraph()
    paragraph_format = para.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(header)
    font = run.font
    font.name = 'Arial'
    font.size = Pt(15)

    para = document.add_paragraph()
    paragraph_format = para.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Perform The Test On Model Performance")
    font = run.font
    font.name = 'Arial'
    font.size = Pt(12)
    table = document.add_table(rows=1, cols=2)
    table2 = document.add_table(rows=1, cols=5)
    table1 = document.add_table(rows=1, cols=2)
    for index, row in df.iterrows():
        # print(' result[i]', result[i])
        y += 5
        pdf.set_xy(20, y)
        pdf.cell(0, 10, row["testName"])

        if row["testName"] == "Classification Report - Validation:" or row["testName"] == "Confusion Matrix - Validation:":

            row_cells = table.add_row().cells
            row_cells[0].text = row["testName"]
            if row["testName"] == "Confusion Matrix - Validation:":
                y += 5
                pdf.set_xy(20, y)
                pdf.cell(0, 10, row["testResult"])
                row_cells = table.add_row().cells
                row_cells[0].text = row["testResult"]
            else:

                dictRes = pd.read_excel(xlsfile_name, sheet_name="Data2")
                # print('dictRes', dictRes)
                y += 5
                pdf.set_xy(20, y)
                pdf.cell(0, 10, "")
                pdf.set_xy(50, y)
                pdf.cell(0, 10, "precision")
                pdf.set_xy(80, y)
                pdf.cell(0, 10, "recall")
                pdf.set_xy(110, y)
                pdf.cell(0, 10, "f1-score")
                pdf.set_xy(140, y)
                pdf.cell(0, 10, "support")

                hdr_cells = table2.rows[0].cells
                hdr_cells[0].text = ''
                hdr_cells[1].text = 'precision'
                hdr_cells[2].text = 'recall'
                hdr_cells[3].text = 'f1-score'
                hdr_cells[4].text = 'support'

                for index2, row2 in dictRes.iterrows():
                    y += 5
                    pdf.set_xy(20, y)
                    pdf.cell(0, 10, row2[0])
                    x2 = 50
                    pdf.set_xy(x2, y)
                    pdf.cell(0, 10, str(row2[1]))
                    x2 += 30
                    pdf.set_xy(x2, y)
                    pdf.cell(0, 10, str(row2[2]))
                    x2 += 30
                    pdf.set_xy(x2, y)
                    pdf.cell(0, 10, str(row2[3]))
                    x2 += 30
                    pdf.set_xy(x2, y)
                    pdf.cell(0, 10, str(row2[4]))

                    row_cells2 = table2.add_row().cells
                    row_cells2[0].text = str(row2[0])
                    row_cells2[1].text = str(row2[1])
                    row_cells2[2].text = str(row2[2])
                    row_cells2[3].text = str(row2[3])
                    row_cells2[4].text = str(row2[4])
        else:
            pdf.set_xy(70, y)
            pdf.cell(0, 10, row["testResult"])

            row_cells1 = table1.add_row().cells
            row_cells1[0].text = row["testName"]
            row_cells1[1].text = row["testResult"]

        y += 5
        pdf.set_xy(20, y)
        pdf.cell(0, 10, "")

    return pdf


def exportTestResultLstFromExcel(header):
    xlsfile_name = os.path.join(BASE_DIR, plot_dir_view + header+".xlsx")

    df = pd.read_excel(xlsfile_name, sheet_name="Data")
    dictRes = pd.read_excel(xlsfile_name, sheet_name="Data2")

    result = df.to_json(orient="records")
    result = json.loads(result)
    clsReport = dictRes.to_json(orient="records")
    clsReport = json.loads(clsReport)
    data = {'result': result, 'clsReport': clsReport}
    return data


def addCommenttoPdf(x, y, pdf, document, comment, header):

    # set style and size of font
    # that you want in the pdf
    pdf.set_font("Arial", size=15)
    pdf.set_xy(x, y)
    pdf.set_text_color(0.0, 0.0, 0.0)
    pdf.multi_cell(0, 10, header, align='C')

    y += 20.0

    pdf.set_xy(x, y)
    pdf.set_font("Arial", size=12)
    pdf.set_text_color(0.0, 0.0, 0.0)
    # pdf.multi_cell(0, 10, "User Comments", align='C')
    # pdf.set_font("Arial", size=10)
    # pdf.set_text_color(0.0, 0.0, 0.0)
    # y += 10
    pdf.multi_cell(0, 10, comment)

    para = document.add_paragraph()
    paragraph_format = para.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(header)
    font = run.font
    font.name = 'Arial'
    font.size = Pt(15)

    paracomment = document.add_paragraph()
    paragraph_format = paracomment.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paracomment.add_run(comment)
    font = run.font
    font.name = 'Arial'
    font.size = Pt(12)

    return pdf


def saveComments(request):
    comments = request.GET['comments']
    modelName = request.GET['modelName']
    method = request.GET['method']
    outputTab = request.GET['outputTab']
    modelSS = ""
    pdfFile = ""
    # print('NT_auc_ginin_output  ', eval(NT_auc_ginin_output))
    if modelName == "K Nearest Neighbors (KNN)":
        modelSS = "KNN"
        pdfFile = file_name + "_KNN.pdf"
    elif modelName == "Support Machine Vector":
        modelSS = "SVM"
        pdfFile = file_name + "_SVM_Output.pdf"
    elif modelName == "Bagging Classifier":
        modelSS = "BC"
        pdfFile = file_name + "_BagCls.pdf"
    elif modelName == "Gradient Boosting":
        modelSS = "GBC"
        pdfFile = file_name + "_Gradient_Boosting.pdf"
    elif modelName == "Multi-Layer Perceptron (MLP)":
        modelName = "Multi-Layer Perceptron"
        modelSS = "MLP"
        pdfFile = file_name + "_MLP.pdf"
    elif modelName == "XGBoost":
        modelSS = "xgboost"
        pdfFile = file_name + "_XGBoost.pdf"
    elif modelName == "Random Forest":
        modelSS = "RF"
        pdfFile = file_name + "_Random_Forest.pdf"

    data = {"is_taken": True}
    saveUserComments(comments, modelSS, method, outputTab)
    # a variable pdf
    pdf = FPDF()
    document = Document()
    # pdf = exportPdfWithComments(pdf, document, modelName, comments, pdfFile, file_name+modelSS + "_NT_roc1.png", file_name+modelSS + "_NT_roc2.png", file_name+"_"+modelSS + "_NT_train_data.png", file_name+"_"+modelSS + "_NT_val_data.png",
    #                             file_name+"_"+modelSS + "_RS_roc1.png", file_name +
    #                             "_"+modelSS + "_RS_roc2.png", file_name
    #                             + "_"+modelSS + "_RS_train_data.png", file_name+"_" +
    #                             modelSS + "_RS_val_data.png",
    #                             file_name+"_"+modelSS + "_GS_roc1.png", file_name+"_"+modelSS + "_GS_roc2.png", file_name+"_"+modelSS + "_GS_train_data.png", file_name+"_"+modelSS + "_GS_val_data.png")
    # pdf.output(os.path.join(
    #     BASE_DIR, plot_dir_view + pdfFile))
    return JsonResponse(data)


def saveUserComments(comments, modelSS, method, outputTab):
    UserCommentsFiles = file_path + file_name + "_UserComments.csv"
    if os.path.exists(UserCommentsFiles):
        df_old = pd.read_csv(UserCommentsFiles)
        if (df_old["modelId"] == modelSS+'-'+str(method)+'-'+str(outputTab)).any():
            df_old.loc[(df_old["modelId"] == modelSS+'-' +
                        str(method)+'-'+str(outputTab)), "comments"] = comments
            df_old.to_csv(UserCommentsFiles, index=False, encoding='utf-8')
        else:
            data = [[modelSS+'-'+str(method)+'-'+str(outputTab),
                     modelSS, comments, method, outputTab]]
            df_new = pd.DataFrame(
                data, columns=['modelId', 'modelSS', 'comments', 'method', 'outputTab'])
            df = pd.concat([df_old, df_new], axis=0)
            df.to_csv(UserCommentsFiles, index=False, encoding='utf-8')
    else:
        data = [[modelSS+'-'+str(method)+'-'+str(outputTab),
                 modelSS, comments, method, outputTab]]
        df = pd.DataFrame(
            data, columns=['modelId', 'modelSS', 'comments', 'method', 'outputTab'])
        df.to_csv(UserCommentsFiles, index=False, encoding='utf-8')


def saveModelChartImage(request):
    
    chartImg = request.GET['chartImg']
    chartType = request.GET['chartType']
    method = request.GET['method']
    outputTab = request.GET['outputTab']
    modelSS = request.GET['modelSS']
    imgPath = request.GET['imgPath']
    imgPath = imgPath[1:]
    print('chartImg ', chartImg, ' chartType ',
          chartType, ' imgPath ', imgPath)
    UserChartFile = file_path + "_Chartimg.csv"
    UserCommentsFiles = file_path + file_name + "_UserComments.csv"
    directory = os.path.join(BASE_DIR, plot_dir_view+user_name+'Chartimgs')

    if os.path.exists(UserChartFile):
        df2 = pd.read_csv(UserChartFile)
        df = pd.read_csv(UserCommentsFiles)

        if not os.path.exists(directory):
            os.makedirs(directory)

        print("modelSS== '"+modelSS+"' and method==" +
              method + " and outputTab== "+outputTab)
        dffilter = df.query(
            "modelSS== '"+modelSS+"' and method==" + method + " and outputTab== "+outputTab)
        chartcomment = ""
        if len(dffilter) > 0:
            chartcomment = dffilter["comments"].values[0]
        destination = plot_dir_view+user_name+'Chartimgs/'+chartImg+'.png'
        data = [[chartType, chartImg, destination,
                chartcomment]]
        dfnew = pd.DataFrame(
            data, columns=['chartType', 'chartImg', 'destination', 'comments'])
        dfmerged = pd.concat([df2, dfnew], axis=0)
        dfmerged.to_csv(UserChartFile, index=False, encoding='utf-8')
        del dfmerged
        del dffilter
        # if (df["chartType"] == chartType).any():
        #     dffilter = df.query("chartType== '"+chartType+"'")
        #     for index, row in dffilter.iterrows():
        #         if os.path.exists(os.path.join(
        #                 BASE_DIR, plot_dir_view+row["imageName"])):
        # Source path
        source = os.path.join(
            BASE_DIR, imgPath)

        # Destination path
        destination = os.path.join(
            BASE_DIR, plot_dir_view+user_name+'Chartimgs/'+chartImg+'.png')
        shutil.copyfile(source, destination)
        print("File copied successfully.")
        #     del dffilter
        del df
        del df2
    else:
        df = pd.read_csv(UserCommentsFiles)
        dffilter = df.query(
            "modelSS== '"+modelSS+"' and method==" + method + " and outputTab== "+outputTab)
        chartcomment = ""
        if len(dffilter) > 0:
            chartcomment = dffilter["comments"].values[0]
        destination = plot_dir_view+user_name+'Chartimgs/'+chartImg+'.png'
        data = [[chartType, chartImg, destination,
                 chartcomment]]
        del dffilter
        dfnew = pd.DataFrame(
            data, columns=['chartType', 'chartImg', 'destination', 'comments'])
        dfnew.to_csv(UserChartFile, index=False, encoding='utf-8')
        del dfnew

        if not os.path.exists(directory):
            os.makedirs(directory)

        source = os.path.join(
            BASE_DIR, imgPath)

        # Destination path
        destination = os.path.join(
            BASE_DIR, plot_dir_view+user_name+'Chartimgs/'+chartImg+'.png')
        shutil.copyfile(source, destination)
        print("File copied successfully.")

    data = {"is_taken": True}
    return JsonResponse(data)


def exportPdfWithComments(pdf, document, modelName, comment, pdffile_name, nt_roc1, nt_roc2, nt_conf1, nt_conf2, rs_roc1, rs_roc2, rs_conf1, rs_conf2, gs_roc1, gs_roc2, gs_conf1, gs_conf2):

    # Add a page
    if(len(comment) > 0):
        pdf.add_page()
        document.add_page_break()
        pdf = addCommenttoPdf(10, 10, pdf, document, comment, modelName)
    # Add a page
    pdf.add_page()
    document.add_page_break()
    pdf = exportPdf(10, 10, pdf, document,  os.path.join(
        BASE_DIR, plot_dir_view, nt_roc1),
        "ROC Curve Validation data", os.path.join(
        BASE_DIR, plot_dir_view, nt_roc2),
        "ROC Curve Test data", modelName + " - No Parameters Tuning")
    pdf.add_page()
    document.add_page_break()
    pdf = exportPdf(10, 10, pdf, document,  os.path.join(
        BASE_DIR, plot_dir_view, nt_conf1),
        "Confusion Matrix Training data", os.path.join(
        BASE_DIR, plot_dir_view, nt_conf2),
        "Confusion Matrix Validation data", modelName + " - No Parameters Tuning")
    pdf.add_page()
    document.add_page_break()
    pdf = exportTestResultPdfFromExcel(
        pdf, document, modelName + " - No Parameters Tuning")

    # Add a page
    pdf.add_page()
    document.add_page_break()
    pdf = exportPdf(10, 10, pdf, document,   os.path.join(
        BASE_DIR, plot_dir_view, rs_roc1),
        "ROC Curve Validation data", os.path.join(
        BASE_DIR, plot_dir_view, rs_roc2),
        "ROC Curve Test data", modelName + " - Random Search")
    pdf.add_page()
    document.add_page_break()
    pdf = exportPdf(10, 10, pdf, document,   os.path.join(
        BASE_DIR, plot_dir_view, rs_conf1),
        "Confusion Matrix Training data", os.path.join(
        BASE_DIR, plot_dir_view, rs_conf2),
        "Confusion Matrix Validation data", modelName + " - Random Search")
    pdf.add_page()
    document.add_page_break()
    pdf = exportTestResultPdfFromExcel(
        pdf, document,  modelName + " - Random Search")

    # Add a page
    pdf.add_page()
    document.add_page_break()
    pdf = exportPdf(10, 10, pdf, document,   os.path.join(
        BASE_DIR, plot_dir_view, gs_roc1),
        "ROC Curve Validation data", os.path.join(
        BASE_DIR, plot_dir_view, gs_roc2),
        "ROC Curve Test data", modelName + " - Grid Search")
    pdf.add_page()
    document.add_page_break()
    pdf = exportPdf(10, 10, pdf, document,   os.path.join(
        BASE_DIR, plot_dir_view, gs_conf1),
        "Confusion Matrix Training data", os.path.join(
        BASE_DIR, plot_dir_view, gs_conf2),
        "Confusion Matrix Validation data", modelName + " - Grid Search")
    pdf.add_page()
    document.add_page_break()
    pdf = exportTestResultPdfFromExcel(
        pdf, document,  modelName + " - Grid Search")

    return pdf


def conceptualsoundness(request):
    try:
        UserChartFile = file_path + "_Chartimg.csv"
        data = {}
        result=[]
        resultDocumentation = []
        resultConcSnd = []
        modelFileExists = False
        processing = os.path.join(BASE_DIR, processingFile_path)
        df = pd.read_csv(processing, na_values='?')

        resultpROCESS = df.to_json(orient="records")
        resultpROCESS = json.loads(resultpROCESS)
        del df
        if(os.path.exists(os.path.join(
                BASE_DIR, plot_dir_view + file_name + "_ModelUsage.pdf"))):
            modelFileExists = True
        if os.path.exists(UserChartFile):
            df = pd.read_csv(UserChartFile)
            result = df.to_json(orient="records")
            result = json.loads(result)
            DocumentationData = file_path + file_name + "_DocumentationData.csv"
            if os.path.exists(DocumentationData):
                df_old = pd.read_csv(DocumentationData)
                idxLst = [*range(1, len(df_old)+1, 1)]
                df_new = pd.DataFrame(
                    idxLst, columns=['docIdx'])
                df = pd.concat([df_old, df_new], axis=1)
                resultDocumentation = df.to_json(orient="records")
                resultDocumentation = json.loads(resultDocumentation)

        report_file_path = os.path.join(BASE_DIR, 'static/csv_files/')
        report_file_name = "CS_"+user_name
        cnfrmsrcFiles = report_file_path + report_file_name + ".csv"
        if os.path.exists(cnfrmsrcFiles):
            df_ConcSnd = pd.read_csv(cnfrmsrcFiles, na_values='?')
            resultConcSnd = df_ConcSnd.to_json(orient="records")
            resultConcSnd = json.loads(resultConcSnd)
            del df_ConcSnd

        data = {
            'imgFiles': result,
            'pdfFile': "/static/media/ValidationReport.pdf",
            'modelDocs': resultDocumentation,
            'modelUsage': modelFileExists,
            'modelUsageFile': plot_dir + file_name + "_ModelUsage.pdf",
            'df': resultpROCESS,
            'resultConcSnd': resultConcSnd,
        }
        return render(request, 'conceptualsoundness.html', data)
    except Exception as e: 
        print('Error is', e,' , ',traceback.print_exc)
        return render(request, 'error.html')


def dataIntegrity(request):
    try:
        UserChartFile = file_path + "_Chartimg.csv"
        data = {}
        resultDocumentation = []
        DocumentationData = file_path + file_name + "_DocumentationData.csv"
        if os.path.exists(DocumentationData):
            df_old = pd.read_csv(DocumentationData)
            idxLst = [*range(1, len(df_old)+1, 1)]
            df_new = pd.DataFrame(
                idxLst, columns=['docIdx'])
            df = pd.concat([df_old, df_new], axis=1)
            resultDocumentation = df.to_json(orient="records")
            resultDocumentation = json.loads(resultDocumentation)
        data = {
            'modelDocs': resultDocumentation,
        }
        return render(request, 'dataIntegrity.html', data)
    except Exception as e:
        print('Error is', e,' , ',traceback.print_exc)
        return render(request, 'error.html')


def viewConcSnd(request):
    report_file_path = os.path.join(BASE_DIR, 'static/csv_files/')
    report_file_name = "CS_"+user_name
    cnfrmsrcFiles = report_file_path + report_file_name + ".csv"
    resultConcSnd = []
    if os.path.exists(cnfrmsrcFiles):
        df_ConcSnd = pd.read_csv(cnfrmsrcFiles, na_values='?')
        resultConcSnd = df_ConcSnd.to_json(orient="records")
        resultConcSnd = json.loads(resultConcSnd)
        del df_ConcSnd

    data = {
        'resultConcSnd': resultConcSnd, 'header': 'Conceptual Soundness'
    }
    return render(request, 'viewConSnd.html', data)


def viewDI(request):
    report_file_path = os.path.join(BASE_DIR, 'static/csv_files/')
    report_file_name = "DI_"+user_name
    cnfrmsrcFiles = report_file_path + report_file_name + ".csv"
    resultConcSnd = []
    if os.path.exists(cnfrmsrcFiles):
        df_ConcSnd = pd.read_csv(cnfrmsrcFiles, na_values='?')
        resultConcSnd = df_ConcSnd.to_json(orient="records")
        resultConcSnd = json.loads(resultConcSnd)
        del df_ConcSnd

    data = {
        'resultConcSnd': resultConcSnd, 'header': 'Data Integrity'
    }
    return render(request, 'viewConSnd.html', data)


def viewModelUsage(request):
    modelUsageReq = file_path + file_name + "_modelUsageReq.csv"
    resultConcSnd = []
    if os.path.exists(modelUsageReq):
        df = pd.read_csv(modelUsageReq)
        df = df.sort_values(by="reqIdx", ascending=True)

        resultConcSnd = df.to_json(orient="records")
        resultConcSnd = json.loads(resultConcSnd)
        del df

    data = {
        'resultConcSnd': resultConcSnd,
    }
    return render(request, 'viewModelUsage.html', data)


def saveCSData(request):
    comment = request.GET['comment']
    reqId = request.GET['reqId']
    title = request.GET['title']
    titleIdx = request.GET['titleIdx']
    report_file_path = os.path.join(BASE_DIR, 'static/csv_files/')
    report_file_name = "CS_"+user_name
    cnfrmsrcFiles = report_file_path + report_file_name + ".csv"
    if os.path.exists(cnfrmsrcFiles):
        df_old = pd.read_csv(cnfrmsrcFiles, encoding='utf-8')
        print('old df ', df_old)
        if (len(reqId) > 0):
            df_old.loc[df_old.reqID == float(reqId), "comment"] = comment
            df_old.to_csv(cnfrmsrcFiles, index=False, encoding='utf-8')
        else:
            maxid = df_old["reqID"].max()+1
            data = [['Comment', title, titleIdx,  maxid, comment,
                     '-',  '-', '-',  '-',  '-', '-']]
            print('data ', data)
            df_new = pd.DataFrame(
                data, columns=['section', 'title', 'titleIdx', 'reqID', 'comment', 'img', 'imgWidth', 'imgHeight', 'imgAlign', 'ImgTitle', 'TitleAlign'])
            df = pd.concat([df_old, df_new], axis=0)
            df.to_csv(cnfrmsrcFiles, index=False, encoding='utf-8')
            reqId = maxid
            # print(unicode(comment, errors='replace'))
    else:
        data = [['Comment', title, titleIdx, 1,
                 comment,  '-',  '-', '-',  '-',  '-', '-']]
        reqId = 1
        df = pd.DataFrame(
            data, columns=['section', 'title', 'titleIdx', 'reqID', 'comment', 'img', 'imgWidth', 'imgHeight', 'imgAlign', 'ImgTitle', 'TitleAlign'])
        df.to_csv(cnfrmsrcFiles, index=False, encoding='utf-8')
    data = {
        'is_taken': True,
        'reqId': str(reqId),
        'titleIdx': int(titleIdx)-1,
    }
    return JsonResponse(data)


def saveDIData(request):
    comment = request.GET['comment']
    reqId = request.GET['reqId']
    title = request.GET['title']
    titleIdx = request.GET['titleIdx']
    report_file_path = os.path.join(BASE_DIR, 'static/csv_files/')
    report_file_name = "DI_"+user_name
    cnfrmsrcFiles = report_file_path + report_file_name + ".csv"
    if os.path.exists(cnfrmsrcFiles):
        df_old = pd.read_csv(cnfrmsrcFiles, encoding='utf-8')
        print('old df ', df_old)
        if (len(reqId) > 0):
            df_old.loc[df_old.reqID == float(reqId), "comment"] = comment
            df_old.to_csv(cnfrmsrcFiles, index=False, encoding='utf-8')
        else:
            maxid = df_old["reqID"].max()+1
            data = [['Comment', title, titleIdx,  maxid, comment,
                     '-',  '-', '-',  '-',  '-', '-']]
            print('data ', data)
            df_new = pd.DataFrame(
                data, columns=['section', 'title', 'titleIdx', 'reqID', 'comment', 'img', 'imgWidth', 'imgHeight', 'imgAlign', 'ImgTitle', 'TitleAlign'])
            df = pd.concat([df_old, df_new], axis=0)
            df.to_csv(cnfrmsrcFiles, index=False, encoding='utf-8')
            reqId = maxid
            # print(unicode(comment, errors='replace'))
    else:
        data = [['Comment', title, titleIdx, 1,
                 comment,  '-',  '-', '-',  '-',  '-', '-']]
        reqId = 1
        df = pd.DataFrame(
            data, columns=['section', 'title', 'titleIdx', 'reqID', 'comment', 'img', 'imgWidth', 'imgHeight', 'imgAlign', 'ImgTitle', 'TitleAlign'])
        df.to_csv(cnfrmsrcFiles, index=False, encoding='utf-8')
    data = {
        'is_taken': True,
        'reqId': str(reqId),
        'titleIdx': int(titleIdx)-1,
    }
    return JsonResponse(data)


def Report(request):
    try:
        UserChartFile = file_path + "_Chartimg.csv"
        data = {}
        resultDocumentation = []
        result = []
        modelFileExists = "true"
        processing = os.path.join(BASE_DIR, processingFile_path)
        df = pd.read_csv(processing, na_values='?', encoding='utf-8')
        savedData = getSavedReportData()
        resultpROCESS = df.to_json(orient="records")
        resultpROCESS = json.loads(resultpROCESS)
        del df
        if(os.path.exists(os.path.join(
                BASE_DIR, plot_dir_view + file_name + "_modelUsageReq1.csv"))):
            dfmodelusage = pd.read_csv(os.path.join(
                BASE_DIR, plot_dir_view + file_name + "_modelUsageReq.csv"), na_values='?', encoding='utf-8')
            if len(dfmodelusage) > 0:
                modelFileExists = "false"
            del dfmodelusage
        print('modelFileExists is ', modelFileExists)
        outputfiles = []
        replication_files = os.path.join(
            BASE_DIR, 'static/replicationFiles/')
        if os.path.exists(replication_files):
            dir_list = os.listdir(os.path.join(
                BASE_DIR, 'static/replicationoutput/'))
            # prints all files
            outputfiles = dir_list
        if os.path.exists(UserChartFile):
            df = pd.read_csv(UserChartFile, encoding='utf-8')
            result = df.to_json(orient="records")
            result = json.loads(result)
            DocumentationData = file_path + file_name + "_DocumentationData.csv"
            if os.path.exists(DocumentationData):
                df_old = pd.read_csv(DocumentationData, encoding='utf-8')
                idxLst = [*range(1, len(df_old)+1, 1)]
                print('idxLst ', idxLst)
                df_new = pd.DataFrame(
                    idxLst, columns=['docIdx'])
                df = pd.concat([df_old, df_new], axis=1)
                resultDocumentation = df.to_json(orient="records")
                resultDocumentation = json.loads(resultDocumentation)
        data = {
            'imgFiles': result,
            'imgReplication': outputfiles,
            'pdfFile': "/static/media/ValidationReport.pdf",
            'policiesLst': getPolicies(),
            'modelDocs': resultDocumentation,
            'ismodelUsage': modelFileExists,
            'modelUsageFile': plot_dir + file_name + "_ModelUsage.pdf",
            'df': resultpROCESS,
            'savedReportData': savedData,
        }
        print('savedData is ', savedData)
        print('modelFileExists is ', modelFileExists)
        report_file_path = os.path.join(BASE_DIR, plot_dir_view)
        report_file_name = "temp_report_"+user_name
        cnfrmsrcFiles = report_file_path + report_file_name + ".csv"
        if os.path.exists(cnfrmsrcFiles):
            os.remove(cnfrmsrcFiles)
        return render(request, 'exportReport3.html', data)
    except Exception as e:
        print('Error is', e)
        print('Error is', traceback.print_exc())
        return render(request, 'error.html')


def ReportTxtEd(request):
    try:
        UserChartFile = file_path + "_Chartimg.csv"
        data = {}
        resultDocumentation = []
        result = []
        lstTbl = []
        isValFindings=False
        modelFileExists = "true"
        processing = os.path.join(BASE_DIR, processingFile_path)
        df = pd.read_csv(processing, na_values='?', encoding='utf-8')
        newTitles=[]
        savedData,newTitles = getSavedReportData()
        resultpROCESS = df.to_json(orient="records")
        resultpROCESS = json.loads(resultpROCESS)
        del df
        validationFindings = file_path + file_name + "_validationFindings.csv"    
        if os.path.exists(validationFindings):
            isValFindings=True
        if(os.path.exists(os.path.join(
                BASE_DIR, plot_dir_view + file_name + "_modelUsageReq1.csv"))):
            dfmodelusage = pd.read_csv(os.path.join(
                BASE_DIR, plot_dir_view + file_name + "_modelUsageReq.csv"), na_values='?', encoding='utf-8')
            if len(dfmodelusage) > 0:
                modelFileExists = "false"
            del dfmodelusage
        tblFile = file_path + user_name+"_Tables.csv"
        if os.path.exists(tblFile):
            df_tbl = pd.read_csv(tblFile)
            resultTbl = df_tbl.to_json(orient="records")
            lstTbl = json.loads(resultTbl)
            resultTbl = ""
            del df_tbl
        outputfiles = []
        replication_files = os.path.join(
            BASE_DIR, 'static/replicationFiles/')
        if os.path.exists(replication_files):
            dir_list = os.listdir(os.path.join(
                BASE_DIR, 'static/replicationoutput/'))
            # prints all files
            outputfiles = dir_list

        scnfiles = []
        scn_files = os.path.join(
            BASE_DIR, 'static/scenarioOutput/')
        if os.path.exists(scn_files):
            dir_list = os.listdir(os.path.join(
                BASE_DIR, 'static/scenarioOutput'))
            # prints all files
            scnfiles = dir_list 
        if os.path.exists(UserChartFile):
            df = pd.read_csv(UserChartFile, encoding='utf-8')
            result = df.to_json(orient="records")
            result = json.loads(result)
            DocumentationData = file_path + file_name + "_DocumentationData.csv"
            if os.path.exists(DocumentationData):
                df_old = pd.read_csv(DocumentationData, encoding='utf-8')
                idxLst = [*range(1, len(df_old)+1, 1)]
                print('idxLst ', idxLst)
                df_new = pd.DataFrame(
                    idxLst, columns=['docIdx'])
                df = pd.concat([df_old, df_new], axis=1)
                resultDocumentation = df.to_json(orient="records")
                resultDocumentation = json.loads(resultDocumentation)
        data = {
            'imgFiles': result,
            'imgReplication': outputfiles,
            'imgScn': scnfiles,
            'pdfFile': "/static/media/ValidationReport.pdf",
            'policiesLst': getPolicies(),
            'modelDocs': resultDocumentation,
            'ismodelUsage': modelFileExists,
            'modelUsageFile': plot_dir + file_name + "_ModelUsage.pdf",
            'df': resultpROCESS,
            'savedReportData': savedData,
            'lstTables': lstTbl,
            'newTitles':newTitles,
            'isValFindings':isValFindings,
        }
        
        report_file_path = os.path.join(BASE_DIR, plot_dir_view)
        report_file_name = "temp_report_"+user_name
        cnfrmsrcFiles = report_file_path + report_file_name + ".csv"
        if os.path.exists(cnfrmsrcFiles):
            os.remove(cnfrmsrcFiles)
        return render(request, 'exportReportTxtEd.html', data)
    except Exception as e:
        print('Error is', e)
        print('Error is', traceback.print_exc())
        return render(request, 'error.html')

def ReportTxtEdVaro(request):
    try:
        UserChartFile = file_path + "_Chartimg.csv"
        gridHeaders=["Assessment Area","Question / Request","MDD References","Request Document Name","Request Date","Response Date","Responsible Party","Status","Notes / Comments from Varo MO","Follow-up/New Questions","Request Date_Follow-up","Response Date_Follow-up"]
        data = {}
        resultDocumentation = []
        result = []
        lstTbl = []
        
        modelFileExists = "true"
        processing = os.path.join(BASE_DIR, processingFile_path)
        df = pd.read_csv(processing, na_values='?', encoding='utf-8')
        newTitles=[]
        savedData,newTitles = getSavedReportData('Varo')
        resultpROCESS = df.to_json(orient="records")
        resultpROCESS = json.loads(resultpROCESS) 
        quetionLogFile = src_files +  "Question_"+user_name + ".csv"
        del df
        qtnresult=[]
        rows=[]
        
        if os.path.exists(quetionLogFile):
            df = pd.read_csv(quetionLogFile)  
            df.fillna("",inplace=True)
            df_new = df.sort_values(by=['section', 'reqID'],ignore_index=True)
             
            df_new=df_new.drop(["emailId","ismailSent"], axis=1)  
            rowSpan=1
            txtSec=""
            # print(df_new)
            for idx, row in df_new.iterrows():   
                arr= str(row['question']).split("\n")   
                rows.append([len(arr),(len(arr)*25) ,(idx+1)])                 
            
            df_new2 = pd.DataFrame(
                rows, columns=['txtrows','rowH','idx'])

             
            df_new3 = pd.concat([df_new, df_new2], axis=1) 
            
            qtnresult = df_new3.to_json(orient="records")
            qtnresult = json.loads(qtnresult) 
        if(os.path.exists(os.path.join(
                BASE_DIR, plot_dir_view + file_name + "_modelUsageReq1.csv"))):
            dfmodelusage = pd.read_csv(os.path.join(
                BASE_DIR, plot_dir_view + file_name + "_modelUsageReq.csv"), na_values='?', encoding='utf-8')
            if len(dfmodelusage) > 0:
                modelFileExists = "false"
            del dfmodelusage
        tblFile = file_path + user_name+"_Tables.csv"
        if os.path.exists(tblFile):
            df_tbl = pd.read_csv(tblFile)
            resultTbl = df_tbl.to_json(orient="records")
            lstTbl = json.loads(resultTbl)
            resultTbl = ""
            del df_tbl
        outputfiles = []
        replication_files = os.path.join(
            BASE_DIR, 'static/replicationFiles/')
        if os.path.exists(replication_files):
            dir_list = os.listdir(os.path.join(
                BASE_DIR, 'static/replicationoutput/'))
            # prints all files
            outputfiles = dir_list

        scnfiles = []
        scn_files = os.path.join(
            BASE_DIR, 'static/scenarioOutput/')
        if os.path.exists(scn_files):
            dir_list = os.listdir(os.path.join(
                BASE_DIR, 'static/scenarioOutput'))
            # prints all files
            scnfiles = dir_list 
        if os.path.exists(UserChartFile):
            df = pd.read_csv(UserChartFile, encoding='utf-8')
            result = df.to_json(orient="records")
            result = json.loads(result)
            DocumentationData = file_path + file_name + "_DocumentationData.csv"
            if os.path.exists(DocumentationData):
                df_old = pd.read_csv(DocumentationData, encoding='utf-8')
                idxLst = [*range(1, len(df_old)+1, 1)]
                print('idxLst ', idxLst)
                df_new = pd.DataFrame(
                    idxLst, columns=['docIdx'])
                df = pd.concat([df_old, df_new], axis=1)
                resultDocumentation = df.to_json(orient="records")
                resultDocumentation = json.loads(resultDocumentation)
        data = {
            'imgFiles': result,
            'imgReplication': outputfiles,
            'imgScn': scnfiles,
            'pdfFile': "/static/media/ValidationReport.pdf",
            'policiesLst': getPolicies(),
            'modelDocs': resultDocumentation,
            'ismodelUsage': modelFileExists,
            'modelUsageFile': plot_dir + file_name + "_ModelUsage.pdf",
            'df': resultpROCESS,
            'savedReportData': savedData,
            'lstTables': lstTbl,
            'newTitles':newTitles,
            'headers':  gridHeaders,
            'qtnresult':qtnresult,
        }
        
        report_file_path = os.path.join(BASE_DIR, plot_dir_view)
        report_file_name = "temp_report_"+user_name
        cnfrmsrcFiles = report_file_path + report_file_name + ".csv"
        if os.path.exists(cnfrmsrcFiles):
            os.remove(cnfrmsrcFiles)
        return render(request, 'exportReportTxtEdVaro.html', data)
    except Exception as e:
        print('Error is', e)
        print('Error is', traceback.print_exc())
        return render(request, 'error.html')

def getPolicies():
    contactFile = file_path + user_name + "_Policies.csv"
    result = []
    if os.path.exists(contactFile):
        df = pd.read_csv(contactFile)
        result = df.to_json(orient="records")
        result = json.loads(result)

    return result


def generateReport(request):

    try:
        csv_file_name = "csvfile_"+user_name
        # a variable pdf
        pdf = FPDF()
        document = Document()
        section = document.sections[0]
        # Changing the orientation to landscape
        section.orientation = WD_ORIENT.LANDSCAPE

        # Printing the new orientation.
        columnPageIdx = -1
        edaPageIdx = -1
        modelsPageIdx = -1

        # pdf.add_page()
        # strsuprscript = "superscript test"
        # pdf.add_font("ArialUnicodeMS", "", font_files +
        #              "ARIALUNI.ttf", uni=True)
        # pdf.set_font('ArialUnicodeMS', '', 11)
        # strsuprscript = strsuprscript+"\u2075" + \
        #     " testtestâµtestÂ¹âµ  b'1.2 Model Description\n\ttest\xc2\xb9\xe2\x81\xb5 erwrwerwerwer wrew'"
        # # print('strsuprscript '+strsuprscript.encode(
        # #     'utf-8', 'replace').decode('latin-1'))
        # pdf.multi_cell(200, 20, strsuprscript, align='L')

        # pdf.output(os.path.join(
        #     BASE_DIR, "static/media/modelReport.pdf"))
        # data = {"is_taken": True}
        # return JsonResponse(data)

        pdf.add_page('P')
        pdf = addTitlenComments(pdf, document)

        pdf.add_page('P')
        addDocumentVesrionHistory(pdf, document)

        # #pdf = addCommentsnImgs(pdf)

        pdf, columnPageIdx, edaPageIdx, modelsPageIdx = addSummarynComments(
            pdf, document, columnPageIdx, edaPageIdx, modelsPageIdx)

        # Add a page

        # pdf.add_page()
        # document.add_page_break()
        # pdf = addImplCtrl(pdf, document)  # to be uncommented

        # Add a page
        # pdf.add_page()
        # document.add_page_break()
        # savefile_name = file_path + csv_file_name + ".csv"
        # df = pd.read_csv(savefile_name, na_values='?')

        # pdf = exportDatatypenCnt(pdf, document, df, "", columnPageIdx)

        # chartFiles = file_path + "_ChartViewd.csv"
        # if os.path.exists(chartFiles):
        #     df = pd.read_csv(chartFiles)

        #     pdf.add_page()
        #     pdf.set_xy(10, 80)
        #     pdf.set_font("Arial",  size=22)
        #     pdf.set_text_color(0.0, 0.0, 0.0)
        #     pdf.set_link(edaPageIdx, pdf.page_no())
        #     pdf.multi_cell(0, 10, "Exploratory Data Analysis", align='C')

        #     document.add_page_break()
        #     paragraphBlank = document.add_paragraph()
        #     paragraph_format = paragraphBlank.paragraph_format
        #     paragraph_format.line_spacing = Pt(180)

        #     paragraph = document.add_paragraph()
        #     paragraph_format = paragraph.paragraph_format
        #     # paragraph_format.line_spacing = Pt(180)
        #     paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #     run = paragraph.add_run("Exploratory Data Analysis")
        #     font = run.font
        #     font.name = 'Arial'
        #     font.size = Pt(22)

        #     chartType = ["Heatmap", "Bar chart", "Stacked Bar chart",
        #                  "Strip Plot", "Distribution", "Box Plot"]
        #     for i in range(len(chartType)):
        #         if (df["chartType"] == chartType[i]).any():
        #             dffilter = df.query("chartType== '"+chartType[i]+"'")
        #             for index, row in dffilter.iterrows():
        #                 if os.path.exists(os.path.join(
        #                         BASE_DIR, plot_dir_view+row["imageName"])):
        #                     pdf.add_page()
        #                     document.add_page_break()
        #                     if(row["chartType"] == "Heatmap"):
        #                         pdf = exportgraphImgPdf(pdf, document, os.path.join(
        #                             BASE_DIR, plot_dir_view+row["imageName"]),  "Correlation on independent variables-Heat map", row["comments"])
        #                     else:
        #                         pdf = exportgraphImgPdf(pdf, document, os.path.join(
        #                             BASE_DIR, plot_dir_view+row["imageName"]), row["chartType"]+" "+row["xaxisval"]+" vs "+row["yaxisval"], row["comments"])
        #             del dffilter
        #     del df

        # pdf.output(os.path.join(
        #     BASE_DIR, "static/media/modelReport.pdf"))
        # data = {"is_taken": True}
        # return JsonResponse(data)

        # pdf.add_page()
        # pdf.set_xy(10, 80)
        # pdf.set_font("Arial",  size=22)
        # pdf.set_text_color(0.0, 0.0, 0.0)
        # pdf.set_link(modelsPageIdx, pdf.page_no())
        # pdf.multi_cell(0, 10, "Modeling", align='C')

        # document.add_page_break()
        # paragraphBlank = document.add_paragraph()
        # paragraph_format = paragraphBlank.paragraph_format
        # paragraph_format.line_spacing = Pt(180)

        # paragraph = document.add_paragraph()
        # paragraph_format = paragraph.paragraph_format
        # # paragraph_format.line_spacing = Pt(180)
        # paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # run = paragraph.add_run("Modeling")
        # font = run.font
        # font.name = 'Arial'
        # font.size = Pt(22)

        # UserCommentsFiles = file_path + file_name + "_UserComments.csv"
        # if os.path.exists(UserCommentsFiles):
        #     df = pd.read_csv(UserCommentsFiles)

        # if os.path.exists(os.path.join(
        #         BASE_DIR, plot_dir_view + file_name + "_KNN.pdf")):
        #     modelName = "K Nearest Neighbors"
        #     modelSS = "KNN"
        #     pdfFile = file_name + "_KNN.pdf"
        #     comments = ""
        #     if os.path.exists(UserCommentsFiles):
        #         if (df["modelSS"] == modelSS).any():
        #             dffilter = df.query("modelSS== '"+modelSS+"'")
        #             for index, row in dffilter.iterrows():
        #                 comments = row["comments"]
        #     pdf = exportPdfWithComments(pdf, document, modelName, comments, pdfFile, file_name+modelSS + "_NT_roc1.png", file_name+modelSS + "_NT_roc2.png", file_name+"_"+modelSS + "_NT_train_data.png", file_name+"_"+modelSS + "_NT_val_data.png",
        #                                 file_name+"_"+modelSS + "_RS_roc1.png", file_name +
        #                                 "_"+modelSS + "_RS_roc2.png", file_name
        #                                 + "_"+modelSS + "_RS_train_data.png", file_name+"_" +
        #                                 modelSS + "_RS_val_data.png",
        #                                 file_name+"_"+modelSS + "_GS_roc1.png", file_name+"_"+modelSS + "_GS_roc2.png", file_name+"_"+modelSS + "_GS_train_data.png", file_name+"_"+modelSS + "_GS_val_data.png")

        # if os.path.exists(os.path.join(
        #         BASE_DIR, plot_dir_view + file_name + "_SVM_Output.pdf")):
        #     modelName = "Support Machine Vector"
        #     modelSS = "SVM"
        #     pdfFile = file_name + "_SVM_Output.pdf"
        #     comments = ""
        #     if os.path.exists(UserCommentsFiles):
        #         if (df["modelSS"] == modelSS).any():
        #             dffilter = df.query("modelSS== '"+modelSS+"'")
        #             for index, row in dffilter.iterrows():
        #                 comments = row["comments"]
        #     pdf = exportPdfWithComments(pdf, document, modelName, comments, pdfFile, file_name+modelSS + "_NT_roc1.png", file_name+modelSS + "_NT_roc2.png", file_name+"_"+modelSS + "_NT_train_data.png", file_name+"_"+modelSS + "_NT_val_data.png",
        #                                 file_name+"_"+modelSS + "_RS_roc1.png", file_name +
        #                                 "_"+modelSS + "_RS_roc2.png", file_name
        #                                 + "_"+modelSS + "_RS_train_data.png", file_name+"_" +
        #                                 modelSS + "_RS_val_data.png",
        #                                 file_name+"_"+modelSS + "_GS_roc1.png", file_name+"_"+modelSS + "_GS_roc2.png", file_name+"_"+modelSS + "_GS_train_data.png", file_name+"_"+modelSS + "_GS_val_data.png")

        # if os.path.exists(os.path.join(
        #         BASE_DIR, plot_dir_view + file_name + "_BagCls.pdf")):
        #     modelName = "Bagging Classifier"
        #     modelSS = "BC"
        #     pdfFile = file_name + "_BagCls.pdf"
        #     comments = ""
        #     if os.path.exists(UserCommentsFiles):
        #         if (df["modelSS"] == modelSS).any():
        #             dffilter = df.query("modelSS== '"+modelSS+"'")
        #             for index, row in dffilter.iterrows():
        #                 comments = row["comments"]
        #     pdf = exportPdfWithComments(pdf, document, modelName, comments, pdfFile, file_name+modelSS + "_NT_roc1.png", file_name+modelSS + "_NT_roc2.png", file_name+"_"+modelSS + "_NT_train_data.png", file_name+"_"+modelSS + "_NT_val_data.png",
        #                                 file_name+"_"+modelSS + "_RS_roc1.png", file_name +
        #                                 "_"+modelSS + "_RS_roc2.png", file_name
        #                                 + "_"+modelSS + "_RS_train_data.png", file_name+"_" +
        #                                 modelSS + "_RS_val_data.png",
        #                                 file_name+"_"+modelSS + "_GS_roc1.png", file_name+"_"+modelSS + "_GS_roc2.png", file_name+"_"+modelSS + "_GS_train_data.png", file_name+"_"+modelSS + "_GS_val_data.png")

        # if os.path.exists(os.path.join(
        #         BASE_DIR, plot_dir_view + file_name + "_Gradient_Boosting.pdf")):
        #     modelName = "Gradient Boosting"
        #     modelSS = "GBC"
        #     pdfFile = file_name + "_Gradient_Boosting.pdf"
        #     comments = ""
        #     if os.path.exists(UserCommentsFiles):
        #         if (df["modelSS"] == modelSS).any():
        #             dffilter = df.query("modelSS== '"+modelSS+"'")
        #             for index, row in dffilter.iterrows():
        #                 comments = row["comments"]
        #     pdf = exportPdfWithComments(pdf, document, modelName, comments, pdfFile, file_name+modelSS + "_NT_roc1.png", file_name+modelSS + "_NT_roc2.png", file_name+"_"+modelSS + "_NT_train_data.png", file_name+"_"+modelSS + "_NT_val_data.png",
        #                                 file_name+"_"+modelSS + "_RS_roc1.png", file_name +
        #                                 "_"+modelSS + "_RS_roc2.png", file_name
        #                                 + "_"+modelSS + "_RS_train_data.png", file_name+"_" +
        #                                 modelSS + "_RS_val_data.png",
        #                                 file_name+"_"+modelSS + "_GS_roc1.png", file_name+"_"+modelSS + "_GS_roc2.png", file_name+"_"+modelSS + "_GS_train_data.png", file_name+"_"+modelSS + "_GS_val_data.png")

        # if os.path.exists(os.path.join(
        #         BASE_DIR, plot_dir_view + file_name + "_MLP.pdf")):
        #     modelName = "Multi-Layer Perceptron (MLP)"
        #     modelName = "Multi-Layer Perceptron"
        #     modelSS = "MLP"
        #     pdfFile = file_name + "_MLP.pdf"
        #     comments = ""
        #     if os.path.exists(UserCommentsFiles):
        #         if (df["modelSS"] == modelSS).any():
        #             dffilter = df.query("modelSS== '"+modelSS+"'")
        #             for index, row in dffilter.iterrows():
        #                 comments = row["comments"]
        #     pdf = exportPdfWithComments(pdf, document, modelName, comments, pdfFile, file_name+modelSS + "_NT_roc1.png", file_name+modelSS + "_NT_roc2.png", file_name+"_"+modelSS + "_NT_train_data.png", file_name+"_"+modelSS + "_NT_val_data.png",
        #                                 file_name+"_"+modelSS + "_RS_roc1.png", file_name +
        #                                 "_"+modelSS + "_RS_roc2.png", file_name
        #                                 + "_"+modelSS + "_RS_train_data.png", file_name+"_" +
        #                                 modelSS + "_RS_val_data.png",
        #                                 file_name+"_"+modelSS + "_GS_roc1.png", file_name+"_"+modelSS + "_GS_roc2.png", file_name+"_"+modelSS + "_GS_train_data.png", file_name+"_"+modelSS + "_GS_val_data.png")

        # if os.path.exists(os.path.join(
        #         BASE_DIR, plot_dir_view + file_name + "_XGBoost.pdf")):
        #     modelName = "XGBoost"
        #     modelSS = "xgboost"
        #     pdfFile = file_name + "_XGBoost.pdf"
        #     comments = ""
        #     if os.path.exists(UserCommentsFiles):
        #         if (df["modelSS"] == modelSS).any():
        #             dffilter = df.query("modelSS== '"+modelSS+"'")
        #             for index, row in dffilter.iterrows():
        #                 comments = row["comments"]
        #     pdf = exportPdfWithComments(pdf, document, modelName, comments, pdfFile, file_name+modelSS + "_NT_roc1.png", file_name+modelSS + "_NT_roc2.png", file_name+"_"+modelSS + "_NT_train_data.png", file_name+"_"+modelSS + "_NT_val_data.png",
        #                                 file_name+"_"+modelSS + "_RS_roc1.png", file_name +
        #                                 "_"+modelSS + "_RS_roc2.png", file_name
        #                                 + "_"+modelSS + "_RS_train_data.png", file_name+"_" +
        #                                 modelSS + "_RS_val_data.png",
        #                                 file_name+"_"+modelSS + "_GS_roc1.png", file_name+"_"+modelSS + "_GS_roc2.png", file_name+"_"+modelSS + "_GS_train_data.png", file_name+"_"+modelSS + "_GS_val_data.png")

        # if os.path.exists(os.path.join(
        #         BASE_DIR, plot_dir_view + file_name + "_Random_Forest.pdf")):
        #     modelName = "Random Forest"
        #     modelSS = "RF"
        #     pdfFile = file_name + "_Random_Forest.pdf"
        #     comments = ""
        #     if os.path.exists(UserCommentsFiles):
        #         if (df["modelSS"] == modelSS).any():
        #             dffilter = df.query("modelSS== '"+modelSS+"'")
        #             for index, row in dffilter.iterrows():
        #                 comments = row["comments"]
        #     pdf = exportPdfWithComments(pdf, document, modelName, comments, pdfFile, file_name+modelSS + "_NT_roc1.png", file_name+modelSS + "_NT_roc2.png", file_name+"_"+modelSS + "_NT_train_data.png", file_name+"_"+modelSS + "_NT_val_data.png",
        #                                 file_name+"_"+modelSS + "_RS_roc1.png", file_name +
        #                                 "_"+modelSS + "_RS_roc2.png", file_name
        #                                 + "_"+modelSS + "_RS_train_data.png", file_name+"_" +
        #                                 modelSS + "_RS_val_data.png",
        #                                 file_name+"_"+modelSS + "_GS_roc1.png", file_name+"_"+modelSS + "_GS_roc2.png", file_name+"_"+modelSS + "_GS_train_data.png", file_name+"_"+modelSS + "_GS_val_data.png")

        pdf = addReferences(pdf)
        pdf.add_page('P')
        document.add_page_break()
        pdf = addDocumentationComments(pdf, document)

        pdf.add_page('P')
        document.add_page_break()
        pdf = addDataQuality(pdf, document)

        pdf.output(os.path.join(
            BASE_DIR, "static/media/ValidationReport.pdf"))

        document.save(os.path.join(
            BASE_DIR, "static/media/demo.docx"))

        reportFilepath = os.path.join(
            BASE_DIR, "static/media/ValidationReport.pdf")
        if os.path.exists(reportFilepath):
            # with open(reportFilepath, 'rb') as fh:
            #     response = HttpResponse(
            #         fh.read(), content_type="application/force-download")
            #     response['Content-Disposition'] = 'attachment; filename=' + \
            #         os.path.basename(reportFilepath)
            #     return response
            # from django.utils.encoding import smart_str

            # # mimetype is replaced by content_type for django 1.7
            # response = HttpResponse(content_type="application/force-download")
            # response['Content-Disposition'] = 'attachment; filename=' + \
            #     os.path.basename(reportFilepath)
            # response['X-Sendfile'] = smart_str(reportFilepath)
            # return response
            data = {"is_taken": True}
            return JsonResponse(data)
        raise Http404
    except Exception as e:
        print(e)
        print(traceback.print_exc())
        data = {"is_taken": True}
        return JsonResponse(data)


def generateReportTxtEd(request):

    try: 
        # a variable pdf
        pdf = MyFPDF()
        document = Document()
        section = document.sections[0]
        # Changing the orientation to landscape
        section.orientation = WD_ORIENT.LANDSCAPE

        # Printing the new orientation.
        columnPageIdx = -1
        edaPageIdx = -1
        modelsPageIdx = -1

        pdf.add_page('P')
        pdf = addTitlenComments(pdf, document)

        pdf.add_page('P')
        addDocumentVesrionHistory(pdf, document)

        # #pdf = addCommentsnImgs(pdf)

        pdf, columnPageIdx, edaPageIdx, modelsPageIdx = addSummarynCommentsHTML(
            pdf, document, columnPageIdx, edaPageIdx, modelsPageIdx)
        pdf = addReferences(pdf)
        pdf.add_page('P')
        document.add_page_break()
        pdf = addDocumentationComments(pdf, document)

        pdf.add_page('P')
        document.add_page_break()
        pdf = addDataQuality(pdf, document)

        pdf.output(os.path.join(
            BASE_DIR, "static/media/ValidationReport.pdf"))

        document.save(os.path.join(
            BASE_DIR, "static/media/demo.docx"))

        reportFilepath = os.path.join(
            BASE_DIR, "static/media/ValidationReport.pdf")
        if os.path.exists(reportFilepath):
            # with open(reportFilepath, 'rb') as fh:
            #     response = HttpResponse(
            #         fh.read(), content_type="application/force-download")
            #     response['Content-Disposition'] = 'attachment; filename=' + \
            #         os.path.basename(reportFilepath)
            #     return response
            # from django.utils.encoding import smart_str

            # # mimetype is replaced by content_type for django 1.7
            # response = HttpResponse(content_type="application/force-download")
            # response['Content-Disposition'] = 'attachment; filename=' + \
            #     os.path.basename(reportFilepath)
            # response['X-Sendfile'] = smart_str(reportFilepath)
            # return response
            data = {"is_taken": True}
            return JsonResponse(data)
        raise Http404
    except Exception as e:
        print(e)
        print(traceback.print_exc())
        data = {"is_taken": False}
        return JsonResponse(data)


def addCommentsnImgs(pdf):
    temp_report_file_path = os.path.join(BASE_DIR, plot_dir_view)
    temp_report_file_name = "temp_report_"+user_name
    temp_cnfrmsrcFiles = temp_report_file_path + temp_report_file_name + ".csv"

    report_file_path = os.path.join(BASE_DIR, plot_dir_view)
    report_file_name = "report_"+user_name
    cnfrmsrcFiles = report_file_path + report_file_name + ".csv"

    if os.path.exists(temp_cnfrmsrcFiles):
        if os.path.exists(cnfrmsrcFiles):
            os.remove(cnfrmsrcFiles)
        os.rename(temp_cnfrmsrcFiles, cnfrmsrcFiles)

    if os.path.exists(cnfrmsrcFiles):
        df = pd.read_csv(cnfrmsrcFiles)
        # ['section',  'reqID', 'comment', 'img', 'imgWidth', 'imgHeight', 'imgAlign', 'ImgTitle', 'TitleAlign']
        dfprint = df.sort_values(by='reqID', ascending=True)
        pdf.add_page()
        pdf.set_xy(10, 10)
        for index, row in dfprint.iterrows():
            if row["section"] == "Comment":
                pdf.set_xy(10, pdf.get_y())
                pdf.set_font("Arial",  size=10)
                pdf.set_text_color(0.0, 0.0, 0.0)
                pdf.multi_cell(0, 5, row["comment"].encode(
                    'latin-1', 'replace').decode('latin-1'), align='L')
            else:
                imgW = (700/4)
                imgW = imgW * float(row["imgWidth"])
                if(row["imgAlign"] == "right"):
                    # print('right corener ', (700/4)*float(row["imgWidth"]))
                    if(175-((700/4)*float(row["imgWidth"])) > 0):
                        pdf.set_xy(
                            175-((700/4)*float(row["imgWidth"])), pdf.get_y())
                    else:
                        pdf.set_xy(10, pdf.get_y())
                elif(row["imgAlign"] == "center"):
                    print('row["imgWidth"] ', row["imgWidth"],
                          '((700/4)*float(row["imgWidth"])) ', ((700/4)*float(row["imgWidth"])))
                    print('center point ',
                          (175-((700/4)*float(row["imgWidth"])))/2)
                    if(((700/4)*float(row["imgWidth"])) < 175):
                        pdf.set_xy(
                            (175-((700/4)*float(row["imgWidth"])))/2, pdf.get_y())
                    else:
                        pdf.set_xy(10, pdf.get_y())
                else:
                    pdf.set_xy(10, pdf.get_y())
                pdf.image(os.path.join(
                    BASE_DIR, row["img"]),  link='', type='',  w=(700/4)*float(row["imgWidth"]), h=(450/4)*float(row["imgHeight"]))

    return pdf


def downloadReport(request):
    reportFilepath = os.path.join(BASE_DIR, "static/media/modelReport.pdf")
    if os.path.exists(reportFilepath):
        with open(reportFilepath, 'rb') as fh:
            response = HttpResponse(
                fh.read(), content_type="application/force-download")
            response['Content-Disposition'] = 'attachment; filename=' + \
                os.path.basename(reportFilepath)
            return response
    raise Http404


def addTitlenComments(pdf, document):
    x, y = 10, 10
    DocumentationData = file_path + file_name + "_ModelInfo.csv"
    if os.path.exists(DocumentationData):
        df = pd.read_csv(DocumentationData, encoding='utf-8')
        # print('df ', df)
        pdf.set_xy(80, y)
        pdf.image(os.path.join(
            BASE_DIR, "static/images/varo-logo.png"),  link='', type='',  w=50, h=20)
        y = y+60
        dfsorted = df.sort_values(by='reqId', ascending=False)
        header = "Varo Money, Inc."
        Title = dfsorted["ModelNm"].values[len(dfsorted)-1]
        SubTitle1 = "Validation Report"
        SubTitle2 = "Model Version " + \
            str(dfsorted["ModelVersion"].values[len(dfsorted)-1])
        SubTitle3 = "Model ID: " + \
            str(dfsorted["ModelID"].values[len(dfsorted)-1])
        pdf.set_xy(x, y)
        pdf.set_font("Arial", 'B', size=26)
        pdf.set_text_color(0.0, 0.0, 0.0)
        pdf.multi_cell(0, 10, header, align='C')
        pdf.set_font("Arial", 'B', size=22)
        y += 30
        pdf.set_xy(x, y)
        pdf.multi_cell(0, 10, Title, align='C')
        y += 10
        pdf.set_xy(x, y)
        pdf.multi_cell(0, 10, str(SubTitle1), align='C')
        y += 10
        pdf.set_xy(x, y)
        pdf.multi_cell(0, 10, str(SubTitle2), align='C')
        y += 10
        pdf.set_xy(x, y)
        pdf.multi_cell(0, 10, str(SubTitle3), align='C')

    return pdf


def addDocumentVesrionHistory(pdf, document):
    x, y = 10, 10
    DocumentationData = file_path + file_name + "_ModelInfo.csv"
    if os.path.exists(DocumentationData):
        df = pd.read_csv(DocumentationData, encoding='utf-8')
        dfsorted = df.sort_values(by='reqId', ascending=False)
        header = "Document Version History"
        pdf.set_xy(x, y)
        pdf.set_font("Arial", size=10)

        pdf.multi_cell(0, 10, header, align='L')

        y = pdf.get_y()+10
        pdf.set_font("Arial", 'B', size=9)
        pdf.set_fill_color(211, 211, 211)
        pdf.set_xy(20, y)
        pdf.cell(20, 5, "Version #", 1, fill=True, align='C')

        pdf.set_xy(40, y)
        pdf.cell(25, 5, "Date", 1, fill=True, align='C')

        pdf.set_xy(65, y)
        pdf.cell(25, 5, "Author", 1, fill=True, align='C')

        pdf.set_xy(90, y)
        pdf.cell(30, 5, "Approved By", 1, fill=True, align='C')

        pdf.set_xy(120, y)
        pdf.cell(0, 5, "Version Description", 1, fill=True, align='C')
        # Start from the first cell. Rows and columns are zero indexed.
        # print('dfsorted ', dfsorted)
        pdf.set_font("Arial", size=9)
        y += 5
        pdf.set_xy(20, y)
        pdf.cell(
            20, 5, str(dfsorted["ModelVersion"].values[len(dfsorted)-1]), 1, align='C')
        pdf.set_xy(40, y)
        pdf.cell(
            25, 5, str(dfsorted["Date"].values[len(dfsorted)-1]), 1, align='C')
        pdf.set_xy(65, y)
        pdf.cell(25, 5, "Varo", 1, align='C')
        pdf.set_xy(90, y)
        pdf.cell(30, 5, "Administrator", 1, align='C')
        pdf.set_xy(120, y)
        pdf.cell(
            0, 5, str(dfsorted["ModelDesc"].values[len(dfsorted)-1]), 1, align='C')

        y += 20
        pdf.set_font("Arial", 'B', size=9)
        pdf.set_xy(20, y)
        pdf.cell(
            0, 5, "Model Identification and Stakeholders Summary")

        y += 10
        pdf.set_fill_color(211, 211, 211)
        pdf.set_font("Arial", 'B', size=9)
        pdf.set_xy(20, y)
        pdf.cell(
            50, 6, "Model Name", 1, fill=True, align='L')
        pdf.set_font("Arial", size=9)
        pdf.set_xy(70, y)
        pdf.cell(
            0, 6, str(dfsorted["ModelNm"].values[len(dfsorted)-1]), 1, align='L')

        y += 6
        pdf.set_fill_color(211, 211, 211)
        pdf.set_font("Arial", 'B', size=9)
        pdf.set_xy(20, y)
        pdf.cell(
            50, 6, "Model Version", 1, fill=True, align='L')
        pdf.set_font("Arial", size=9)
        pdf.set_xy(70, y)
        pdf.cell(
            0, 6, str(dfsorted["ModelVersion"].values[len(dfsorted)-1]), 1, align='L')

        y += 6
        pdf.set_fill_color(211, 211, 211)
        pdf.set_font("Arial", 'B', size=9)
        pdf.set_xy(20, y)
        pdf.cell(
            50, 6, "Model ID", 1, fill=True, align='L')
        pdf.set_font("Arial", size=9)
        pdf.set_xy(70, y)
        pdf.cell(
            0, 6, str(dfsorted["ModelID"].values[len(dfsorted)-1]), 1, align='L')

        y += 6
        pdf.set_fill_color(211, 211, 211)
        pdf.set_font("Arial", 'B', size=9)
        pdf.set_xy(20, y)
        pdf.cell(
            50, 6, "Model Developer/Vendor:", 1, fill=True, align='L')
        pdf.set_font("Arial", size=9)
        pdf.set_xy(70, y)
        pdf.cell(
            0, 6, str(dfsorted["ModelDev"].values[len(dfsorted)-1]), 1, align='L')

        y += 6
        pdf.set_fill_color(211, 211, 211)
        pdf.set_font("Arial", 'B', size=9)
        pdf.set_xy(20, y)
        pdf.cell(
            50, 6, "Business Unit/Department:", 1, fill=True, align='L')
        pdf.set_font("Arial", size=9)
        pdf.set_xy(70, y)
        pdf.cell(
            0, 6, str(dfsorted["ModelDept"].values[len(dfsorted)-1]), 1, align='L')

        y += 6
        pdf.set_fill_color(211, 211, 211)
        pdf.set_font("Arial", 'B', size=9)
        pdf.set_xy(20, y)
        pdf.cell(
            50, 6, "Model Owner (Position):", 1, fill=True, align='L')
        pdf.set_font("Arial", size=9)
        pdf.set_xy(70, y)
        pdf.cell(
            0, 6, str(dfsorted["ModelOwner"].values[len(dfsorted)-1]), 1, align='L')

    return pdf


def addSummarynCommentsHTML(pdf, document, columnPageIdx, edaPageIdx, modelsPageIdx):

    # SummaryDataFiles = file_path + file_name + "_SummaryData.csv"
    temp_report_file_path = os.path.join(BASE_DIR, plot_dir_view)
    temp_report_file_name = "temp_report_"+user_name
    temp_cnfrmsrcFiles = temp_report_file_path + temp_report_file_name + ".csv"

    report_file_path = os.path.join(BASE_DIR, plot_dir_view)

    report_file_name = "report_"+user_name
    cnfrmsrcFiles = report_file_path + report_file_name + ".csv"
    linkArr = []
    pageNoArr = []
    valFindingsLinkIdx = -1
    if os.path.exists(temp_cnfrmsrcFiles):
        if os.path.exists(cnfrmsrcFiles):
            os.remove(cnfrmsrcFiles)
        os.rename(temp_cnfrmsrcFiles, cnfrmsrcFiles)

    if os.path.exists(cnfrmsrcFiles):
        df_old = pd.read_csv(cnfrmsrcFiles, encoding='utf-8')
        df = df_old.sort_values(
            by=['titleIdx', 'subtitleIdx', 'subsubtitleIdx', 'reqID'], ascending=True)
        # print('sorted df ', df)
        pdf, linkArr, valFindingsLinkIdx, pageNoArr, columnPageIdx, edaPageIdx, modelsPageIdx = addTableofContents(
            pdf, df, document, valFindingsLinkIdx, columnPageIdx, edaPageIdx, modelsPageIdx)
        # pdf.set_left_margin(32)
        pdf.add_page('P')
        document.add_page_break()
        x, y = 10, pdf.get_y()
        exeSummTtlIdx = ""
        exeSummTtl = ""

        for index, row in df.iterrows():
            x, y = 10, pdf.get_y()
            if(str(row["title"]) == "Executive Summary"):
                exeSummTtl = str(row["title"])
                dffilter = df.loc[df.title == exeSummTtl]
                maxIdx = float(dffilter["subtitleIdx"].max())

                maxIdx = maxIdx+0.1
                exeSummTtlIdx = str(round(maxIdx, 2))
                if(len(dffilter) == 1):
                    exeSummTtlIdx = "1.1"
                del(dffilter)
            # if((exeSummTtl != str(row["title"]) and exeSummTtl != "")):
            #     pdf, new_y = addValFinfings(
            #         pdf, exeSummTtlIdx, x, y, valFindingsLinkIdx)
            #     exeSummTtl = ""
            #     if(new_y>y):
            #         y = new_y+10
            #         pdf.set_xy(x, y)
            #         print('new_y 1', y)
            # pdf.add_font("comic sans ms;", "",
            #              font_files + "ComicSansMS3.ttf", uni=True)
            if row["section"] == "Comment":
                if(str(row["comment"]) != "nan"):
                    y = pdf.get_y()
                    print('y at start is ',y)
                    pdf.set_xy(x, y)
                    pdf.add_font("ArialUnicodeMS", "",
                                 font_files + "ARIALUNI.ttf", uni=True)
                    pdf.set_font('ArialUnicodeMS', '', 9)
                    # ##pdf.set_font("Arial",  size=9)
                    pdf.set_text_color(0.0, 0.0, 0.0)
                    pdf.set_left_margin(10)
                    pdf.set_link(linkArr[index], y, pdf.page_no())
                    encdStr = str(row["comment"]).replace(
                        'src="/static', 'src="static').replace('height="24"', 'height="20"')
                    encdStr = str(encdStr.encode('utf-8'), 'utf-8')
                    commentstr = ""
                    import re
                    tblLst = re.findall(
                        '<div class="appTblsss" id="', encdStr) 
                    for match in re.finditer('<div class="appTblsss" id="', encdStr):
                        print('match is ', match)

                    tbl_index = -1
                    if '<div class="appTblsss" id="'.lower() in encdStr.lower():
                        tbl_index = encdStr.index(
                            '<div class="appTblsss" id="')
                        tblid_index = encdStr.index(
                            '"><table style=', tbl_index)
                        # print('tblid_index ', tblid_index)
                        # print(' tableId is ', encdStr[(tbl_index+27):tblid_index])
                        table_ID = encdStr[(tbl_index+27):tblid_index]
                        tblend_index = encdStr.index(table_ID+'End"')
                        tblend_index = tblend_index+17+len(table_ID)
                    # print('tblend_index is ', tblend_index)
                    itr = 0
                    # print('coment after replacing the table is ',
                    #       encdStr[:tbl_index] + "" + encdStr[tblend_index:])
                    # print('len of encdStr ', len(encdStr))
                    while itr < len(encdStr):
                        # for itr in range(len(encdStr)):
                        # print('itr is ', itr)
                        if tbl_index == itr:
                            y = pdf.get_y()
                            x = 10
                            pdf.set_xy(x, y)
                            pdf.set_font('Arial', '', 9)
                            pdf.write_html(commentstr) 
                            commentstr = ""
                            print('table_ID is ',table_ID)
                            addTabletoRpt(pdf, table_ID)
                            itr = tblend_index
                            print('added here  1', commentstr)
                        elif(checkSymbol(encdStr[itr:itr+1]) == False):
                            commentstr += encdStr[itr:itr+1]

                        else:
                            # print('commentstr ', commentstr)
                            pdf.set_font('Arial', '', 9)
                            pdf.write_html(commentstr)
                            y = pdf.get_y()
                            x = pdf.get_x()
                            # print('x before cell', pdf.get_x(), pdf.get_y())
                            pdf.set_xy(x, y)
                            pdf.set_font('ArialUnicodeMS', '', 9)
                            pdf.cell(2, 5, encdStr[itr:itr+1])
                            y = pdf.get_y()
                            x = pdf.get_x()
                            # print('x,y  ', pdf.get_x(), pdf.get_y())
                            # print('added here  2',commentstr)
                            commentstr = ""
                        itr += 1
                    if len(encdStr) == itr and commentstr != "":
                        # print('commentst at end is ', commentstr)
                        print('y at srite html is ',y)
                        pdf.set_font('Arial', '', 9)
                        pdf.write_html(commentstr)
                        y = pdf.get_y()
                        x = pdf.get_x()
                        # print('added here  3',commentstr,' y is ',y)
                        commentstr = ""
                        
                    # if(len(df) == 1 and exeSummTtlIdx == "1.1"):
                    #     pdf, new_y = addValFinfings(
                    #         pdf, exeSummTtlIdx, x, y, valFindingsLinkIdx)
                    #     exeSummTtl = ""
                    y = pdf.get_y()+5
                    pdf.set_xy(x, y)
                    pdf.multi_cell(0, 5, "", align='L')
    
            # if((exeSummTtl == str(row["title"]) and index==len(df)-1 and str(row["title"]) == "Executive Summary")):
            #     pdf, new_y = addValFinfings(
            #         pdf, exeSummTtlIdx, x, y, valFindingsLinkIdx)
            #     exeSummTtl = ""
            #     if new_y>y:
            #         y = new_y+10
            #         pdf.set_xy(x, y)
            #         print('new_y', y)
    return pdf, columnPageIdx, edaPageIdx, modelsPageIdx


def addTabletoRpt(pdf, tblName):
    tblFile = file_path + user_name+"_Tables.csv"
    if os.path.exists(tblFile):
        df_tbl = pd.read_csv(tblFile)
        print('tblName in addTabletoRpt ', tblName)
        df_filter = df_tbl.query("tableName =='" + tblName + "' ")
        # print('df_filter for ', tblName, ' is ', df_filter)
        tableType = df_filter["tableType"].values[0]
        del df_tbl, df_filter
        print('tableType is ',tableType)
        if(tableType == "DataTypenCnt"):
            exportDatatypenCnt(pdf)
        elif (tableType == "DataDesc"):
            exportviewNumData(tableType, pdf)
        elif (tableType == "DataMean"):
            exportviewNumData(tableType, pdf)
        elif (tableType == "DataMedian"):
            exportviewNumData(tableType, pdf)
        elif(tableType == "NumVarDIst"):
            exportdist_numevari_catvar(tblName, pdf)
        elif(tableType == "VIFData"):
            exportVIFData(pdf)
        elif(tableType == "TarvsCat"):
            exportCT(tblName, pdf)
    else:
        if(tblName == "ValFinding"):
            exportValFinfings(pdf)

        x = pdf.get_x()+5
        y = pdf.get_y()+5
        pdf.set_xy(x, y)
        pdf.multi_cell(0, 5, "", align='L')
    return pdf


def exportviewNumData(strType, pdf):
    from statsmodels import robust
    savefile_name = file_path + "csvfile_"+user_name + ".csv"
    df = pd.read_csv(savefile_name, na_values='?')
    num_cols = [c for i, c in enumerate(
        df.columns) if df.dtypes[i] not in [np.object]]
    x_numeric = pd.DataFrame(df, columns=num_cols)
    if(strType == "DataDesc"):
        desc = df.describe()
        x = pdf.get_x()
        y = pdf.get_y()
        pdf.set_xy(x, y)

        pdf.set_font("Arial", size=10)
        # pdf.set_text_color(255, 255, 255)

        pdf.set_xy(20, y)
        pdf.cell(40, 5, "", 1, fill=True)
        pdf.set_xy(60, y)
        pdf.cell(20, 5, "count", 1, fill=True)
        pdf.set_xy(80, y)
        pdf.cell(20, 5, "min", 1, fill=True)
        pdf.set_xy(100, y)
        pdf.cell(20, 5, "min", 1, fill=True)
        pdf.set_xy(120, y)
        pdf.cell(20, 5, "mean", 1, fill=True)
        pdf.set_xy(140, y)
        pdf.cell(20, 5, "std", 1, fill=True)
        pdf.set_xy(160, y)
        pdf.cell(20, 5, "25%", 1, fill=True)
        pdf.set_xy(180, y)
        pdf.cell(20, 5, "50%", 1, fill=True)
        pdf.set_xy(200, y)
        pdf.cell(20, 5, "75%", 1, fill=True)
        # Start from the first cell. Rows and columns are zero indexed.

        result = dict(df.dtypes)
        pdf.set_font("Arial", size=10)
        pdf.set_text_color(0.0, 0.0, 0.0)
        for recs, vals in dict(desc).items():
            if(y > 265):
                pdf.add_page()
                y = 10
            y += 5
            pdf.set_xy(20, y)
            pdf.cell(40, 5, recs, 1)
            pdf.set_xy(60, y)
            pdf.cell(20, 5, str(vals["count"]), 1)
            pdf.set_xy(80, y)
            pdf.cell(20, 5, str(vals["min"]), 1)
            pdf.set_xy(100, y)
            pdf.cell(20, 5, str(vals["min"]), 1)
            pdf.set_xy(120, y)
            pdf.cell(20, 5, str(vals["mean"]), 1)
            pdf.set_xy(140, y)
            pdf.cell(20, 5, str(vals["std"]), 1)
            pdf.set_xy(160, y)
            pdf.cell(20, 5, str(vals["25%"]), 1)
            pdf.set_xy(180, y)
            pdf.cell(20, 5, str(vals["50%"]), 1)
            pdf.set_xy(200, y)
            pdf.cell(20, 5, str(vals["75%"]), 1)
    elif(strType == "DataMean"):

        mean_ad = x_numeric.mad().round(decimals=3)
        # print('mean_ad is ', mean_ad)
        mean_adresult = mean_ad.to_json(orient='index')
        mean_adresult = json.loads(mean_adresult)
        y = pdf.get_y()
        x = pdf.get_x()
        pdf.set_xy(x, y)

        pdf.set_font("Arial", size=10)
        # pdf.set_text_color(255, 255, 255)

        pdf.set_xy(20, y)
        pdf.cell(80, 5, "Column", 1, fill=True)
        pdf.set_xy(100, y)
        pdf.cell(0, 5, "Value", 1, fill=True)
        for key in mean_adresult:
            value = mean_adresult[key]
            if(y > 265):
                pdf.add_page()
                y = 10
            y += 5
            pdf.set_xy(20, y)
            pdf.cell(80, 5, key, 1)
            pdf.set_xy(100, y)
            pdf.cell(0, 5, str(value), 1)
    elif(strType == "DataMedian"):
        median_ad = x_numeric.apply(robust.mad).round(decimals=3)
        # print(mean_ad)
        median_adresult = median_ad.to_json(orient='index')
        median_adresult = json.loads(median_adresult)
        y = pdf.get_y()
        x = pdf.get_x()
        pdf.set_xy(x, y)

        pdf.set_font("Arial", size=10)
        pdf.set_xy(20, y)
        pdf.cell(80, 5, "Column", 1, fill=True)
        pdf.set_xy(100, y)
        pdf.cell(0, 5, "Value", 1, fill=True)
        for key in median_adresult:
            value = median_adresult[key]
            if(y > 265):
                pdf.add_page()
                y = 10
            y += 5
            pdf.set_xy(20, y)
            pdf.cell(80, 5, key, 1)
            pdf.set_xy(100, y)
            pdf.cell(0, 5, str(value), 1)
    return pdf


def exportdist_numevari_catvar(tableName, pdf):
    tblFile = file_path + user_name+"_Tables.csv"
    arrdescData = ''

    if os.path.exists(tblFile):
        df = pd.read_csv(tblFile)
        dffilter = df.query("tableName== '"+tableName +
                            "' and tableType== 'NumVarDIst'")
        if(len(dffilter) > 0):
            var1 = dffilter["var1"].values[0]
            var2 = dffilter["var2"].values[0]
        savefile_withoutnull = file_path + file_name + ".csv"

        df = pd.read_csv(savefile_withoutnull, na_values='?')
        cat_cols = [c for i, c in enumerate(
            df.columns) if df.dtypes[i] in [np.object]]
        num_cols = [c for i, c in enumerate(
            df.columns) if df.dtypes[i] not in [np.object]]
        if(var1 != False):
            cat_var = var1  # cat_cols[i]
            num_var = var2  # num_cols[j]
        else:
            cat_var = cat_cols[0]  # cat_cols[i]
            num_var = num_cols[0]  # num_cols[j]
        dist_num_cat = df.groupby(cat_var)[num_var].describe()
        result = dist_num_cat.to_json(orient='index')
        result = json.loads(result)

        x = pdf.get_x()
        y = pdf.get_y()
        pdf.set_xy(x, y)

        pdf.set_font("Arial", size=10)
        pdf.set_xy(20, y)
        pdf.cell(40, 5, "", 1, fill=True)
        irow = 0
        for col in dist_num_cat.columns:
            pdf.set_xy((60+irow), y)
            pdf.cell(20, 5, str(col), 1, fill=True)
            irow += 20

        for key in result:
            value = result[key]
            if(y > 265):
                pdf.add_page()
                y = 10
            y += 5
            pdf.set_xy(20, y)
            pdf.cell(40, 5, key, 1)
            for key2 in value:
                value2 = value[key2]
                pdf.set_xy((60+irow), y)
                pdf.cell(20, 5, str(value2), 1)
                irow += 20
    return pdf


def exportVIFData(pdf):
    print('inside exportvif ')
    from statsmodels.stats.outliers_influence import variance_inflation_factor

    savefile_x_final = file_path + "csvfile_"+user_name + "_x_final.csv"
    if os.path.exists(savefile_x_final):
        savefile_x_scaled = savefile_x_final
    else:
        savefile_x_scaled = file_path + "csvfile_"+user_name + "_x_scaled.csv"
    savefile_x_keep = file_path + "csvfile_"+user_name + "_x_keep.csv"

    if os.path.exists(savefile_x_scaled):
        x_scaled_df = pd.read_csv(savefile_x_scaled, na_values='?')
        x_keep = pd.read_csv(savefile_x_keep, na_values='?')

        targetVarFile = file_path + "csvfile_"+user_name + "_targetVar.txt"
        file1 = open(targetVarFile, "r")  # write mode
        targetVar = file1.read()
        file1.close()
        x_keep = x_keep.drop(targetVar, axis=1)
        x_scaled_df = x_scaled_df.drop(targetVar, axis=1)

        vif_data = pd.DataFrame()
        vif_data["feature"] = x_scaled_df.columns

        # calculating VIF for each feature
        vif_data["VIF"] = [variance_inflation_factor(x_scaled_df.values, i)
                           for i in range(len(x_scaled_df.columns))]

        vif_data = vif_data.sort_values(
            "VIF", ascending=False)
        x = pdf.get_x()
        y = pdf.get_y()
        pdf.set_xy(x, y)

        pdf.set_font("Arial", size=10)
        pdf.set_xy(20, y)
        pdf.cell(80, 5, "Column", 1, fill=True)
        pdf.set_xy(100, y)
        pdf.cell(0, 5, "VIF", 1, fill=True)

        for index, row in vif_data.iterrows():
            if(y > 265):
                pdf.add_page()
                y = 10
            y += 5
            pdf.set_xy(20, y)
            pdf.cell(80, 5, row["feature"], 1)
            pdf.set_xy(100, y)
            pdf.cell(0, 5, str(row["VIF"]), 1)

    return pdf


def exportCT(tableName, pdf):
    tblFile = file_path + user_name+"_Tables.csv"
    arrdescData = ''

    if os.path.exists(tblFile):
        df = pd.read_csv(tblFile)
        dffilter = df.query("tableName== '"+tableName +
                            "' and tableType== 'TarvsCat'")
        if(len(dffilter) > 0):
            var1 = dffilter["var1"].values[0]
            var2 = dffilter["var2"].values[0]
    csvfile = file_path + file_name + ".csv"
    df = pd.read_csv(csvfile, na_values='?')
    dfCRossTab = pd.crosstab(df[var1], df[var2], rownames=[
                             var1], colnames=[var2])
    resultCrossTab = dfCRossTab.to_json(orient='index')
    resultCrossTab = json.loads(resultCrossTab)

    x = pdf.get_x()
    y = pdf.get_y()
    pdf.set_xy(x, y)

    pdf.set_font("Arial", size=10)
    pdf.set_xy(20, y)
    pdf.cell(40, 5, var1, 1, fill=True)
    irow = 0
    for key in resultCrossTab:
        value = resultCrossTab[key]
        for key2 in value:
            pdf.set_xy((60+irow), y)
            pdf.cell(20, 5, str(key2), 1, fill=True)
            irow += 20
        break
    y = y+5
    pdf.set_xy(20, y)
    pdf.cell(0, 5, var2, 1, fill=True)
    for key in resultCrossTab:
        irow = 0
        value = resultCrossTab[key]
        x = pdf.get_x()
        y = pdf.get_y()
        if(y > 265):
            pdf.add_page()
            y = 10
        y += 5
        pdf.set_xy(x, y)
        pdf.set_xy(20, y)
        pdf.cell(40, 5, key, 1, fill=True)

        for key2 in value:
            val1 = value[key2]
            pdf.set_xy((60+irow), y)
            pdf.cell(20, 5, str(val1), 1, fill=True)
            irow += 20
    return pdf

def exportValFinfings(pdf):
    print('export val findings ')
    validationFindings = file_path + file_name + "_validationFindings.csv"
    cellHeight = 0
    x = pdf.get_x()
    y = pdf.get_y()
    pdf.set_xy(x, y)
    if os.path.exists(validationFindings):
        df = pd.read_csv(validationFindings)
        df = df.sort_values(by="reqId", ascending=True) 
        if(y > 150):
            pdf.add_page('P')
            y = pdf.get_y()
        pdf.set_xy(x, y)
        pdf.set_font("Arial", "B", size=9)
        pdf.set_text_color(0.0, 0.0, 0.0) 
        pdf.set_font("Arial",  size=9)
        pdf.set_fill_color(211, 211, 211)
        y = pdf.get_y()+10
        pdf.set_xy(20, y)
        pdf.cell(20, 5, "Finding ID#",
                 1, fill=True, align='L')

        pdf.set_xy(40, y)
        pdf.cell(35, 5, "Assessment Area",
                 1, fill=True, align='L')
        pdf.set_xy(75, y)
        pdf.cell(85, 5, "Description",
                 1, fill=True, align='L')
        pdf.set_xy(160, y)
        pdf.cell(30, 5, "Risk Level", 1,
                 fill=True, align='L')
        y = pdf.get_y()+5
        initheight = y
        
        pdf.set_font("Arial",  size=8)
        for index, row in df.iterrows():
            pdf.set_xy(20, y)
            pdf.multi_cell(20, 4, str(
                row["findingsId"]), 0, fill=False)

            pdf.set_xy(40, y)
            pdf.multi_cell(35, 4, str(row["Assessment"]).encode(
                'latin-1', 'replace').decode('latin-1'), 0, fill=False)

            pdf.set_xy(75, y)
            pdf.multi_cell(85, 4, str(row["Desc"]).encode(
                'latin-1', 'replace').decode('latin-1'), 0, fill=False)
            if(pdf.get_y()>cellHeight):
                cellHeight = pdf.get_y()
            pdf.set_xy(160, y)
            pdf.cell(30, 4, str(
                row["Risk_Level"]), 0, fill=False)
            pdf.rect(20, y, 20, cellHeight-y)
            pdf.rect(40, y, 35, cellHeight-y)
            pdf.rect(75, y, 85, cellHeight-y)
            pdf.rect(160, y, 30, cellHeight-y)
            y = cellHeight
            if(len(str(row["Response"])) > 0 and str(row["Response"]) != "-"):
                pdf.set_xy(20, y)
                pdf.multi_cell(170, 4, str(row["Response"]).encode(
                    'latin-1', 'replace').decode('latin-1'), 1, fill=False)
                y = pdf.get_y()
            # print(' y is ', y, ' cellHeight ', cellHeight)

        if(str(df["Comment"].values[0]) != "-" and str(df["Comment"].values[0]) != "nan"):
            y = y+10
            pdf.set_xy(20, y)
            pdf.multi_cell(170, 4, str(df["Comment"].values[0]).encode(
                'latin-1', 'replace').decode('latin-1'), 0, fill=False)
            y = pdf.get_y()

         
    return pdf

def checkSymbol(charStr):
    try:
        a = charStr.encode('latin-1')
        return False
    except Exception as e:
        return True


def addSummarynComments(pdf, document, columnPageIdx, edaPageIdx, modelsPageIdx):

    # SummaryDataFiles = file_path + file_name + "_SummaryData.csv"
    temp_report_file_path = os.path.join(BASE_DIR, plot_dir_view)
    temp_report_file_name = "temp_report_"+user_name
    temp_cnfrmsrcFiles = temp_report_file_path + temp_report_file_name + ".csv"

    report_file_path = os.path.join(BASE_DIR, plot_dir_view)

    report_file_name = "report_"+user_name
    cnfrmsrcFiles = report_file_path + report_file_name + ".csv"
    linkArr = []
    pageNoArr = []
    valFindingsLinkIdx = -1
    if os.path.exists(temp_cnfrmsrcFiles):
        if os.path.exists(cnfrmsrcFiles):
            os.remove(cnfrmsrcFiles)
        os.rename(temp_cnfrmsrcFiles, cnfrmsrcFiles)

    if os.path.exists(cnfrmsrcFiles):
        df_old = pd.read_csv(cnfrmsrcFiles, encoding='utf-8')
        df = df_old.sort_values(
            by=['titleIdx', 'subtitleIdx', 'subsubtitleIdx', 'reqID'], ascending=True)
        # print('sorted df ', df)
        pdf, linkArr, valFindingsLinkIdx, pageNoArr, columnPageIdx, edaPageIdx, modelsPageIdx = addTableofContents(
            pdf, df, document, valFindingsLinkIdx, columnPageIdx, edaPageIdx, modelsPageIdx)
        # pdf.set_left_margin(32)
        pdf.add_page('P')
        document.add_page_break()
        x, y = 10, pdf.get_y()
        exeSummTtlIdx = ""
        exeSummTtl = ""
        titleAdded = ""
        subtitleAdded = ""
        subsubtitleAdded = ""
        for index, row in df.iterrows():
            if(str(row["title"]) == "Executive Summary"):
                exeSummTtl = str(row["title"])
                dffilter = df.loc[df.title == exeSummTtl]
                maxIdx = float(dffilter["subtitleIdx"].max())

                maxIdx = maxIdx+0.1
                exeSummTtlIdx = str(round(maxIdx, 2))
                if(len(dffilter) == 1):
                    exeSummTtlIdx = "1.1"
                del(dffilter)
            if((exeSummTtl != str(row["title"]) and exeSummTtl != "") or exeSummTtlIdx == "1.1"):
                pdf, new_y = addValFinfings(
                    pdf, exeSummTtlIdx, x, y, valFindingsLinkIdx)
                exeSummTtl = ""
                y = new_y+10

            if((str(row["subtitleIdx"]) == '0' or str(row["subtitleIdx"]) == '0.0') and str(row["subsubtitleIdx"]) == '0'):
                if(row["title"] != titleAdded):
                    if(y > 150):
                        pdf.add_page('P')
                        y = pdf.get_y()
                    pdf.set_xy(x, y)
                    pdf.set_font("Arial", "B", size=9)
                    pdf.set_text_color(0.0, 0.0, 0.0)
                    pdf.set_link(linkArr[index], y, pdf.page_no())
                    strsuprscript = str(row["titleIdx"]) + \
                        " " + str(row["title"])
                    pdf.multi_cell(200, 20, strsuprscript, align='L')
                    pageNoArr[index] = pdf.page_no()
                    titleAdded = str(row["title"])

                paraTitle = document.add_paragraph()
                paragraph_format = paraTitle.paragraph_format
                paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = paraTitle.add_run(str(row["title"]).encode(
                    'latin-1', 'replace').decode('latin-1'))
                run.bold = True
                font = run.font
                font.name = 'Arial'
                font.size = Pt(9)
                print('to be uncommented')
            elif((row["subtitle"] != subtitleAdded) and (str(row["subtitleIdx"]) != '0' or str(row["subtitleIdx"]) != '0.0') and str(row["subsubtitleIdx"]) == '0'):
                y = pdf.get_y()
                pdf.set_xy(x, y)
                pdf.set_font("Arial",  "B", size=9)
                pdf.set_text_color(0.0, 0.0, 0.0)
                pdf.set_link(linkArr[index], y, pdf.page_no())
                pdf.multi_cell(0, 10, str(row["subtitleIdx"]) + " " + str(row["subtitle"]).encode(
                    'latin-1', 'replace').decode('latin-1'), align='L')
                print('str(row["subtitle"]) is ', str(row["subtitle"]))
                subtitleAdded = str(row["subtitle"])
                parasubTitle = document.add_paragraph()
                paragraph_format = parasubTitle.paragraph_format
                paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = parasubTitle.add_run(str(row["subtitle"]).encode(
                    'latin-1', 'replace').decode('latin-1'))
                run.bold = True
                font = run.font
                font.name = 'Arial'
                font.size = Pt(9)
                print('to be uncommented')
            elif((row["subsubtitle"] != subsubtitleAdded) and str(row["subsubtitleIdx"]) != '0'):
                y = pdf.get_y()
                pdf.set_xy(x, y)
                pdf.set_font("Arial",  "B", size=9)
                pdf.set_text_color(0.0, 0.0, 0.0)
                pdf.multi_cell(0, 10, "     "+str(row["subsubtitleIdx"]) + " " + str(row["subsubtitle"]).encode(
                    'latin-1', 'replace').decode('latin-1'), align='L')
                subsubtitleAdded = str(row["subsubtitle"])
                parasubTitle = document.add_paragraph()
                paragraph_format = parasubTitle.paragraph_format
                paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = parasubTitle.add_run(str(row["subsubtitle"]).encode(
                    'latin-1', 'replace').decode('latin-1'))
                run.bold = True
                font = run.font
                font.name = 'Arial'
                font.size = Pt(9)
                print('to be uncommented')
            if row["section"] == "Comment":
                if(str(row["comment"]) != "nan"):
                    y = pdf.get_y()
                    pdf.set_xy(x, y)
                    pdf.add_font("NotoSans-Regular", "", font_files +
                                 "NotoSans-Regular.ttf", uni=True)
                    pdf.set_font('NotoSans-Regular', '', 9)
                    # pdf.set_font("Arial",  size=9)
                    pdf.set_text_color(0.0, 0.0, 0.0)
                    pdf.set_left_margin(10)
                    pdf.set_link(linkArr[index], y, pdf.page_no())
                    print('comment us ', str(row["comment"]).replace("\t", "\u0020\u0020\u0020\u0020").replace(str(row["title"])+"\n", "").replace(
                        str(row["subtitle"])+"\n", "").replace(str(row["subsubtitle"])+"\n", "").encode(
                        'utf-8', 'replace').decode('utf-8'))
                    pdf.multi_cell(0, 5, str(row["comment"]).replace("\t", "\u0020\u0020\u0020\u0020").replace(str(row["title"])+"\n", "").replace(
                        str(row["subtitle"])+"\n", "").replace(str(row["subsubtitle"])+"\n", "").encode(
                        'utf-8', 'replace').decode('utf-8'),  align='L')
                    # print('coment is ', (str(row["comment"])))
                    # pdf.write_html(str(row["comment"]).replace(
                    #     'src="/static', 'src="static'))
                    y = pdf.get_y()
                    pdf.set_xy(x, y)
                    pdf.set_font("Arial",  size=8)
                    pdf.set_text_color(0.0, 0.0, 0.0)
                    pdf.multi_cell(0, 5, "", align='L')

                    # paracomments = document.add_paragraph()
                    # paragraph_format = paracomments.paragraph_format
                    # paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # run = paracomments.add_run(str(row["comment"]).encode(
                    #     'latin-1', 'replace').decode('latin-1'))
                    # font = run.font
                    # font.name = 'Arial'
                    # font.size = Pt(9)
            else:
                imgW = (700/4)
                imgW = imgW * float(row["imgWidth"])
                if(row["imgAlign"] == "right"):
                    # print('right corener ', (700/4)*float(row["imgWidth"]))
                    if(175-((700/4)*float(row["imgWidth"])) > 0):
                        pdf.set_xy(
                            175-((700/4)*float(row["imgWidth"])), pdf.get_y())
                    else:
                        pdf.set_xy(10, pdf.get_y())
                elif(row["imgAlign"] == "center"):
                    if(((700/4)*float(row["imgWidth"])) < 175):
                        pdf.set_xy(
                            (175-((700/4)*float(row["imgWidth"])))/2, pdf.get_y())
                    else:
                        pdf.set_xy(10, pdf.get_y())
                else:
                    pdf.set_xy(10, pdf.get_y())

                if(pdf.get_y() > 150):
                    pdf.add_page('P')
                    y = pdf.get_y()
                    pdf.set_xy(10, pdf.get_y())

                pdf.image(os.path.join(
                    BASE_DIR, row["img"]),  link='', type='',  w=(700/4)*float(row["imgWidth"]), h=(450/4)*float(row["imgHeight"]))
                # y = pdf.get_y()+((450/4)*float(row["imgHeight"]))
                # pdf.set_xy(10, y)
                y = pdf.get_y()
                print('to be uncommented')
    return pdf, columnPageIdx, edaPageIdx, modelsPageIdx


def addTableofContents(pdf, df, document, valFindingsLinkIdx, columnPageIdx, edaPageIdx, modelsPageIdx):
    pdf.set_left_margin(32)
    pdf.add_page()
    document.add_page_break()
    x, y = 10, pdf.get_y()
    pdf.set_xy(x, y)
    pdf.set_font("Arial", size=9)
    pdf.set_text_color(0.0, 0.0, 0.0)
    pdf.cell(0, 5, "Table of Contents", align='L')
    linkArr = [0] * len(df)
    pageNoArr = [0] * len(df)
    exeSummTtlIdx = ""
    exeSummTtl = ""
    titleAdded = 0
    # print('df', df["title"], df["titleIdx"], df["subtitle"])

    # print('df_new is', df)
    for index, row in df.iterrows():
        print('idxs ', index,str(row["title"]), str(
            row["title"]) == "Executive Summary")
        if((str(row["subtitleIdx"]) ==
                '0' or str(row["subtitleIdx"]) ==
                '0.0') and str(row["subsubtitleIdx"]) == '0'):
            if(str(row["title"]) == "Executive Summary"):
                exeSummTtl = str(row["title"])
                dffilter = df.loc[df.title == exeSummTtl]
                # print('dffilter ', dffilter)
                maxIdx = dffilter["subtitleIdx"].max()
                # print('maxIdx ', maxIdx)
                maxIdx = maxIdx+0.1
                # print('maxIdx + 0.1 ', round(maxIdx, 2))
                exeSummTtlIdx = str(round(maxIdx, 2))
                if(len(dffilter) == 1):
                    exeSummTtlIdx = "1.1"
                del(dffilter)
                print('exeSummTtlIdx ',exeSummTtlIdx)
            if((exeSummTtl != str(row["title"]) and exeSummTtl != "") or (exeSummTtl == str(row["title"]) and index==len(df)-1) ):
                # exeSummTtl = ""
                # y = pdf.get_y()+7
                # pdf.set_xy(x, y)
                # pdf.set_font("Arial", size=9)
                # pdf.set_text_color(0.0, 0.0, 0.0)
                # to_page = pdf.add_link()
                # valFindingsLinkIdx = to_page
                # pdf.cell(0, 5, "    "+exeSummTtlIdx +
                #          " Validation Findings", align='L', link=to_page)
                print('uncommented s')
            if(row["title"] != titleAdded):
                y = pdf.get_y()+7
                pdf.set_xy(x, y)
                pdf.set_font("Arial", size=9)
                pdf.set_text_color(0.0, 0.0, 0.0)
                to_page = pdf.add_link()
                linkArr[index] = to_page
                pdf.cell(0, 5, str(row["titleIdx"]) + " "+str(row["title"]).encode(
                    'utf-8', 'replace').decode('utf-8'), align='L', link=to_page)
                paraTitle = document.add_paragraph()
                paragraph_format = paraTitle.paragraph_format
                paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = paraTitle.add_run(str(row["title"]).encode(
                    'utf-8', 'replace').decode('utf-8'))
                run.bold = True
                font = run.font
                font.name = 'Arial'
                font.size = Pt(9)
                titleAdded = str(row["title"])

                

        elif(str(row["subtitleIdx"]) != '0.0' and str(row["subtitleIdx"]) != '0' and (str(row["subsubtitleIdx"]) == '0.0' or str(row["subsubtitleIdx"]) == '0') and str(row["subtitle"]) !=
                'NaN'):
            y = pdf.get_y()+7
            pdf.set_xy(x, y)
            pdf.set_font("Arial",  size=9)
            pdf.set_text_color(0.0, 0.0, 0.0)
            to_page = pdf.add_link()
            linkArr[index] = to_page
            pdf.cell(0, 5, "    "+str(row["subtitleIdx"])+" "+str(row["subtitle"]).encode(
                'utf-8', 'replace').decode('utf-8'), align='L', link=to_page)
            parasubTitle = document.add_paragraph()
            paragraph_format = parasubTitle.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = parasubTitle.add_run(str(row["subtitle"]).encode(
                'utf-8', 'replace').decode('utf-8'))
            run.bold = True
            font = run.font
            font.name = 'Arial'
            font.size = Pt(9)
 
            # if((exeSummTtl == str(row["title"]) and index==len(df)-1)):
            #     exeSummTtl = ""
            #     y = pdf.get_y()+7
            #     pdf.set_xy(x, y)
            #     pdf.set_font("Arial", size=9)
            #     pdf.set_text_color(0.0, 0.0, 0.0)
            #     to_page = pdf.add_link()
            #     valFindingsLinkIdx = to_page
            #     pdf.cell(0, 5, "    "+exeSummTtlIdx +
            #              " Validation Findings", align='L', link=to_page)

    # maxTtlIdx = df['titleIdx'].max()+1
    # y = pdf.get_y()+7
    # pdf.set_xy(x, y)
    # pdf.set_font("Arial", size=9)
    # pdf.set_text_color(0.0, 0.0, 0.0)
    # to_page = pdf.add_link()
    # columnPageIdx = to_page
    # pdf.cell(0, 5, str(maxTtlIdx)+" Variables by type and count",
    #          align='L', link=to_page)

    # maxTtlIdx = maxTtlIdx+1
    # y = pdf.get_y()+7
    # pdf.set_xy(x, y)
    # pdf.set_font("Arial", size=9)
    # pdf.set_text_color(0.0, 0.0, 0.0)
    # to_page = pdf.add_link()
    # edaPageIdx = to_page
    # pdf.cell(0, 5, str(maxTtlIdx)+" Exploratory Data Analysis",
    #          align='L', link=to_page)

    # maxTtlIdx = maxTtlIdx+1
    # y = pdf.get_y()+7
    # pdf.set_xy(x, y)
    # pdf.set_font("Arial", size=9)
    # pdf.set_text_color(0.0, 0.0, 0.0)
    # to_page = pdf.add_link()
    # modelsPageIdx = to_page
    # pdf.cell(0, 5, str(maxTtlIdx)+" Modeling",
    #          align='L', link=to_page)
    return pdf, linkArr, valFindingsLinkIdx, pageNoArr, columnPageIdx, edaPageIdx, modelsPageIdx


def addReferences(pdf):
    pdf.set_left_margin(32)
    pdf.add_page()
    x, y = 10, pdf.get_y()
    pdf.set_xy(x, y)
    pdf.set_font("Arial", size=9)
    pdf.set_text_color(0.0, 0.0, 0.0)
    pdf.cell(0, 5, "References", align='L')
    contactFile = file_path + user_name + "_RptRef.csv"
    y = y+20
    if os.path.exists(contactFile):
        df = pd.read_csv(contactFile)
        for index, row in df.iterrows():
            pdf.set_xy(20, y)
        pdf.cell(20, 5, str(row["Srno"])+" "+row["policy"],
                 0, fill=False, align='L')

        pdf.set_xy(80, y)
        pdf.cell(50, 5, row["reference"],
                 0, fill=False, align='L')
        y = pdf.get_y()+5
    return pdf


def addValFinfings(pdf, exeSummTtlIdx, x, y, valFindingsLinkIdx):
    validationFindings = file_path + file_name + "_validationFindings.csv"
    cellHeight = 0
    if os.path.exists(validationFindings):
        df = pd.read_csv(validationFindings)
        df = df.sort_values(by="reqId", ascending=True)

        print(' y on start findins', y)
        if(y > 150):
            pdf.add_page('P')
            y = pdf.get_y()
        pdf.set_xy(x, y)
        pdf.set_font("Arial", "B", size=9)
        pdf.set_text_color(0.0, 0.0, 0.0)
        pdf.set_link(valFindingsLinkIdx, y, pdf.page_no())
        pdf.cell(
            0, 5, exeSummTtlIdx + " validation findings", align='L')
        pdf.set_font("Arial",  size=9)
        pdf.set_fill_color(211, 211, 211)
        y = pdf.get_y()+10
        pdf.set_xy(20, y)
        pdf.cell(20, 5, "Finding ID#",
                 1, fill=True, align='L')

        pdf.set_xy(40, y)
        pdf.cell(35, 5, "Assessment Area",
                 1, fill=True, align='L')
        pdf.set_xy(75, y)
        pdf.cell(85, 5, "Description",
                 1, fill=True, align='L')
        pdf.set_xy(160, y)
        pdf.cell(30, 5, "Risk Level", 1,
                 fill=True, align='L')
        y = pdf.get_y()+5
        initheight = y
        
        pdf.set_font("Arial",  size=8)
        for index, row in df.iterrows():
            pdf.set_xy(20, y)
            pdf.multi_cell(20, 4, str(
                row["findingsId"]), 0, fill=False)

            pdf.set_xy(40, y)
            pdf.multi_cell(35, 4, str(row["Assessment"]).encode(
                'latin-1', 'replace').decode('latin-1'), 0, fill=False)

            pdf.set_xy(75, y)
            pdf.multi_cell(85, 4, str(row["Desc"]).encode(
                'latin-1', 'replace').decode('latin-1'), 0, fill=False)
            if(pdf.get_y()>cellHeight):
                cellHeight = pdf.get_y()
            pdf.set_xy(160, y)
            pdf.cell(30, 4, str(
                row["Risk_Level"]), 0, fill=False)
            pdf.rect(20, y, 20, cellHeight-y)
            pdf.rect(40, y, 35, cellHeight-y)
            pdf.rect(75, y, 85, cellHeight-y)
            pdf.rect(160, y, 30, cellHeight-y)
            y = cellHeight
            if(len(str(row["Response"])) > 0 and str(row["Response"]) != "-"):
                pdf.set_xy(20, y)
                pdf.multi_cell(170, 4, str(row["Response"]).encode(
                    'latin-1', 'replace').decode('latin-1'), 1, fill=False)
                y = pdf.get_y()
            # print(' y is ', y, ' cellHeight ', cellHeight)

        if(str(df["Comment"].values[0]) != "-" and str(df["Comment"].values[0]) != "nan"):
            y = y+10
            pdf.set_xy(20, y)
            pdf.multi_cell(170, 4, str(df["Comment"].values[0]).encode(
                'latin-1', 'replace').decode('latin-1'), 0, fill=False)
            y = pdf.get_y()

        if cellHeight < y:
            cellHeight = y
    return pdf, cellHeight


def addDocumentationComments(pdf, document):
    x, y = 20, 10
    DocumentationDataFiles = file_path + file_name + "_DocumentationData.csv"
    pdf.set_xy(x, y)
    pdf.set_font("Arial",  size=22)
    pdf.set_text_color(0.0, 0.0, 0.0)
    pdf.multi_cell(0, 10, "Appendix 1 : Documentation", align='L')
    y = 30
    if os.path.exists(DocumentationDataFiles):
        df = pd.read_csv(DocumentationDataFiles)
        # '', ''
        table = document.add_table(rows=1, cols=2)

        for index, row in df.iterrows():
            pdf.set_font("Arial", size=9)
            y += 5
            pdf.set_xy(20, y)
            pdf.cell(0, 10, str(row["doc"]))
            pdf.set_xy(100, y)
            pdf.cell(0, 10, str(row["doc_file"]))

            _cells = table.add_row().cells
            _cells[0].text = str(row["doc"])
            _cells[1].text = str(row["doc_file"])

    return pdf


def addDataQuality(pdf, document):
    x, y = 20, 10
    cnfrmsrc_file_path = os.path.join(BASE_DIR, 'static/cnfrmsrc_files/')
    cnfrmsrc_file_name = "cnfrmsrc_"+user_name
    cnfrmsrcFiles = cnfrmsrc_file_path + cnfrmsrc_file_name + ".csv"
    pdf.set_xy(x, y)
    pdf.set_font("Arial",  size=22)
    pdf.set_text_color(0.0, 0.0, 0.0)
    pdf.multi_cell(
        0, 10, "Appendix 2 : Data Quality", align='L')

    if os.path.exists(cnfrmsrcFiles):
        df = pd.read_csv(cnfrmsrcFiles)

        pdf.set_font("Arial",  size=9)
        pdf.set_fill_color(211, 211, 211)
        y = pdf.get_y()+10
        pdf.set_xy(20, y)
        pdf.cell(40, 5, "Column Name",
                 1, fill=True, align='L')

        pdf.set_xy(60, y)
        pdf.cell(45, 5, "Datasource",
                 1, fill=True, align='L')
        pdf.set_xy(105, y)
        pdf.cell(85, 5, "Data Quality",
                 1, fill=True, align='L')

        y = pdf.get_y()+5
        initheight = y
        cellHeight = 0
        pdf.set_font("Arial",  size=8)
        for index, row in df.iterrows():
            pdf.set_xy(20, y)
            pdf.multi_cell(40, 4, str(
                row["colName"]), 0, fill=False)

            pdf.set_xy(60, y)
            pdf.multi_cell(45, 4, str(row["srcName"]).encode(
                'latin-1', 'replace').decode('latin-1'), 0, fill=False)

            pdf.set_xy(105, y)
            pdf.multi_cell(85, 4, str(row["dataQuality"]).encode(
                'latin-1', 'replace').decode('latin-1'), 0, fill=False)
            cellHeight = pdf.get_y()

            pdf.rect(20, y, 40, cellHeight-y)
            pdf.rect(60, y, 45, cellHeight-y)
            pdf.rect(105, y, 85, cellHeight-y)
            y = cellHeight

    return pdf


def addImplCtrl(pdf, document):
    x, y = 20, 10
    cnfrmsrc_file_path = os.path.join(BASE_DIR, 'static/cnfrmsrc_files/')
    cnfrmsrc_file_name = "ImpCtrl_"+user_name
    cnfrmsrcFiles = cnfrmsrc_file_path + cnfrmsrc_file_name + ".csv"
    # pdf.set_xy(x, y)
    # pdf.set_font("Arial",  size=22)
    # pdf.set_text_color(0.0, 0.0, 0.0)
    # pdf.multi_cell(0, 10, "Documentation", align='C')
    y = 30
    if os.path.exists(cnfrmsrcFiles):
        df = pd.read_csv(cnfrmsrcFiles)
        # '', ''
        table = document.add_table(rows=1, cols=2)
        pdf.set_xy(x, y)
        pdf.set_font("Arial",  size=9)
        pdf.set_text_color(0.0, 0.0, 0.0)
        pdf.multi_cell(0, 10, df["reportComment"].values[0], align='L')
        for index, row in df.iterrows():
            pdf.set_font("Arial", size=9)
            y += 5
            pdf.set_xy(20, y)
            pdf.cell(0, 10, str(row["section"]))
            pdf.set_xy(100, y)
            pdf.cell(0, 10, str(row["reqRessepon"]))

            _cells = table.add_row().cells
            _cells[0].text = str(row["section"])
            _cells[1].text = str(row["reqRessepon"])
            _cells[0].width = Inches(2.0)

    return pdf


def exportgraphImgPdf(pdf, document, graph, header, comments=""):
    if(str(comments) == "nan"):
        comments = ""
    if(len(comments) > 0):
        x, y = 10, 10
    else:
        x, y = 10, 50
    # set style and size of font
    # that you want in the pdf
    pdf.set_font("Arial", size=15)
    pdf.set_xy(x, y)
    pdf.set_text_color(0.0, 0.0, 0.0)
    pdf.multi_cell(0, 10, header, align='C')

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    # paragraph_format.line_spacing = Pt(180)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(header)
    font = run.font
    font.name = 'Arial'
    font.size = Pt(15)

    # if(len(comments) > 0):
    y = pdf.get_y()+5.0
    pdf.set_font("Arial", size=10)
    pdf.set_xy(x, y)
    pdf.set_text_color(0.0, 0.0, 0.0)
    pdf.multi_cell(0, 5, comments.encode(
        'latin-1', 'replace').decode('latin-1'), align='L')

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(comments)
    font = run.font
    font.name = 'Arial'
    font.size = Pt(10)

    y = pdf.get_y()+20.0
    pdf.set_xy(20, y)
    pdf.image(graph,  link='', type='', w=700/4, h=450/4)

    document.add_picture(graph, width=Inches(6.0), height=Inches(4.5))

    return pdf


def Summary(request):
    return render(request, 'Summary.html', {'pdfFile': "/static/media/modelReport.pdf", })


def modelValReq(request):
    try:
        destination_path = os.path.join(BASE_DIR, 'static/document_files/')
        MCD_path = os.path.join(BASE_DIR, 'static/modelCode/')
        reqId = 0
        resultDocumentation={"noData":"noData"}
        if request.method == 'POST':
            # for f in os.listdir(file_path):
            #     os.remove(os.path.join(file_path, f))
            txt_MD = request.POST['txt_MD']
            rbTarget_MD = request.POST.get('rbTarget_MD')
            reqId = saveModelInfo(request.POST.get('txtModelNm'),
                                  request.POST.get('txtModelVersion'),
                                  request.POST.get('txtModelID'),
                                  request.POST.get('txtModelDesc'),
                                  request.POST.get('txtModelDate'),
                                  request.POST.get('txtModelDev'),
                                  request.POST.get('txtModelDept'),
                                  request.POST.get('txtModelOwner'),
                                  request.POST.get('Model_Type'))

            saveModelRiskInfo('Materiality', request.POST.get(
                'optMateriality'), reqId)
            saveModelRiskInfo(
                'Reliance', request.POST.get('optReliance'), reqId)
            saveModelRiskInfo('Intrinsic_Risk',
                              request.POST.get('optIntrinsic_Risk'), reqId)
            if rbTarget_MD == "file":
                fl_MD = request.FILES.get('fl_MD', 'none')
                if fl_MD != 'none':
                    fs = FileSystemStorage()
                    savefile_name = destination_path + fl_MD.name
                    if os.path.exists(savefile_name):
                        os.remove(savefile_name)
                    fs.save(savefile_name, fl_MD)
                    saveModelDocuments('Model Development',
                                       'file', fl_MD.name, reqId, request.POST['txt_MD_Comment'])
            elif rbTarget_MD == "link":
                saveModelDocuments('Model Development', 'link',
                                   txt_MD, reqId, request.POST['txt_MD_Comment'])

            if request.POST.get('rbTarget_UM') == "file":
                fl_UM = request.FILES.get('fl_UM', 'none')
                if fl_UM != 'none':
                    fs = FileSystemStorage()
                    savefile_name = destination_path + fl_UM.name
                    if os.path.exists(savefile_name):
                        os.remove(savefile_name)
                    fs.save(savefile_name, fl_UM)
                    saveModelDocuments(
                        'User Manual', 'file', fl_UM.name, reqId, request.POST['txt_UM_Comment'])
            elif request.POST.get('rbTarget_UM') == "link":
                saveModelDocuments('User Manual', 'link',
                                   request.POST['txt_UM'], reqId, request.POST['txt_UM_Comment'])

            if request.POST.get('rbTarget_MDT') == "file":
                fl_MDT = request.FILES.get('fl_MDT', 'none')
                if(fl_MDT != 'none'):
                    fs = FileSystemStorage()
                    savefile_name = destination_path + fl_MDT.name
                    if os.path.exists(savefile_name):
                        os.remove(savefile_name)
                    fs.save(savefile_name, fl_MDT)
                    saveModelDocuments('Model Data', 'file',
                                       fl_MDT.name, reqId, request.POST['txt_MDT_Comment'])
            elif request.POST.get('rbTarget_MDT') == "link":
                saveModelDocuments('Model Data', 'link',
                                   request.POST['txt_MDT'], reqId, request.POST['txt_MDT_Comment'])

            if request.POST.get('rbTarget_MCD') == "file":
                fl_MCD = request.FILES.get('fl_MCD', 'none')
                if fl_MCD != 'none':
                    fs = FileSystemStorage()
                    savefile_name = MCD_path + fl_MCD.name
                    if os.path.exists(savefile_name):
                        os.remove(savefile_name)
                    fs.save(savefile_name, fl_MCD)
                    saveModelDocuments('Model Code', 'file',
                                       fl_MCD.name, reqId, request.POST['txt_MCD_Comment'])
            elif request.POST.get('rbTarget_MCD') == "link":
                saveModelDocuments('Model Code', 'link',
                                   request.POST['txt_MCD'], reqId, request.POST['txt_MCD_Comment'])

            if request.POST.get('rbTarget_UT') == "file":
                fl_UT = request.FILES.get('fl_UT', 'none')
                if fl_UT != 'none':
                    fs = FileSystemStorage()
                    savefile_name = destination_path + fl_UT.name
                    if os.path.exists(savefile_name):
                        os.remove(savefile_name)
                    fs.save(savefile_name, fl_UT)
                    saveModelDocuments(
                        'User Acceptance Testing', 'file', fl_UT.name, reqId, request.POST['txt_UT_Comment'])
            elif request.POST.get('rbTarget_UT') == "link":
                saveModelDocuments('User Acceptance Testing',
                                   'link', request.POST['txt_UT'], reqId, request.POST['txt_UT_Comment'])

            if request.POST.get('rbTarget_TM') == "file":
                fl_TM = request.FILES.get('fl_TM', 'none')
                if fl_TM != 'none':
                    fs = FileSystemStorage()
                    savefile_name = destination_path + fl_TM.name
                    if os.path.exists(savefile_name):
                        os.remove(savefile_name)
                    fs.save(savefile_name, fl_TM)
                    saveModelDocuments('Technical Manual',
                                       'file', fl_TM.name, reqId, request.POST['txt_TM_Comment'])
            elif request.POST.get('rbTarget_TM') == "link":
                saveModelDocuments('Technical Manual',
                                   'link', request.POST['txt_TM'], reqId, request.POST['txt_TM_Comment'])

            if request.POST.get('rbTarget_OD') == "file":
                fl_OD = request.FILES.get('fl_OD', 'none')
                if fl_OD != 'none':
                    fs = FileSystemStorage()
                    savefile_name = destination_path + fl_OD.name
                    if os.path.exists(savefile_name):
                        os.remove(savefile_name)
                    fs.save(savefile_name, fl_OD)
                    saveModelDocuments('Onboarding Documents',
                                       'file', fl_OD.name, reqId, request.POST['txt_OD_Comment'])
            elif request.POST.get('rbTarget_TM') == "link":
                saveModelDocuments('Onboarding Documents',
                                   'link', request.POST['txt_OD'], reqId, request.POST['txt_OD_Comment'])
         
       

        DocumentationData = file_path + file_name + "_ModelInfo.csv"

        today = date.today().strftime("%m/%d/%Y")
        result = '[{"ModelNm": "", "ModelVersion":"", "ModelID":"", "ModelDesc":"", "ModelDate":"' + \
            today+'", "ModelDev":"", "ModelDept":"", "ModelOwner":"", "Model_Type": "Vendor"}]'
        result = json.loads(result)
        reqId = 0
        if os.path.exists(DocumentationData):
            df = pd.read_csv(DocumentationData)
            reqId = df["reqId"].max()
            dffilter = df.query("reqId== "+str(reqId))
            result = dffilter.to_json(orient="records")
            result = json.loads(result)
            del df, dffilter

        DocumentationData = file_path + file_name + "_ModelRiskInfo.csv"
        resultRiskInfo = []
        if os.path.exists(DocumentationData):
            df = pd.read_csv(DocumentationData)
            dffilter = df.query("reqId== "+str(reqId))
            resultRiskInfo = dffilter.to_json(orient="records")
            resultRiskInfo = json.dumps(resultRiskInfo)
            del df, dffilter
            print('resultRiskInfo ', resultRiskInfo)

        DocumentationData = file_path + file_name + "_DocumentationData.csv"
        if os.path.exists(DocumentationData):
            dffilter = pd.read_csv(DocumentationData)
            # dffilter = df.query("reqId== "+str(reqId))
            resultDocumentation = dffilter.to_json(orient="records")
            resultDocumentation = json.loads(resultDocumentation)
            print('resultDocumentation ', resultDocumentation)
        _isDisabled="disabled"
        processing = os.path.join(BASE_DIR, processingFile_path)
        df_old_proc = pd.read_csv(processing) 
        statusReq=df_old_proc.loc[df_old_proc.Idx == 1, "Status"] 
        del df_old_proc
        if(statusReq == "Not done").any():
            _isDisabled=""
        return render(request, 'modelValReq.html', {'isDisabled':_isDisabled,'modelInfo': result, 'RiskInfo': resultRiskInfo, 'Documentation': resultDocumentation})
    except Exception as e:
        print(e)
        print('stacktrace is ',traceback.print_exc())
        return render(request, 'error.html')


def saveModelDocuments(doc, doc_type, doc_file, reqId, docref=""):
    DocumentationData = file_path + file_name + "_DocumentationData.csv"
    print('doc_type ', doc_type, ' doc_type ', doc_type)
    if(len(doc_file) > 0):
        if os.path.exists(DocumentationData):
            df_old = pd.read_csv(DocumentationData)
            if ((df_old["doc"] == doc) & (
                    df_old['reqId'] == reqId)).any():
                df_old.loc[(df_old['doc'] == doc) & (
                    df_old['reqId'] == reqId),
                    "doc_file"] = doc_file
                df_old.loc[(df_old['doc'] == doc),
                           "doc_type"] = doc_type
                df_old.loc[(df_old['doc'] == doc),
                           "docref"] = docref
                df_old.to_csv(DocumentationData, index=False)
            else:
                data = [[doc, doc_type, doc_file, reqId, docref]]
                df_new = pd.DataFrame(
                    data, columns=['doc', 'doc_type', 'doc_file', 'reqId', 'docref'])
                df = pd.concat([df_old, df_new], axis=0)
                df.to_csv(DocumentationData, index=False)
                del df
                del df_new
            del df_old
        else:
            data = [[doc, doc_type, doc_file, reqId, docref]]
            df = pd.DataFrame(
                data, columns=['doc', 'doc_type', 'doc_file', 'reqId', 'docref'])
            df.to_csv(DocumentationData, index=False)
            del df


def saveModelRiskInfo(risk_type, risk_val, reqId):
    DocumentationData = file_path + file_name + "_ModelRiskInfo.csv"
    if os.path.exists(DocumentationData):
        df_old = pd.read_csv(DocumentationData)
        if ((df_old["risk_type"] == risk_type) & (
                df_old['reqId'] == reqId)).any():
            df_old.loc[(df_old["risk_type"] == risk_type) & (
                df_old['reqId'] == reqId),
                "risk_val"] = risk_val
            df_old.to_csv(DocumentationData, index=False)
        else:
            data = [[risk_type, risk_val, 'reqId']]
            df_new = pd.DataFrame(
                data, columns=['risk_type', 'risk_val', 'reqId'])
            df = pd.concat([df_old, df_new], axis=0)
            df.to_csv(DocumentationData, index=False)
            del df
            del df_new
        del df_old
    else:
        data = [[risk_type, risk_val, reqId]]
        df = pd.DataFrame(
            data, columns=['risk_type', 'risk_val', 'reqId'])
        df.to_csv(DocumentationData, index=False)
        del df


def saveModelInfo(ModelNm, ModelVersion, ModelID, ModelDesc, ModelDate, ModelDev, ModelDept, ModelOwner, Model_Type):
    DocumentationData = file_path + file_name + "_ModelInfo.csv"
    reqId = 0
    today = date.today().strftime("%m/%d/%Y")
    if os.path.exists(DocumentationData):
        df_old = pd.read_csv(DocumentationData)
        if (df_old["ModelNm"] == ModelNm).any():
            df_old.loc[(df_old['ModelNm'] == ModelNm),
                       "ModelVersion"] = ModelVersion
            df_old.loc[(df_old['ModelNm'] == ModelNm),
                       "ModelID"] = ModelID
            df_old.loc[(df_old['ModelNm'] == ModelNm),
                       "ModelDesc"] = ModelDesc
            df_old.loc[(df_old['ModelNm'] == ModelNm),
                       "ModelDate"] = ModelDate
            df_old.loc[(df_old['ModelNm'] == ModelNm),
                       "ModelDev"] = ModelDev
            df_old.loc[(df_old['ModelNm'] == ModelNm),
                       "ModelDept"] = ModelDept
            df_old.loc[(df_old['ModelNm'] == ModelNm),
                       "ModelOwner"] = ModelOwner
            df_old.loc[(df_old['ModelNm'] == ModelNm),
                       "Model_Type"] = Model_Type
            df_old.to_csv(DocumentationData, index=False)
            reqId = df_old["reqId"].max()
        else:
            maxid = df_old["reqId"].max()+1
            reqId = maxid
            data = [[reqId, ModelNm, ModelVersion, ModelID, ModelDesc,
                    ModelDate, ModelDev, ModelDept, ModelOwner, Model_Type, today]]
            df_new = pd.DataFrame(
                data, columns=['reqId', 'ModelNm', 'ModelVersion', 'ModelID', 'ModelDesc', 'ModelDate', 'ModelDev', 'ModelDept', 'ModelOwner', 'Model_Type', 'Date'])
            df = pd.concat([df_old, df_new], axis=0)
            df.to_csv(DocumentationData, index=False)
            del df
            del df_new
            if(len(ModelNm) > 0):
                processing = os.path.join(
                    BASE_DIR, processingFile_path)
                df_old_proc = pd.read_csv(processing)
                df_old_proc["Status"] = "Not done"
                df_old_proc.loc[df_old_proc.Idx == 1, "Status"] = "Done"
                df_old_proc.to_csv(processing, index=False)
                del df_old_proc
        del df_old
    else:
        data = [['1', ModelNm, ModelVersion, ModelID, ModelDesc,
                ModelDate, ModelDev, ModelDept, ModelOwner, Model_Type, today]]
        df = pd.DataFrame(
            data, columns=['reqId', 'ModelNm', 'ModelVersion', 'ModelID', 'ModelDesc', 'ModelDate', 'ModelDev', 'ModelDept', 'ModelOwner', 'Model_Type', 'Date'])
        df.to_csv(DocumentationData, index=False)
        if(len(ModelNm) > 0):
            processing = os.path.join(BASE_DIR, processingFile_path)
            df_old_proc = pd.read_csv(processing)
            df_old_proc["Status"] = "Not done"
            df_old_proc.loc[df_old_proc.Idx == 1, "Status"] = "Done"
            df_old_proc.to_csv(processing, index=False)
            del df_old_proc
        reqId = 1
        del df

    return reqId


def saveSummaryData(request):
    comments = request.GET['comments']
    title = request.GET['title']
    subTitle = request.GET['subTitle']
    SummaryDataFiles = file_path + file_name + "_SummaryData.csv"
    if os.path.exists(SummaryDataFiles):
        df_old = pd.read_csv(SummaryDataFiles)
        if len(subTitle) > 0:
            # dffilter = df_old.query(
            #     "title== '"+title+"' and subTitle='" + subTitle + "' ")
            dffilter = df_old.loc[(df_old['title'] == title) & (
                df_old['subTitle'] == subTitle)]
            if (df_old["subTitle"] == subTitle).any():
                df_old.loc[(df_old['title'] == title) & (
                    df_old['subTitle'] == subTitle), "comments"] = comments
                df_old.to_csv(SummaryDataFiles, index=False)
            else:
                data = [[title, comments, subTitle]]
                df_new = pd.DataFrame(
                    data, columns=['title', 'comments', 'subTitle'])
                df = pd.concat([df_old, df_new], axis=0)
                df.to_csv(SummaryDataFiles, index=False)
            del dffilter
        else:
            if (df_old["title"] == title).any():
                df_old.loc[(df_old['title'] == title) & (
                    df_old['subTitle'] == "-"), "comments"] = comments
                df_old.to_csv(SummaryDataFiles, index=False)
            else:
                data = [[title, comments, '-']]
                df_new = pd.DataFrame(
                    data, columns=['title', 'comments', 'subTitle'])
                df = pd.concat([df_old, df_new], axis=0)
                df.to_csv(SummaryDataFiles, index=False)
    else:
        data = [[title, comments, '-']]
        df = pd.DataFrame(
            data, columns=['title', 'comments', 'subTitle'])
        df.to_csv(SummaryDataFiles, index=False)

    return JsonResponse({'is_taken': True})


def getEmails():
    contactFile = file_path + user_name + "_Contacts.csv"
    result = {
        '': '',
    }
    if os.path.exists(contactFile):
        df = pd.read_csv(contactFile)
        result = df.to_json(orient="records")
        result = json.loads(result)

    return result


def valFindings(request):
    try:
        validationFindings = file_path + file_name + "_validationFindings.csv"
        data = {'List': ''}

        # Returns the current local date
        today = date.today().strftime("%m/%d/%Y")
        result = []
        if os.path.exists(validationFindings):
            df = pd.read_csv(validationFindings)

            # df = df["findingsId"]
            for index, row in df.iterrows():
                print('str(row["Response"])  ', row["Response"])
                if (str(row["Response"]) != "-"):
                    result.append(
                        {'val': row["findingsId"], 'bgColor': 'green', 'color': 'white'})
                else:
                    result.append(
                        {'val': row["findingsId"], 'bgColor': 'white', 'color': 'black'})

            data = {'List': result, 'today': today, 'emailLst': getEmails()}
        return render(request, 'valFindings.html', data)
    except Exception as e:
        print(e)
        return render(request, 'error.html')


def valFindingsResp(request):
    try:
        validationFindings = file_path + file_name + "_validationFindings.csv"
        data = {'List': ''}
        if os.path.exists(validationFindings):
            df = pd.read_csv(validationFindings)
            df = df["findingsId"]
            result = df.to_json(orient="records")
            result = json.loads(result)
            data = {'List': result}
        return render(request, 'valFindingsResp.html', data)
    except Exception as e:
        print(e)
        return render(request, 'error.html')


def getvalFindings(request):
    validationFindings = file_path + file_name + "_validationFindings.csv"
    findingsId = request.GET['findingsId']
    if os.path.exists(validationFindings):
        df = pd.read_csv(validationFindings)
        dffilter = df.query("findingsId== '"+findingsId+"'")
        result = dffilter.to_json(orient="records")
        result = json.loads(result)
        data = {'findingData': result}
        del dffilter, df 
    return JsonResponse(data)


def documantation(request):
    try:
        DocumentationData = file_path + file_name + "_ModelInfo.csv"

        today = date.today().strftime("%m/%d/%Y")
        result = '[{"ModelNm": "", "ModelVersion":"", "ModelID":"", "ModelDesc":"", "ModelDate":"' + \
            today+'", "ModelDev":"", "ModelDept":"", "ModelOwner":"", "Model_Type": "Vendor"}]'
        result = json.loads(result)
        reqId = 0
        resultDocumentation=[]
        if os.path.exists(DocumentationData):
            df = pd.read_csv(DocumentationData)
            reqId = df["reqId"].max()
            dffilter = df.query("reqId== "+str(reqId))
            result = dffilter.to_json(orient="records")
            result = json.loads(result)
            del df, dffilter

        DocumentationData = file_path + file_name + "_DocumentationData.csv"
        if os.path.exists(DocumentationData):
            dffilter = pd.read_csv(DocumentationData)
            # dffilter = df.query("reqId== "+str(reqId))
            resultDocumentation = dffilter.to_json(orient="records")
            resultDocumentation = json.loads(resultDocumentation)

        validationFindings = file_path + file_name + "_DocComments.csv"
      

        # Returns the current local date
        today = date.today().strftime("%m/%d/%Y")
        commresult = []
        if os.path.exists(validationFindings):
            df = pd.read_csv(validationFindings) 
            for index, row in df.iterrows():
                print('str(row["Response"])  ', row["Response"])
                if (str(row["Response"]) != "-"):
                    commresult.append(
                        {'val': row["commentsId"], 'bgColor': 'green', 'color': 'white'})
                else:
                    commresult.append(
                        {'val': row["commentsId"], 'bgColor': 'white', 'color': 'black'})

            print('commresult ',commresult)
        return render(request, 'documentation2.html', {'Documentation': resultDocumentation,'List': commresult, 'today': today, 'emailLst': getEmails()})
    except Exception as e:
        print('stacktrac ',traceback.print_exc())
        return render(request, 'error.html')


def saveDocumentationData(request):
    DocumentationData = file_path + file_name + "_DocumentationData.csv"
    content = request.GET['doc_type']
    json_dictionary = json.loads(content)
    for colval in json_dictionary:
        print(' colval ', colval)
        doc = colval["doc_type"]
        doc_file = colval["doc_file"]
        doc_comment = colval["doc_comment"]
        if os.path.exists(DocumentationData):
            df_old = pd.read_csv(DocumentationData)
            if (df_old["doc"] == doc).any():
                df_old.loc[(df_old['doc'] == doc),
                           "doc_file"] = doc_file
                df_old.loc[(df_old['doc'] == doc),
                           "docref"] = doc_comment
                df_old.to_csv(DocumentationData, index=False)
            else:
                data = [[doc, 'link', doc_file, doc_comment]]
                df_new = pd.DataFrame(
                    data, columns=['doc', 'doc_type', 'doc_file', 'docref'])
                df = pd.concat([df_old, df_new], axis=0)
                df.to_csv(DocumentationData, index=False)
        else:
            data = [[doc, 'link', doc_file, doc_comment]]
            df = pd.DataFrame(
                data, columns=['doc', 'doc_type', 'doc_file', 'docref'])
            df.to_csv(DocumentationData, index=False)

    return JsonResponse({'is_taken': True})


def saveReportImage(request):
    # {'':imgHeight,'':imgWidth,'':$("#selectImg").val(),
    imgAlign = request.GET['imgAlign']
    ImgTitle = request.GET['ImgTitle']
    TitleAlign = request.GET['TitleAlign']
    imgHeight = request.GET['imgHeight']
    imgWidth = request.GET['imgWidth']
    img = request.GET['img']
    title = request.GET['title']
    titleIdx = request.GET['titleIdx']
    subtitle = request.GET['subtitle']
    subtitleIdx = request.GET['subtitleIdx']
    subsubtitle = request.GET['subsubtitle']
    subsubtitleIdx = request.GET['subsubtitleIdx']
    report_file_path = os.path.join(BASE_DIR, plot_dir_view)
    report_file_name = "temp_report_"+user_name
    cnfrmsrcFiles = report_file_path + report_file_name + ".csv"
    report_file_name = "report_"+user_name
    savedReport = report_file_path + report_file_name + ".csv"
    savedData = ""
    if(os.path.exists(savedReport)):
        cnfrmsrcFiles = savedReport
    if os.path.exists(cnfrmsrcFiles):
        df_old = pd.read_csv(cnfrmsrcFiles)
        maxid = df_old["reqID"].max()+1
        reqId = maxid
        data = [['Image', title, titleIdx, subtitle, subtitleIdx, subsubtitle, subsubtitleIdx, maxid, '-', img, imgWidth,
                imgHeight, imgAlign, ImgTitle, TitleAlign]]
        df_new = pd.DataFrame(
            data, columns=['section', 'title', 'titleIdx', 'subtitle', 'subtitleIdx', 'subsubtitle', 'subsubtitleIdx', 'reqID', 'comment', 'img', 'imgWidth', 'imgHeight', 'imgAlign', 'ImgTitle', 'TitleAlign'])
        df = pd.concat([df_old, df_new], axis=0)
        df.to_csv(cnfrmsrcFiles, index=False)
    else:
        data = [['Image', title, titleIdx, subtitle, subtitleIdx, subsubtitle, subsubtitleIdx,   1, '-', img, imgWidth,
                imgHeight, imgAlign, ImgTitle, TitleAlign]]
        df = pd.DataFrame(
            data, columns=['section', 'title', 'titleIdx', 'subtitle', 'subtitleIdx', 'subsubtitle',  'subsubtitleIdx', 'reqID', 'comment', 'img', 'imgWidth', 'imgHeight', 'imgAlign', 'ImgTitle', 'TitleAlign'])
        df.to_csv(cnfrmsrcFiles, index=False)
        reqId = 1
    savedData = getSavedReportData()
    data = {
        'is_taken': True,
        'reqId': str(reqId),
        'savedReportData': savedData,
    }
    return JsonResponse(data)


def saveReportComment(request):
    comment = request.GET['comment']
    reqId = request.GET['reqId']
    title = request.GET['title']
    titleIdx = request.GET['titleIdx']
    subtitle = request.GET['subtitle']
    subtitleIdx = request.GET['subtitleIdx']
    subsubtitle = request.GET['subsubtitle']
    subsubtitleIdx = request.GET['subsubtitleIdx']
    report_file_path = os.path.join(BASE_DIR, plot_dir_view)
    report_file_name = "temp_report_"+user_name
    cnfrmsrcFiles = report_file_path + report_file_name + ".csv"
    report_file_name = "report_"+user_name
    savedReport = report_file_path + report_file_name + ".csv"
    savedData = ""
    if(os.path.exists(savedReport)):
        cnfrmsrcFiles = savedReport
    if os.path.exists(cnfrmsrcFiles):
        df_old = pd.read_csv(cnfrmsrcFiles, encoding='utf-8')
        print('old df ', df_old)
        if(float(subsubtitleIdx) != 0):
            dffilter = df_old.query("subsubtitleIdx ==" + subsubtitleIdx)
        elif(float(subsubtitleIdx) == 0 and float(subtitleIdx) != 0):
            dffilter = df_old.query("subtitleIdx ==" + subtitleIdx)
        else:
            dffilter = df_old.query("titleIdx ==" + titleIdx)
        print('reqID ', reqId)
        if(len(reqId) > 0):
            df_old.loc[df_old.reqID == float(
                reqId), "comment"] = comment
            df_old.to_csv(cnfrmsrcFiles, index=False, encoding='utf-8')
        # elif (len(dffilter) > 0):
        #     if(int(subsubtitleIdx) != 0):
        #         df_old.loc[df_old.subsubtitleIdx == float(
        #             subsubtitleIdx), "comment"] = comment
        #     elif(int(subsubtitleIdx) == 0 and int(subtitleIdx) != 0):
        #         df_old.loc[df_old.subtitleIdx == float(
        #             subtitleIdx), "comment"] = comment
        #     else:
        #         df_old.loc[df_old.titleIdx == float(
        #             titleIdx), "comment"] = comment

        #     df_old.to_csv(cnfrmsrcFiles, index=False, encoding='utf-8')
        else:
            maxid = df_old["reqID"].max()+1
            data = [['Comment', title, titleIdx, subtitle, subtitleIdx, subsubtitle,  subsubtitleIdx, maxid, comment,
                    '-',  '-', '-',  '-',  '-', '-']]
            print('data ', data)
            df_new = pd.DataFrame(
                data, columns=['section', 'title', 'titleIdx', 'subtitle', 'subtitleIdx', 'subsubtitle', 'subsubtitleIdx', 'reqID', 'comment', 'img', 'imgWidth', 'imgHeight', 'imgAlign', 'ImgTitle', 'TitleAlign'])
            df = pd.concat([df_old, df_new], axis=0)
            df.to_csv(cnfrmsrcFiles, index=False, encoding='utf-8')
            reqId = maxid
            # print(unicode(comment, errors='replace'))
    else:
        data = [['Comment', title, titleIdx, subtitle, subtitleIdx, subsubtitle, subsubtitleIdx, 1,
                comment,  '-',  '-', '-',  '-',  '-', '-']]
        reqId = 1
        df = pd.DataFrame(
            data, columns=['section', 'title', 'titleIdx', 'subtitle', 'subtitleIdx', 'subsubtitle', 'subsubtitleIdx', 'reqID', 'comment', 'img', 'imgWidth', 'imgHeight', 'imgAlign', 'ImgTitle', 'TitleAlign'])
        df.to_csv(cnfrmsrcFiles, index=False, encoding='utf-8')
    newTitles=[]
    savedData,newTitles = getSavedReportData()
    print('savedData ', savedData)
    data = {
        'is_taken': True,
        'reqId': str(reqId),
        'titleIdx': int(titleIdx)-1,
        'savedReportData': savedData,
    }
    return JsonResponse(data)


def saveReportCommentTxtEd(request):
    body_unicode = request.body.decode('utf-8')
    body = json.loads(body_unicode)
    comment = body['comment']
    reqId = body['reqId']
    title = body['title']
    titleIdx = body['titleIdx']
    subtitle = body['subtitle']
    subtitleIdx = body['subtitleIdx']
    subsubtitle = body['subsubtitle']
    subsubtitleIdx = body['subsubtitleIdx']
    report_file_path = os.path.join(BASE_DIR, plot_dir_view)
    report_file_name = "temp_report_"+user_name
    cnfrmsrcFiles = report_file_path + report_file_name + ".csv"
    report_file_name = "report_"+user_name
    savedReport = report_file_path + report_file_name + ".csv"
    savedData = "" 
    if(os.path.exists(savedReport)):
        cnfrmsrcFiles = savedReport
    if os.path.exists(cnfrmsrcFiles):
        df_old = pd.read_csv(cnfrmsrcFiles, encoding='utf-8')
        print('subsubtitleIdx ', subsubtitleIdx, ' subtitleIdx ',  subtitleIdx, ' titleIdx ', titleIdx)
        
        if(str(subsubtitleIdx) != "0"):
            dffilter = df_old.query(
                "subsubtitleIdx =='" + str(subsubtitleIdx) + "'")
        elif(str(subsubtitleIdx) == "0" and float(subtitleIdx) != 0):
            # print('inside if')
            dffilter = df_old.query("subtitleIdx ==" + str(subtitleIdx))
            # print('len(dffilter) ', len(dffilter))
        else:
            dffilter = df_old.query("titleIdx ==" + str(titleIdx)+" and subtitleIdx ==0")
        # print('reqID ', reqId, ' len(dffilter) ', len(dffilter))
        if(len(reqId) > 0):
            df_old.loc[df_old.reqID == float(
                reqId), "comment"] = comment
            df_old.to_csv(cnfrmsrcFiles, index=False, encoding='utf-8')
        elif (len(dffilter) > 0):
            if(str(subsubtitleIdx) != "0"):
                df_old.loc[df_old.subsubtitleIdx ==
                           subsubtitleIdx, "comment"] = comment
            elif(str(subsubtitleIdx) == "0" and float(subtitleIdx) != 0):
                df_old.loc[df_old.subtitleIdx == float(
                    subtitleIdx), "comment"] = comment
            else:
                df_old.loc[df_old.titleIdx == float(
                    titleIdx), "comment"] = comment

            df_old.to_csv(cnfrmsrcFiles, index=False, encoding='utf-8')
        else:
            maxid = df_old["reqID"].max()+1
            data = [['Comment', title, titleIdx, subtitle, subtitleIdx, subsubtitle,  subsubtitleIdx, maxid, comment,
                    '-',  '-', '-',  '-',  '-', '-']]

            df_new = pd.DataFrame(
                data, columns=['section', 'title', 'titleIdx', 'subtitle', 'subtitleIdx', 'subsubtitle', 'subsubtitleIdx', 'reqID', 'comment', 'img', 'imgWidth', 'imgHeight', 'imgAlign', 'ImgTitle', 'TitleAlign'])
            df = pd.concat([df_old, df_new], axis=0)
            df.to_csv(cnfrmsrcFiles, index=False, encoding='utf-8')
            reqId = maxid
            # print(unicode(comment, errors='replace'))
    else:
        data = [['Comment', title, titleIdx, subtitle, subtitleIdx, subsubtitle, subsubtitleIdx, 1,
                comment,  '-',  '-', '-',  '-',  '-', '-']]
        reqId = 1
        df = pd.DataFrame(
            data, columns=['section', 'title', 'titleIdx', 'subtitle', 'subtitleIdx', 'subsubtitle', 'subsubtitleIdx', 'reqID', 'comment', 'img', 'imgWidth', 'imgHeight', 'imgAlign', 'ImgTitle', 'TitleAlign'])
        df.to_csv(cnfrmsrcFiles, index=False, encoding='utf-8')
    newTtl=[]
    savedData,newTtl = getSavedReportData()
    # print('savedData ', savedData)
    data = {
        'is_taken': True,
        'reqId': str(reqId),
        'titleIdx': int(titleIdx)-1,
        'savedReportData': savedData,
    }
    # print('data ', data)
    return JsonResponse(data)


def getSavedReportData(template="General"): 
    # SummaryDataFiles = file_path + file_name + "_SummaryData.csv" 
    temp_report_file_path = os.path.join(BASE_DIR, plot_dir_view)
    temp_report_file_name = "temp_report_"+user_name
    if template=="Varo":
        temp_report_file_path= os.path.join(BASE_DIR, 'static/reportTemplates/')
        temp_report_file_name = "varo_report"
    temp_cnfrmsrcFiles = temp_report_file_path + temp_report_file_name + ".csv"
    divMain = ""
    newTitles=[]
    divSection = ""
    report_file_path = os.path.join(BASE_DIR, plot_dir_view)
    staticTitles='"Executive Summary" ,"Model Assessment","Model Performance & Testing","Implementation and Controls","Governance and Oversight"'
    report_file_name = "report_"+user_name
    cnfrmsrcFiles = report_file_path + report_file_name + ".csv"
    print('os.path.exists(temp_cnfrmsrcFiles) is ',temp_cnfrmsrcFiles ,',',os.path.exists(temp_cnfrmsrcFiles))
    if os.path.exists(temp_cnfrmsrcFiles):
        if not os.path.exists(cnfrmsrcFiles):
            # os.remove(cnfrmsrcFiles)
        # os.rename(temp_cnfrmsrcFiles, cnfrmsrcFiles)
            shutil.copyfile(temp_cnfrmsrcFiles, cnfrmsrcFiles)
    ms_report_file_path = os.path.join(BASE_DIR, 'static/reportTemplates/')
    ms_report_file_name = "report_master.csv"
    df_ms = pd.read_csv(ms_report_file_path+ms_report_file_name, encoding='utf-8')
    if os.path.exists(cnfrmsrcFiles):
        df_old = pd.read_csv(cnfrmsrcFiles, encoding='utf-8')
        df = df_old.sort_values(
            by=['titleIdx', 'subtitleIdx', 'subsubtitleIdx', 'reqID'], ascending=True)
        
        titleAdded = ""
        titleAddedIdx=""
        subtitles_not_added ="" 
        idx=0 
        for index, row in df.iterrows():  
            if((str(row["subtitleIdx"]) == '0' or str(row["subtitleIdx"]) == '0.0') and (str(row["subsubtitleIdx"]) == '0' or str(row["subsubtitleIdx"]) == '0.0')):

                if(row["title"] != titleAdded): 
                    dfms_subcols = df_ms.query("title =='" + str(row["title"])+"'") 
                    dfms_subcols=dfms_subcols["subtitle"].tolist()
                    df_subcols = df.query("title =='" + str(row["title"])+"'") 
                    # if(titleAdded!=""):
                    #     for ttl in subtitles_not_added:
                    #         if not (ttl=='-'):
                    #             parattl="'"+titleAdded+"','"+ttl+"','"+titleAddedIdx+"'"
                    #             divSection = divSection+'<div style="display: flex; justify-content: flex-start;"><div style="width:20px;"><i class="fa fa-edit" style="margin-right:5px;cursor:pointer;"  title="Edit comment" onclick="getSection('+ parattl+')"></i> </div><div id="div_' + str(idx+1000) + '" >&nbsp;&nbsp;&nbsp;&nbsp;<b>'+ ttl.replace("'","") +'</b></div></div>'
                    #             divSection = divSection+"<br>" 
                        # parattl="'"+titleAdded+"','-1','"+titleAddedIdx+"'"
                        # divSection = divSection+'<div style="display: flex; justify-content: flex-start;"><div style="width:20px;"><i class="fa fa-edit" style="margin-right:5px;cursor:pointer;"  title="Edit comment" onclick="getSection('+ parattl+')"></i> </div><div id="div_sub_new' + str(titleAddedIdx) + '" >&nbsp;&nbsp;&nbsp;&nbsp;<b>Add New</b></div></div>'
                        # divSection = divSection+"<br>" 

                    # if(len(df_subcols)>0):
                    #     df_subcols=df_subcols["subtitle"].tolist()
                    #     subtitles_not_added=list(set(dfms_subcols).difference(set(df_subcols))) 
                    
                    titleAdded=str(row["title"])
                    titleAddedIdx=str(row["titleIdx"])
                    if(staticTitles.find(str(row["title"])) == -1):
                        newTitles.append(str(row["title"])) 
                
            if row["section"] == "Comment":
                if(str(row["comment"]) != "nan"): 
                    if((str(row["subtitleIdx"]) != '0' and str(row["subtitleIdx"]) != '0.0') and (str(row["subsubtitleIdx"]) == '0' or str(row["subsubtitleIdx"]) == '0.0')):
                        print( (str(row["comment"])) ,", ",  (str(row["subtitle"])))
                        print(len(str(row["comment"])) ,", ", len(str(row["subtitle"])))
                        if(len(str(row["comment"])) >=len(str(row["subtitle"]))+17 and len(str(row["comment"])) <=len(str(row["subtitle"]))+22):
                            divSection = divSection+"<div style='display: flex; justify-content: flex-start;'><div style='width:40px;'><i class='fa fa-edit' style='margin-left:16px;margin-right:5px;cursor:pointer;'  title='Edit comment' onclick='getData("+str(row["reqID"]) +")'></i></div>&nbsp;&nbsp;&nbsp;&nbsp;<div id='div_" + str(row["reqID"]) + "' ondblclick='getData("+str(row["reqID"]) +")'>"+  str(row["comment"]).replace("\t", "\u0020\u0020\u0020\u0020").replace(str(row["title"])+"\n", "").replace(
                            str(row["subtitle"])+"\n", "").replace(str(row["subsubtitle"])+"\n", "").replace("\n", "") +"</div></div>"
                            divSection = divSection+"<br>"
                        else:
                            divSection = divSection+"<div style='display: flex; justify-content: flex-start;'><div style='width:17px;'><i id='tggl_" + str(row["reqID"]) + "' class='fa fa-plus-square-o' style='margin-right:5px;cursor:pointer;' onclick='toggleHeight("+str(row["reqID"]) +",this.id)'></i></div><div style='width:20px;'><i class='fa fa-edit' style='margin-right:5px;cursor:pointer;'  title='Edit comment' onclick='getData("+str(row["reqID"]) +")'></i> </div>&nbsp;&nbsp;&nbsp;&nbsp;<div id='div_" + str(row["reqID"]) + "' style='height:20px;overflow:hidden;'>"+  str(row["comment"]).replace("\t", "\u0020\u0020\u0020\u0020").replace(str(row["title"])+"\n", "").replace(
                            str(row["subtitle"])+"\n", "").replace(str(row["subsubtitle"])+"\n", "").replace("\n", "") +"</div></div>"
                            divSection = divSection+"<br>"
                    elif((str(row["subsubtitleIdx"]) != '0' and str(row["subsubtitleIdx"]) != '0.0')):
                        if(len(str(row["comment"])) >=len(str(row["subsubtitle"]))+19 and len(str(row["comment"])) <=len(str(row["subsubtitle"]))+24):
                            divSection = divSection+"<div style='display: flex; justify-content: flex-start;'><div style='width:40px;'><i class='fa fa-edit' style='margin-left:16px;margin-right:5px;cursor:pointer;'  title='Edit comment' onclick='getData("+str(row["reqID"]) +")'></i></div>&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;<div id='div_" + str(row["reqID"]) + "' ondblclick='getData("+str(row["reqID"]) +")'>"+  str(row["comment"]).replace("\t", "\u0020\u0020\u0020\u0020").replace(str(row["title"])+"\n", "").replace(
                            str(row["subtitle"])+"\n", "").replace(str(row["subsubtitle"])+"\n", "").replace("\n", "") +"</div></div>"
                            divSection = divSection+"<br>"
                        else:
                            divSection = divSection+"<div style='display: flex; justify-content: flex-start;'><div style='width:17px;'><i id='tggl_" + str(row["reqID"]) + "' class='fa fa-plus-square-o' style='margin-right:5px;cursor:pointer;' onclick='toggleHeight("+str(row["reqID"]) +",this.id)'></i></div><div style='width:20px;'><i class='fa fa-edit' style='margin-right:5px;cursor:pointer;'  title='Edit comment' onclick='getData("+str(row["reqID"]) +")'></i> </div>&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;<div id='div_" + str(row["reqID"]) + "' style='height:20px;overflow:hidden;'>"+  str(row["comment"]).replace("\t", "\u0020\u0020\u0020\u0020").replace(str(row["title"])+"\n", "").replace(
                            str(row["subtitle"])+"\n", "").replace(str(row["subsubtitle"])+"\n", "").replace("\n", "") +"</div></div>"
                            divSection = divSection+"<br>"
                    else:
                        
                        if(len(str(row["comment"])) >=len(str(row["title"]))+15 and len(str(row["comment"])) <=len(str(row["title"]))+19):
                            divSection = divSection+"<div style='display: flex; justify-content: flex-start;'><div style='width:40px;'><i class='fa fa-edit' style='margin-left:16px;margin-right:5px;cursor:pointer;'  title='Edit comment' onclick='getData("+str(row["reqID"]) +")'></i></div><div id='div_" + str(row["reqID"]) + "' style='height:20px;overflow:hidden;'>"+  str(row["comment"]).replace("\t", "\u0020\u0020\u0020\u0020").replace(str(row["title"])+"\n", "").replace(
                                str(row["subtitle"])+"\n", "").replace(str(row["subsubtitle"])+"\n", "").replace("\n", "") +"</div></div>"
                            divSection = divSection+"<br>" 
                        else:
                            divSection = divSection+"<div style='display: flex; justify-content: flex-start;'><div style='width:17px;'><i id='tggl_" + str(row["reqID"]) + "' class='fa fa-plus-square-o' style='margin-right:5px;cursor:pointer;' onclick='toggleHeight("+str(row["reqID"]) +",this.id)'></i></div><div style='width:20px;'><i class='fa fa-edit' style='margin-right:5px;cursor:pointer;'  title='Edit comment' onclick='getData("+str(row["reqID"]) +")'></i> </div><div id='div_" + str(row["reqID"]) + "' style='height:20px;overflow:hidden;'>"+  str(row["comment"]).replace("\t", "\u0020\u0020\u0020\u0020").replace(str(row["title"])+"\n", "").replace(
                                str(row["subtitle"])+"\n", "").replace(str(row["subsubtitle"])+"\n", "").replace("\n", "") +"</div></div>"
                            divSection = divSection+"<br>" 
                        
            # if(idx==len(df)-1): 
            #     for ttl in subtitles_not_added:
            #         if not (ttl=='-'):
            #             parattl="'"+titleAdded+"','"+ttl+"','"+titleAddedIdx+"'"
            #             divSection = divSection+'<div style="display: flex; justify-content: flex-start;"><div style="width:20px;"><i class="fa fa-edit" style="margin-right:5px;cursor:pointer;"  title="Edit comment" onclick="getSection('+ parattl+')"></i> </div><div id="div_' + str(idx+1000) + '" >&nbsp;&nbsp;&nbsp;&nbsp;<b>'+ ttl.replace("'","") +'</b></div></div>'
            #             divSection = divSection+"<br>" 
            idx=idx+1

            
    
    
    # if(not(template=="Varo")):
    #     for index, row in df_ms.iterrows(): 
    #         if((str(row["subtitle"]) == '-' and  str(row["subsubtitle"]) == '-')):            
    #             dffilter = df.query("title =='" + str(row["title"])+"'") 
    #             if((len(dffilter)==0)):
    #                 # print('str(row["title"]) ', str(row["title"]))  
    #                 parattl='"'+str(row["title"])+'","",""'
    #                 divSection = divSection+"<div><i class='fa fa-edit' style='margin-right:5px;cursor:pointer;cursor:pointer;'  title='Edit comment' onclick='getSection(" + parattl+")'></i><b>"+ str(row["title"]) +"</b></div>"
    #                 divSection = divSection+"<br>"
#    add new title
    # parattl='"-1","",""'
    # divSection = divSection+"<div><i class='fa fa-edit' style='margin-right:5px;cursor:pointer;cursor:pointer;'  title='Edit comment' onclick='getSection("+parattl +")'></i><b>Add New</b></div>"
    # divSection = divSection+"<br>"     
    return divSection,newTitles


def getReportComment(request):
    comment = request.GET['comment']
    reqId = request.GET['reqId']
    title = request.GET['title']
    titleIdx = request.GET['titleIdx']
    subtitle = request.GET['subtitle']
    subtitleIdx = request.GET['subtitleIdx']
    subsubtitle = request.GET['subsubtitle']
    subsubtitleIdx = request.GET['subsubtitleIdx']
    report_file_path = os.path.join(BASE_DIR, plot_dir_view)
    report_file_name = "temp_report_"+user_name
    cnfrmsrcFiles = report_file_path + report_file_name + ".csv"
    report_file_name = "report_"+user_name
    savedReport = report_file_path + report_file_name + ".csv"
    if(os.path.exists(savedReport)):
        cnfrmsrcFiles = savedReport
    if os.path.exists(cnfrmsrcFiles):
        df_old = pd.read_csv(cnfrmsrcFiles, encoding='utf-8')
        print('old df ', df_old)
        if (len(reqId) > 0):
            df_old.loc[df_old.reqID == float(reqId), "comment"] = comment
            df_old.to_csv(cnfrmsrcFiles, index=False, encoding='utf-8')
        else:
            maxid = df_old["reqID"].max()+1
            data = [['Comment', title, titleIdx, subtitle, subtitleIdx, subsubtitle,  subsubtitleIdx, maxid, comment,
                    '-',  '-', '-',  '-',  '-', '-']]
            print('data ', data)
            df_new = pd.DataFrame(
                data, columns=['section', 'title', 'titleIdx', 'subtitle', 'subtitleIdx', 'subsubtitle', 'subsubtitleIdx', 'reqID', 'comment', 'img', 'imgWidth', 'imgHeight', 'imgAlign', 'ImgTitle', 'TitleAlign'])
            df = pd.concat([df_old, df_new], axis=0)
            df.to_csv(cnfrmsrcFiles, index=False, encoding='utf-8')
            reqId = maxid
            # print(unicode(comment, errors='replace'))
    else:
        data = [['Comment', title, titleIdx, subtitle, subtitleIdx, subsubtitle, subsubtitleIdx, 1,
                comment,  '-',  '-', '-',  '-',  '-', '-']]
        reqId = 1
        df = pd.DataFrame(
            data, columns=['section', 'title', 'titleIdx', 'subtitle', 'subtitleIdx', 'subsubtitle', 'subsubtitleIdx', 'reqID', 'comment', 'img', 'imgWidth', 'imgHeight', 'imgAlign', 'ImgTitle', 'TitleAlign'])
        df.to_csv(cnfrmsrcFiles, index=False, encoding='utf-8')
    data = {
        'is_taken': True,
        'reqId': str(reqId),
        'titleIdx': int(titleIdx)-1,
    }
    return JsonResponse(data)


def deleteReportComment(request):
    reqID = request.GET['reqID']
    titleIdx = request.GET['titleIdx'] 
    subtitleIdx = request.GET['subtitleIdx'] 
    subsubtitleIdx = request.GET['subsubtitleIdx']  
    report_file_path = os.path.join(BASE_DIR, plot_dir_view)
    report_file_name = "temp_report_"+user_name
    cnfrmsrcFiles = report_file_path + report_file_name + ".csv"
    report_file_name = "report_"+user_name
    savedReport = report_file_path + report_file_name + ".csv"
    data = {
            'is_taken': False
        }
    if(os.path.exists(savedReport)):
        cnfrmsrcFiles = savedReport
    print(titleIdx,subtitleIdx,subsubtitleIdx)
    if os.path.exists(cnfrmsrcFiles):
        if(len(reqID)>0):
            df = pd.read_csv(cnfrmsrcFiles)
            df = df.set_index("reqID")
            df.drop(float(reqID), axis=0, inplace=True)
            df = df.reset_index()
            df.to_csv(cnfrmsrcFiles, index=False)
            del df
        else:
            df_old = pd.read_csv(cnfrmsrcFiles, encoding='utf-8')
            # print('subsubtitleIdx ', subsubtitleIdx, ' subtitleIdx ',  subtitleIdx, ' titleIdx ', titleIdx)
            if(str(subsubtitleIdx) != "0"):
                dffilter = df_old.query(
                    "subsubtitleIdx =='" + str(subsubtitleIdx) + "'")
            elif(str(subsubtitleIdx) == "0" and float(subtitleIdx) != 0):
                # print('inside if')
                dffilter = df_old.query("subtitleIdx ==" + str(subtitleIdx))
                # print('len(dffilter) ', len(dffilter))
            else:
                dffilter = df_old.query("titleIdx ==" + str(titleIdx))
            del df_old
            if len(dffilter)>0:
                reqID =dffilter['reqID'].values[0]
                df = pd.read_csv(cnfrmsrcFiles)
                df = df.set_index("reqID")
                df.drop(float(reqID), axis=0, inplace=True)
                df = df.reset_index()
                df.to_csv(cnfrmsrcFiles, index=False)
                del df
                newTtl=[]
        savedData,newTtl = getSavedReportData() 
        data = {
            'is_taken': True,
            'reqId': '',
            'titleIdx': int(titleIdx)-1,
            'savedReportData': savedData,
        }
    return JsonResponse(data)


def savevalFindings(request):
    Desc = request.GET['Desc']
    Risk_Level = request.GET['Risk_Level']
    Assessment = request.GET['Assessment']
    findingsId = request.GET['findingsId']
    Level = request.GET['Lvl']
    letters = ""
    words = Assessment.split()
    if(len(words)>1):
        for i in range(2):
            letters = letters + words[i][0]
    else:
        letters = Assessment[:2]
    print('letters is ',letters)
    validationFindings = file_path + file_name + "_validationFindings.csv"
    today = date.today().strftime("%m/%d/%Y")
    if os.path.exists(validationFindings):
        df_old = pd.read_csv(validationFindings)
        if ((df_old["Assessment"] == Assessment)).any():
            df_old.loc[(df_old['Assessment'] == Assessment),
                       "Desc"] = Desc
            df_old.to_csv(validationFindings, index=False)
        else:
            findingsId = letters + str(len(df_old)+1)
            maxId = df_old["reqId"].max()+1
            data = [[Assessment, Risk_Level, Desc,
                    findingsId, '-', '-', maxId, today, '-',Level]]
            df_new = pd.DataFrame(
                data, columns=['Assessment', 'Risk_Level', 'Desc', 'findingsId', 'Response', 'EmailId', 'reqId', 'Date', 'Comment','Level'])
            df = pd.concat([df_old, df_new], axis=0)
            df.to_csv(validationFindings, index=False)
            del df
            del df_new
        del df_old
    else:
        findingsId = letters+"1"
        data = [[Assessment, Risk_Level, Desc,
                findingsId, '-', '-', 1, today, '-']]
        df = pd.DataFrame(
            data, columns=['Assessment', 'Risk_Level', 'Desc', 'findingsId', 'Response', 'EmailId', 'reqId', 'Date', 'Comment'])
        df.to_csv(validationFindings, index=False)
        del df
    data = {
        'is_taken': True,
        'findingsId': findingsId
    }
    return JsonResponse(data)

def saveDocComments(request):
    Desc = request.GET['Desc'] 
    Section = request.GET['Section']
    commentsId = request.GET['commentsId'] 
    DocSel= request.GET['DocSel'] 
    
    DocComments = file_path + file_name + "_DocComments.csv"
    today = date.today().strftime("%m/%d/%Y")
    if os.path.exists(DocComments):
        df_old = pd.read_csv(DocComments, encoding='utf-8')
        if ((df_old["Section"] == Section)).any():
            df_old.loc[(df_old['Section'] == Section),
                       "Desc"] = Desc
            df_old.to_csv(DocComments, index=False, encoding='utf-8')
        else:
            commentsId = "Comment" + str(len(df_old)+1)
            maxId = df_old["reqId"].max()+1
            data = [[Section,DocSel,  Desc,
                    commentsId, '-', '-', maxId, today]]
            df_new = pd.DataFrame(
                data, columns=['Section','DocSel',   'Desc', 'commentsId', 'Response', 'EmailId', 'reqId', 'Date'])
            df = pd.concat([df_old, df_new], axis=0)
            df.to_csv(DocComments, index=False, encoding='utf-8')
            del df
            del df_new
        del df_old
    else:
        commentsId = "Comment1"
        data = [[Section, DocSel,  Desc,
                commentsId, '-', '-', 1, today]]
        df = pd.DataFrame(
            data, columns=['Section', 'DocSel','Desc', 'commentsId', 'Response', 'EmailId', 'reqId', 'Date'])
        df.to_csv(DocComments, index=False)
        del df
    data = {
        'is_taken': True,
        'findingsId': commentsId
    }
    return JsonResponse(data)

def getDocComments(request):
    DocComments = file_path + file_name + "_DocComments.csv"
    commentsId = request.GET['commentsId']
    # DocSel= request.GET['DocSel']
    print(' commentsId ',commentsId)
    if os.path.exists(DocComments):
        df = pd.read_csv(DocComments)
        dffilter = df.query("commentsId== '"+commentsId+"'")# and DocSel=="+DocSel+"
        result = dffilter.to_json(orient="records")
        result = json.loads(result)
        data = {'findingData': result}
        del dffilter, df
        print('data ', data)
    return JsonResponse(data)


def docCommResp(request):
    try:
        validationFindings = file_path + file_name + "_DocComments.csv"
        data = {'List': ''}
        if os.path.exists(validationFindings):
            df = pd.read_csv(validationFindings)
            df = df["commentsId"]
            result = df.to_json(orient="records")
            result = json.loads(result)
            data = {'List': result}
        return render(request, 'docCommResp.html', data)
    except Exception as e:
        print(e)
        return render(request, 'error.html')

def saveDocCommentsResp(request):
    Resp = request.GET['Resp']
    commentsId = request.GET['commentsId']
    validationFindings = file_path + file_name + "_DocComments.csv"
    if os.path.exists(validationFindings):
        df_old = pd.read_csv(validationFindings)
        if ((df_old["commentsId"] == commentsId)).any():
            df_old.loc[(df_old['commentsId'] == commentsId),
                       "Response"] = Resp
            df_old.to_csv(validationFindings, index=False)
        del df_old
    data = {
        'is_taken': True
    }
    return JsonResponse(data)

def savevalFindingsResp(request):
    Resp = request.GET['Resp']
    findingsId = request.GET['findingsId']
    validationFindings = file_path + file_name + "_validationFindings.csv"
    if os.path.exists(validationFindings):
        df_old = pd.read_csv(validationFindings)
        if ((df_old["findingsId"] == findingsId)).any():
            df_old.loc[(df_old['findingsId'] == findingsId),
                       "Response"] = Resp
            df_old.to_csv(validationFindings, index=False)
        del df_old
    data = {
        'is_taken': True
    }
    return JsonResponse(data)


def exportDoc():

    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    document.add_picture(os.path.join(
        BASE_DIR, 'static/media/output.png'), width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    document.save(os.path.join(
        BASE_DIR, "static/media/demo.docx"))


def modelUsage(request):
    try:
        catLst = ['Business Requirement', 'Explanation of Model Output', 'Business Requirement Met', 'Model Performance Monitoring ',
                  'Model Maintenance (Frequency of Failure, Run-time etc)', 'Model Dependencies', 'Violations of Model Assumptions']
        data = {'emailLst': getEmails(), 'catLst': catLst}
        return render(request, 'modelUsage.html', data)
    except Exception as e:
        print('error is ', e)
        return render(request, 'error.html')


def saveModelUsageReq(request):
    email = request.GET['email']
    categories = request.GET['categories']
    json_dictionary = json.loads(categories)
    modelUsageReq = file_path + file_name + "_modelUsageReq.csv"

    if os.path.exists(modelUsageReq):
        os.remove(modelUsageReq)
    if os.path.exists(modelUsageReq):
        for colval in json_dictionary:
            for attribute, value in colval.items():
                df_old = pd.read_csv(modelUsageReq)
                colName = value
                if (df_old["categories"] == colName).any():
                    print('already exists')
                else:
                    maxId = df_old["reqIdx"].max()+1
                    data = [
                        [maxId, 'req_'+str(maxId), email, colName, '-', '-']]
                    df_new = pd.DataFrame(
                        data, columns=['reqIdx', 'reqId', 'email', 'categories', 'resp', 'comments'])
                    df = pd.concat([df_old, df_new], axis=0)
                    df.to_csv(modelUsageReq, index=False)
                    del df
                    del df_new
                    del df_old
        sendModelUsage(email)
    else:
        for colval in json_dictionary:
            for attribute, value in colval.items():
                colName = value
                print('colName ', colName)
                if os.path.exists(modelUsageReq):
                    df_old = pd.read_csv(modelUsageReq)
                    maxId = df_old["reqIdx"].max()+1
                    data = [
                        [maxId, 'req_'+str(maxId), email, colName, '-', '-']]
                    df_new = pd.DataFrame(
                        data, columns=['reqIdx', 'reqId',  'email', 'categories', 'resp', 'comments'])
                    df = pd.concat([df_old, df_new], axis=0)
                    df.to_csv(modelUsageReq, index=False)
                    del df
                    del df_new
                    del df_old
                else:
                    data = [[1, 'req_1', email, colName, '-', '-']]
                    df_new = pd.DataFrame(
                        data, columns=['reqIdx', 'reqId', 'email', 'categories', 'resp', 'comments'])
                    df_new.to_csv(modelUsageReq, index=False)
                    del df_new
        sendModelUsage(email)
    data = {
        'is_taken': True,
    }
    return JsonResponse(data)


def sendModelUsage(emailId):
    try:
        import smtplib
        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText
        mail_content = """Hello,
        Please click link below to responde the model usage.
        """+app_url + """modelUsageResp
        Thank You
        """
        # The mail addresses and password
        sender_address = 'modvaladm@gmail.com'
        sender_pass = 'modelVal#12'
        receiver_address = emailId
        # Setup the MIME
        message = MIMEMultipart()
        message['From'] = sender_address
        message['To'] = receiver_address
        # The subject line
        message['Subject'] = 'Model usage.'

        # The body and the attachments for the mail
        message.attach(MIMEText(mail_content, 'plain'))

        # Create SMTP session for sending the mail
        # use gmail with port
        session = smtplib.SMTP('smtp.gmail.com', 587)
        session.starttls()  # enable security
        # login with mail_id and password
        session.login(sender_address, sender_pass)
        text = message.as_string()
        session.sendmail(sender_address, receiver_address, text)
        session.quit()
        print('Mail Sent')
        data = {'is_taken': True}
    except Exception as e:
        print(e)
        print("Error: unable to send email")
        data = {'is_taken': False}
    return JsonResponse(data)


def modelUsageResp(request):
    try:
        modelUsageReq = file_path + file_name + "_modelUsageReq.csv"
        data = {'List': ''}
        if os.path.exists(modelUsageReq):
            df = pd.read_csv(modelUsageReq)
            result = df.to_json(orient="records")
            result = json.loads(result)
            data = {'List': result}
        return render(request, 'modelUsageResp.html', data)
    except Exception as e:
        print(e)
        return render(request, 'error.html')


def saveModelUsageResp(request):
    feqModel = request.GET['feqModel']
    comments = request.GET['comments']
    users = request.GET['users']
    comments_dictionary = json.loads(comments)
    users_dictionary = json.loads(users)
    modelUsageReq = file_path + file_name + "_modelUsageReq.csv"
    modelUsers = file_path + file_name + "_modelUsers.csv"

    if os.path.exists(modelUsageReq):
        for colval in comments_dictionary:
            for attribute, value in colval.items():
                # print('attribute, value ', attribute, value)
                df_old = pd.read_csv(modelUsageReq)
                if (df_old["reqId"] == attribute).any():
                    df_old.loc[(df_old['reqId'] == attribute),
                               "resp"] = value
                    df_old.to_csv(modelUsageReq, index=False)
                del df_old

    if os.path.exists(modelUsers):
        os.remove(modelUsers)

    newpd = pd.DataFrame.from_dict(users_dictionary)
    # print('new ', newpd)
    newpd.to_csv(modelUsers, index=False)
    del newpd
    exportModelUsage()
    data = {
        'is_taken': True,
    }
    return JsonResponse(data)


def exportModelUsage():
    modelUsageReq = file_path + file_name + "_modelUsageReq.csv"
    pdf = FPDF()
    pdf.add_page()
    x, y = 10, 10
    if os.path.exists(modelUsageReq):
        df = pd.read_csv(modelUsageReq)
        df = df.sort_values(by="reqIdx", ascending=True)
        pdf.set_xy(x, y)
        pdf.set_font("Arial", "B", size=9)
        pdf.set_text_color(0.0, 0.0, 0.0)
        pdf.cell(
            0, 5, "Model Usage", align='L')
        pdf.set_font("Arial",  size=9)
        pdf.set_fill_color(211, 211, 211)
        y = pdf.get_y()+10
        pdf.set_xy(20, y)
        pdf.cell(60, 5, "Section",
                 1, fill=True, align='L')

        pdf.set_xy(80, y)
        pdf.cell(100, 5, "Comments",
                 1, fill=True, align='L')
        y = pdf.get_y()+5
        cellHeight = 0
        pdf.set_font("Arial",  size=8)
        for index, row in df.iterrows():
            pdf.set_xy(20, y)
            pdf.multi_cell(60, 4, str(
                row["categories"]), 0, fill=False)
            cellHeight = pdf.get_y()
            pdf.set_xy(80, y)
            pdf.multi_cell(100, 4, str(row["resp"]).encode(
                'latin-1', 'replace').decode('latin-1'), 0, fill=False)
            if(pdf.get_y() > cellHeight):
                cellHeight = pdf.get_y()
            pdf.rect(20, y, 60, cellHeight-y)
            pdf.rect(80, y, 100, cellHeight-y)
            y = cellHeight
            print(' y is ', y, ' cellHeight ', cellHeight)
    pdf.output(os.path.join(
        BASE_DIR, plot_dir_view + file_name + "_ModelUsage.pdf"))


def openFile(request):
    try:
        fileName = request.GET['fileName']
        return render(request, 'openFile.html', {'fileNm': fileName})
    except Exception as e:
        print(e)
        return render(request, 'error.html')


def showPlotnComm(request):
    try:
        ImgId = request.GET['ImgId']
        UserChartFile = file_path + "_Chartimg.csv"
        comments = ""
        graphpath = ""
        print('ImgId is ', ImgId)
        print('ImgId is ', len(ImgId.strip()))
        if os.path.exists(UserChartFile):
            df = pd.read_csv(UserChartFile)
            print('df is ', df)
            dffilter = df.query("chartImg == '" + ImgId.strip() + "'")
            print('dffilter is ', dffilter)
            if(len(dffilter) > 0):
                graphpath = dffilter["destination"].values[0]
                comments = dffilter["comments"].values[0]

        return render(request, 'showPlotnComment.html', {'graphpath': "/"+graphpath, 'comments': comments})
    except Exception as e:
        print(e)
        return render(request, 'error.html')


def runFile(request):
    try:
        # import subprocess
        # cmd = 'python runscriptupdated.py'

        # # p = subprocess.Popen(cmd, stdout=subprocess.PIPE, shell=True)
        # # out, err = p.communicate()
        # # result = out.split('\n')
        # # for lin in result:
        # #     if not lin.startswith('#'):
        # #         print(lin)
        # print('BASE_DIR is ', BASE_DIR)
        # scriptFile = os.path.join(
        #     BASE_DIR, "modelvalidation/runscriptupdated.py")
        # subprocess.call("python "+scriptFile, shell=True)
        return render(request, 'test.html')
    except Exception as e:
        print('error is  ', e)
        print('traceback is  ', traceback.print_exc())
        return render(request, 'error.html')


def replication(request):
    try:
        codefile = os.path.join(BASE_DIR, 'static/modelCode/')
        replication_files = os.path.join(
            BASE_DIR, 'static/replicationFiles/')
        DocumentationData = file_path + file_name + "_DocumentationData.csv"
        outputfiles = []
        userComments = ""
        UserCommentsFiles = file_path + file_name + "_ReplicationComments.csv"
        htlp_Data = []
        gridDttypes=[]
        scripCode=""
        if os.path.exists(DocumentationData):
            df = pd.read_csv(DocumentationData)
            dffilter = df.query("doc == 'Model Code'")
            resultDocumentation = dffilter["doc_file"].values[0]
            codefile = codefile+resultDocumentation
            replication_name = user_name+'_edited_'+resultDocumentation
            replicationCode = replication_files + replication_name
            if os.path.exists(UserCommentsFiles):
                dfuserComm = pd.read_csv(UserCommentsFiles)
                if (dfuserComm["codeFile"] == resultDocumentation).any():
                    userComments = dfuserComm.loc[(
                        dfuserComm["codeFile"] == resultDocumentation)]["comments"].values[0]
                del dfuserComm
            if os.path.exists(replicationCode):
                codefile = replicationCode
                dir_list = os.listdir(os.path.join(
                    BASE_DIR, 'static/replicationoutput/'))
                # prints all files
                outputfiles = dir_list
                   
                    
                file1 = open(codefile, "r")  # write mode
                scripCode = file1.read()
                file1.close()
                savefile_name = file_path + "labeled_csvfile_"+file_name + ".csv"
                idx=1
                if os.path.exists(savefile_name):
                    print('get output cols')
                    gridDttypesrr= getValidNumCols()+getValidCatCols()
                    gridDttypesrr.append('dd_label')
                    for i in gridDttypesrr:
                        gridDttypes.append({'colName': i, 'chkId': idx})
                        idx = idx + 1

        return render(request, 'replication.html', {'data': scripCode, 'helpData': htlp_Data, 'userComments': userComments, 'scriptFile': codefile, 'imgFiles': outputfiles,'cols':gridDttypes})
    except Exception as e:
        print(e)
        print('traceback is  ', traceback.print_exc())
        return render(request, 'error.html')


def saveReplicationData(request):
    try:
        replication_files = os.path.join(
            BASE_DIR, 'static/replicationFiles/')
        body_unicode = request.body.decode('utf-8')
        body = json.loads(body_unicode)
        content = body['comment']
        DocumentationData = file_path + file_name + "_DocumentationData.csv"
        if os.path.exists(DocumentationData):
            df = pd.read_csv(DocumentationData)
            dffilter = df.query("doc == 'Model Code'")
            resultDocumentation = dffilter["doc_file"].values[0]
        replication_name = user_name+'_edited_'+resultDocumentation
        replicationCode = replication_files + replication_name
        print('path is ', replicationCode, os.path.exists(replicationCode))
        if os.path.exists(replicationCode):
            file1 = open(replicationCode, "w")  # write mode
            file1.write(content)
            file1.close()
        else:
            file1 = open(replicationCode, "w+")  # write mode
            file1.write(content)
            file1.close()
        data = {
            'is_taken': True,
            'codeFile': replicationCode,
        }
        return JsonResponse(data)
    except Exception as e:
        print(e)
        print('traceback is  ', traceback.print_exc())


def runReplicationFile(request):
    try:
        import subprocess
        from subprocess import Popen, PIPE
        replicationFile = request.GET['replicationFile'] 
        cmd = 'python runscriptupdated.py' 
        dir = os.path.join(BASE_DIR, 'static/replicationoutput/')
        for f in os.listdir(dir):
            os.remove(os.path.join(dir, f))
        # subprocess.call("python "+replicationFile, shell=True)
        
        cmdTxt="python "+replicationFile
        

        proc = subprocess.Popen(cmdTxt, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        (stdout, stderr) = proc.communicate()
        
        outputfiles = []
        gridDttypes = []
        if os.path.exists(replicationFile):
            dir_list = os.listdir(os.path.join(
                BASE_DIR, 'static/replicationoutput/'))
            # prints all files
            outputfiles = dir_list
        if(len(str(stderr))<=3):
            savefile_name = file_path + "labeled_csvfile_"+file_name + ".csv"  
            idx=1
            if os.path.exists(savefile_name):
                print('get output cols')
                gridDttypesrr= getValidNumCols()+getValidCatCols()
                gridDttypesrr.append('dd_label')
                for i in gridDttypesrr:
                    gridDttypes.append({'colName': i, 'chkId': idx})
                    idx = idx + 1
            
            
            print('gridDttypes ',gridDttypes)
            data = {
                'is_taken': True,
                'imgFiles': outputfiles,
                'cols':gridDttypes
            }
        else:
                data = {
                'is_taken': False,
                'error': str(stderr),
                'cols':gridDttypes
            }
        return JsonResponse(data)
    except Exception as e:
        print('error is  ', e)
        print('traceback is  ', traceback.print_exc())
        data = {
            'is_taken': False,
            'error':traceback.print_exc()
        }
        return JsonResponse(data)

def getValidCatCols():
    csv_file_name = "csvfile_"+user_name
    savefile_withoutnull = file_path + csv_file_name + "_catcols.csv" 
    df = pd.read_csv(savefile_withoutnull)
     
    # x_keep = pd.read_csv(savefile_x_keep)
    cat_cols=[]
    cat_cols =  df["colName"].tolist()
    del df
    return cat_cols

def getValidNumCols():
    csv_file_name = "csvfile_"+user_name
    savefile_withoutnull = file_path + csv_file_name + "_numcols.csv" 
    df = pd.read_csv(savefile_withoutnull)
     
    # x_keep = pd.read_csv(savefile_x_keep)
    num_cols=[]
    num_cols =  df["colName"].tolist()
    del df
    return num_cols

def sendReportMail(request):
    try:
        import smtplib
        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText
        from email.mime.application import MIMEApplication
        from email.mime.base import MIMEBase
        from email import encoders

        emailId = request.GET['emailId']
        subject = request.GET['emailSub']
        mailtxt = request.GET['emailTxt']
        reportFilepath = os.path.join(
            BASE_DIR, "static/media/ValidationReport.pdf")
        # The mail addresses and password
        sender_address = 'modvaladm@gmail.com'
        sender_pass = 'modelVal#12'
        # ['n.bawaskar@prescio.com', 'nilesh.bawaskar@prescio.com']
        recipients = emailId.split(",")
        # Setup the MIME
        message = MIMEMultipart()
        message['From'] = sender_address
        message['To'] = ", ".join(recipients)  # receiver_address
        # The subject line
        message['Subject'] = subject

        # The body and the attachments for the mail
        message.attach(MIMEText(mailtxt, 'plain'))

        reportpdf = MIMEApplication(open(reportFilepath, 'rb').read())
        attachment = open(reportFilepath, "rb")

        # instance of MIMEBase and named as p
        filename = "ValidationReport.pdf"
        reportpdf = MIMEBase('application', 'octet-stream')

        # To change the payload into encoded form
        reportpdf.set_payload((attachment).read())

        # encode into base64
        encoders.encode_base64(reportpdf)

        reportpdf.add_header('Content-Disposition',
                             "attachment; filename= %s" % filename)

        # attach the instance 'p' to instance 'msg'
        message.attach(reportpdf)

        # Create SMTP session for sending the mail
        # use gmail with port
        session = smtplib.SMTP('smtp.gmail.com', 587)
        session.starttls()  # enable security
        # login with mail_id and password
        session.login(sender_address, sender_pass)
        text = message.as_string()
        session.sendmail(sender_address, recipients, text)
        session.quit()
        print('Mail Sent')
        data = {'is_taken': True}
    except Exception as e:
        print(e)
        print("Error: unable to send email")
        data = {'is_taken': False}
    return JsonResponse(data)


def saveReplicationComments(request):
    try:
        comments = request.GET['comments']
        DocumentationData = file_path + file_name + "_DocumentationData.csv"
        codeFile = ""
        if os.path.exists(DocumentationData):
            df = pd.read_csv(DocumentationData)
            dffilter = df.query("doc == 'Model Code'")
            resultDocumentation = dffilter["doc_file"].values[0]
            codeFile = resultDocumentation

        UserCommentsFiles = file_path + file_name + "_ReplicationComments.csv"
        if os.path.exists(UserCommentsFiles):
            df_old = pd.read_csv(UserCommentsFiles)
            if (df_old["codeFile"] == codeFile).any():
                df_old.loc[(df_old["codeFile"] == codeFile),
                           "comments"] = comments
                df_old.to_csv(UserCommentsFiles, index=False, encoding='utf-8')
            else:
                data = [[codeFile, comments]]
                df_new = pd.DataFrame(
                    data, columns=['codeFile', 'comments'])
                df = pd.concat([df_old, df_new], axis=0)
                df.to_csv(UserCommentsFiles, index=False, encoding='utf-8')
        else:
            data = [[codeFile, comments]]
            df = pd.DataFrame(
                data, columns=['codeFile', 'comments'])
            df.to_csv(UserCommentsFiles, index=False, encoding='utf-8')
        data = {"is_taken": True}
        return JsonResponse(data)
    except Exception as e:
        print(e)
        print('traceback is  ', traceback.print_exc())


def saveCRComments(request):
    try:
        comments = request.GET['comments']
        section = request.GET['section']
        DocumentationData = file_path + file_name + "_DocumentationData.csv"
        codeFile = ""
        if os.path.exists(DocumentationData):
            df = pd.read_csv(DocumentationData)
            dffilter = df.query("doc == 'Model Code'")
            resultDocumentation = dffilter["doc_file"].values[0]
            codeFile = resultDocumentation

        UserCommentsFiles = file_path + file_name + "_CRComments.csv"
        if os.path.exists(UserCommentsFiles):
            df_old = pd.read_csv(UserCommentsFiles)
            if (df_old["reqId"] == codeFile+"-"+section).any():
                df_old.loc[(df_old["reqId"] == codeFile+"-"+section),
                           "comments"] = comments
                df_old.to_csv(UserCommentsFiles, index=False, encoding='utf-8')
            else:
                data = [[codeFile, comments, section, codeFile+"-"+section]]
                df_new = pd.DataFrame(
                    data, columns=['codeFile', 'comments', 'section', 'reqId'])
                df = pd.concat([df_old, df_new], axis=0)
                df.to_csv(UserCommentsFiles, index=False, encoding='utf-8')
        else:
            data = [[codeFile, comments, section, codeFile+"-"+section]]
            df = pd.DataFrame(
                data, columns=['codeFile', 'comments', 'section', 'reqId'])
            df.to_csv(UserCommentsFiles, index=False, encoding='utf-8')
        data = {"is_taken": True}
        return JsonResponse(data)
    except Exception as e:
        print(e)
        print('traceback is  ', traceback.print_exc())


def getCRComments(request):
    try:
        section = request.GET['section']
        DocumentationData = file_path + file_name + "_DocumentationData.csv"
        codeFile = ""
        comment = ""
        if os.path.exists(DocumentationData):
            df = pd.read_csv(DocumentationData)
            dffilter = df.query("doc == 'Model Code'")
            resultDocumentation = dffilter["doc_file"].values[0]
            codeFile = resultDocumentation

        UserCommentsFiles = file_path + file_name + "_CRComments.csv"
        if os.path.exists(UserCommentsFiles):
            df_old = pd.read_csv(UserCommentsFiles)
            if (df_old["reqId"] == codeFile+"-"+section).any():
                comment = df_old.loc[(
                    df_old["reqId"] == codeFile+"-"+section)]["comments"].values[0]
        print('comment is ', comment)
        data = {"is_taken": True, 'comment': comment}
        return JsonResponse(data)
    except Exception as e:
        print(e)
        print('traceback is  ', traceback.print_exc())


def getAllCRComments(request):
    try:
        DocumentationData = file_path + file_name + "_DocumentationData.csv"
        codeFile = ""
        result = []
        if os.path.exists(DocumentationData):
            df = pd.read_csv(DocumentationData)
            dffilter = df.query("doc == 'Model Code'")
            resultDocumentation = dffilter["doc_file"].values[0]
            codeFile = resultDocumentation
        headersData = [[1, 'Comment'],
                       [2, 'Consistency'],
                       [3, 'Design'],
                       [4, 'Error and Warning Handling'],
                       [5, 'Functionality'],
                       [6, 'Hard Coding'],
                       [7, 'Repeated Codes'],
                       [8, 'Run-time Algorithms'],
                       [9, 'Safe Parallel Programming']]
        dfHeader = pd.DataFrame(
            headersData, columns=['section', 'header'])
        UserCommentsFiles = file_path + file_name + "_CRComments.csv"
        if os.path.exists(UserCommentsFiles):
            df = pd.read_csv(UserCommentsFiles)
            df = df.merge(dfHeader, left_on='section', right_on='section')
            print('df ', df)
            result = df.to_json(orient="records")
            result = json.loads(result)
            del df
        data = {"is_taken": True, 'comment': result}
        return JsonResponse(data)
    except Exception as e:
        print(e)
        print('traceback is  ', traceback.print_exc())


def getCRHelpData(request):
    result = []
    help_File = os.path.join(BASE_DIR, "static/media/CR_Help.csv")
    df = pd.read_csv(help_File, encoding='utf-8')
    result = df.to_json(orient="records")
    result = json.loads(result)
    del df
    return JsonResponse({"is_taken": True, 'CRHelpData': result})


def exportDatatypenCnt(pdf):

    csv_file_name = "csvfile_"+user_name
    savefile_name = file_path + csv_file_name + ".csv"
    df = pd.read_csv(savefile_name, na_values='?')
    # y += 20.0
    x = pdf.get_x()
    y = pdf.get_y()
    pdf.set_xy(x, y)

    pdf.set_font("Arial", size=10)
    # pdf.set_text_color(255, 255, 255)

    pdf.set_xy(20, y)
    pdf.cell(0, 5, "Column Name", 1, fill=True)
    pdf.set_xy(100, y)
    pdf.cell(0, 5, "Not-Null Count", 1)
    pdf.set_xy(130, y)
    pdf.cell(0, 5, "Column Data type", 1)
    # Start from the first cell. Rows and columns are zero indexed.

    result = dict(df.dtypes)
    pdf.set_font("Arial", size=10)
    pdf.set_text_color(0.0, 0.0, 0.0)
    for key, value in result.items():
        # gridDttypes.append(
        #     {'colName': key, 'dataType': value, 'notnull': df[key].count()})
        # print('y is ', y, ' - ', key)
        if(y > 265):
            pdf.add_page()
            y = 10
        y += 5
        pdf.set_xy(20, y)
        pdf.cell(80, 5, key, 1)
        pdf.set_xy(100, y)
        pdf.cell(30, 5, str(df[key].count())+" non-null", 1)
        pdf.set_xy(130, y)
        pdf.cell(0, 5, str(value), 1)

    return pdf


def copyModelData(request):  
    DocumentationData = file_path + file_name + "_ModelInfo.csv"
    bkdir =""
    if os.path.exists(DocumentationData):
        df = pd.read_csv(DocumentationData, encoding='utf-8')        
        Title = df["ModelNm"].values[len(df)-1]  
        del df
        Title=Title.replace(" ","_")
        bkdir = os.path.join(BASE_DIR, 'static/archived/'+Title)
        if not os.path.exists(bkdir):
            os.makedirs(bkdir)
    
    data_dir=["media","csv_files","cnfrmsrc_files","modelCode","replicationFiles","replicationoutput","scenarioFiles","scenarioOutput","scenarioScripts","document_files"]
    for d in data_dir:
        dir = os.path.join(BASE_DIR, 'static/'+d)
        destination = os.path.join(bkdir,d)
        shutil.copytree(dir,destination)
        # for f in os.listdir(dir):
            
        #     source = os.path.join(dir, f)

        #     # Destination path
        #     destination = os.path.join(bkdir,f)
        #     shutil.copyfile(source, destination)
            # os.remove(os.path.join(dir, f))
    deleteModelData(request)
    
    return JsonResponse({"is_taken": True})

def deleteModelData(request):   
    # data_dir=["cnfrmsrc_files"]
    data_dir=["media","csv_files","cnfrmsrc_files","modelCode","replicationFiles","replicationoutput","scenarioFiles","scenarioOutput","scenarioScripts","document_files"]
    for d in data_dir:
        folder = os.path.join(BASE_DIR, 'static/'+d)
        for filename in os.listdir(folder):
            file_path = os.path.join(folder, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
                
            except Exception as e:
                print('Failed to delete %s. Reason: %s' % (file_path, e))
        processing = os.path.join(BASE_DIR, processingFile_path)
        df_old_proc = pd.read_csv(processing)
        df_old_proc["Status"] = "Not done" 
        df_old_proc.to_csv(processing, index=False, encoding='utf-8')
        del df_old_proc
    return JsonResponse({"is_taken": True})

def uploadImage(request):
    if request.method == 'POST':  
        # myfile = request.FILES['myfile']
        filename = request.POST.get('filenm','none')
        print(' filename ',filename)
        files = request.FILES
        myfile = files.get('filename', None)
        print('myfile ',myfile,' filename ',filename)
        if myfile=="":
            res = JsonResponse({'data':'Invalid Request'})
            return res
        else:     
            fs = FileSystemStorage() 
            fs.save(os.path.join(BASE_DIR, plot_dir_view+user_name+'Chartimgs/',filename), myfile)   
            UserChartFile = file_path + "_Chartimg.csv" 
            directory = os.path.join(BASE_DIR, plot_dir_view+user_name+'Chartimgs')
            destination = plot_dir_view+user_name+'Chartimgs/'+filename
            if os.path.exists(UserChartFile):
                df2 = pd.read_csv(UserChartFile) 

                if not os.path.exists(directory):
                    os.makedirs(directory)   
                data = [["External Image", filename, destination,""]]
                dfnew = pd.DataFrame(
                    data, columns=['chartType', 'chartImg', 'destination', 'comments'])
                dfmerged = pd.concat([df2, dfnew], axis=0)
                dfmerged.to_csv(UserChartFile, index=False, encoding='utf-8')
                del dfmerged 
                del df2
            else:  
                data = [["External Image", filename, destination, ""]] 
                dfnew = pd.DataFrame(
                    data, columns=['chartType', 'chartImg', 'destination', 'comments'])
                dfnew.to_csv(UserChartFile, index=False, encoding='utf-8')
                del dfnew  
            res = JsonResponse({"is_taken": True,'data':'Uploaded Successfully','ddlVal':destination,'ddlTxt':'External Image-'+filename})            
            return res    

class MyFPDF(FPDF, HTMLMixin):
    pass
